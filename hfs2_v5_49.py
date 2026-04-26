"""
HfS₂ 박막 산화도 분석기  v5.0  (1단계 개선)
─────────────────────────────────────────────────────
수정사항:
  1. 파일명 파싱 정확도 향상 (Nday_MRH_조건명 완전 지원)
  2. 황색 잔존 비율 수정 (H채널 → 실제 각도 변환 후 비교)
  3. H/S/I 채널 그래프 탭 하나에 통합 + 더블클릭 확대
  4. ROI 복사 + 드래그 이동 지원 (신규 ROI 그리기도 유지)
  6. 라이트 테마 (흰색 계통)
─────────────────────────────────────────────────────
pip install opencv-python pillow matplotlib tkinterdnd2
"""

import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import re as _re
import numpy as np
from PIL import Image, ImageTk, ImageDraw
import cv2, os, csv, io, datetime, glob, json, base64, subprocess, traceback, tempfile, shutil
import matplotlib
# ★ TkAgg 백엔드: FigureCanvasTkAgg와 매칭. "Agg"는 화면 렌더링이 안 되는
#   파일 전용 백엔드라 tkinter 팝업에서 빈 화면으로 나타남.
matplotlib.use("TkAgg")
import matplotlib.pyplot as plt
import matplotlib.font_manager as _fm
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from matplotlib.backends.backend_tkagg import NavigationToolbar2Tk

# ══════════════════════════════════════════════
#  한글 폰트
# ══════════════════════════════════════════════
def _setup_korean_font():
    installed = {f.name for f in _fm.fontManager.ttflist}
    for c in ["Malgun Gothic","NanumGothic","NanumBarunGothic",
              "Apple SD Gothic Neo","AppleGothic",
              "Noto Sans KR","Noto Sans CJK KR","Gulim","Dotum"]:
        if c in installed:
            matplotlib.rc("font", family=c)
            matplotlib.rcParams["axes.unicode_minus"] = False
            return True
    for pat in ["C:/Windows/Fonts/malgun*.ttf",
                "C:/Windows/Fonts/NanumGothic*.ttf",
                "C:/Windows/Fonts/gulim*.tt*",
                "/usr/share/fonts/**/*othi*.tt*"]:
        hits = glob.glob(pat, recursive=True)
        if hits:
            prop = _fm.FontProperties(fname=hits[0])
            matplotlib.rc("font", family=prop.get_name())
            matplotlib.rcParams["axes.unicode_minus"] = False
            return True
    return False

_KO = False   # 기본값: 영어 (한국어로 전환하려면 UI 버튼 사용)
def _L(ko, en): return ko if _KO else en

def set_lang(ko: bool):
    """언어 전환 — True=한국어, False=영어"""
    global _KO
    _KO = ko

try:
    from tkinterdnd2 import TkinterDnD, DND_FILES
    _DND = True
except ImportError:
    _DND = False

# ══════════════════════════════════════════════
#  라이트 테마
# ══════════════════════════════════════════════
BG      = "#f4f6fa"
PANEL   = "#ffffff"
PANEL2  = "#eef0f5"
CARD    = "#ffffff"
CARD2   = "#f8f9fc"
BORDER  = "#d1d5e0"
TXT     = "#1a1d2e"
SUB     = "#6b7280"
ACCENT  = "#3b82f6"
GREEN   = "#16a34a"
AMBER   = "#d97706"
RED     = "#dc2626"
TEAL    = "#0d9488"
PURPLE  = "#7c3aed"
GOLD    = "#b45309"
BTN     = "#e5e7eb"
BTN_H   = "#d1d5db"
SHADOW  = "#c8ccd8"

MF  = ("Segoe UI", 9)
MFB = ("Segoe UI", 9, "bold")
LF  = ("Segoe UI", 8)
TF  = ("Segoe UI", 10, "bold")

COND_COLORS = ["#3b82f6","#16a34a","#d97706","#7c3aed",
               "#dc2626","#0d9488","#b45309","#db2777"]
EXTRA_TARGET_COLORS = ["#0ea5e9", "#f97316", "#ec4899", "#14b8a6"]
TARGET_COLOR_PALETTE = COND_COLORS + EXTRA_TARGET_COLORS
PRED_MAX_TARGETS = 8
_IMG_EXTS = {".png",".jpg",".jpeg",".bmp",".tiff",".tif"}

# ══════════════════════════════════════════════
#  1. 파일명 파싱 (개선)
# ══════════════════════════════════════════════
def parse_filename_tags(name: str) -> tuple[str, str]:
    """
    0day_35RH_NativeHFS2.jpg  → ('0',  'NativeHfS2-35%RH')
    28day_70RH_NativeHFS2.jpg → ('28', 'NativeHfS2-70%RH')
    7days_70RH_PMMA.png       → ('7',  'PMMA-70%RH')
    sample_14d_50RH_Al2O3.jpg → ('14', 'Al2O3-50%RH')
    """
    stem = _re.sub(r'\.(jpg|jpeg|png|bmp|tiff|tif)$', '',
                   name, flags=_re.IGNORECASE)

    # ① day 추출 — Nday / Ndays / Nd (언더스코어/하이픈 경계 인식)
    day = ""
    for pat in [r'(?:^|[_\-])(\d+)\s*days?(?:[_\-]|$)',
                r'(?:^|[_\-])(\d+)[dD](?:[_\-]|$)',
                r'(\d+)\s*days?',
                r'(\d+)[dD]']:
        m = _re.search(pat, stem, _re.IGNORECASE)
        if m:
            day = m.group(1)
            stem = stem[:m.start()] + "_" + stem[m.end():]
            break

    # ② RH 추출
    rh = ""
    m2 = _re.search(r'(\d+)\s*%?\s*[Rr][Hh]', stem)
    if m2:
        rh = m2.group(1)
        stem = stem[:m2.start()] + "_" + stem[m2.end():]

    # ③ 나머지 정리
    cond = _re.sub(r'[_\-\s]+', ' ', stem).strip()
    # 표준화: HFS2 / hfs2 → HfS2
    cond = _re.sub(r'[Hh][Ff][Ss]2', 'HfS2', cond)
    # 앞뒤 공백 제거
    cond = cond.strip()

    if rh:
        cond = f"{cond}-{rh}%RH" if cond else f"{rh}%RH"

    return day, cond


# ══════════════════════════════════════════════
#  2. 색상 분석 (HSI + 황색도 수정)
# ══════════════════════════════════════════════
def rgb_to_hsi(arr: np.ndarray):
    """RGB uint8 → H(0~255), S(0~255), I(0~255)"""
    img = arr.astype(np.float64) / 255.0
    R, G, B = img[:,:,0], img[:,:,1], img[:,:,2]
    I = (R + G + B) / 3.0
    mn = np.minimum(np.minimum(R, G), B)
    S = np.where(I > 1e-8, 1.0 - mn / (I + 1e-10), 0.0)
    num = 0.5 * ((R-G) + (R-B))
    den = np.sqrt((R-G)**2 + (R-B)*(G-B)) + 1e-10
    theta = np.arccos(np.clip(num/den, -1.0, 1.0))
    H = np.where(B <= G, theta, 2*np.pi - theta) / (2*np.pi)
    u8 = lambda a: (a * 255).clip(0, 255).astype(np.uint8)
    return u8(H), u8(S), u8(I)


def compute_yellow_ratio(rgb: np.ndarray, mask: np.ndarray,
                          h_lo_deg: float = 35.0,
                          h_hi_deg: float = 75.0,
                          s_thresh: float = 0.10) -> float:
    """
    황색 잔존 비율
    - H 범위: 35~75° (연한 황색까지 포괄)
    - S 임계값: 0.10 (사진 조명 변동 대응, 흰색/베이지는 S≈0.02~0.04로 자동 제외)
    - 마스크 내 픽셀 중 황색 비율 반환 (0~1)
    """
    m = mask.astype(bool)
    if m.sum() == 0:
        return 0.0

    img = rgb.astype(np.float64) / 255.0
    R, G, B = img[:,:,0][m], img[:,:,1][m], img[:,:,2][m]
    I = (R + G + B) / 3.0 + 1e-10
    mn = np.minimum(np.minimum(R, G), B)
    S_n = 1.0 - mn / I                          # 0~1

    num = 0.5 * ((R-G) + (R-B))
    den = np.sqrt((R-G)**2 + (R-B)*(G-B)) + 1e-10
    theta = np.arccos(np.clip(num/den, -1.0, 1.0))
    H_deg = np.where(B <= G, theta, 2*np.pi - theta) * (180.0 / np.pi)  # 0~360°

    yellow = ((H_deg >= h_lo_deg) & (H_deg <= h_hi_deg) & (S_n >= s_thresh))
    return float(yellow.sum()) / m.sum()


def compute_yellowness_index(rgb: np.ndarray, mask: np.ndarray) -> float:
    """
    YI (Yellowness Index, ASTM E313)
    YI = 100 * (1.28R - 1.06B) / G
    미산화 시편: YI ≈ 50~110 / 산화 시편: YI ≈ 20~35
    ROI 마스크 내 평균값 반환
    """
    m = mask.astype(bool)
    if m.sum() == 0:
        return 0.0
    img = rgb.astype(np.float64) / 255.0
    R = img[:,:,0][m]
    G = img[:,:,1][m]
    B = img[:,:,2][m]
    G_safe = np.where(G > 0.01, G, 0.01)
    YI = 100.0 * (1.28*R - 1.06*B) / G_safe
    return float(np.mean(YI))


def compute_s_mean(rgb: np.ndarray, mask: np.ndarray) -> float:
    """마스크 내 채도(S) 평균 — 0~255"""
    _, S, _ = rgb_to_hsi(rgb)
    vals = S[mask.astype(bool)]
    return float(np.mean(vals)) if len(vals) > 0 else np.nan


# ══════════════════════════════════════════════
#  5. 추가 컬러모델 + 영상처리 지표
# ══════════════════════════════════════════════

def compute_lab_metrics(rgb: np.ndarray, mask: np.ndarray) -> dict:
    """
    CIE Lab 색공간 기반 지표
    - L*  : 밝기 (산화 시 증가 → 흰색으로)
    - a*  : 적-녹 축 (HfS2 산화에서는 크게 변하지 않음)
    - b*  : 황-청 축 (미산화: +50~60, 산화: +4~15) ★ 핵심 지표
    반환: {'L': float, 'a': float, 'b': float}
    """
    m = mask.astype(bool)
    if m.sum() == 0:
        return {"L": np.nan, "a": np.nan, "b": np.nan}
    bgr  = cv2.cvtColor(rgb, cv2.COLOR_RGB2BGR)
    lab  = cv2.cvtColor(bgr, cv2.COLOR_BGR2Lab).astype(np.float32)
    # OpenCV Lab 범위: L(0~255), a/b(0~255, 중심=128)
    # 실제 L*(0~100), a*b*(-128~127) 로 변환
    L_real = lab[:,:,0][m] / 255.0 * 100.0
    a_real = lab[:,:,1][m] - 128.0
    b_real = lab[:,:,2][m] - 128.0
    return {
        "L": float(np.mean(L_real)),
        "a": float(np.mean(a_real)),
        "b": float(np.mean(b_real)),
    }


def compute_delta_e(lab_curr: dict, lab_ref: dict) -> float:
    """
    ΔE (CIE76 색차) — 현재 이미지와 기준(0일차) 이미지 간 색 변화량
    ΔE = √((ΔL*)² + (Δa*)² + (Δb*)²)
    산화 진행 시 ΔE 증가 (0=변화없음, 10+=뚜렷한 변화)
    """
    if any(np.isnan(v) for v in [lab_curr["L"], lab_ref["L"]]):
        return np.nan
    dL = lab_curr["L"] - lab_ref["L"]
    da = lab_curr["a"] - lab_ref["a"]
    db = lab_curr["b"] - lab_ref["b"]
    return float(np.sqrt(dL**2 + da**2 + db**2))


def compute_glcm_metrics(rgb: np.ndarray, mask: np.ndarray) -> dict:
    """
    GLCM (Gray-Level Co-occurrence Matrix) 텍스처 지표
    - contrast   : 대비 (산화 → 표면 균일 → 감소)
    - energy     : 에너지/균일도 (산화 → 균일해짐 → 증가)
    - homogeneity: 균질성 (산화 → 증가)
    - correlation: 상관관계
    반환: {'contrast':float, 'energy':float,
           'homogeneity':float, 'correlation':float}
    """
    try:
        from skimage.feature import graycomatrix, graycoprops
    except ImportError:
        return {k: np.nan for k in
                ("contrast","energy","homogeneity","correlation")}

    m = mask.astype(bool)
    if m.sum() < 100:
        return {k: np.nan for k in
                ("contrast","energy","homogeneity","correlation")}

    # ROI 영역만 크롭 후 그레이스케일
    rows, cols = np.where(m)
    r0,r1 = rows.min(), rows.max()+1
    c0,c1 = cols.min(), cols.max()+1
    crop = rgb[r0:r1, c0:c1]
    gray = cv2.cvtColor(crop, cv2.COLOR_RGB2GRAY)

    # 레벨 줄여서 계산 속도 향상 (256→64)
    gray_q = (gray // 4).astype(np.uint8)
    gcm = graycomatrix(gray_q, distances=[1],
                       angles=[0, np.pi/4, np.pi/2, 3*np.pi/4],
                       levels=64, symmetric=True, normed=True)
    return {
        "contrast":    float(graycoprops(gcm,"contrast").mean()),
        "energy":      float(graycoprops(gcm,"energy").mean()),
        "homogeneity": float(graycoprops(gcm,"homogeneity").mean()),
        "correlation": float(graycoprops(gcm,"correlation").mean()),
    }


def compute_all_metrics(rgb: np.ndarray, mask: np.ndarray,
                        ref_lab: dict | None = None) -> dict:
    """
    모든 지표를 한 번에 계산해서 dict 반환
    ref_lab: 0일차 이미지의 lab_metrics (ΔE 계산용, None이면 skip)
    """
    lab  = compute_lab_metrics(rgb, mask)
    glcm = compute_glcm_metrics(rgb, mask)
    dE   = compute_delta_e(lab, ref_lab) if ref_lab else np.nan
    return {
        # Lab
        "lab_L":  lab["L"],
        "lab_a":  lab["a"],
        "lab_b":  lab["b"],   # ★ 핵심: 황색도
        "delta_e": dE,
        # GLCM
        "glcm_contrast":    glcm["contrast"],
        "glcm_energy":      glcm["energy"],
        "glcm_homogeneity": glcm["homogeneity"],
        "glcm_correlation": glcm["correlation"],
    }


# ══════════════════════════════════════════════
#  세그먼트 통계
# ══════════════════════════════════════════════
def seg_stats(ch: np.ndarray, mask: np.ndarray,
              roi: tuple, rows: int, cols: int,
              min_pix: int = 10) -> list:
    x0, y0, x1, y1 = roi
    sh = (y1-y0) / rows
    sw = (x1-x0) / cols
    out = []
    for r in range(rows):
        for c in range(cols):
            ry0 = y0 + int(r*sh)
            ry1 = y0 + int((r+1)*sh) if r < rows-1 else y1
            rx0 = x0 + int(c*sw)
            rx1 = x0 + int((c+1)*sw) if c < cols-1 else x1
            cm   = mask[ry0:ry1, rx0:rx1]
            vals = ch[ry0:ry1, rx0:rx1][cm]
            base = {"seg":r*cols+c, "row":r, "col":c,
                    "ry0":ry0, "ry1":ry1, "rx0":rx0, "rx1":rx1}
            if len(vals) < min_pix:
                out.append({**base, "mode":np.nan, "mean":np.nan,
                             "std":np.nan, "pixels":0})
            else:
                out.append({**base,
                    "mode":  float(np.bincount(vals.astype(np.int32)).argmax()),
                    "mean":  float(np.mean(vals)),
                    "std":   float(np.std(vals)),
                    "pixels":int(cm.sum())})
    return out


def roi_to_mask(shape, roi):
    h, w = shape[:2]
    x0, y0, x1, y1 = roi
    m = np.zeros((h, w), bool)
    m[y0:y1, x0:x1] = True
    return m


# ══════════════════════════════════════════════
#  자동 ROI 추정 — HfS₂ 시편 vs 흰 종이 분리
# ══════════════════════════════════════════════
def _find_corner_far_from_curved_side(contour, curved_mid):
    """곡률 큰 변의 중심점 (curved_mid) 에서 가장 먼 polygon vertex 찾기.

    approxPolyDP 로 시편 다각형 근사 후 각 vertex 까지의 유클리드 거리 측정.
    가장 먼 vertex 반환.

    Returns: (vx, vy) or None
    """
    if curved_mid is None:
        return None
    try:
        peri = cv2.arcLength(contour, True)
        approx = None
        # 적당한 단순함의 polygon — 너무 디테일하면 실제 corner 가 아닌
        # 곡선 위 점이 잡힐 수 있음
        for eps_factor in (0.03, 0.02, 0.04, 0.01):
            cand = cv2.approxPolyDP(contour, eps_factor * peri, True)
            if 3 <= len(cand) <= 12:
                approx = cand
                break
        if approx is None:
            return None
        pts = approx.reshape(-1, 2).astype(float)
        cmx, cmy = curved_mid
        max_d = -1.0
        farthest = None
        for p in pts:
            d = (p[0] - cmx) ** 2 + (p[1] - cmy) ** 2
            if d > max_d:
                max_d = d
                farthest = (float(p[0]), float(p[1]))
        return farthest
    except Exception:
        pass
    return None


def _find_most_curved_side_midpoint(contour, sx, sy, sbw, sbh):
    """시편 contour 에서 곡률이 가장 큰 (가장 원형에 가까운) 쪽의 중점.

    bbox 중심에서 4 사분면(상/하/좌/우) 으로 contour 점들을 분류한 뒤,
    각 사분면 점 분포에 PCA 적용 → 작은 eigenvalue / 큰 eigenvalue 비율
    이 클수록 원형에 가까움 (0 = 완벽 직선, 1 = 원).

    Returns: (cx, cy) or None — 곡률 큰 사이드의 중점
    """
    try:
        pts = contour.reshape(-1, 2).astype(float)
        if len(pts) < 8:
            return None
        cx_bb = sx + sbw / 2.0
        cy_bb = sy + sbh / 2.0

        sides = {"top": [], "bottom": [], "left": [], "right": []}
        for p in pts:
            dx = p[0] - cx_bb
            dy = p[1] - cy_bb
            if abs(dx) > abs(dy):
                sides["right" if dx > 0 else "left"].append(p)
            else:
                sides["bottom" if dy > 0 else "top"].append(p)

        worst_score = -1.0  # 곡률 가장 큰 점수 (작은/큰 eigvalue 비율 가장 큼)
        worst_mid = None
        total = len(pts)
        for name, side_pts in sides.items():
            if len(side_pts) < 5:
                continue
            arr = np.array(side_pts)
            mean = arr.mean(axis=0)
            centered = arr - mean
            cov = (centered.T @ centered) / max(1, len(centered) - 1)
            try:
                eigvals = np.linalg.eigvalsh(cov)
            except Exception:
                continue
            eigvals = np.sort(eigvals)
            if eigvals[1] < 1.0:
                continue
            curvature = float(eigvals[0] / eigvals[1])  # 0 ≈ 직선, 1 ≈ 원형
            # 사이드 점 수가 너무 적으면 신뢰도 낮음 → 패널티 (점수 깎음)
            if len(side_pts) < total * 0.10:
                curvature -= 0.10
            if curvature > worst_score:
                worst_score = curvature
                worst_mid = (float(mean[0]), float(mean[1]))
        # 충분한 곡률 있을 때만 인정 (직선이면 무의미)
        if worst_mid is not None and worst_score > 0.04:
            return worst_mid
    except Exception:
        pass
    return None


# cond 별 학습된 ROI 면적 비율 (이미지 면적 대비).
# 사용자 수동 정답(roi_mod.db, 33장) cond별 평균값:
#   NativeHfS2-35%RH n=5  mean=0.128  median=0.129
#   NativeHfS2-70%RH n=11 mean=0.150  median=0.160
#   Al2O3HfS2-70%RH  n=12 mean=0.172  median=0.177
#   PMMA HfS2-70%RH  n=5  mean=0.151  median=0.136
# 여기 없는 cond 는 DEFAULT_AREA_RATIO 사용.
COND_AREA_RATIOS = {
    "NativeHfS2-35%RH": 0.13,
    "NativeHfS2-70%RH": 0.15,
    "Al2O3HfS2-70%RH": 0.17,
    "PMMA HfS2-70%RH": 0.15,
}
DEFAULT_AREA_RATIO = 0.15


def _resolve_area_ratio(cond) -> float:
    if not cond:
        return DEFAULT_AREA_RATIO
    return COND_AREA_RATIOS.get(str(cond).strip(), DEFAULT_AREA_RATIO)


def auto_detect_roi(rgb: np.ndarray,
                    paper_v_thresh: int = 215,
                    paper_s_thresh: int = 25,
                    target_area_ratio: float = None,
                    max_specimen_fraction: float = 0.70,
                    edge_margin_ratio: float = 0.05,
                    bias_ratio: float = 0.0,
                    min_area_ratio: float = 0.03,
                    paper_inside_ratio: float = 0.02,
                    cond=None) -> tuple:
    """
    HfS₂ 시편 사진에서 자동 ROI 추정.

    사용자 수동 정답(roi_mod.db) 분석 기반:
      - ROI 면적 ≈ cond 별 학습 비율 (기본 22%, COND_AREA_RATIOS)
      - ROI 중심 = 시편 mass-centroid (사용자 의도: "시편 중심부 안전영역")
      - 가로/세로 비 = (시편 bbox AR) ** 0.5  — 정사각형 쪽으로 완화
      - 이미지 가장자리 5% 안쪽 보장
      - paper mask: 이미지별 V/S 적응 임계 + convex hull 로 광택 반사 흡수

    알고리즘:
    1) HSV V/S 적응 임계 → paper 마스크
    2) non_paper 형태학 정리 (close → open) + 가장 큰 contour
    3) convex hull 로 시편 bbox 산출 (광택 반사 hole 회복)
    4) ROI 면적 = img_area × _resolve_area_ratio(cond), AR = sqrt(sbw/sbh)
    5) ROI 중심 = mass-centroid (cv2.moments). pointPolygonTest 로 내부 보장.
    6) edge_margin_ratio (5%) 안쪽으로 밀어넣기
    7) 품질 평가: 면적 / 가장자리 근접 / paper 비율

    Parameters
    ----------
    target_area_ratio : float | None
        None 이면 cond 기반 자동. 명시값이 있으면 그 값을 강제.
    bias_ratio : float
        과거 호환용 인자. 0.0 (기본) 이면 mass-centroid 그대로 사용.
    cond : str | None
        cond 문자열. COND_AREA_RATIOS 룩업에 사용.

    Returns
    -------
    roi    : (x0, y0, x1, y1)
    flag   : 'good' | 'warn_small' | 'warn_off' | 'warn_paper' | 'failed'
    reason : 한 줄 한국어 설명
    """
    h, w = rgb.shape[:2]
    img_area = h * w
    default_roi = (w // 4, h // 4, 3 * w // 4, 3 * h // 4)

    if target_area_ratio is None:
        target_area_ratio = _resolve_area_ratio(cond)

    try:
        hsv = cv2.cvtColor(rgb, cv2.COLOR_RGB2HSV)
        S = hsv[:, :, 1]
        V = hsv[:, :, 2]
        # 적응 임계 (보수): V 92퍼센타일이 기본보다 *조금만* 높으면 채택,
        # 너무 크면 (시편을 paper 로 흡수) 거부. 70%RH 광택 케이스 보호용.
        v_hi = paper_v_thresh
        v_p92 = int(np.percentile(V, 92))
        if paper_v_thresh < v_p92 < paper_v_thresh + 20:
            v_hi = v_p92
        paper = (V > v_hi) & (S < paper_s_thresh)
        non_paper = (~paper).astype(np.uint8) * 255

        # 형태학 정리 — close 를 먼저 해 시편 내부 hole 흡수
        k_clean = max(7, min(h, w) // 80)
        if k_clean % 2 == 0:
            k_clean += 1
        kern = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (k_clean, k_clean))
        non_paper = cv2.morphologyEx(non_paper, cv2.MORPH_CLOSE, kern)
        non_paper = cv2.morphologyEx(non_paper, cv2.MORPH_OPEN, kern)

        contours, _ = cv2.findContours(
            non_paper, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        if not contours:
            return default_roi, "failed", "시편 영역을 찾지 못함 — 수동 설정 필요"
        c = max(contours, key=cv2.contourArea)

        # convex hull — 시편 내부 hole(광택 반사) 흡수.
        # 단, hull 이 이미지의 90% 이상이면 거부 (시편 검출 실패의 fallback 으로 보호)
        hull = cv2.convexHull(c)
        if cv2.contourArea(hull) < 0.9 * img_area:
            c = hull

        sx, sy, sbw, sbh = cv2.boundingRect(c)
        if sbw < 10 or sbh < 10:
            return default_roi, "failed", "시편 영역이 너무 작음"

        # 목표 ROI 크기 — 면적 비율 + 시편 모양 AR^0.5 완화
        target_area = img_area * target_area_ratio
        raw_aspect = sbw / sbh if sbh > 0 else 1.0
        aspect = raw_aspect ** 0.5  # 정사각형 쪽으로 완화 (사용자 수동 평균 AR≈1.07)
        roi_w = int((target_area * aspect) ** 0.5)
        roi_h = int((target_area / aspect) ** 0.5)
        roi_w = min(roi_w, int(sbw * max_specimen_fraction))
        roi_h = min(roi_h, int(sbh * max_specimen_fraction))
        roi_w = max(roi_w, int(min(w, sbw) * 0.18))
        roi_h = max(roi_h, int(min(h, sbh) * 0.18))

        # ROI 중심 — 시편 bbox 중심과 mass-centroid 의 평균
        # (사용자 수동 패턴 분석: 순수 bbox 중심보다 약간 mass 쪽,
        #  순수 mass-centroid 보다 약간 bbox 쪽이 가장 잘 맞음.
        #  비대칭 시편에서도 두 점 평균이 안전영역 중심에 가까움)
        cx_base = sx + sbw / 2
        cy_base = sy + sbh / 2
        M = cv2.moments(c)
        if M["m00"] > 0:
            cx_mass = M["m10"] / M["m00"]
            cy_mass = M["m01"] / M["m00"]
        else:
            cx_mass, cy_mass = cx_base, cy_base
        # contour 외부에 떨어지는 thin-shell 케이스 가드
        if cv2.pointPolygonTest(c, (float(cx_mass), float(cy_mass)), False) < 0:
            cx_mass, cy_mass = cx_base, cy_base
        cx = (cx_base + cx_mass) / 2
        cy = (cy_base + cy_mass) / 2
        # bias_ratio > 0 이면 과거 호환 모드: 곡률 큰 변에서 먼 꼭지점 쪽으로 추가 이동
        if bias_ratio > 0:
            curved_pt = _find_most_curved_side_midpoint(c, sx, sy, sbw, sbh)
            far_vertex = _find_corner_far_from_curved_side(c, curved_pt)
            if far_vertex is not None:
                cx = cx + (far_vertex[0] - cx) * bias_ratio
                cy = cy + (far_vertex[1] - cy) * bias_ratio

        # 가장자리 안전 margin — 안쪽으로 밀어넣기
        margin_x = int(w * edge_margin_ratio)
        margin_y = int(h * edge_margin_ratio)
        x0 = int(cx - roi_w / 2)
        y0 = int(cy - roi_h / 2)
        x1 = x0 + roi_w
        y1 = y0 + roi_h
        if x0 < margin_x:
            shift = margin_x - x0
            x0 += shift; x1 += shift
        if x1 > w - margin_x:
            shift = x1 - (w - margin_x)
            x0 -= shift; x1 -= shift
        if y0 < margin_y:
            shift = margin_y - y0
            y0 += shift; y1 += shift
        if y1 > h - margin_y:
            shift = y1 - (h - margin_y)
            y0 -= shift; y1 -= shift
        x0 = max(0, x0); y0 = max(0, y0)
        x1 = min(w, x1); y1 = min(h, y1)
        if x1 <= x0 + 1 or y1 <= y0 + 1:
            return default_roi, "failed", "ROI 너비/높이 0"
        roi = (int(x0), int(y0), int(x1), int(y1))

        # 품질 평가
        roi_area = (x1 - x0) * (y1 - y0)
        area_ratio = roi_area / img_area
        if area_ratio < min_area_ratio:
            return roi, "warn_small", f"ROI 면적 작음 ({area_ratio*100:.0f}%)"
        # 가장자리 너무 가까운지 (margin 의 절반 이내) — 시편이 가장자리에 있으면 위험
        edge_warn = (x0 < margin_x * 0.5 or y0 < margin_y * 0.5 or
                     x1 > w - margin_x * 0.5 or y1 > h - margin_y * 0.5)
        if edge_warn:
            return roi, "warn_off", "ROI 가 이미지 가장자리에 가까움 — 시편 위치 확인 필요"
        roi_paper = paper[y0:y1, x0:x1]
        paper_in_roi = float(roi_paper.mean()) if roi_paper.size > 0 else 0.0
        if paper_in_roi > paper_inside_ratio:
            return roi, "warn_paper", f"ROI 안 흰 배경 {paper_in_roi*100:.1f}% 포함"
        return roi, "good", "정상 자동 추정"
    except Exception as e:
        return default_roi, "failed", f"오류: {e}"


def evaluate_roi_quality(rgb: np.ndarray, roi: tuple,
                         paper_v_thresh: int = 215,
                         paper_s_thresh: int = 25,
                         min_area_ratio: float = 0.05,
                         edge_margin_ratio: float = 0.025,
                         paper_inside_ratio: float = 0.02) -> tuple:
    """주어진 ROI 에 대해 품질만 평가 (자동 추정 X — DB 로드 등에 사용).
    Returns (flag, reason) — flag ∈ good / warn_small / warn_off / warn_paper / failed
    """
    if rgb is None or roi is None:
        return "failed", "이미지 또는 ROI 없음"
    h, w = rgb.shape[:2]
    x0, y0, x1, y1 = roi
    if x1 <= x0 or y1 <= y0:
        return "failed", "ROI 너비/높이 0"
    img_area = h * w
    roi_area = (x1 - x0) * (y1 - y0)
    area_ratio = roi_area / img_area
    if area_ratio < min_area_ratio:
        return "warn_small", f"ROI 면적 작음 ({area_ratio*100:.0f}%)"
    # 가장자리 근접 검사 — ROI 가 이미지 끝에 닿으면 잘못 설정
    margin_x = int(w * edge_margin_ratio)
    margin_y = int(h * edge_margin_ratio)
    if x0 < margin_x or y0 < margin_y or x1 > w - margin_x or y1 > h - margin_y:
        return "warn_off", "ROI 가 이미지 가장자리에 가까움 — 시편 위치 확인 필요"
    try:
        hsv = cv2.cvtColor(rgb, cv2.COLOR_RGB2HSV)
        paper = (hsv[:, :, 2] > paper_v_thresh) & (hsv[:, :, 1] < paper_s_thresh)
        sub = paper[y0:y1, x0:x1]
        paper_in_roi = float(sub.mean()) if sub.size > 0 else 0.0
        if paper_in_roi > paper_inside_ratio:
            return "warn_paper", f"ROI 안 흰 배경 {paper_in_roi*100:.0f}% 포함"
    except Exception as e:
        return "failed", f"품질 평가 오류: {e}"
    return "good", "정상"


def _roi_to_ratio(roi: tuple, w: int, h: int) -> tuple:
    return (roi[0] / w, roi[1] / h, roi[2] / w, roi[3] / h)


def _ratio_iou(a: tuple, b: tuple) -> float:
    """두 비율 ROI 의 IoU."""
    ax0, ay0, ax1, ay1 = a
    bx0, by0, bx1, by1 = b
    ix0 = max(ax0, bx0); iy0 = max(ay0, by0)
    ix1 = min(ax1, bx1); iy1 = min(ay1, by1)
    iw = max(0.0, ix1 - ix0); ih = max(0.0, iy1 - iy0)
    inter = iw * ih
    aa = max(0.0, ax1 - ax0) * max(0.0, ay1 - ay0)
    bb = max(0.0, bx1 - bx0) * max(0.0, by1 - by0)
    uni = aa + bb - inter
    return inter / uni if uni > 0 else 0.0


def check_roi_group_consistency(images: list, iou_thresh: float = 0.70,
                                snap_to_median: bool = False) -> dict:
    """같은 cond 그룹 내 ROI 비율 일관성 점검.

    각 그룹의 ROI 비율 중앙값과 IoU < iou_thresh 인 이미지를 outlier 로 마킹.

    snap_to_median=True 이면 outlier 의 ROI 를 그룹 중앙값 비율로 덮어씀
    (`roi_source` 가 'db' 또는 'manual' 인 이미지는 보호).

    Returns dict {idx: (is_outlier: bool, group_iou: float)}
    """
    out = {}
    # cond 별 그룹화
    groups = {}
    for i, img in enumerate(images):
        cond = (img.get("cond") or "").strip()
        if not cond or img.get("roi") is None or img.get("rgb") is None:
            out[i] = (False, 1.0)
            continue
        groups.setdefault(cond, []).append(i)
    for cond, idxs in groups.items():
        if len(idxs) < 2:
            for i in idxs:
                out[i] = (False, 1.0)
            continue
        ratios = []
        for i in idxs:
            img = images[i]
            h, w = img["rgb"].shape[:2]
            ratios.append(_roi_to_ratio(img["roi"], w, h))
        # 중앙값
        med = tuple(float(np.median([r[k] for r in ratios])) for k in range(4))
        for i, r in zip(idxs, ratios):
            iou = _ratio_iou(r, med)
            is_out = iou < iou_thresh
            out[i] = (is_out, iou)
            if (snap_to_median and is_out and
                    images[i].get("roi_source") not in ("db", "manual")):
                h, w = images[i]["rgb"].shape[:2]
                images[i]["roi"] = (int(med[0] * w), int(med[1] * h),
                                    int(med[2] * w), int(med[3] * h))
                images[i]["roi_source"] = "auto_snapped"
    return out


def _border_color_for_roi(flag, has_roi: bool, defaults: dict,
                          inconsistent: bool = False) -> str:
    """카드 테두리 색상. defaults = {'green':..,'amber':..,'red':..,'purple':..,'border':..}"""
    # 우선순위: failed > warn > inconsistent > good > none
    if flag == "failed":
        return defaults["red"]
    if flag in ("warn_small", "warn_off", "warn_paper"):
        return defaults["amber"]
    if inconsistent and has_roi:
        return defaults["purple"]
    if flag in ("manual", "good"):
        return defaults["green"]
    return defaults["green"] if has_roi else defaults["border"]


_ROI_FLAG_LABEL = {
    "manual":     ("✓", "사용자 확인", "Manual"),
    "good":       ("✓", "자동 OK",     "Auto OK"),
    "warn_small": ("⚠", "면적 작음",   "Too small"),
    "warn_off":   ("⚠", "끝쪽 설정",   "Near edge"),
    "warn_paper": ("⚠", "배경 포함",   "Includes paper"),
    "failed":     ("✗", "추정 실패",   "Detection failed"),
}


# ══════════════════════════════════════════════
#  고급 분석 백엔드 — Histogram + FFT + Ensemble
# ══════════════════════════════════════════════

def adv_hist_signature(rgb: np.ndarray, mask: np.ndarray,
                       bins: int = 64) -> np.ndarray:
    """
    마스크 내 b* 채널 히스토그램 → 정규화된 시그니처 벡터 반환.
    b* 범위 -30 ~ +80 을 bins개 구간으로 나눔.
    Wasserstein 거리 계산에 사용.
    """
    m = mask.astype(bool)
    if m.sum() == 0:
        return np.ones(bins) / bins          # fallback: uniform

    bgr = cv2.cvtColor(rgb, cv2.COLOR_RGB2BGR)
    lab = cv2.cvtColor(bgr, cv2.COLOR_BGR2Lab).astype(np.float32)
    b_star = (lab[:, :, 2][m] - 128.0)      # OpenCV b: 0~255 → -128~127

    hist, _ = np.histogram(b_star, bins=bins, range=(-30.0, 80.0))
    total = hist.sum()
    if total == 0:
        return np.ones(bins) / bins
    return hist.astype(np.float64) / total   # 확률 분포


def adv_wasserstein_dist(p: np.ndarray, q: np.ndarray) -> float:
    """
    1D Wasserstein (Earth Mover's) 거리.
    두 히스토그램 p, q (같은 bins 수의 확률 분포).
    scipy 없는 환경 대비: CDF 차이의 L1-norm으로 직접 계산.
    """
    # CDF 차이의 절댓값 합 = W1 거리 (같은 support 가정)
    cdf_p = np.cumsum(p)
    cdf_q = np.cumsum(q)
    return float(np.sum(np.abs(cdf_p - cdf_q))) / len(p)


def adv_wasserstein_estimate(target_hist: np.ndarray,
                              pool: list,
                              bins: int = 64) -> dict:
    """
    참조 DB pool의 각 이미지와 Wasserstein 거리 계산 →
    거리 역수 가중 평균으로 day 추정.
    반환: {est_day, confidence, scores [(dist, img), ...]}
    """
    scores = []
    for img in pool:
        ref_hist = img.get("_adv_hist")
        if ref_hist is None:
            continue
        d = adv_wasserstein_dist(target_hist, ref_hist)
        scores.append((d, img))

    if not scores:
        return {"est_day": None, "confidence": 0, "scores": []}

    scores.sort(key=lambda x: x[0])
    top3 = scores[:3]

    day_weights = []
    for dist, img in top3:
        try:
            day_weights.append((float(img["day"]), 1.0 / (dist + 1e-6)))
        except Exception:
            pass

    if not day_weights:
        return {"est_day": None, "confidence": 0, "scores": scores}

    total_w = sum(w for _, w in day_weights)
    est_day = sum(d * w for d, w in day_weights) / total_w
    conf    = max(0.0, min(100.0, 100.0 - scores[0][0] * 500))
    return {"est_day": est_day, "confidence": conf, "scores": scores}


def adv_fft_features(rgb: np.ndarray, mask: np.ndarray) -> dict:
    """
    ROI 그레이스케일 이미지 → 2D FFT 주파수 도메인 분석.
    반환:
      power_map   : 2D log-power 스펙트럼 (시각화용, shape H×W)
      hf_ratio    : 고주파 에너지 비율  (0~1, 산화 → 표면 거칠어짐 → 증가)
      radial_mean : 방사형 평균 스펙트럼 (1D array, bins=64)
      entropy     : 스펙트럼 엔트로피 (분포 불균일도)
    """
    m = mask.astype(bool)
    gray = cv2.cvtColor(rgb, cv2.COLOR_RGB2GRAY).astype(np.float32)

    # 마스크 영역만 사용 (나머지는 평균값으로 채워 경계 artifact 최소화)
    mean_val = float(gray[m].mean()) if m.sum() > 0 else 128.0
    gray_m = np.full_like(gray, mean_val)
    gray_m[m] = gray[m]

    # 2D FFT + 중심 이동
    F  = np.fft.fft2(gray_m)
    Fs = np.fft.fftshift(F)
    power = np.abs(Fs) ** 2
    log_power = np.log1p(power)

    H, W = power.shape
    cy, cx = H // 2, W // 2

    # 방사형 거리 맵
    yy, xx = np.ogrid[:H, :W]
    r_map = np.sqrt((yy - cy) ** 2 + (xx - cx) ** 2)
    max_r = np.sqrt(cx ** 2 + cy ** 2)

    # 고주파 임계: 전체 반경의 40% 이상을 고주파로 정의
    hf_thresh = max_r * 0.40
    dc_thresh = max_r * 0.02          # DC 성분 제외
    mask_hf = (r_map > hf_thresh)
    mask_all = (r_map > dc_thresh)    # DC 제외 전체

    total_e = power[mask_all].sum()
    hf_e    = power[mask_hf].sum()
    hf_ratio = float(hf_e / (total_e + 1e-12))

    # 방사형 평균 스펙트럼 (64개 구간)
    r_bins  = 64
    r_edges = np.linspace(0, max_r, r_bins + 1)
    radial  = np.zeros(r_bins)
    for i in range(r_bins):
        ring = (r_map >= r_edges[i]) & (r_map < r_edges[i + 1])
        if ring.sum() > 0:
            radial[i] = float(power[ring].mean())
    radial_norm = radial / (radial.sum() + 1e-12)

    # 스펙트럼 엔트로피
    eps = 1e-12
    entropy = float(-np.sum(radial_norm * np.log(radial_norm + eps)))

    return {
        "power_map":   log_power,
        "hf_ratio":    hf_ratio,
        "radial_mean": radial_norm,
        "entropy":     entropy,
    }


def adv_fft_estimate(target_fft: dict,
                     pool: list) -> dict:
    """
    참조 DB의 FFT 피처와 타겟의 FFT 피처를 비교 → day 추정.
    거리 = 0.5*|hf_ratio 차| + 0.3*|entropy 차| + 0.2*radial cosine 거리.
    반환: {est_day, confidence, scores [(dist, img), ...]}
    """
    t_hf   = target_fft["hf_ratio"]
    t_ent  = target_fft["entropy"]
    t_rad  = target_fft["radial_mean"]

    # 정규화 기준 (pool 전체 범위)
    hf_vals  = [img["_adv_fft"]["hf_ratio"]  for img in pool
                if img.get("_adv_fft")]
    ent_vals = [img["_adv_fft"]["entropy"]    for img in pool
                if img.get("_adv_fft")]
    if not hf_vals:
        return {"est_day": None, "confidence": 0, "scores": []}

    hf_r   = max(max(hf_vals)  - min(hf_vals),  1e-6)
    ent_r  = max(max(ent_vals) - min(ent_vals), 1e-6)

    scores = []
    for img in pool:
        ff = img.get("_adv_fft")
        if ff is None:
            continue
        d_hf  = abs(t_hf  - ff["hf_ratio"])  / hf_r
        d_ent = abs(t_ent - ff["entropy"])    / ent_r
        # 방사형 코사인 거리
        dot   = np.dot(t_rad, ff["radial_mean"])
        n1    = np.linalg.norm(t_rad)
        n2    = np.linalg.norm(ff["radial_mean"])
        cos_d = 1.0 - dot / (n1 * n2 + 1e-12)

        dist = 0.5 * d_hf + 0.3 * d_ent + 0.2 * cos_d
        scores.append((dist, img))

    if not scores:
        return {"est_day": None, "confidence": 0, "scores": []}

    scores.sort(key=lambda x: x[0])
    top3 = scores[:3]

    day_weights = []
    for dist, img in top3:
        try:
            day_weights.append((float(img["day"]), 1.0 / (dist + 1e-6)))
        except Exception:
            pass

    if not day_weights:
        return {"est_day": None, "confidence": 0, "scores": scores}

    total_w = sum(w for _, w in day_weights)
    est_day = sum(d * w for d, w in day_weights) / total_w
    conf    = max(0.0, min(100.0, 100.0 - scores[0][0] * 300))
    return {"est_day": est_day, "confidence": conf, "scores": scores}


def adv_ensemble(knn_day: float | None, knn_conf: float,
                 wass_day: float | None, wass_conf: float,
                 fft_day: float | None,  fft_conf: float,
                 spatial_day: float | None = None, spatial_conf: float = 0,
                 kinetic_day: float | None = None, kinetic_conf: float = 0) -> dict:
    """
    최대 5개 추정치를 신뢰도 가중 평균으로 앙상블.
    신뢰도 20% 미만 추정치는 제외.
    반환: {est_day, confidence, weights: {knn, wass, fft, spatial, kinetic}}
    """
    candidates = []
    for name, day, conf in [("knn",     knn_day,     knn_conf),
                             ("wass",    wass_day,    wass_conf),
                             ("fft",     fft_day,     fft_conf),
                             ("spatial", spatial_day, spatial_conf),
                             ("kinetic", kinetic_day, kinetic_conf)]:
        if day is not None and conf >= 20.0:
            candidates.append((name, day, conf))

    if not candidates:
        return {"est_day": None, "confidence": 0,
                "weights": {"knn": 0, "wass": 0, "fft": 0,
                            "spatial": 0, "kinetic": 0}}

    total_c = sum(c for _, _, c in candidates)
    est_day = sum(d * c for _, d, c in candidates) / total_c
    days    = [d for _, d, _ in candidates]
    spread  = max(days) - min(days) if len(days) > 1 else 0
    consistency_bonus = max(0.0, 10.0 - spread * 2.0)
    ens_conf = min(100.0, total_c / len(candidates) + consistency_bonus)

    weights = {"knn": 0.0, "wass": 0.0, "fft": 0.0,
               "spatial": 0.0, "kinetic": 0.0}
    for name, _, conf in candidates:
        weights[name] = conf / total_c

    return {"est_day": est_day, "confidence": ens_conf, "weights": weights}


# ══════════════════════════════════════════════
#  고급 분석 — Spatial Pattern Features
# ══════════════════════════════════════════════

def adv_spatial_features(rgb: np.ndarray, mask: np.ndarray,
                          roi: tuple, rows: int = 3, cols: int = 3) -> dict:
    """
    세그먼트 분할 기반 공간 산화 패턴 피처 추출.
    반환:
      entropy       : 세그먼트 간 b* 분포의 공간 엔트로피 (값이 클수록 불균일)
      boundary_grad : 경계 세그먼트 vs 중심 세그먼트 b* 차이 (산화 전파 방향)
      anisotropy    : 가로/세로 방향 분산 비율 (이방성 산화 지표)
      seg_means     : 세그먼트별 b* 평균 (rows×cols 배열)
      seg_stds      : 세그먼트별 b* 표준편차
    """
    m = mask.astype(bool)
    if m.sum() == 0:
        return {"entropy": 0, "boundary_grad": 0, "anisotropy": 1,
                "seg_means": np.zeros((rows, cols)),
                "seg_stds":  np.zeros((rows, cols))}

    bgr = cv2.cvtColor(rgb, cv2.COLOR_RGB2BGR)
    lab = cv2.cvtColor(bgr, cv2.COLOR_BGR2Lab).astype(np.float32)
    b_star = lab[:, :, 2] - 128.0   # b* 전체 맵

    x0, y0, x1, y1 = roi
    H = y1 - y0
    W = x1 - x0
    sh = H / rows
    sw = W / cols

    seg_means = np.full((rows, cols), np.nan)
    seg_stds  = np.full((rows, cols), np.nan)

    for r in range(rows):
        for c in range(cols):
            ry0 = y0 + int(r * sh)
            ry1 = y0 + int((r + 1) * sh) if r < rows - 1 else y1
            rx0 = x0 + int(c * sw)
            rx1 = x0 + int((c + 1) * sw) if c < cols - 1 else x1
            seg_m = m[ry0:ry1, rx0:rx1]
            vals  = b_star[ry0:ry1, rx0:rx1][seg_m]
            if len(vals) >= 5:
                seg_means[r, c] = float(np.mean(vals))
                seg_stds[r, c]  = float(np.std(vals))

    valid = seg_means[~np.isnan(seg_means)]
    if len(valid) == 0:
        return {"entropy": 0, "boundary_grad": 0, "anisotropy": 1,
                "seg_means": seg_means, "seg_stds": seg_stds}

    # 공간 엔트로피 — b* 값의 세그먼트 간 분산
    entropy = float(np.nanstd(seg_means))

    # 경계-중심 그래디언트
    center_mask = np.zeros((rows, cols), bool)
    cr, cc = rows // 2, cols // 2
    center_mask[max(0, cr-1):cr+2, max(0, cc-1):cc+2] = True
    boundary_mask = ~center_mask

    center_vals   = seg_means[center_mask & ~np.isnan(seg_means)]
    boundary_vals = seg_means[boundary_mask & ~np.isnan(seg_means)]

    if len(center_vals) > 0 and len(boundary_vals) > 0:
        boundary_grad = float(np.mean(boundary_vals) - np.mean(center_vals))
    else:
        boundary_grad = 0.0

    # 이방성 (가로 vs 세로 방향 b* 분산 비율)
    row_vars = [np.nanvar(seg_means[r, :]) for r in range(rows)
                if not np.all(np.isnan(seg_means[r, :]))]
    col_vars = [np.nanvar(seg_means[:, c]) for c in range(cols)
                if not np.all(np.isnan(seg_means[:, c]))]
    row_var = float(np.mean(row_vars)) if row_vars else 0
    col_var = float(np.mean(col_vars)) if col_vars else 0
    anisotropy = row_var / (col_var + 1e-6) if col_var > 1e-6 else 1.0

    return {
        "entropy":       entropy,
        "boundary_grad": boundary_grad,
        "anisotropy":    anisotropy,
        "seg_means":     seg_means,
        "seg_stds":      seg_stds,
    }


def adv_spatial_estimate(target_spatial: dict,
                          pool: list,
                          rows: int = 3, cols: int = 3) -> dict:
    """
    공간 피처 유사도로 day 추정.
    거리 = 0.4*|entropy차| + 0.4*|boundary_grad차| + 0.2*|anisotropy차|
    반환: {est_day, confidence, scores}
    """
    t_ent = target_spatial["entropy"]
    t_bg  = target_spatial["boundary_grad"]
    t_ani = target_spatial["anisotropy"]

    ent_vals = [img["_adv_spatial"]["entropy"]       for img in pool
                if img.get("_adv_spatial")]
    bg_vals  = [img["_adv_spatial"]["boundary_grad"] for img in pool
                if img.get("_adv_spatial")]
    if not ent_vals:
        return {"est_day": None, "confidence": 0, "scores": []}

    ent_r = max(max(ent_vals) - min(ent_vals), 1e-6)
    bg_r  = max(max(abs(v) for v in bg_vals) * 2, 1e-6)

    scores = []
    for img in pool:
        sf = img.get("_adv_spatial")
        if sf is None:
            continue
        d_ent = abs(t_ent - sf["entropy"])       / ent_r
        d_bg  = abs(t_bg  - sf["boundary_grad"]) / bg_r
        d_ani = abs(t_ani - sf["anisotropy"])    / (abs(sf["anisotropy"]) + 1)
        dist  = 0.4 * d_ent + 0.4 * d_bg + 0.2 * d_ani
        scores.append((dist, img))

    if not scores:
        return {"est_day": None, "confidence": 0, "scores": []}

    scores.sort(key=lambda x: x[0])
    top3 = scores[:3]

    day_weights = []
    for dist, img in top3:
        try:
            day_weights.append((float(img["day"]), 1.0 / (dist + 1e-6)))
        except Exception:
            pass

    if not day_weights:
        return {"est_day": None, "confidence": 0, "scores": scores}

    total_w = sum(w for _, w in day_weights)
    est_day = sum(d * w for d, w in day_weights) / total_w
    conf    = max(0.0, min(100.0, 100.0 - scores[0][0] * 400))
    return {"est_day": est_day, "confidence": conf, "scores": scores}


# ══════════════════════════════════════════════
#  고급 분석 — Kinetic 물리 모델
# ══════════════════════════════════════════════

def adv_kinetic_fit(pool: list) -> dict:
    """
    참조 DB에서 조건(cond)별로 b*(t) = b_inf + (b0 - b_inf)*exp(-k*t) 피팅.
    3개 미만 데이터 조건도 선형 fallback으로 처리.
    반환: {cond: {k, b0, b_inf, r2, b_min, b_max, t_max, fit_method}}
    """
    from scipy.optimize import curve_fit

    cond_data: dict[str, list] = {}
    for img in pool:
        cond = img.get("cond", "unknown")
        try:
            day = float(img["day"])
            b   = float(img.get("lab", {}).get("b", np.nan))
            # b=0.0이면 미분석 이미지 → 제외 (1.0 임계값)
            if not np.isnan(b) and b > 1.0:
                cond_data.setdefault(cond, []).append((day, b))
        except Exception:
            pass

    result = {}
    for cond, pairs in cond_data.items():
        pairs.sort(key=lambda x: x[0])
        ts = np.array([p[0] for p in pairs], dtype=float)
        bs = np.array([p[1] for p in pairs], dtype=float)
        b_min, b_max = float(bs.min()), float(bs.max())
        t_max = float(ts.max())

        def decay(t, b0, b_inf, k):
            return b_inf + (b0 - b_inf) * np.exp(-k * t)

        if len(pairs) >= 3:
            try:
                # b_inf 하한 -20 → 매우 산화된 샘플도 수렴 가능
                popt, _ = curve_fit(
                    decay, ts, bs,
                    p0=[b_max, max(-5, b_min - 2), 0.15],
                    bounds=([0, -20, 1e-6], [150, 80, 20]),
                    maxfev=8000)
                b0_f, b_inf_f, k_f = popt
                bs_pred = decay(ts, *popt)
                ss_res = np.sum((bs - bs_pred) ** 2)
                ss_tot = np.sum((bs - np.mean(bs)) ** 2)
                r2 = max(0.0, 1.0 - ss_res / (ss_tot + 1e-12))
                result[cond] = {
                    "k": float(k_f), "b0": float(b0_f),
                    "b_inf": float(b_inf_f), "r2": float(r2),
                    "b_min": b_min, "b_max": b_max, "t_max": t_max,
                    "fit_method": "exponential",
                }
                continue
            except Exception:
                pass

        # fallback: 데이터 부족 or 피팅 실패 → 선형 근사
        if len(pairs) >= 2:
            try:
                coef = np.polyfit(ts, bs, 1)
                b0_lin = float(np.polyval(coef, 0))
                slope  = float(coef[0])
                # 선형에서 k 추정: dI/dt ≈ -k * I → k ≈ -slope / mean(bs)
                k_lin  = max(1e-4, -slope / (np.mean(bs) + 1e-6))
                b_inf_lin = max(-10.0, b_min - 2.0)
                bs_pred = np.polyval(coef, ts)
                ss_res = np.sum((bs - bs_pred) ** 2)
                ss_tot = np.sum((bs - np.mean(bs)) ** 2)
                r2_lin = max(0.0, 1.0 - ss_res / (ss_tot + 1e-12))
                result[cond] = {
                    "k": k_lin, "b0": b0_lin,
                    "b_inf": b_inf_lin, "r2": r2_lin * 0.7,  # fallback 패널티
                    "b_min": b_min, "b_max": b_max, "t_max": t_max,
                    "fit_method": "linear_fallback",
                }
            except Exception:
                pass

    return result


def adv_kinetic_estimate(target_b: float,
                          kinetic_params: dict,
                          target_cond: str = "") -> dict:
    """
    타겟 b*와 피팅 파라미터로 day 역산.
    ratio 범위 초과 시 외삽(extrapolation) 또는 경계값으로 처리.
    """
    _none = {"est_day": None, "confidence": 0,
             "cond_used": "", "model_params": {}, "fail_reason": ""}

    if not kinetic_params:
        return {**_none, "fail_reason": "참조 DB 피팅 결과 없음 (데이터 부족)"}
    if np.isnan(target_b):
        return {**_none, "fail_reason": "타겟 b* 값이 NaN"}

    # 조건 선택: 직접 매칭 → 부분 매칭 → 최고 R²
    best_cond = None
    if target_cond:
        if target_cond in kinetic_params:
            best_cond = target_cond
        else:
            for cond in kinetic_params:
                tc = target_cond.lower()
                cc = cond.lower()
                if tc in cc or cc in tc:
                    best_cond = cond
                    break
    if best_cond is None:
        best_cond = max(kinetic_params, key=lambda c: kinetic_params[c]["r2"])

    p  = kinetic_params[best_cond]
    b0 = p["b0"];  b_inf = p["b_inf"];  k = p["k"];  r2 = p["r2"]
    b_min = p.get("b_min", b_inf);  t_max = p.get("t_max", 28.0)

    try:
        denom = b0 - b_inf
        if abs(denom) < 0.5:
            # b0 ≈ b_inf → 모델이 상수 → 추정 불가
            return {**_none, "fail_reason": "모델 b0≈b_inf (데이터 분산 없음)",
                    "cond_used": best_cond, "model_params": p}

        ratio = (target_b - b_inf) / denom

        # ── 정상 범위 (0 < ratio <= 1) ──────────────────
        if 0 < ratio <= 1.0:
            est_day = float(-np.log(ratio) / (k + 1e-12))
            est_day = float(np.clip(est_day, 0, 365))
            # 신뢰도: R² 기반 + ratio 중간일수록 안정
            stability = 1.0 - abs(ratio - 0.5) * 0.8
            conf = float(max(0, min(100, r2 * 100 * stability)))
            return {
                "est_day": est_day, "confidence": conf,
                "cond_used": best_cond, "model_params": p,
                "fail_reason": "",
            }

        # ── ratio > 1: 타겟이 b0보다 높음 (신선한 초기) → 0일 추정 ──
        if ratio > 1.0:
            est_day = 0.0
            conf = float(max(0, min(60, r2 * 60)))  # 외삽이므로 신뢰도 제한
            return {
                "est_day": est_day, "confidence": conf,
                "cond_used": best_cond, "model_params": p,
                "fail_reason": "b*가 초기값보다 높음 (ratio>1) → Day 0 추정",
            }

        # ── ratio <= 0: 타겟이 b_inf보다 낮음 (포화 초과) → 외삽 ──
        # t_max에서의 b* 기준으로 선형 외삽
        # b*(t) ≈ b_inf + small_residual → t 매우 큰 값
        # 안전하게: t_max 기준으로 남은 감쇠를 역산
        b_at_tmax = b_inf + denom * np.exp(-k * t_max)
        if target_b >= b_at_tmax * 0.5:
            # b_at_tmax보다 약간 낮음 → t_max 근방 외삽
            ratio_ext = max(1e-6, (target_b - b_inf) / denom)
            est_day = float(-np.log(ratio_ext) / (k + 1e-12))
            est_day = float(np.clip(est_day, t_max, 365))
            conf = float(max(0, min(40, r2 * 40)))  # 외삽 신뢰도 낮음
        else:
            # 완전 산화 → t_max의 1.5배 추정
            est_day = float(min(365, t_max * 1.5))
            conf = float(max(0, min(25, r2 * 25)))
        return {
            "est_day": est_day, "confidence": conf,
            "cond_used": best_cond, "model_params": p,
            "fail_reason": "b*가 포화값보다 낮음 (ratio≤0) → 외삽 추정",
        }

    except Exception as ex:
        return {**_none, "fail_reason": f"역산 오류: {ex}",
                "cond_used": best_cond, "model_params": p}


def adv_precompute_pool(pool: list, bins: int = 64,
                         rows: int = 3, cols: int = 3) -> None:
    """
    참조 DB pool의 각 이미지에 _adv_hist, _adv_fft, _adv_spatial 캐시 삽입.
    (없는 경우에만 계산 — 중복 방지)
    rgb, mask, roi 키가 있어야 함.
    """
    for img in pool:
        rgb  = img.get("rgb")
        mask = img.get("mask")
        roi  = img.get("roi")
        if rgb is None or mask is None:
            continue
        if roi is None:
            roi = (0, 0, rgb.shape[1], rgb.shape[0])
        if img.get("_adv_hist") is None:
            img["_adv_hist"]    = adv_hist_signature(rgb, mask, bins)
            img["_adv_fft"]     = adv_fft_features(rgb, mask)
            img["_adv_spatial"] = adv_spatial_features(rgb, mask, roi,
                                                        rows, cols)


def make_thumb(rgb, tw, th, roi=None):
    pil = Image.fromarray(rgb)
    pil.thumbnail((tw, th), Image.LANCZOS)
    if roi is not None:
        sx = pil.width  / rgb.shape[1]
        sy = pil.height / rgb.shape[0]
        x0, y0, x1, y1 = roi
        draw = ImageDraw.Draw(pil)
        draw.rectangle([int(x0*sx), int(y0*sy),
                        int(x1*sx), int(y1*sy)],
                       outline=(59, 130, 246), width=2)
    return pil


# ══════════════════════════════════════════════
#  DnD 경로 파싱
# ══════════════════════════════════════════════
def parse_drop_paths(raw: str) -> list:
    pairs = _re.findall(r'\{([^}]+)\}|([^\s{}]+)', raw)
    return [a.strip() or b.strip() for a,b in pairs if (a or b).strip()]


# ══════════════════════════════════════════════
#  DB 저장 / 로드  (SQLite3 내장)
# ══════════════════════════════════════════════
import sqlite3, json as _json, pickle as _pickle

_DB_VERSION = 2   # 스키마 변경 시 증가

def _db_open_safe(path: str, timeout: float = 15.0) -> sqlite3.Connection:
    """WSL/네트워크 마운트 호환 SQLite connection 헬퍼.

    \\\\wsl$\\... 같은 9p 마운트에서 WAL 모드는 .shm/.wal 파일 lock 충돌
    유발 → 안전한 DELETE 모드 (기본 rollback journal) 사용.
    busy_timeout 으로 OS-레벨 lock 도 15초까지 대기.
    """
    pre_size = os.path.getsize(path) if os.path.exists(path) else -1
    leftovers = [ext for ext in ("-wal", "-shm", "-journal")
                 if os.path.exists(path + ext)]
    print(f"[db-open] path={path} pre_size={pre_size} leftovers={leftovers}")
    try:
        con = sqlite3.connect(path, timeout=timeout)
    except Exception as e:
        print(f"[db-open] FAIL connect: {type(e).__name__}: {e}")
        raise
    try:
        con.execute(f"PRAGMA busy_timeout={int(timeout * 1000)}")
        jm = con.execute("PRAGMA journal_mode=DELETE").fetchone()
        sync = con.execute("PRAGMA synchronous=NORMAL").fetchone()
        print(f"[db-open] ok: journal_mode={jm} synchronous={sync}")
    except Exception as e:
        print(f"[db-open] FAIL pragma: {type(e).__name__}: {e}")
        try: con.close()
        except Exception: pass
        raise
    return con


def db_init(path: str) -> sqlite3.Connection:
    """DB 파일을 열고 테이블이 없으면 생성.

    통합 DB (v2): images + raman_data + 매칭 컬럼(images.raman_id) 한 파일에.
    구 DB (v1): images 만 있던 경우 raman_id 컬럼 자동 ALTER.
    """
    con = _db_open_safe(path)
    con.execute(f"""
        CREATE TABLE IF NOT EXISTS meta (
            key   TEXT PRIMARY KEY,
            value TEXT
        )""")
    con.execute("""
        CREATE TABLE IF NOT EXISTS images (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            name        TEXT,
            day         TEXT,
            cond        TEXT,
            roi_x0      INTEGER, roi_y0 INTEGER,
            roi_x1      INTEGER, roi_y1 INTEGER,
            s_mean      REAL, yellow_ratio REAL, yellowness_idx REAL,
            lab_L       REAL, lab_a REAL, lab_b REAL,
            delta_e     REAL,
            glcm_con    REAL, glcm_eng REAL,
            glcm_hom    REAL, glcm_cor REAL,
            stats_json  TEXT,
            rgb_blob    BLOB,
            thumb_blob  BLOB,
            raman_id    INTEGER,
            saved_at    TEXT DEFAULT (datetime('now','localtime'))
        )""")
    # 라만 데이터 테이블 — 통합 DB 의 일부. id 보존(저장→로드 매칭).
    con.execute("""
        CREATE TABLE IF NOT EXISTS raman_data (
            id           INTEGER PRIMARY KEY,
            cond         TEXT,
            day          TEXT,
            peak         REAL,
            norm_peak    REAL,
            peak_shift   REAL,
            peak_range   TEXT,
            spectrum_json TEXT,
            saved_at     TEXT DEFAULT (datetime('now','localtime'))
        )""")
    # v1 → v2 마이그레이션: images 에 raman_id 가 없으면 추가
    cols = [r[1] for r in con.execute("PRAGMA table_info(images)").fetchall()]
    if "raman_id" not in cols:
        try:
            con.execute("ALTER TABLE images ADD COLUMN raman_id INTEGER")
            print("[db-init] migrated: added images.raman_id column")
        except Exception as e:
            print(f"[db-init] WARN ALTER images.raman_id failed: {e}")
    con.execute(f"INSERT OR IGNORE INTO meta VALUES ('version','{_DB_VERSION}')")
    con.commit()
    return con


def _rgb_to_blob(rgb: np.ndarray) -> bytes:
    pil = Image.fromarray(rgb)
    buf = io.BytesIO()
    pil.save(buf, format="PNG", optimize=True)
    return buf.getvalue()


def _blob_to_rgb(blob: bytes) -> np.ndarray:
    return np.array(Image.open(io.BytesIO(blob)).convert("RGB"))


def _migrate_eval_target_schema(con):
    """eval_target 테이블 마이그레이션.

    구 스키마: (id INTEGER PRIMARY KEY, rgb_blob BLOB, roi TEXT, saved_at TEXT)
    신 스키마: (target_id INTEGER PRIMARY KEY, name TEXT, rgb_blob BLOB,
               roi TEXT, color TEXT, cond_hint TEXT, result_json TEXT,
               saved_at TEXT)

    Returns: 'created' | 'already_migrated' | 'migrated_with_backup'
    """
    cur = con.execute(
        "SELECT name FROM sqlite_master "
        "WHERE type='table' AND name='eval_target'")
    if not cur.fetchone():
        con.execute(
            "CREATE TABLE eval_target ("
            "target_id INTEGER PRIMARY KEY, "
            "name TEXT, rgb_blob BLOB, roi TEXT, "
            "color TEXT, cond_hint TEXT, "
            "result_json TEXT, saved_at TEXT)")
        con.commit()
        return "created"
    cols = {row[1]
            for row in con.execute("PRAGMA table_info(eval_target)")}
    if {"target_id", "name", "color"}.issubset(cols):
        return "already_migrated"
    # 백업 + 재구성 — 기존 backup 이 있으면 먼저 정리 (가장 최신 backup 만 유지)
    con.execute("DROP TABLE IF EXISTS eval_target_v1_backup")
    con.execute("ALTER TABLE eval_target RENAME TO eval_target_v1_backup")
    con.execute(
        "CREATE TABLE eval_target ("
        "target_id INTEGER PRIMARY KEY, "
        "name TEXT, rgb_blob BLOB, roi TEXT, "
        "color TEXT, cond_hint TEXT, "
        "result_json TEXT, saved_at TEXT)")
    # 기존 단일 row 를 첫 target 으로 복사
    try:
        rows = list(con.execute(
            "SELECT rgb_blob, roi, saved_at FROM eval_target_v1_backup"))
    except Exception:
        rows = []
    for i, row in enumerate(rows[:1], start=1):
        blob = row[0]
        roi  = row[1] if len(row) > 1 else None
        saved_at = row[2] if len(row) > 2 else None
        con.execute(
            "INSERT INTO eval_target VALUES (?,?,?,?,?,?,?,?)",
            (i, f"Target #{i}", blob, roi,
             COND_COLORS[0], "", None, saved_at))
    con.commit()
    return "migrated_with_backup"


def db_save_all(path: str, images: list) -> int:
    """
    이미지들을 DB에 저장 (upsert: name+cond+day 기준).
    분석 결과(s_mean 등)가 NaN 이어도 저장 — ROI 보존용 checkpoint 가능.
    저장된 행 수 반환.
    """
    saveable = [img for img in images if img.get("rgb") is not None]
    n_with_raman = sum(1 for img in saveable
                       if img.get("raman_id") is not None)
    print(f"[db-save-all] start: path={path} total_images={len(images)} "
          f"saveable={len(saveable)} with_raman_id={n_with_raman}")
    if not saveable:
        print(f"[db-save-all] WARN: no saveable images (all rgb is None)")
        return 0

    con = db_init(path)
    print(f"[db-save-all] db_init ok, file_size={os.path.getsize(path)}")
    count = 0
    for img in saveable:
        roi = img.get("roi")
        lab  = img.get("lab",  {})
        glcm = img.get("glcm", {})

        # stats: tkVar 제외하고 직렬화 가능한 것만
        stats_safe = {}
        for ch, segs in img.get("stats", {}).items():
            stats_safe[ch] = segs   # seg_stats는 순수 dict 리스트

        # rgb 블롭
        rgb_blob   = _rgb_to_blob(img["rgb"])
        # 썸네일 블롭
        th = img.get("thumb")
        thumb_blob = _rgb_to_blob(np.array(th.convert("RGB"))) if th else b""

        con.execute("""
            INSERT OR REPLACE INTO images
            (name, day, cond,
             roi_x0, roi_y0, roi_x1, roi_y1,
             s_mean, yellow_ratio, yellowness_idx,
             lab_L, lab_a, lab_b, delta_e,
             glcm_con, glcm_eng, glcm_hom, glcm_cor,
             stats_json, rgb_blob, thumb_blob,
             raman_id,
             saved_at)
            VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,
                    datetime('now','localtime'))
        """, (
            img["name"], img["day"], img["cond"],
            roi[0] if roi else None, roi[1] if roi else None,
            roi[2] if roi else None, roi[3] if roi else None,
            img.get("s_mean", np.nan),
            img.get("yellow_ratio", np.nan),
            img.get("yellowness_idx", np.nan),
            lab.get("L", np.nan), lab.get("a", np.nan), lab.get("b", np.nan),
            img.get("delta_e", np.nan),
            glcm.get("contrast", np.nan), glcm.get("energy", np.nan),
            glcm.get("homogeneity", np.nan), glcm.get("correlation", np.nan),
            _json.dumps(stats_safe, ensure_ascii=False),
            rgb_blob, thumb_blob,
            img.get("raman_id"),
        ))
        count += 1

    print(f"[db-save-all] inserts done: count={count}, committing...")
    try:
        con.commit()
    except Exception as e:
        print(f"[db-save-all] FAIL commit: {type(e).__name__}: {e}")
        try: con.close()
        except Exception: pass
        raise
    try:
        con.close()
    except Exception as e:
        print(f"[db-save-all] WARN close: {type(e).__name__}: {e}")
    final_size = os.path.getsize(path) if os.path.exists(path) else -1
    print(f"[db-save-all] ok: count={count} file_size={final_size}")
    return count


# LOAD 경로의 local temp 복사 디렉토리 추적 — 프로세스 종료 시 일괄 정리
_LOAD_TEMP_DIRS: list = []


def _cleanup_load_temp_dirs():
    """atexit hook — 프로세스 종료 시 LOAD 임시 디렉토리 정리."""
    for d in _LOAD_TEMP_DIRS:
        try:
            shutil.rmtree(d, ignore_errors=True)
        except Exception:
            pass


import atexit as _atexit
_atexit.register(_cleanup_load_temp_dirs)


def _db_open_read(path: str, timeout: float = 15.0) -> sqlite3.Connection:
    """읽기 전용 open — 9p / Windows UNC / 외부 도구 lock 모두 회피.

    LOAD 경로 전용. 시도 순서:
      1) file URI ?mode=ro (pathlib.Path(path).resolve().as_uri()) —
         SQLite 가 read-only 모드로 열어 외부 도구 lock 영향 안 받음.
      2) src.db → OS local temp 복사 → temp 일반 open. 모든 lock 회피.
         con close 후 tmp_dir 정리는 프로세스 종료 시 atexit hook 으로.
      3) 일반 connect — PRAGMA journal_mode 변경 없이 busy_timeout 만.
    """
    pre_size = os.path.getsize(path) if os.path.exists(path) else -1
    leftovers = [ext for ext in ("-wal", "-shm", "-journal")
                 if os.path.exists(path + ext)]
    print(f"[db-open-read] path={path} pre_size={pre_size} leftovers={leftovers}")

    # 시도 1: file URI ?mode=ro
    try:
        import pathlib
        uri = pathlib.Path(path).resolve().as_uri() + "?mode=ro"
        con = sqlite3.connect(uri, uri=True, timeout=timeout)
        con.execute(f"PRAGMA busy_timeout={int(timeout * 1000)}")
        print(f"[db-open-read] ok (uri mode=ro): {uri}")
        return con
    except Exception as e:
        print(f"[db-open-read] URI mode=ro failed: {type(e).__name__}: {e}; "
              f"trying local temp copy")

    # 시도 2: local temp 복사 후 open (모든 lock 회피)
    tmp_dir = None
    try:
        tmp_dir = tempfile.mkdtemp(prefix="hfs2_load_")
        tmp_path = os.path.join(tmp_dir, "load.db")
        shutil.copy2(path, tmp_path)
        # 동반 lock 파일도 복사 (best-effort — 잠겨있어도 무시)
        for ext in ("-wal", "-shm", "-journal"):
            src_ext = path + ext
            if os.path.exists(src_ext):
                try: shutil.copy2(src_ext, tmp_path + ext)
                except OSError: pass
        con = sqlite3.connect(tmp_path, timeout=timeout)
        con.execute(f"PRAGMA busy_timeout={int(timeout * 1000)}")
        _LOAD_TEMP_DIRS.append(tmp_dir)
        print(f"[db-open-read] ok (local temp copy): {tmp_path}")
        return con
    except Exception as e:
        print(f"[db-open-read] local temp copy failed: "
              f"{type(e).__name__}: {e}; trying plain connect")
        if tmp_dir:
            try: shutil.rmtree(tmp_dir, ignore_errors=True)
            except Exception: pass

    # 시도 3: 일반 connect (PRAGMA journal_mode 변경 안 함)
    con = sqlite3.connect(path, timeout=timeout)
    con.execute(f"PRAGMA busy_timeout={int(timeout * 1000)}")
    print(f"[db-open-read] ok (plain, no pragma change)")
    return con


def db_load_all(path: str) -> list:
    """DB에서 전체 이미지 레코드를 dict 리스트로 반환 (rgb_blob + raman_id 포함).

    read-only 로 열어 9p 마운트의 SQLite lock 회피.
    raman_id 컬럼이 구 DB 에 없으면 None 으로 fallback (ALTER 안 함).
    """
    con = _db_open_read(path)
    cols = [r[1] for r in con.execute("PRAGMA table_info(images)").fetchall()]
    has_raman_col = "raman_id" in cols
    select_raman = "raman_id" if has_raman_col else "NULL AS raman_id"
    print(f"[db-load-all] path={path} raman_id_col={has_raman_col}")
    rows = con.execute(f"""
        SELECT name, day, cond,
               roi_x0, roi_y0, roi_x1, roi_y1,
               s_mean, yellow_ratio, yellowness_idx,
               lab_L, lab_a, lab_b, delta_e,
               glcm_con, glcm_eng, glcm_hom, glcm_cor,
               stats_json, rgb_blob, thumb_blob, {select_raman}, saved_at
        FROM images ORDER BY cond, CAST(day AS REAL)
    """).fetchall()
    n_with_raman = sum(1 for r in rows if r[21] is not None)
    print(f"[db-load-all] rows={len(rows)} with_raman_id={n_with_raman}")
    con.close()

    result = []
    for r in rows:
        (name, day, cond,
         x0, y0, x1, y1,
         s_mean, yr, yi,
         lab_L, lab_a, lab_b, de,
         gc, ge, gh, gcr,
         stats_json, rgb_blob, thumb_blob, raman_id, saved_at) = r

        roi = (x0,y0,x1,y1) if x0 is not None else None
        rgb = _blob_to_rgb(rgb_blob) if rgb_blob else None
        H_ch,S_ch,I_ch = rgb_to_hsi(rgb) if rgb is not None else (None,None,None)
        th  = Image.open(io.BytesIO(thumb_blob)).convert("RGB") \
              if thumb_blob else None

        result.append({
            "name": name, "day": day, "cond": cond,
            "rgb":  rgb,
            "hsi":  (H_ch,S_ch,I_ch),
            "roi":  roi,
            "mask": roi_to_mask(rgb.shape, roi) if (rgb is not None and roi) else None,
            "thumb": th,
            "stats": _json.loads(stats_json) if stats_json else {},
            "s_mean":       float(s_mean) if s_mean is not None else np.nan,
            "yellow_ratio": float(yr)     if yr     is not None else np.nan,
            "yellowness_idx": float(yi)   if yi     is not None else np.nan,
            "lab":  {"L": float(lab_L or 0),
                     "a": float(lab_a or 0),
                     "b": float(lab_b or 0)},
            "delta_e": float(de) if de is not None else np.nan,
            "glcm": {"contrast":    float(gc  or 0),
                     "energy":      float(ge  or 0),
                     "homogeneity": float(gh  or 0),
                     "correlation": float(gcr or 0)},
            "raman_id": int(raman_id) if raman_id is not None else None,
            "auto_parsed": False,
            "saved_at": saved_at,
            "_from_db": True,
        })
    return result


def db_load_raman_all(path: str) -> list:
    """통합 DB 에서 raman_data 전체 로드 (id 보존). 빈 테이블이면 [] 반환.

    read-only 모드로 열어 9p lock 회피.
    """
    out = []
    try:
        con = _db_open_read(path)
        try:
            cur = con.execute(
                "SELECT name FROM sqlite_master "
                "WHERE type='table' AND name='raman_data'")
            if not cur.fetchone():
                print(f"[db-load-raman-all] no raman_data table in {path}")
                return out
            rows = con.execute(
                "SELECT id, cond, day, peak, norm_peak, peak_shift, "
                "peak_range, spectrum_json FROM raman_data ORDER BY id"
            ).fetchall()
            print(f"[db-load-raman-all] {path} rows={len(rows)}")
            for (rid, cond, day, peak, norm_peak,
                 peak_shift, peak_range, spec_j) in rows:
                spec = _json.loads(spec_j) if spec_j else None
                out.append({
                    "_id":       int(rid),
                    "cond":      cond or "",
                    "day":       day  or "",
                    "peak":      float(peak or 0),
                    "norm_peak": float(norm_peak or 0),
                    "peak_shift":float(peak_shift or 0),
                    "peak_range":str(peak_range or ""),
                    "spectrum":  spec,
                })
        finally:
            con.close()
    except Exception as e:
        print(f"[db-load-raman] WARN: {type(e).__name__}: {e}")
    return out


def db_save_raman_all(con: sqlite3.Connection, raman_list: list) -> int:
    """이미 열린 connection 에 라만 데이터 일괄 저장. id 보존(있으면 사용).

    raman_list 의 각 dict 가 가진 _id (int) 가 있으면 그 id 로 INSERT,
    없으면 자동 할당 후 dict 에 _id 채워 넣음. 매칭 보존을 위함.
    저장 후 _id 가 새로 부여된 경우 호출자가 매칭 (img['raman_id']) 갱신 필요.
    """
    con.execute("DELETE FROM raman_data")
    n = 0
    for r in raman_list:
        spec_j = (_json.dumps(r["spectrum"]) if r.get("spectrum") else None)
        rid = r.get("_id")
        if rid is not None:
            con.execute(
                "INSERT INTO raman_data"
                " (id, cond, day, peak, norm_peak, peak_shift,"
                "  peak_range, spectrum_json, saved_at)"
                " VALUES (?,?,?,?,?,?,?,?,datetime('now','localtime'))",
                (int(rid), r.get("cond",""), r.get("day",""),
                 float(r.get("peak", 0)),
                 float(r.get("norm_peak") or 0),
                 float(r.get("peak_shift") or 0),
                 str(r.get("peak_range","")),
                 spec_j))
        else:
            cur = con.execute(
                "INSERT INTO raman_data"
                " (cond, day, peak, norm_peak, peak_shift,"
                "  peak_range, spectrum_json, saved_at)"
                " VALUES (?,?,?,?,?,?,?,datetime('now','localtime'))",
                (r.get("cond",""), r.get("day",""),
                 float(r.get("peak", 0)),
                 float(r.get("norm_peak") or 0),
                 float(r.get("peak_shift") or 0),
                 str(r.get("peak_range","")),
                 spec_j))
            r["_id"] = cur.lastrowid
        n += 1
    return n


# ══════════════════════════════════════════════
#  4. 사용자 설정 (API 키 등) 영구 저장
# ══════════════════════════════════════════════
def _settings_dir() -> str:
    if os.name == "nt":
        base = os.environ.get("APPDATA") or os.path.expanduser("~")
        return os.path.join(base, "hfs2_analyzer")
    return os.path.join(os.path.expanduser("~"), ".config", "hfs2_analyzer")


def _settings_path() -> str:
    return os.path.join(_settings_dir(), "settings.json")


def load_settings() -> dict:
    try:
        with open(_settings_path(), encoding="utf-8") as f:
            data = _json.load(f)
            return data if isinstance(data, dict) else {}
    except (FileNotFoundError, OSError, ValueError):
        return {}


def save_settings(data: dict) -> bool:
    try:
        os.makedirs(_settings_dir(), exist_ok=True)
        with open(_settings_path(), "w", encoding="utf-8") as f:
            _json.dump(data, f, ensure_ascii=False, indent=2)
        return True
    except (OSError, ValueError):
        return False


# ══════════════════════════════════════════════
#  5. ROI 선택 팝업 (드래그 이동 + 복사 지원)
# ══════════════════════════════════════════════
class ROISelector(tk.Toplevel):
    """
    - 드래그: 새 ROI 그리기
    - 박스 안쪽 드래그: 이동
    - ← → 방향키 / 버튼: 이미지 전환 (현재 ROI 자동 저장)
    - [⊕ 전체복사]: 비율 기준으로 다른 이미지에 ROI 적용
                   → 이미지 크기/해상도 달라도 화면상 같은 위치
    """
    HANDLE = 10

    def __init__(self, parent, img_entry, on_confirm,
                 on_copy_all=None,
                 images=None,
                 current_idx=0):
        super().__init__(parent)
        self.configure(bg=PANEL)
        self.resizable(True, True)

        self._images     = images or [img_entry]
        self._cur_idx    = current_idx
        self._on_confirm = on_confirm
        self._on_copy    = on_copy_all

        self._rgb        = img_entry["rgb"]
        self._roi        = img_entry.get("roi")
        self._mode       = "idle"
        self._drag_start = None
        self._roi_at_drag= None
        self._scale      = 1.0
        self._ox = self._oy = 0

        self._build()
        self._update_title()
        sw = parent.winfo_screenwidth()
        sh = parent.winfo_screenheight()
        ww, wh = int(sw*0.84), int(sh*0.84)
        self.geometry(f"{ww}x{wh}+{(sw-ww)//2}+{(sh-wh)//2}")
        self.grab_set()
        self.focus_set()   # ★ 방향키 수신을 위해 팝업에 포커스

        # 방향키: 팝업 + 캔버스 모두 바인딩
        for w in (self, self._cv):
            w.bind("<Left>",  lambda e: self._switch_image(-1))
            w.bind("<Right>", lambda e: self._switch_image(+1))
            w.bind("<Prior>", lambda e: self._switch_image(-1))
            w.bind("<Next>",  lambda e: self._switch_image(+1))

    def _update_title(self):
        n   = len(self._images)
        img = self._images[self._cur_idx]
        self.title(f"ROI  [{self._cur_idx+1}/{n}]  —  {img['name']}")

    # ─── UI ─────────────────────────────────────
    def _build(self):
        hdr = tk.Frame(self, bg=PANEL2)
        hdr.pack(fill="x")

        left_hdr = tk.Frame(hdr, bg=PANEL2)
        left_hdr.pack(side="left", fill="y")

        tk.Label(left_hdr,
                 text=_L("  드래그: ROI  |  박스 안쪽: 이동  |  ← →: 이미지 전환","  Drag: ROI  |  Inside drag: Move  |  ←→: Switch image"),
                 bg=PANEL2, fg=TEAL, font=MFB).pack(side="left", padx=10, pady=8)

        nav = tk.Frame(left_hdr, bg=PANEL2)
        nav.pack(side="left", padx=4)

        self._prev_btn = tk.Button(nav, text=_L("◀ 이전","◀ Prev"),
            command=lambda: self._switch_image(-1),
            bg=BTN, fg=TXT, font=MF,
            relief="flat", padx=8, pady=4, cursor="hand2")
        self._prev_btn.pack(side="left", padx=2)

        self._nav_label = tk.Label(nav, text="1/1",
            bg=PANEL2, fg=SUB, font=MFB, width=6)
        self._nav_label.pack(side="left", padx=2)

        self._next_btn = tk.Button(nav, text=_L("다음 ▶","Next ▶"),
            command=lambda: self._switch_image(+1),
            bg=BTN, fg=TXT, font=MF,
            relief="flat", padx=8, pady=4, cursor="hand2")
        self._next_btn.pack(side="left", padx=2)

        self._info = tk.StringVar(value=_L("ROI를 드래그하여 선택한다","Drag to select ROI"))
        tk.Label(hdr, textvariable=self._info,
                 bg=PANEL2, fg=SUB, font=LF).pack(side="left", padx=6)

        for txt, cmd, bg_, fg_ in [
            ("✔ Confirm",    self._confirm,  ACCENT, "white"),
            ("⊕ Copy All", self._copy_all, GREEN,  "white"),
            ("↺ Reset",  self._reset,    BTN,    TXT),
            ("✕ Cancel",    self.destroy,   BTN,    TXT),
        ]:
            tk.Button(hdr, text=txt, command=cmd,
                      bg=bg_, fg=fg_, font=MF,
                      relief="flat", padx=10, pady=5,
                      cursor="hand2").pack(side="right", padx=3, pady=6)

        self._cv = tk.Canvas(self, bg="#e8eaf0",
                             cursor="crosshair", highlightthickness=0)
        self._cv.pack(fill="both", expand=True)
        self._cv.bind("<ButtonPress-1>",   self._press)
        self._cv.bind("<B1-Motion>",       self._drag)
        self._cv.bind("<ButtonRelease-1>", self._release)
        self._cv.bind("<Motion>",          self._hover)
        self._cv.bind("<Configure>",       lambda e: self.after(50, self._render))
        # ★ 캔버스 클릭 후 포커스 되돌리기 — 방향키 유지
        self._cv.bind("<ButtonRelease-1>", lambda e: self.focus_set(), add="+")

        self._refresh_nav()

    def _refresh_nav(self):
        n = len(self._images)
        i = self._cur_idx
        self._nav_label.configure(text=f"{i+1}/{n}")
        self._prev_btn.configure(
            state="normal" if i > 0   else "disabled",
            fg=TXT        if i > 0   else BORDER)
        self._next_btn.configure(
            state="normal" if i < n-1 else "disabled",
            fg=TXT        if i < n-1 else BORDER)

    # ─── 이미지 전환 ────────────────────────────
    def _switch_image(self, delta):
        new_idx = self._cur_idx + delta
        if new_idx < 0 or new_idx >= len(self._images):
            return

        # 현재 ROI 임시 저장
        self._images[self._cur_idx]["_pending_roi"] = self._roi

        # 새 이미지로 전환
        self._cur_idx = new_idx
        new_img = self._images[new_idx]
        self._rgb  = new_img["rgb"]
        self._roi  = new_img.get("_pending_roi") or new_img.get("roi")
        self._mode = "idle"
        self._drag_start = None

        self._update_title()
        self._refresh_nav()
        self._render()

        has_roi = self._roi is not None
        self._info.set(
            _L(f"✔ ROI 있음 ({self._roi[2]-self._roi[0]}×{self._roi[3]-self._roi[1]}px)  — 드래그로 수정 가능",
               f"✔ ROI {self._roi[2]-self._roi[0]}×{self._roi[3]-self._roi[1]}px  — drag to modify")
            if has_roi else _L("ROI 없음  —  드래그하여 ROI 선택","No ROI  —  drag to select ROI"))

        self.focus_set()   # ★ 전환 후 포커스 유지

    # ─── 렌더 ───────────────────────────────────
    def _render(self):
        cw = self._cv.winfo_width()
        ch = self._cv.winfo_height()
        if cw < 2 or ch < 2: return
        ih, iw = self._rgb.shape[:2]
        self._scale = min(cw/iw, ch/ih)
        dw = int(iw * self._scale)
        dh = int(ih * self._scale)
        self._ox = (cw - dw) // 2
        self._oy = (ch - dh) // 2

        pil = Image.fromarray(self._rgb).resize((dw, dh), Image.LANCZOS)
        self._tk_img = ImageTk.PhotoImage(pil)
        self._cv.delete("all")
        self._cv.create_image(self._ox, self._oy, anchor="nw",
                              image=self._tk_img)
        if self._roi:
            self._draw_roi(self._roi)

    def _i2c(self, ix, iy):
        return self._ox + ix*self._scale, self._oy + iy*self._scale

    def _c2i(self, cx, cy):
        ih, iw = self._rgb.shape[:2]
        ix = int((cx - self._ox) / self._scale)
        iy = int((cy - self._oy) / self._scale)
        return max(0, min(ix, iw)), max(0, min(iy, ih))

    def _draw_roi(self, roi):
        x0,y0,x1,y1 = roi
        cx0,cy0 = self._i2c(x0, y0)
        cx1,cy1 = self._i2c(x1, y1)

        cw = self._cv.winfo_width()
        ch = self._cv.winfo_height()
        for rect, stipple in [
            ((0,0,cw,cy0),     "gray50"),
            ((0,cy1,cw,ch),    "gray50"),
            ((0,cy0,cx0,cy1),  "gray50"),
            ((cx1,cy0,cw,cy1), "gray50"),
        ]:
            self._cv.create_rectangle(*rect, fill="#000000",
                                      stipple=stipple, outline="")

        self._cv.create_rectangle(cx0,cy0,cx1,cy1, outline=ACCENT, width=2)
        w = x1-x0; h = y1-y0
        self._cv.create_text(cx0+6, cy0+4, anchor="nw",
                             text=f" {w}\u00d7{h}px ",
                             fill="white", font=("Consolas",9,"bold"))
        hh = self.HANDLE
        for cx,cy in [(cx0,cy0),(cx1,cy0),(cx0,cy1),(cx1,cy1)]:
            self._cv.create_oval(cx-hh,cy-hh,cx+hh,cy+hh,
                                 fill=ACCENT, outline="white", width=1)

    def _inside_roi(self, cx, cy):
        if not self._roi: return False
        x0,y0,x1,y1 = self._roi
        cx0,cy0 = self._i2c(x0,y0)
        cx1,cy1 = self._i2c(x1,y1)
        return cx0 <= cx <= cx1 and cy0 <= cy <= cy1

    # ─── 마우스 ─────────────────────────────────
    def _hover(self, event):
        if self._roi and self._inside_roi(event.x, event.y):
            self._cv.config(cursor="fleur")
        else:
            self._cv.config(cursor="crosshair")

    def _press(self, event):
        self._drag_start = (event.x, event.y)
        if self._roi and self._inside_roi(event.x, event.y):
            self._mode = "move"
            self._roi_at_drag = self._roi
        else:
            self._mode = "draw"
            self._roi_at_drag = None

    def _drag(self, event):
        if not self._drag_start: return
        sx, sy = self._drag_start
        dx = event.x - sx
        dy = event.y - sy

        if self._mode == "move" and self._roi_at_drag:
            dix = int(dx / self._scale)
            diy = int(dy / self._scale)
            ox0,oy0,ox1,oy1 = self._roi_at_drag
            ih, iw = self._rgb.shape[:2]
            w = ox1-ox0; h = oy1-oy0
            nx0 = max(0, min(ox0+dix, iw-w))
            ny0 = max(0, min(oy0+diy, ih-h))
            self._roi = (nx0, ny0, nx0+w, ny0+h)
            self._render()
            self._info.set(_L(f"이동 중  ({self._roi[0]},{self._roi[1]})", f"Moving  ({self._roi[0]},{self._roi[1]})"))

        elif self._mode == "draw":
            self._render()
            ix0,iy0 = self._c2i(sx, sy)
            ix1,iy1 = self._c2i(event.x, event.y)
            tmp = (min(ix0,ix1),min(iy0,iy1),max(ix0,ix1),max(iy0,iy1))
            self._draw_roi(tmp)
            self._info.set(
                _L(f"선택 중  {tmp[2]-tmp[0]}×{tmp[3]-tmp[1]}px",
                   f"Drawing  {tmp[2]-tmp[0]}×{tmp[3]-tmp[1]}px"))

    def _release(self, event):
        if not self._drag_start: return
        sx, sy = self._drag_start

        if self._mode == "draw":
            ix0,iy0 = self._c2i(sx, sy)
            ix1,iy1 = self._c2i(event.x, event.y)
            if abs(ix1-ix0) > 5 and abs(iy1-iy0) > 5:
                self._roi = (min(ix0,ix1),min(iy0,iy1),
                             max(ix0,ix1),max(iy0,iy1))
                self._render()
                w=self._roi[2]-self._roi[0]; h=self._roi[3]-self._roi[1]
                self._info.set(
                    _L(f"✔ {w}×{h}px  ({self._roi[0]},{self._roi[1]})  — [✔ 확정] 또는 [⊕ 전체복사]",
                       f"✔ {w}×{h}px  ({self._roi[0]},{self._roi[1]})  — [✔ Confirm] or [⊕ Copy All]"))

        elif self._mode == "move":
            w=self._roi[2]-self._roi[0]; h=self._roi[3]-self._roi[1]
            self._info.set(
                _L(f"✔ 이동완료  {w}×{h}px  ({self._roi[0]},{self._roi[1]})",
                   f"✔ Moved  {w}×{h}px  ({self._roi[0]},{self._roi[1]})"))

        self._mode = "idle"
        self._drag_start = None

    def _reset(self):
        self._roi = None
        self._images[self._cur_idx].pop("_pending_roi", None)
        self._render()
        self._info.set(_L(_L("ROI 초기화됨","ROI reset"),"ROI cleared"))

    def _confirm(self):
        if not self._roi:
            messagebox.showwarning(_L("주의","Warning"),_L("ROI를 먼저 선택한다.","Please select ROI first."),parent=self)
            return
        self._images[self._cur_idx]["_pending_roi"] = self._roi
        self._flush_pending()
        self.destroy()

    def _copy_all(self):
        """
        현재 ROI를 비율(0~1)로 변환해서 다른 이미지에 적용.
        이미지 크기/해상도가 달라도 화면상 같은 위치에 ROI 생성.
        """
        if not self._roi:
            messagebox.showwarning(_L("주의","Warning"),_L("ROI를 먼저 선택한다.","Please select ROI first."),parent=self)
            return
        self._images[self._cur_idx]["_pending_roi"] = self._roi
        if self._on_copy:
            self._on_copy(self._roi)
        self._flush_pending()
        self.destroy()

    def _flush_pending(self):
        for img in self._images:
            roi = img.pop("_pending_roi", None)
            if roi is not None:
                self._on_confirm(roi, img)


def styled_ax(ax, bg=PANEL):
    ax.set_facecolor(bg)
    for sp in ax.spines.values():
        sp.set_color(BORDER)
    ax.tick_params(colors=SUB)
    ax.grid(True, color=BORDER, linewidth=0.6, linestyle="--")


# ══════════════════════════════════════════════
#  메인 앱
# ══════════════════════════════════════════════
_Base = TkinterDnD.Tk if _DND else tk.Tk

class App(_Base):
    TW, TH = 140, 115

    def __init__(self):
        super().__init__()
        self.title(_L("HfS₂ 산화도 분석기  v5.0","HfS₂ Oxidation Analyzer  v5.0"))
        self.configure(bg=BG)
        self.geometry("1600x940")
        self.minsize(1300, 760)

        self.images   = []
        self.sel_idx  = -1
        self.rows_var = tk.IntVar(value=10)
        self.cols_var = tk.IntVar(value=10)
        self.ch_var   = tk.StringVar(value="S")
        self._refs    = {}

        # ── 다중 평가 대상 (Evaluation 탭) ──────────
        # 각 target dict 형식:
        #   {tid, name, rgb, roi, roi_flag, roi_reason, roi_source,
        #    color, thumb, cond_hint, result}
        self._pred_targets: list = []
        self._pred_sel_tid = None

        # ── 사용자 설정 임계값 ──────────────────────
        # 황색 판정 (HSI)
        self.cfg_h_lo    = tk.DoubleVar(value=35.0)   # H 하한 (°)
        self.cfg_h_hi    = tk.DoubleVar(value=75.0)   # H 상한 (°)
        self.cfg_s_thresh= tk.DoubleVar(value=0.10)   # S 최소 채도 (0~1)

        # S채널 색상 판정
        self.cfg_s_good  = tk.IntVar(value=80)        # S ≥ 이값 → 녹색(양호)
        self.cfg_s_warn  = tk.IntVar(value=40)        # S ≥ 이값 → 주황(경고)

        # YI 색상 판정
        self.cfg_yi_good = tk.IntVar(value=60)        # YI ≥ → 녹색
        self.cfg_yi_warn = tk.IntVar(value=35)        # YI ≥ → 주황

        # Lab b* 색상 판정
        self.cfg_b_good  = tk.IntVar(value=40)        # b* ≥ → 녹색
        self.cfg_b_warn  = tk.IntVar(value=20)        # b* ≥ → 주황

        # 날짜 추정 가중치 (합계 = 1.0)
        self.cfg_w_b     = tk.DoubleVar(value=0.45)   # b* 가중치
        self.cfg_w_s     = tk.DoubleVar(value=0.30)   # S 가중치
        self.cfg_w_yi    = tk.DoubleVar(value=0.25)   # YI 가중치

        # 세그먼트
        self.cfg_min_pix = tk.IntVar(value=10)        # 유효 세그먼트 최소 픽셀

        self._presets = [
            ("Native-35%", COND_COLORS[0]),
            ("Native-70%", COND_COLORS[1]),
            ("PMMA-70%",   COND_COLORS[2]),
            ("Al₂O₃-70%", COND_COLORS[3]),
        ]

        self._settings = load_settings()
        self._build()
        self.bind_all("<Control-v>", lambda e: self._paste())
        self.bind_all("<Control-V>", lambda e: self._paste())

    # ─────────────────────────────────────────
    #  다중 평가 대상 — property getter (READ only)
    # ─────────────────────────────────────────
    @property
    def _pred_rgb(self):
        """호환용 — 첫 target 의 rgb 반환. setter 정의 없음."""
        if self._pred_targets:
            return self._pred_targets[0].get("rgb")
        return None

    @property
    def _pred_roi(self):
        """호환용 — 첫 target 의 roi 반환. setter 정의 없음."""
        if self._pred_targets:
            return self._pred_targets[0].get("roi")
        return None

    # ─────────────────────────────────────────
    #  헬퍼
    # ─────────────────────────────────────────
    def _pred_max_tid(self) -> int:
        """현재 사용 중인 최대 tid 반환 (없으면 0)"""
        if not self._pred_targets:
            return 0
        return max(int(t.get("tid", 0)) for t in self._pred_targets)

    def _pred_get_target_by_tid(self, tid):
        """tid 로 target 검색 (없으면 None)"""
        for t in self._pred_targets:
            if t.get("tid") == tid:
                return t
        return None

    def _pred_assign_color(self) -> str:
        """현재 사용 중이 아닌 첫 빈 색상 반환"""
        used = {t.get("color") for t in self._pred_targets}
        for c in TARGET_COLOR_PALETTE:
            if c not in used:
                return c
        # 8개 모두 사용 중이면 PRED_MAX_TARGETS 가드로 도달 불가 (방어 코드)
        return TARGET_COLOR_PALETTE[len(self._pred_targets)
                                    % len(TARGET_COLOR_PALETTE)]

    # ─────────────────────────────────────────
    #  UI
    # ─────────────────────────────────────────
    def _build(self):
        # 툴바
        tb = tk.Frame(self, bg=PANEL2,
                      highlightbackground=BORDER, highlightthickness=1)
        tb.pack(fill="x")

        tk.Label(tb, text="⬡ HfS₂  v5",
                 bg=PANEL2, fg=ACCENT,
                 font=("Segoe UI",12,"bold")).pack(side="left", padx=16, pady=10)

        for txt,cmd,acc in [
            ("📂 Add",         self._load,             False),
            ("📋 Paste",       self._paste,            False),
            ("▶ Analyze All",  self._run_all,          True),
            ("🗄 DB Save",     self._db_save,          False),
            ("📂 DB Load",     self._db_load,          False),
            ("📦 Load All",    self._load_all_db,      False),
            ("💾 CSV",          self._export_csv,       False),
            ("📄 Report",      self._generate_report,  False),
            ("🗑 Clear",       self._clear,            False),
        ]:
            tk.Button(tb, text=txt, command=cmd,
                      bg=ACCENT if acc else BTN,
                      fg="white" if acc else TXT,
                      font=MF, relief="flat",
                      padx=10, pady=5, cursor="hand2",
                      activebackground=BTN_H).pack(
                      side="left", padx=3, pady=8)

        if _DND:
            tk.Label(tb, text=_L("  📂 드래그앤드롭 지원","  📂 Drag & Drop supported"),
                     bg=PANEL2, fg=TEAL, font=LF).pack(side="left", padx=8)

        # 파라미터
        pf = tk.Frame(tb, bg=PANEL2)
        pf.pack(side="right", padx=14)
        for lbl,var,lo,hi in [(_L("행","Rows"),self.rows_var,1,20),
                               (_L("열","Cols"),self.cols_var,1,20)]:
            tk.Label(pf,text=lbl,bg=PANEL2,fg=SUB,font=LF).pack(
                side="left",padx=(8,1))
            tk.Spinbox(pf,textvariable=var,from_=lo,to=hi,width=3,
                       bg=PANEL,fg=TXT,relief="flat",font=LF).pack(side="left")

        # 언어 전환 버튼
        lang_f = tk.Frame(tb, bg=PANEL2)
        lang_f.pack(side="right", padx=8)
        tk.Label(lang_f, text=_L("언어:","Lang:"), bg=PANEL2, fg=SUB,
                 font=LF).pack(side="left", padx=(4,2))
        self._lang_btn_ko = tk.Button(
            lang_f, text="한국어",
            command=lambda: self._set_lang(True),
            bg=ACCENT if _KO else BTN,
            fg="white" if _KO else TXT,
            font=("Segoe UI",8,"bold"),
            relief="flat", padx=6, pady=3, cursor="hand2")
        self._lang_btn_ko.pack(side="left", padx=1)
        self._lang_btn_en = tk.Button(
            lang_f, text="English",
            command=lambda: self._set_lang(False),
            bg=BTN if _KO else ACCENT,
            fg=TXT if _KO else "white",
            font=("Segoe UI",8,"bold"),
            relief="flat", padx=6, pady=3, cursor="hand2")
        self._lang_btn_en.pack(side="left", padx=1)

        # 본문
        body = tk.Frame(self, bg=BG)
        body.pack(fill="both", expand=True, padx=6, pady=5)

        # 좌: 이미지 목록 (사용자 요청 — 더 크게)
        left = tk.Frame(body, bg=PANEL,
                        highlightbackground=BORDER, highlightthickness=1,
                        width=440)
        left.pack(side="left", fill="y", padx=(0,5))
        left.pack_propagate(False)
        self._build_list_panel(left)

        # 우: 탭
        right = tk.Frame(body, bg=BG)
        right.pack(side="left", fill="both", expand=True)
        self._build_tabs(right)

        # 상태바
        self.sv = tk.StringVar(value=_L("이미지를 추가하고 ROI를 선택한 후 분석한다.","Add images, set ROI, then run analysis."))
        tk.Label(self, textvariable=self.sv,
                 bg=PANEL2, fg=SUB, font=LF,
                 anchor="w",
                 highlightbackground=BORDER,
                 highlightthickness=1).pack(
                 side="bottom", fill="x", padx=0, pady=0)

        if _DND:
            self.drop_target_register(DND_FILES)
            self.dnd_bind("<<Drop>>",      self._on_drop)
            self.dnd_bind("<<DragEnter>>",
                lambda e: self.sv.set(_L("📂 여기에 파일을 놓으세요!","📂 Drop files here!")))
            self.dnd_bind("<<DragLeave>>",
                lambda e: self.sv.set(""))

    # ─────────────────────────────────────────
    #  이미지 목록 패널
    # ─────────────────────────────────────────
    def _build_list_panel(self, parent):
        # 헤더
        hf = tk.Frame(parent, bg=PANEL2,
                      highlightbackground=BORDER, highlightthickness=1)
        hf.pack(fill="x")
        tk.Label(hf, text=_L("  📋 이미지 목록","  📋 Image List"),
                 bg=PANEL2, fg=TXT, font=MFB).pack(side="left",pady=6,padx=8)

        # 자동 ROI 재실행 (사용자가 직접 설정한 것은 보호)
        tk.Button(hf, text=_L("🤖 자동 ROI","🤖 Auto ROI"),
                  command=self._auto_roi_all_unmanual,
                  bg=PANEL2, fg=ACCENT, font=("Segoe UI",8,"bold"),
                  relief="flat", padx=8, pady=2, cursor="hand2",
                  activebackground=PANEL).pack(side="left", padx=2, pady=4)
        # 전체 삭제 버튼 (Raman 데이터는 보존)
        tk.Button(hf, text=_L("🗑 전체 삭제","🗑 Delete All"),
                  command=self._delete_all_images,
                  bg=PANEL2, fg=RED, font=("Segoe UI",8,"bold"),
                  relief="flat", padx=8, pady=2, cursor="hand2",
                  activebackground=PANEL).pack(side="left", padx=2, pady=4)

        self._roi_stat = tk.StringVar(value="ROI 0/0")
        self._roi_ok_lbl = tk.Label(hf, textvariable=self._roi_stat,
                                    bg=PANEL2, fg=AMBER, font=MFB)
        self._roi_ok_lbl.pack(side="right", padx=10)

        # 안내
        gf = tk.Frame(parent, bg=CARD2,
                      highlightbackground=BORDER, highlightthickness=1)
        gf.pack(fill="x", padx=6, pady=(6,2))

        tk.Label(gf, text=_L("💡 사용법","💡 How to use"),
                 bg=CARD2, fg=ACCENT, font=MFB).pack(anchor="w", padx=8, pady=(5,1))
        tk.Label(gf,
                 text=_L("1. 이미지 추가 → 자동 ROI 적용 (DB 로드는 ROI 보존)\n"
                         "2. 카드 테두리 색으로 품질 확인:\n"
                         "   🟢 녹색=OK  🟠 주황=검토  🟣 보라=조건 불일치  🔴 빨강=실패\n"
                         "3. 주황·보라·빨강만 🎯 눌러 수동 보정\n"
                         "4. Run [▶ Analyze All]",
                         "1. Add images → auto ROI (DB load preserves ROI)\n"
                         "2. Card border = quality:\n"
                         "   🟢 OK  🟠 review  🟣 group mismatch  🔴 failed\n"
                         "3. Fix flagged ones with 🎯\n"
                         "4. Run [▶ Analyze All]"),
                 bg=CARD2, fg=TXT, font=LF,
                 justify="left").pack(anchor="w", padx=10, pady=(0,2))
        tk.Label(gf, text=_L("✦ 파일명: Nday_MRH_조건명.jpg → 자동인식","✦ Filename: Nday_MRH_Cond.jpg → auto-parsed"),
                 bg=CARD2, fg=TEAL,
                 font=("Segoe UI",7)).pack(anchor="w", padx=10, pady=(0,5))

        # 프리셋
        tk.Label(gf, text=_L("조건 프리셋:","Condition Presets:"),
                 bg=CARD2, fg=SUB, font=LF).pack(anchor="w", padx=10)
        pf2 = tk.Frame(gf, bg=CARD2)
        pf2.pack(fill="x", padx=8, pady=(2,6))
        for lbl,col in self._presets:
            tk.Button(pf2, text=lbl,
                      command=lambda l=lbl: self._apply_preset(l),
                      bg=BTN, fg=col, font=("Segoe UI",7,"bold"),
                      relief="flat", cursor="hand2",
                      padx=3, pady=2).pack(side="left", padx=1)

        # 스크롤 목록
        sc = tk.Frame(parent, bg=PANEL)
        sc.pack(fill="both", expand=True, padx=4, pady=4)
        self._lc = tk.Canvas(sc, bg=PANEL, highlightthickness=0)
        vsb = tk.Scrollbar(sc, orient="vertical",
                           command=self._lc.yview)
        self._lc.configure(yscrollcommand=vsb.set)
        vsb.pack(side="right", fill="y")
        self._lc.pack(fill="both", expand=True)
        self._lf = tk.Frame(self._lc, bg=PANEL)
        _w = self._lc.create_window((0,0), window=self._lf, anchor="nw")
        self._lf.bind("<Configure>",
            lambda e: self._lc.configure(
                scrollregion=self._lc.bbox("all")))
        self._lc.bind("<Configure>",
            lambda e: self._lc.itemconfig(_w, width=e.width))

    # ─────────────────────────────────────────
    #  탭
    # ─────────────────────────────────────────
    def _build_tabs(self, parent):
        tbar = tk.Frame(parent, bg=PANEL2,
                        highlightbackground=BORDER,
                        highlightthickness=1)
        tbar.pack(fill="x")

        self._tbts = {}
        self._tfs  = {}
        self._atab = tk.StringVar(value="detail")

        for key,lbl in [
                ("detail",   _L("🔍 ROI·상세",    "🔍 ROI·Detail")),
                ("compare",  _L("🔲 조건 비교",    "🔲 Condition Grid")),
                ("chart",    _L("📈 차트",           "📈 Charts")),
                ("color",    _L("🔬 컬러 분석",     "🔬 Color Analysis")),
                ("raman",    _L("📡 Raman 분석",    "📡 Raman Analysis")),
                ("predict",  _L("🎯 평가 대상",     "🎯 Evaluation")),
                ("advanced", _L("🧪 고급 분석",     "🧪 Advanced")),
                ("settings", _L("⚙ 설정",            "⚙ Settings"))]:
            b = tk.Button(tbar, text=lbl,
                          command=lambda k=key: self._switch(k),
                          bg=ACCENT if key=="detail" else PANEL2,
                          fg="white" if key=="detail" else TXT,
                          font=MF, relief="flat",
                          padx=11, pady=8, cursor="hand2")
            b.pack(side="left")
            self._tbts[key] = b
            f = tk.Frame(parent, bg=BG)
            self._tfs[key] = f

        self._tfs["detail"].pack(fill="both", expand=True)
        self._build_detail()
        self._build_compare()
        self._build_chart()
        self._build_color_tab()
        self._build_raman_tab()
        self._build_predict_tab()
        self._build_advanced_tab()
        self._build_settings_tab()

    def _switch(self, key):
        for k,f in self._tfs.items(): f.pack_forget()
        self._tfs[key].pack(fill="both", expand=True)
        for k,b in self._tbts.items():
            b.configure(bg=ACCENT if k==key else PANEL2,
                        fg="white" if k==key else TXT)
        self._atab.set(key)
        if key=="detail":  self._refresh_orig(); self._refresh_hsi()
        elif key=="compare": self._refresh_compare()
        elif key=="chart":   self._refresh_charts()
        elif key=="color":   self._refresh_color_tab()
        elif key=="raman":   self._refresh_raman_tab()
        elif key=="predict": pass
        elif key=="advanced": self._adv_on_switch()
        elif key=="settings": pass

    # ─────────────────────────────────────────
    #  상세 탭
    # ─────────────────────────────────────────
    def _build_detail(self):
        f = self._tfs["detail"]

        # 좌: 원본 + ROI
        left = tk.Frame(f, bg=PANEL,
                        highlightbackground=BORDER, highlightthickness=1)
        left.pack(side="left", fill="both", expand=True, padx=(0,4))

        lhdr = tk.Frame(left, bg=PANEL2,
                        highlightbackground=BORDER, highlightthickness=1)
        lhdr.pack(fill="x")
        tk.Label(lhdr, text=_L("  🖼 원본  —  ROI 미리보기","  🖼 Original  —  ROI Preview"),
                 bg=PANEL2, fg=TXT, font=MFB).pack(side="left", pady=6, padx=8)

        self._roi_info = tk.StringVar(value=_L("ROI 미선택","No ROI selected"))
        tk.Label(lhdr, textvariable=self._roi_info,
                 bg=PANEL2, fg=AMBER, font=LF).pack(side="left", padx=6)

        tk.Button(lhdr, text=_L("🎯 ROI 선택","🎯 Set ROI"),
                  command=self._open_roi,
                  bg=ACCENT, fg="white", font=MFB,
                  relief="flat", padx=10, pady=4,
                  cursor="hand2").pack(side="right", padx=6, pady=5)
        tk.Button(lhdr, text=_L("✕ ROI 제거","✕ Clear ROI"),
                  command=self._clear_roi,
                  bg=BTN, fg=RED, font=MF,
                  relief="flat", padx=8, pady=4,
                  cursor="hand2").pack(side="right", padx=2)

        self._orig_cv = tk.Canvas(left, bg=PANEL,
                                  highlightthickness=0,
                                  height=200)   # 고정 높이로 축소
        self._orig_cv.pack(fill="x", padx=4, pady=4)
        self._orig_cv.bind("<Configure>", lambda e: self._refresh_orig())

        # 우: PanedWindow — 채널뷰/세그먼트/지표요약/히트맵 드래그 크기 조절
        right_outer = tk.Frame(f, bg=BG, width=500)
        right_outer.pack(side="right", fill="both")
        right_outer.pack_propagate(False)

        # 채널 선택 (고정 영역 — PanedWindow 밖)
        ch_bar = tk.Frame(right_outer, bg=PANEL2,
                          highlightbackground=BORDER, highlightthickness=1)
        ch_bar.pack(fill="x")
        row1 = tk.Frame(ch_bar, bg=PANEL2)
        row1.pack(fill="x", padx=6, pady=(4,1))
        tk.Label(row1, text="HSI:",
                 bg=PANEL2, fg=SUB, font=LF, width=5,
                 anchor="w").pack(side="left")
        for ch, col in [("H", TEAL), ("S", PURPLE), ("I", AMBER)]:
            tk.Radiobutton(
                row1, text=ch, variable=self.ch_var, value=ch,
                bg=PANEL2, fg=col, selectcolor=PANEL2,
                activebackground=PANEL2, font=MFB,
                command=self._refresh_hsi
            ).pack(side="left", padx=4)

        row2 = tk.Frame(ch_bar, bg=PANEL2)
        row2.pack(fill="x", padx=6, pady=(1,4))
        tk.Label(row2, text=_L("Lab/기타:","Lab/Other:"),
                 bg=PANEL2, fg=SUB, font=LF, width=5,
                 anchor="w").pack(side="left")
        for ch, col, tip in [
            ("b*",  "#d97706", "Lab b*(황색도)"),
            ("L*",  "#6b7280", "Lab L*(밝기)"),
            ("ΔE",  "#dc2626", "색차(0일기준)"),
            ("YI",  "#16a34a", "황색도지수"),
        ]:
            tk.Radiobutton(
                row2, text=ch, variable=self.ch_var, value=ch,
                bg=PANEL2, fg=col, selectcolor=PANEL2,
                activebackground=PANEL2, font=MFB,
                command=self._refresh_hsi
            ).pack(side="left", padx=4)

        # PanedWindow (수직 방향, 드래그 크기 조절)
        right_pw = ttk.PanedWindow(right_outer, orient="vertical")
        right_pw.pack(fill="both", expand=True)

        # ── 패널1: 채널 뷰 ─────────────────────
        p1 = tk.Frame(right_pw, bg=PANEL,
                      highlightbackground=BORDER, highlightthickness=1)
        right_pw.add(p1, weight=2)
        hsi_hdr = tk.Frame(p1, bg=PANEL2,
                           highlightbackground=BORDER, highlightthickness=1)
        hsi_hdr.pack(fill="x")
        tk.Label(hsi_hdr, text=_L("  📊 채널 뷰 (더블클릭: 확대)","  📊 Channel View (dbl-click: enlarge)"),
                 bg=PANEL2, fg=TXT, font=MFB).pack(side="left", pady=3, padx=8)
        tk.Label(hsi_hdr, text=_L("⇕ 경계 드래그로 높이 조절","⇕ Drag border to resize"),
                 bg=PANEL2, fg=SUB, font=("Segoe UI",6)).pack(side="right", padx=6)
        self._hsi_cv = tk.Canvas(p1, bg=PANEL, highlightthickness=0)
        self._hsi_cv.pack(fill="both", expand=True, padx=4, pady=4)
        self._hsi_cv.bind("<Configure>", lambda e: self._refresh_hsi())
        self._hsi_cv.bind("<Double-Button-1>", lambda e: self._hsi_popup())

        # ── 패널2: 세그먼트 통계 ───────────────
        p2 = tk.Frame(right_pw, bg=PANEL,
                      highlightbackground=BORDER, highlightthickness=1)
        right_pw.add(p2, weight=2)
        seg_hdr = tk.Frame(p2, bg=PANEL2,
                           highlightbackground=BORDER, highlightthickness=1)
        seg_hdr.pack(fill="x")
        tk.Label(seg_hdr, text=_L("  세그먼트 통계","  Segment Stats"),
                 bg=PANEL2, fg=TXT, font=MFB).pack(side="left", pady=4, padx=8)
        tk.Label(seg_hdr,
                 text=_L("S평균/편차 · b* · YI · 황색% · ΔE · L*","S-avg/std · b* · YI · YR% · ΔE · L*"),
                 bg=PANEL2, fg=SUB, font=("Segoe UI",6)).pack(side="left", padx=2)

        seg_tree_f = tk.Frame(p2, bg=PANEL)
        seg_tree_f.pack(fill="both", expand=True, padx=2, pady=2)
        cols = ("seg","mean","std","b","yi","yr","de","L")
        self.tree = ttk.Treeview(seg_tree_f, columns=cols,
                                 show="headings")
        for c,h,w in [
            ("seg","Seg",34),("mean","S-avg",52),("std","S-std",48),
            ("b","b*",48),("yi","YI",44),("yr","YR%",50),
            ("de","ΔE",44),("L","L*",44),
        ]:
            self.tree.heading(c, text=h)
            self.tree.column(c, width=w, anchor="center")
        sty = ttk.Style()
        sty.theme_use("clam")
        sty.configure("Treeview",
                      background=PANEL, foreground=TXT,
                      fieldbackground=PANEL, rowheight=18, font=LF)
        sty.configure("Treeview.Heading",
                      background=PANEL2, foreground=SUB, font=LF)
        sty.map("Treeview", background=[("selected", ACCENT)],
                foreground=[("selected","white")])
        seg_vsb = tk.Scrollbar(seg_tree_f, orient="vertical",
                               command=self.tree.yview)
        self.tree.configure(yscrollcommand=seg_vsb.set)
        seg_vsb.pack(side="right", fill="y")
        self.tree.pack(fill="both", expand=True)

        # ── 패널3: 지표 요약 ───────────────────
        p3 = tk.Frame(right_pw, bg=PANEL,
                      highlightbackground=BORDER, highlightthickness=1)
        right_pw.add(p3, weight=3)
        stat_hdr = tk.Frame(p3, bg=PANEL2,
                            highlightbackground=BORDER, highlightthickness=1)
        stat_hdr.pack(fill="x")
        tk.Label(stat_hdr, text=_L("  📋 지표 요약","  📋 Metric Summary"),
                 bg=PANEL2, fg=TXT, font=MFB).pack(side="left", pady=4, padx=8)
        tk.Button(stat_hdr, text="⤢",
                  command=lambda: self._stat_popup(),
                  bg=BTN, fg=TXT, font=LF,
                  relief="flat", padx=6, pady=2,
                  cursor="hand2").pack(side="right", padx=2, pady=4)
        tk.Button(stat_hdr, text="📋",
                  command=lambda: self._copy_stat_text(),
                  bg=BTN, fg=TXT, font=LF,
                  relief="flat", padx=6, pady=2,
                  cursor="hand2").pack(side="right", padx=2, pady=4)

        stat_scroll_f = tk.Frame(p3, bg=PANEL)
        stat_scroll_f.pack(fill="both", expand=True, padx=2, pady=2)
        stat_vsb = tk.Scrollbar(stat_scroll_f, orient="vertical")
        self._stat_text = tk.Text(
            stat_scroll_f, wrap="word",
            bg=CARD2, fg=TXT,
            font=("Segoe UI",8),
            relief="flat", padx=8, pady=6,
            highlightbackground=BORDER,
            highlightthickness=1,
            cursor="xterm",
            yscrollcommand=stat_vsb.set,
            state="disabled")
        stat_vsb.configure(command=self._stat_text.yview)
        stat_vsb.pack(side="right", fill="y")
        self._stat_text.pack(fill="both", expand=True)
        self._stat_text.bind("<Control-c>",
            lambda e: self._copy_text(self._stat_text))
        self._stat_text.bind("<Control-C>",
            lambda e: self._copy_text(self._stat_text))

        # ── 패널4: 히트맵 ──────────────────────
        p4 = tk.Frame(right_pw, bg=PANEL,
                      highlightbackground=BORDER, highlightthickness=1)
        right_pw.add(p4, weight=2)
        hm_hdr = tk.Frame(p4, bg=PANEL2,
                          highlightbackground=BORDER, highlightthickness=1)
        hm_hdr.pack(fill="x")
        tk.Label(hm_hdr, text=_L("  히트맵 (이미지×세그먼트)",
                                  "  Heatmap (Image×Segment)"),
                 bg=PANEL2, fg=TXT, font=MFB).pack(side="left", pady=4, padx=8)
        self._hm_metric = tk.StringVar(value="S")
        for lbl, val, col in [
            ("S", "S",  PURPLE),
            ("b*","b",  "#d97706"),
            ("YI","YI", GREEN),
            ("ΔE","dE", RED),
        ]:
            tk.Radiobutton(
                hm_hdr, text=lbl, variable=self._hm_metric, value=val,
                bg=PANEL2, fg=col, selectcolor=PANEL2,
                activebackground=PANEL2, font=LF,
                command=self._update_heatmap
            ).pack(side="left", padx=3)

        self.hm_fig = plt.Figure(figsize=(4.5, 2.6), facecolor=PANEL)
        self.hm_cv  = FigureCanvasTkAgg(self.hm_fig, master=p4)
        self.hm_cv.get_tk_widget().pack(
            fill="both", expand=True, padx=2, pady=2)

    # ─────────────────────────────────────────
    #  비교 탭
    # ─────────────────────────────────────────
    def _build_compare(self):
        f = self._tfs["compare"]

        # ── 범례 패널 (상단 고정) ──────────────────
        legend_f = tk.Frame(f, bg=PANEL2,
                            highlightbackground=BORDER, highlightthickness=1)
        legend_f.pack(fill="x")
        tk.Label(legend_f, text=_L("  📋 표시 항목 설명:","  📋 Legend:"),
                 bg=PANEL2, fg=TXT, font=MFB).pack(side="left", pady=5, padx=8)

        legend_items = [
            ("S=",   PURPLE,
             _L("S채널 평균 (HSI 채도, 0~255). 높을수록 색 진함/미산화",
                "S-ch avg (HSI Saturation, 0-255). Higher=more colored/pristine")),
            ("Y%=",  GOLD,
             _L("황색 픽셀 비율 (H:35~75°·S≥임계값 기준, 0~100%)",
                "Yellow pixel ratio (H:35~75°·S≥thresh, 0-100%)")),
            ("YI=",  GREEN,
             _L("Yellowness Index (ASTM E313: 100×(1.28R-1.06B)/G)",
                "Yellowness Index (ASTM E313: 100×(1.28R-1.06B)/G)")),
            ("b*=",  "#d97706",
             _L("Lab b* 황색도 (+양수=황색, 0=무채색). 핵심지표★",
                "Lab b* yellowness (+pos=yellow, 0=achromatic). Key★")),
            ("ΔE=",  RED,
             _L("0일차 기준 CIE76 색차 (<3 작음, >10 큰 변화)",
                "CIE76 ΔE from day-0 (<3 minor, >10 major)")),
        ]
        for prefix, col, tip in legend_items:
            item_f = tk.Frame(legend_f, bg=PANEL2)
            item_f.pack(side="left", padx=6)
            tk.Label(item_f, text=prefix, bg=PANEL2, fg=col,
                     font=("Segoe UI",8,"bold")).pack(side="left")
            tk.Label(item_f, text=tip, bg=PANEL2, fg=SUB,
                     font=("Segoe UI",7),
                     wraplength=160, justify="left").pack(side="left")

        # 색상 기준
        color_f = tk.Frame(legend_f, bg=PANEL2)
        color_f.pack(side="right", padx=8)
        for lbl, col in [
            (_L("■ 미산화","■ Pristine"), GREEN),
            (_L("■ 경계","■ Boundary"),   AMBER),
            (_L("■ 산화","■ Oxidized"),   RED),
        ]:
            tk.Label(color_f, text=lbl, bg=PANEL2, fg=col,
                     font=("Segoe UI",7,"bold")).pack(side="left", padx=4)

        # ── 비교 그리드 (스크롤) ───────────────────
        tk.Label(f, text=_L("  조건 × 날짜 비교 그리드  (더블클릭: 상세)",
                             "  Condition × Day Grid  (dbl-click: detail)"),
                 bg=BG, fg=SUB, font=MFB).pack(anchor="w", pady=(4,2))
        outer = tk.Frame(f, bg=BG)
        outer.pack(fill="both", expand=True)
        self._cmp_cv = tk.Canvas(outer, bg=BG, highlightthickness=0)
        hsc = tk.Scrollbar(outer, orient="horizontal",
                           command=self._cmp_cv.xview)
        vsc = tk.Scrollbar(outer, orient="vertical",
                           command=self._cmp_cv.yview)
        self._cmp_cv.configure(xscrollcommand=hsc.set,
                               yscrollcommand=vsc.set)
        hsc.pack(side="bottom", fill="x")
        vsc.pack(side="right", fill="y")
        self._cmp_cv.pack(fill="both", expand=True)
        self._cmp_fr = tk.Frame(self._cmp_cv, bg=BG)
        _w = self._cmp_cv.create_window(
            (0,0), window=self._cmp_fr, anchor="nw")
        self._cmp_fr.bind("<Configure>",
            lambda e: self._cmp_cv.configure(
                scrollregion=self._cmp_cv.bbox("all")))

    # ─────────────────────────────────────────
    #  1. DB 저장 / 로드
    # ─────────────────────────────────────────
    # ─────────────────────────────────────────
    #  통합 DB 로드 (이미지 + Raman + 평가대상 한꺼번에)
    # ─────────────────────────────────────────
    def _generate_report(self):
        """보고서 생성 — 영문/한글 동시, 워드(.docx) 또는 HTML 선택"""
        fmt_win = tk.Toplevel(self)
        fmt_win.title("Select Report Format")
        fmt_win.configure(bg=PANEL)
        fmt_win.geometry("340x170")
        fmt_win.grab_set()
        chosen = [None]

        tk.Label(fmt_win, text="Select output format:",
                 bg=PANEL, fg=TXT, font=MFB).pack(pady=(20,10))
        bf = tk.Frame(fmt_win, bg=PANEL)
        bf.pack()

        def pick(fmt):
            chosen[0] = fmt
            fmt_win.destroy()

        tk.Button(bf, text="📝  Word (.docx)",
                  command=lambda: pick("docx"),
                  bg=ACCENT, fg="white", font=MF,
                  relief="flat", padx=14, pady=7,
                  cursor="hand2").pack(side="left", padx=6)
        tk.Button(bf, text="🌐  HTML",
                  command=lambda: pick("html"),
                  bg=BTN, fg=TXT, font=MF,
                  relief="flat", padx=14, pady=7,
                  cursor="hand2").pack(side="left", padx=6)

        tk.Label(fmt_win,
                 text="영문(EN) + 한글(KO) 두 파일이 함께 생성된다.",
                 bg=PANEL, fg=SUB,
                 font=("Segoe UI",8)).pack(pady=(8,4))

        self.wait_window(fmt_win)
        if chosen[0] is None:
            return

        fmt = chosen[0]
        ext = ".docx" if fmt == "docx" else ".html"
        ftypes = ([("Word Document","*.docx")] if fmt == "docx"
                  else [("HTML Report","*.html")])

        path_en = filedialog.asksaveasfilename(
            title="Save English Report As",
            defaultextension=ext,
            filetypes=ftypes + [("All","*.*")])
        if not path_en:
            return

        # 한글 파일 경로: 같은 폴더, _KO 접미사
        base, ext2 = os.path.splitext(path_en)
        path_ko = base + "_KO" + ext2

        self._set_status("📄 Generating reports (EN + KO)...")
        self.update_idletasks()

        try:
            # ── 공통 데이터 수집 ──────────────────
            analyzed = []
            for img in self.images:
                if img.get("s_mean") is None:
                    continue
                th = img.get("thumb")
                tb64 = ""
                if th is not None:
                    try:
                        buf = io.BytesIO()
                        th.save(buf, format="PNG")
                        tb64 = base64.b64encode(buf.getvalue()).decode()
                    except Exception:
                        pass
                def fv(v, d=2):
                    try: return f"{float(v):.{d}f}"
                    except: return "-"
                analyzed.append({
                    "name": img.get("name",""),
                    "cond": img.get("cond",""),
                    "day":  img.get("day",""),
                    "b":  fv(img.get("lab",{}).get("b")),
                    "s":  fv(img.get("s_mean"),0),
                    "yi": fv(img.get("yellowness_idx"),0),
                    "de": fv(img.get("delta_e")),
                    "yr": fv(float(img.get("yellow_ratio",0))*100,1)+"%",
                    "tb64": tb64,
                })

            raman_data = getattr(self, "_raman_data", [])
            eval_ctx   = getattr(self, "_last_eval_ctx", {}) or {}

            eval_comment = pseudo_result = ""
            for attr, name in [("_pred_comment","ec"),
                                ("_pseudo_res_txt","pr")]:
                if hasattr(self, attr):
                    try:
                        w = getattr(self, attr)
                        w.configure(state="normal")
                        val = w.get("1.0","end").strip()
                        w.configure(state="disabled")
                        if name == "ec": eval_comment = val
                        else: pseudo_result = val
                    except Exception:
                        pass

            chart_b64 = {}
            chart_figs = {}
            for k, fd in getattr(self, "_pred_figs", {}).items():
                if fd.get("fig"): chart_figs[k] = fd["fig"]
            for k, attr in [("pseudo_reg_fig","_pseudo_reg_fig"),
                             ("pseudo_spec_fig","_pseudo_spec_fig")]:
                fig = getattr(self, attr, None)
                if fig: chart_figs[k] = fig
            for k, cell in getattr(self, "_raman_charts", {}).items():
                if cell.get("fig"): chart_figs[k] = cell["fig"]

            for k, fig in chart_figs.items():
                try:
                    buf = io.BytesIO()
                    fig.savefig(buf, format="png", dpi=180,
                                facecolor=PANEL, bbox_inches="tight")
                    chart_b64[k] = base64.b64encode(buf.getvalue()).decode()
                except Exception:
                    pass

            gen_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
            args = (analyzed, raman_data, eval_ctx,
                    chart_b64, eval_comment, pseudo_result, gen_time)

            # ── 영문 생성 ─────────────────────────
            if fmt == "html":
                with open(path_en, "w", encoding="utf-8") as f:
                    f.write(self._build_report_html(*args, lang="en"))
                with open(path_ko, "w", encoding="utf-8") as f:
                    f.write(self._build_report_html(*args, lang="ko"))
                import webbrowser
                webbrowser.open(path_ko)   # 한글 버전 먼저 열기
            else:
                self._build_report_docx(path_en, *args, lang="en")
                self._build_report_docx(path_ko, *args, lang="ko")

            self._set_status(f"📄 Reports saved (EN + KO)")
            tip = "\n\nTip: 브라우저 인쇄 → PDF로 저장 가능." if fmt == "html" else ""
            messagebox.showinfo("보고서 생성 완료",
                f"영문: {os.path.basename(path_en)}\n"
                f"한글: {os.path.basename(path_ko)}\n\n"
                f"  이미지 {len(analyzed)}개 | Raman {len(raman_data)}건 "
                f"| 차트 {len(chart_b64)}개" + tip)

        except Exception as ex:
            self._set_status("Report generation failed.")
            messagebox.showerror("Report Error", f"Failed:\n{ex}")


    def _report_flowchart_svg(self, ko=False) -> str:
        """분석 전체 흐름을 SVG 순서도로 반환"""
        svg = """
<div style="overflow-x:auto;margin:16px 0">
<svg viewBox="0 0 820 780" xmlns="http://www.w3.org/2000/svg"
     style="width:100%;max-width:820px;font-family:Arial,sans-serif">
  <defs>
    <marker id="arr" markerWidth="10" markerHeight="7"
            refX="9" refY="3.5" orient="auto">
      <polygon points="0 0,10 3.5,0 7" fill="#2e75b6"/>
    </marker>
    <marker id="arr2" markerWidth="10" markerHeight="7"
            refX="9" refY="3.5" orient="auto">
      <polygon points="0 0,10 3.5,0 7" fill="#2e8b57"/>
    </marker>
  </defs>

  <!-- ── STEP 1: 입력 ── -->
  <rect x="260" y="10" width="300" height="52" rx="26" ry="26"
        fill="#1f3864" stroke="none"/>
  <text x="410" y="30" text-anchor="middle" fill="white" font-size="13" font-weight="bold">INPUT</text>
  <text x="410" y="50" text-anchor="middle" fill="#aac4e8" font-size="11">Target Image  +  Reference Image Set</text>

  <!-- arrow -->
  <line x1="410" y1="62" x2="410" y2="92"
        stroke="#2e75b6" stroke-width="2" marker-end="url(#arr)"/>

  <!-- ── STEP 2: ROI / 전처리 ── -->
  <rect x="280" y="92" width="260" height="48" rx="8"
        fill="#2e75b6" stroke="none"/>
  <text x="410" y="110" text-anchor="middle" fill="white" font-size="12" font-weight="bold">Region of Interest (ROI)</text>
  <text x="410" y="128" text-anchor="middle" fill="#d0e8ff" font-size="10">User-defined crop  |  Full image fallback</text>

  <!-- arrow -->
  <line x1="410" y1="140" x2="410" y2="170"
        stroke="#2e75b6" stroke-width="2" marker-end="url(#arr)"/>

  <!-- ── STEP 3: 색상 지표 추출 (분기 3개) ── -->
  <rect x="240" y="170" width="340" height="48" rx="8"
        fill="#1a6496" stroke="none"/>
  <text x="410" y="189" text-anchor="middle" fill="white" font-size="12" font-weight="bold">Colorimetric Feature Extraction</text>
  <text x="410" y="207" text-anchor="middle" fill="#d0e8ff" font-size="10">CIE Lab (b*)  |  HSI S-ch  |  YI  |  ΔE</text>

  <!-- 3갈래 화살표 -->
  <line x1="410" y1="218" x2="410" y2="238"
        stroke="#2e75b6" stroke-width="1.5"/>
  <line x1="410" y1="238" x2="165" y2="238"
        stroke="#2e75b6" stroke-width="1.5"/>
  <line x1="410" y1="238" x2="655" y2="238"
        stroke="#2e75b6" stroke-width="1.5"/>
  <line x1="165" y1="238" x2="165" y2="265"
        stroke="#2e75b6" stroke-width="1.5" marker-end="url(#arr)"/>
  <line x1="410" y1="238" x2="410" y2="265"
        stroke="#2e75b6" stroke-width="1.5" marker-end="url(#arr)"/>
  <line x1="655" y1="238" x2="655" y2="265"
        stroke="#2e75b6" stroke-width="1.5" marker-end="url(#arr)"/>

  <!-- 지표 박스 3개 -->
  <rect x="80"  y="265" width="170" height="44" rx="6" fill="#0f4c8a" stroke="none"/>
  <text x="165" y="283" text-anchor="middle" fill="white" font-size="11" font-weight="bold">Lab b*  /  ΔE</text>
  <text x="165" y="300" text-anchor="middle" fill="#aac4e8" font-size="9">Yellow–White axis</text>

  <rect x="325" y="265" width="170" height="44" rx="6" fill="#0f4c8a" stroke="none"/>
  <text x="410" y="283" text-anchor="middle" fill="white" font-size="11" font-weight="bold">S-channel</text>
  <text x="410" y="300" text-anchor="middle" fill="#aac4e8" font-size="9">Chromatic purity</text>

  <rect x="570" y="265" width="170" height="44" rx="6" fill="#0f4c8a" stroke="none"/>
  <text x="655" y="283" text-anchor="middle" fill="white" font-size="11" font-weight="bold">Yellowness Index</text>
  <text x="655" y="300" text-anchor="middle" fill="#aac4e8" font-size="9">ASTM E313 scalar</text>

  <!-- 합류 -->
  <line x1="165" y1="309" x2="165" y2="330" stroke="#2e75b6" stroke-width="1.5"/>
  <line x1="410" y1="309" x2="410" y2="330" stroke="#2e75b6" stroke-width="1.5"/>
  <line x1="655" y1="309" x2="655" y2="330" stroke="#2e75b6" stroke-width="1.5"/>
  <line x1="165" y1="330" x2="655" y2="330" stroke="#2e75b6" stroke-width="1.5"/>
  <line x1="410" y1="330" x2="410" y2="355" stroke="#2e75b6" stroke-width="2" marker-end="url(#arr)"/>

  <!-- ── STEP 4: 거리 계산 ── -->
  <rect x="240" y="355" width="340" height="48" rx="8"
        fill="#155724" stroke="none"/>
  <text x="410" y="374" text-anchor="middle" fill="white" font-size="12" font-weight="bold">Weighted Euclidean Distance</text>
  <text x="410" y="392" text-anchor="middle" fill="#b8f0c8" font-size="10">d = √(w_b·Δb*² + w_S·ΔS² + w_YI·ΔYI²)</text>

  <!-- 2갈래 -->
  <line x1="410" y1="403" x2="410" y2="423" stroke="#2e8b57" stroke-width="1.5"/>
  <line x1="410" y1="423" x2="220" y2="423" stroke="#2e8b57" stroke-width="1.5"/>
  <line x1="410" y1="423" x2="600" y2="423" stroke="#2e8b57" stroke-width="1.5"/>
  <line x1="220" y1="423" x2="220" y2="448" stroke="#2e8b57" stroke-width="1.5" marker-end="url(#arr2)"/>
  <line x1="600" y1="423" x2="600" y2="448" stroke="#2e8b57" stroke-width="1.5" marker-end="url(#arr2)"/>

  <!-- ── STEP 5A: 날짜 추정 ── -->
  <rect x="100" y="448" width="240" height="80" rx="8"
        fill="#1e4d2b" stroke="#2e8b57" stroke-width="1.5"/>
  <text x="220" y="468" text-anchor="middle" fill="white" font-size="12" font-weight="bold">Date Estimation</text>
  <text x="220" y="486" text-anchor="middle" fill="#b8f0c8" font-size="9">Top-k matching</text>
  <text x="220" y="500" text-anchor="middle" fill="#b8f0c8" font-size="9">Inv-dist weighted avg</text>
  <text x="220" y="516" text-anchor="middle" fill="#ffd700" font-size="10" font-weight="bold">→ Est. Day + Confidence</text>

  <!-- ── STEP 5B: Pseudo-Raman ── -->
  <rect x="460" y="448" width="240" height="80" rx="8"
        fill="#1e4d2b" stroke="#2e8b57" stroke-width="1.5"/>
  <text x="580" y="468" text-anchor="middle" fill="white" font-size="12" font-weight="bold">Pseudo-Raman</text>
  <text x="580" y="486" text-anchor="middle" fill="#b8f0c8" font-size="9">4-metric regression ensemble</text>
  <text x="580" y="500" text-anchor="middle" fill="#b8f0c8" font-size="9">R²-weighted averaging</text>
  <text x="580" y="516" text-anchor="middle" fill="#ffd700" font-size="10" font-weight="bold">→ A₁g peak + 95% CI</text>

  <!-- 합류 -->
  <line x1="220" y1="528" x2="220" y2="560" stroke="#2e8b57" stroke-width="1.5"/>
  <line x1="580" y1="528" x2="580" y2="560" stroke="#2e8b57" stroke-width="1.5"/>
  <line x1="220" y1="560" x2="580" y2="560" stroke="#2e8b57" stroke-width="1.5"/>
  <line x1="410" y1="560" x2="410" y2="585" stroke="#2e8b57" stroke-width="2" marker-end="url(#arr2)"/>

  <!-- ── STEP 6: 종합 판정 ── -->
  <rect x="200" y="585" width="420" height="52" rx="8"
        fill="#7b2d00" stroke="#ff8c00" stroke-width="1.5"/>
  <text x="410" y="605" text-anchor="middle" fill="white" font-size="13" font-weight="bold">Integrated Oxidation Assessment</text>
  <text x="410" y="623" text-anchor="middle" fill="#ffd9b0" font-size="10">Stage I–IV classification  |  Metric comparison  |  Recommendations</text>

  <!-- arrow -->
  <line x1="410" y1="637" x2="410" y2="665" stroke="#ff8c00" stroke-width="2" marker-end="url(#arr)"/>

  <!-- ── STEP 7: 출력 ── -->
  <rect x="200" y="665" width="420" height="52" rx="26" ry="26"
        fill="#1f3864" stroke="none"/>
  <text x="410" y="685" text-anchor="middle" fill="white" font-size="13" font-weight="bold">OUTPUT</text>
  <text x="410" y="705" text-anchor="middle" fill="#aac4e8" font-size="10">Analysis Report (HTML / Word)  |  Charts  |  Final Opinion</text>

  <!-- 범례 -->
  <rect x="20" y="700" width="12" height="12" fill="#1a6496"/>
  <text x="36" y="711" fill="#555" font-size="10">Color feature extraction</text>
  <rect x="20" y="720" width="12" height="12" fill="#155724"/>
  <text x="36" y="731" fill="#555" font-size="10">Estimation / Regression</text>
  <rect x="20" y="740" width="12" height="12" fill="#7b2d00"/>
  <text x="36" y="751" fill="#555" font-size="10">Assessment &amp; Output</text>
</svg>
</div>
<p style="font-size:12px;color:#555;margin-top:8px">
<strong>Figure 0.</strong> Complete analysis pipeline.
Starting from an input photographic image, colorimetric features are extracted (CIE Lab b*, HSI S-channel, Yellowness Index, ΔE), combined into a weighted Euclidean distance metric, and fed into two parallel estimation branches:
(1) <em>Date Estimation</em> via inverse-distance weighted k-NN matching against a reference image database,
and (2) <em>Pseudo-Raman Regression</em> via R²-weighted ensemble of four univariate linear regressions.
The results are integrated into a final oxidation stage assessment (Stage I–IV) and exported as a structured report.
</p>
"""
        if ko:
            svg = svg.replace(
                "<strong>Figure 0.</strong> Complete analysis pipeline.",
                "<strong>그림 0.</strong> 전체 분석 파이프라인 (Analysis Pipeline)."
            ).replace(
                "Starting from an input photographic image, colorimetric features are extracted (CIE Lab b*, HSI S-channel, Yellowness Index, ΔE), combined into a weighted Euclidean distance metric, and fed into two parallel estimation branches:",
                "입력 이미지에서 색상 특성(CIE Lab b*, HSI S채널(S-channel), 황색지수(YI), 색차(ΔE))을 추출하고, 가중 유클리드 거리로 결합하여 두 병렬 추정 경로로 전달한다:"
            ).replace(
                "(1) <em>Date Estimation</em> via inverse-distance weighted k-NN matching against a reference image database,",
                "(1) <em>날짜 추정 (Date Estimation)</em>: 참조 이미지 DB 대비 역거리 가중 k-NN 매칭,"
            ).replace(
                "and (2) <em>Pseudo-Raman Regression</em> via R²-weighted ensemble of four univariate linear regressions.",
                "(2) <em>Pseudo-Raman 회귀 (Regression)</em>: 4개 단변량 선형회귀 R²-가중 앙상블."
            ).replace(
                "The results are integrated into a final oxidation stage assessment (Stage I–IV) and exported as a structured report.",
                "결과를 통합하여 최종 산화 단계 판정(1~4단계)을 수행하고 구조화된 보고서로 출력한다."
            )
        return svg


    def _report_intro_ko(self) -> str:
        return (
            "본 보고서는 제어된 습도 환경에서 HfS₂(이황화 하프늄) 박막의 산화 진행을 "
            "다채널 색상계 이미지 분석과 Pseudo-Raman 스펙트럼 추정을 결합하여 종합 분석한 결과물이다.\n\n"
            "분석 프레임워크는 세 가지 상호보완적 색상 공간 모델 — "
            "CIE Lab, HSI, ASTM 황색지수(YI) — 과 지도학습(supervised) 회귀 앙상블을 통합하여, "
            "단일 광학 이미지만으로 산화 상태와 노출 일수를 추정한다.\n\n"
            "HfS₂는 전이금속 다이칼코게나이드(TMD) 계열 소재로, 주변 습도에 의한 산화 분해에 민감한다. "
            "표면 황(S) 원자가 산소로 치환되면서 HfO₂로 변환되고, 이 과정에서 "
            "특유의 금색/호박색 색조 소실, 채도 감소, CIE Lab b* 좌표의 양수(황색)에서 "
            "0 근방(무채색 백색)으로의 이동이 나타난다. "
            "본 분석은 이러한 변화를 정량적으로 측정한다."
        )

    def _report_color_model_ko(self) -> str:
        return (
            "CIE Lab 색상 공간 — b* 좌표 (노란도 축)\n"
            "CIE Lab은 기기 독립적 지각 균일 색상 공간이다. b* 축은 음수(파랑)에서 0(무채색)을 거쳐 "
            "양수(노랑)로 이어집니다. 신선한 HfS₂는 b*=45~65(강한 황색)이며, "
            "산화가 진행되면 b*=5~20(HfO₂의 근백색)으로 감소한다. "
            "b*는 HfS₂ 산화도의 가장 민감한 단일 지표이다.\n\n"
            "HSI 채도 채널 — S-채널 (S-channel)\n"
            "색조-채도-강도(HSI) 분해는 색도(색의 선명도)를 휘도와 분리한다. "
            "S-채널(0~255)은 황색의 선명도를 반영한다: "
            "신선 S=150~200, 심한 산화 S<20. "
            "조명 변화에 대해 Raw RGB 채널보다 강건한 산화 지표이다.\n\n"
            "ASTM E313 황색지수 — YI (Yellowness Index)\n"
            "YI는 순백 기준 대비 황색 편차를 정량화하는 산업 표준 스칼라 값이다. "
            "HfS₂: YI=60~120(신선) → YI=20~35(산화). "
            "초기 단계 산화 감지에 특히 민감한다.\n\n"
            "CIE 색차 — ΔE (Delta E)\n"
            "ΔE는 동일 실험 조건 Day-0 참조 이미지와의 CIE Lab 공간 유클리드 거리이다. "
            "ΔE<3: 인지 불가 차이; ΔE>10: 주요 지각 변화. "
            "초기 상태 대비 누적 산화 지수로 활용된다."
        )

    def _report_estimation_method_ko(self) -> str:
        return (
            "날짜 추정 알고리즘 (Date Estimation Algorithm)\n"
            "노출 일수는 3차원 지표 공간 {b*, S-채널, YI}에서 가중 k-최근접 이웃(k-NN) 방식으로 추정한다. "
            "가중 유클리드 거리: d = √(w_b·Δb*² + w_S·ΔS² + w_YI·ΔYI²) "
            "기본 가중치: w_b=0.45, w_S=0.30, w_YI=0.25 "
            "(각 지표의 산화 민감도 차이를 반영). "
            "추정 일수 = 상위 k개 참조 일수의 역거리 가중 평균, "
            "신뢰도 = max(0, 100 - dist_min × 200).\n\n"
            "Pseudo-Raman 회귀 (Pseudo-Raman Regression)\n"
            "직접 라만 측정 없이 이미지 지표로부터 A₁g 피크 강도(정규화)를 추정한다. "
            "4개 단변량 선형회귀(b*→라만, S→라만, YI→라만, ΔE→라만)를 독립 학습하고, "
            "R²-가중 앙상블로 결합한다: "
            "est_peak = Σ(pred_m × w_m × R²_m) / Σ(w_m × R²_m). "
            "337 cm⁻¹ 근방의 A₁g 모드는 HfS₂의 주요 라만 특성 피크이며, "
            "정규화 강도는 산화와 단조 감소한다. "
            "95% 신뢰 구간은 회귀 앙상블의 잔차 표준오차에서 계산된다."
        )

    def _report_chart_guide_ko(self) -> str:
        return (
            "지표 유사도 레이더 차트 (Metric Similarity Radar Chart)\n"
            "4개 지표(b*, S-채널, YI, 황색비율)에 대해 대상 이미지(빨간 실선)와 "
            "Top-3 참조 후보(점선)의 정규화 다각형을 표시한다. "
            "다각형이 클수록 산화 전 상태에 가깝다. "
            "중심 방향으로 수축할수록 고산화를 의미한다.\n\n"
            "거리 타임라인 (Distance Timeline)\n"
            "대상 이미지와 모든 참조 이미지 간의 가중 유클리드 거리를 노출 일수별로 표시한다. "
            "거리가 낮을수록 유사도가 높다. "
            "빨간 점선 = 추정 노출 일수(역거리 가중 평균). "
            "낮은 거리 포인트들이 특정 일수 근방에 밀집할수록 신뢰도가 높다.\n\n"
            "회귀 차트 (Regression: Image Metrics → Raman Peak)\n"
            "4개 서브패널이 각 이미지 지표와 정규화 A₁g 라만 피크 간의 선형 관계를 보여준다. "
            "주황 점선 = 적합 회귀선, 빨간 삼각형 = 대상 이미지 지표가 회귀선과 교차하는 점(예측값). "
            "R² 값이 클수록 해당 지표의 예측 신뢰도가 높다.\n\n"
            "추정 스펙트럼 + 95% 신뢰 구간 (Estimated Spectrum + 95% CI)\n"
            "앙상블 추정 피크비율로 가중된 두 인접 참조 스펙트럼의 보간으로 재구성한 라만 스펙트럼이다. "
            "음영 밴드 = 회귀 불확도에서 유도된 95% 신뢰 구간. "
            "~337 cm⁻¹의 A₁g 주석은 HfS₂ 및 HfO₂ 문헌값과 직접 비교를 가능하게 한다."
        )

    def _report_intro(self) -> str:
        return (
            "This report presents a comprehensive optical analysis of HfS2 thin films "
            "under controlled humidity conditions, using multi-channel colorimetric imaging "
            "combined with Pseudo-Raman spectrum estimation. The analytical framework "
            "integrates three complementary color-space models (CIE Lab, HSI, and ASTM "
            "Yellowness Index) with a supervised regression ensemble to estimate oxidation "
            "state and aging day from a single photographic image.\n\n"
            "HfS2 belongs to the family of transition metal dichalcogenides (TMDs) and "
            "exhibits strong sensitivity to oxidative degradation: under ambient humidity, "
            "surface sulfur atoms are progressively replaced by oxygen, converting the "
            "material toward HfO2. This reaction is accompanied by pronounced optical "
            "changes -- loss of the characteristic golden/amber hue, reduction in "
            "saturation, and a shift in the CIE Lab b* coordinate from high positive "
            "values (yellow) toward near-zero (achromatic white). The present analysis "
            "quantifies these changes systematically."
        )

    def _report_color_model(self) -> str:
        return (
            "CIE Lab Color Space (b* coordinate)\n"
            "The CIE Lab model is a device-independent, perceptually uniform color space. "
            "The b* axis spans from negative values (blue) through zero (achromatic) to "
            "positive values (yellow). For pristine HfS2, b* typically ranges from 45-65. "
            "Progressive oxidation drives b* toward 5-20, indicative of near-white HfO2. "
            "b* is the single most sensitive indicator of HfS2 oxidation.\n\n"
            "HSI Saturation Channel (S-channel)\n"
            "The Hue-Saturation-Intensity decomposition isolates chromatic purity from "
            "luminance. S-channel (0-255) reflects vividness of yellow coloration: "
            "pristine films show S = 150-200; heavily oxidized samples fall below 20.\n\n"
            "ASTM E313 Yellowness Index (YI)\n"
            "YI is an industry-standard scalar calibrated to quantify yellowness deviation "
            "from a pure white reference. For HfS2: YI = 60-120 (pristine) -> YI = 20-35 "
            "(oxidized). YI is particularly sensitive to early-stage oxidation.\n\n"
            "CIE DeltaE (Color Difference)\n"
            "DeltaE measures the Euclidean distance in CIE Lab space between each sample "
            "and the Day-0 reference of the same condition. DeltaE < 3: imperceptible; "
            "DeltaE > 10: major perceptual change. DeltaE serves as a cumulative "
            "oxidation index relative to initial state."
        )

    def _report_estimation_method(self) -> str:
        return (
            "Date Estimation Algorithm\n"
            "The aging day is estimated via a weighted k-nearest-neighbor scheme in "
            "three-dimensional metric space {b*, S-channel, YI}. Weighted Euclidean "
            "distance: d = sqrt(w_b*(b*_t-b*_r)^2 + w_S*(S_t-S_r)^2 + w_YI*(YI_t-YI_r)^2). "
            "Default weights: w_b=0.45, w_S=0.30, w_YI=0.25. Estimated day = "
            "inverse-distance weighted average of top-N reference days.\n\n"
            "Pseudo-Raman Regression\n"
            "In absence of a direct Raman measurement, the A1g peak intensity (normalized) "
            "is estimated from image metrics using a linear regression ensemble. Four "
            "univariate regressions (b*->Raman, S->Raman, YI->Raman, DeltaE->Raman) "
            "combined via R2-weighted averaging: "
            "est_peak = Sum(pred_m * w_m * R2_m) / Sum(w_m * R2_m). "
            "The A1g mode near 337 cm-1 is the primary Raman fingerprint of HfS2; "
            "its normalized intensity decreases monotonically with oxidation. "
            "A 95% confidence interval is computed from the pooled residual standard "
            "error of the regression ensemble."
        )

    def _report_chart_guide(self) -> str:
        """전체 차트 안내 요약 (보고서 6절용)"""
        return (
            "This section provides interpretation guidance for each chart generated "
            "by the HfS2 oxidation analysis pipeline. Each chart addresses a different "
            "aspect of the oxidation assessment and should be read in conjunction with "
            "the numerical results in Section 5.\n\n"
            "Figure 1 — Metric Similarity Radar: Shows how closely the target image "
            "matches the top reference candidates across four color metrics simultaneously. "
            "A polygon close to the outer boundary indicates a pristine film; collapse "
            "toward the center indicates advanced oxidation.\n\n"
            "Figure 2 — Distance Timeline: Reveals which reference days are most similar "
            "to the target, and how similarity evolves with aging time. The red dashed "
            "line shows the estimated aging day; a narrow cluster of low-distance points "
            "around this line confirms high estimation confidence.\n\n"
            "Figure 3 — Regression (Image Metrics to Raman Peak): Validates the "
            "color-to-Raman mapping relationship. High R2 across all four sub-panels "
            "confirms that image color is a reliable proxy for Raman state.\n\n"
            "Figure 4 — Estimated Spectrum + 95% CI: Provides a physical Raman spectrum "
            "estimate without requiring an actual spectrometer. The confidence band "
            "reflects the uncertainty in the regression ensemble; a narrow band indicates "
            "a well-constrained prediction.\n\n"
            "Figures 5-7 — Raman Reference Charts: Show the evolution of A1g peak "
            "intensity across all reference conditions and days, providing the empirical "
            "basis for the Pseudo-Raman regression."
        )

    def _chart_desc_radar(self, eval_ctx, ko=False) -> str:
        """Metric Similarity Radar 상세 설명 (실측값 포함)"""
        if not eval_ctx or not eval_ctx.get("target"):
            return ""
        def fv(v, d=2):
            try: return f"{float(v):.{d}f}"
            except: return "N/A"
        tgt = eval_ctx.get("target", {})
        top = eval_ctx.get("top", [])
        t_b  = fv(tgt.get("lab",{}).get("b"))
        t_s  = fv(tgt.get("s_mean"),0)
        t_yi = fv(tgt.get("yellowness_idx"),0)
        t_yr = fv(float(tgt.get("yellow_ratio",0))*100,1)

        ref_lines = []
        for rank, (dist, img) in enumerate(top[:3], 1):
            if ko:
                ref_lines.append(
                    f"  참조 #{rank} ({img.get('cond','')} Day {img.get('day','')}): "
                    f"b*={fv(img.get('lab',{}).get('b'))}  "
                    f"S={fv(img.get('s_mean'),0)}  "
                    f"YI={fv(img.get('yellowness_idx'),0)}  "
                    f"dist={dist:.4f}")
            else:
                ref_lines.append(
                    f"  Reference #{rank} ({img.get('cond','')} Day {img.get('day','')}): "
                    f"b*={fv(img.get('lab',{}).get('b'))}  "
                    f"S={fv(img.get('s_mean'),0)}  "
                    f"YI={fv(img.get('yellowness_idx'),0)}  "
                    f"dist={dist:.4f}")

        if ko:
            return (
                "━━━ 그래프 개념 및 정의 ━━━\n"
                "지표 유사도 레이더 차트(Metric Similarity Radar Chart)는 "
                "다차원 색상 지표를 한 화면에서 비교하기 위한 방사형 시각화이다. "
                "HfS₂ 박막은 산화가 진행될수록 황색이 사라지고 무색·회색으로 변하는데, "
                "이 변화가 CIE Lab b*, HSI S-채널, 황색지수(YI), 황색비율(YR%) 등 4개 지표에 "
                "동시에 반영된다. 레이더의 '면적'이 넓을수록 박막이 신선하고, "
                "면적이 좁아질수록(중심으로 수축) 산화가 진행된 상태이다.\n\n"
                "━━━ 각 축의 의미 ━━━\n"
                "  ① b*(노란도)  : CIE Lab 색공간의 황색-청색 축\n"
                "                  신선: 45–65 │ 산화: 5–20\n"
                "                  → 산화 시 HfO₂로 전환되면서 노란색 소멸\n"
                "  ② S-채널(채도): HSI 모델의 채도 성분 (0–255)\n"
                "                  신선: 150–200 │ 산화: 20 미만\n"
                "                  → 무색 산화물 형성 시 채도 급감\n"
                "  ③ YI(황색지수): ASTM E313 기준 황색도 지수\n"
                "                  신선: 60–120 │ 산화: 20–35\n"
                "                  → 색상 변화의 절대적 크기를 반영\n"
                "  ④ YR%(황색비율): 마스크 내 황색 픽셀 비율 (%)\n"
                "                  신선: 25–33% │ 산화: 5% 미만\n"
                "                  → 표면의 황색 잔존 면적\n\n"
                "━━━ 추정 연산 순서 ━━━\n"
                "  [1단계] 대상 이미지의 4개 지표 측정\n"
                f"          b*={t_b}  S={t_s}  YI={t_yi}  YR={t_yr}%\n\n"
                "  [2단계] 참조 DB 전체 범위로 정규화 (z-score)\n"
                "          정규화값 = (측정값 - 참조DB 평균) / 참조DB 표준편차\n\n"
                "  [3단계] 가중 유클리드 거리 계산 (각 참조 이미지와 비교)\n"
                "          dist = √(wb·(Δb*)² + ws·(ΔS)² + wyi·(ΔYI)²)\n"
                "          기본 가중치: b*×0.45, S×0.30, YI×0.25\n\n"
                "  [4단계] 거리 오름차순 정렬 → Top-3 추출\n\n"
                "  [5단계] 역거리 가중 평균으로 날짜 추정\n"
                "          est_day = Σ(day_i × 1/dist_i) / Σ(1/dist_i)\n\n"
                "━━━ 이 차트에서 볼 점 ━━━\n"
                "  • 빨간 실선(대상)과 참조 점선이 많이 겹칠수록 → 유사도 높음 → 신뢰도 ↑\n"
                "  • 대상 다각형이 참조보다 전반적으로 작다면 → 참조보다 더 산화된 상태\n"
                "  • 특정 축만 작다면 → 해당 지표만 특이하게 반응한 것\n\n"
                "━━━ 대상 이미지 실측값 ━━━\n"
                f"  Lab b*(노란도)   = {t_b}   (신선: 45–65 │ 산화: 5–20)\n"
                f"  S-채널(채도)     = {t_s}   (신선: 150–200 │ 산화: <20)\n"
                f"  황색지수 YI      = {t_yi}   (신선: 60–120 │ 산화: 20–35)\n"
                f"  황색비율 YR%     = {t_yr}%\n\n"
                "━━━ Top-3 참조 이미지 비교 ━━━\n"
                + "\n".join(ref_lines)
            )
        return (
            "━━━ CONCEPT & DEFINITION ━━━\n"
            "The Metric Similarity Radar Chart is a radial visualization for "
            "comparing multi-dimensional color metrics in a single view. "
            "As HfS₂ thin film oxidizes, its yellow color fades to colorless/grey, "
            "and this change is simultaneously reflected in four metrics: "
            "CIE Lab b*, HSI S-channel, Yellowness Index (YI), and Yellow Ratio (YR%). "
            "A larger radar area indicates a fresher film; "
            "contraction toward the center indicates oxidation.\n\n"
            "━━━ AXIS DEFINITIONS ━━━\n"
            "  ① b* (yellowness) : CIE Lab yellow-blue axis\n"
            "                      Pristine: 45–65 │ Oxidized: 5–20\n"
            "                      → Yellow disappears as HfO₂ forms\n"
            "  ② S-channel       : HSI saturation component (0–255)\n"
            "                      Pristine: 150–200 │ Oxidized: <20\n"
            "                      → Colorless oxide dramatically reduces saturation\n"
            "  ③ YI              : Yellowness Index per ASTM E313\n"
            "                      Pristine: 60–120 │ Oxidized: 20–35\n"
            "                      → Absolute magnitude of color shift\n"
            "  ④ YR% (yellow ratio): % of mask pixels classified as yellow\n"
            "                      Pristine: 25–33% │ Oxidized: <5%\n"
            "                      → Remaining yellow surface area\n\n"
            "━━━ ESTIMATION SEQUENCE ━━━\n"
            "  [Step 1] Measure 4 metrics from target image\n"
            f"           b*={t_b}  S={t_s}  YI={t_yi}  YR={t_yr}%\n\n"
            "  [Step 2] Normalize by reference DB range (z-score)\n"
            "           norm = (measured - DB_mean) / DB_std\n\n"
            "  [Step 3] Compute weighted Euclidean distance to each reference\n"
            "           dist = √(wb·(Δb*)² + ws·(ΔS)² + wyi·(ΔYI)²)\n"
            "           Weights: b*×0.45, S×0.30, YI×0.25\n\n"
            "  [Step 4] Sort ascending → extract Top-3\n\n"
            "  [Step 5] Inverse-distance weighted mean → day estimate\n"
            "           est_day = Σ(day_i / dist_i) / Σ(1 / dist_i)\n\n"
            "━━━ WHAT TO LOOK FOR ━━━\n"
            "  • Target (red) closely overlapping reference → high similarity → high confidence\n"
            "  • Target smaller than all references → more oxidized than references\n"
            "  • Only one axis contracted → that single metric responded unusually\n\n"
            "━━━ MEASURED VALUES — TARGET IMAGE ━━━\n"
            f"  Lab b*           = {t_b}   (pristine: 45–65 │ oxidized: 5–20)\n"
            f"  S-channel        = {t_s}   (pristine: 150–200 │ oxidized: <20)\n"
            f"  Yellowness Index = {t_yi}   (pristine: 60–120 │ oxidized: 20–35)\n"
            f"  Yellow Ratio     = {t_yr}%\n\n"
            "━━━ TOP-3 REFERENCE COMPARISON ━━━\n"
            + "\n".join(ref_lines)
        )

    def _chart_desc_timeline(self, eval_ctx, ko=False) -> str:
        """Distance Timeline 상세 설명 (실측값 포함)"""
        if not eval_ctx:
            return ""
        def fv(v, d=2):
            try: return f"{float(v):.{d}f}"
            except: return "N/A"
        est_day = eval_ctx.get("est_day")
        conf    = eval_ctx.get("confidence", 0)
        top     = eval_ctx.get("top", [])
        scores  = eval_ctx.get("scores", [])
        n_refs  = len(scores)

        top3_lines = []
        for rank, (dist, img) in enumerate(top[:3], 1):
            top3_lines.append(
                f"  #{rank}: {img.get('cond','')} Day {img.get('day','')}  "
                f"dist={dist:.4f}  weight={1/(dist+1e-6):.1f}")

        days_spread = ""
        if top:
            days = []
            for _, img in top[:3]:
                try: days.append(float(img["day"]))
                except: pass
            if days:
                days_spread = (f"Top-3 날짜 범위: {min(days):.0f}–{max(days):.0f}일"
                               if ko else
                               f"Top-3 day spread: {min(days):.0f}–{max(days):.0f} days")

        # 신뢰도 해석
        if conf >= 80:
            conf_interp = ("높음 (High) — 복수의 근접 매칭, 추정 신뢰 가능"
                           if ko else "High — multiple close matches, reliable estimate")
        elif conf >= 50:
            conf_interp = ("중간 (Medium) — 합리적 추정, 추가 참조 데이터 권장"
                           if ko else "Medium — reasonable estimate, more references recommended")
        else:
            conf_interp = ("낮음 (Low) — 참조 커버리지 부족, 주의 필요"
                           if ko else "Low — sparse reference coverage, treat with caution")

        if ko:
            return (
                "━━━ 그래프 개념 및 정의 ━━━\n"
                "거리 타임라인(Distance Timeline)은 대상 이미지와 참조 DB 내 각 이미지 간의 "
                "색상 유사도를 시간축(노출 일수)으로 표현한 그래프이다. "
                "날짜별로 얼마나 비슷한 색상이 있는지를 한눈에 보여주어, "
                "대상 이미지가 DB 내 어느 날짜 구간에 가장 잘 맞는지 파악한다.\n\n"
                "━━━ 그래프 읽는 방법 ━━━\n"
                "  x축: 노출 일수 (Day) — 참조 이미지의 실제 측정 날짜\n"
                "  y축: 가중 유클리드 거리 — 낮을수록 더 유사\n"
                "  각 선: 실험 조건 하나 (예: Native-70%, PMMA-70% 등)\n"
                "  빨간 수직 점선: 추정 날짜 (역거리 가중 평균)\n"
                "  → 점선 근처에서 여러 조건의 거리 곡선이 골(valley)을 이루면 → 신뢰도 높음\n"
                "  → 골이 얕고 넓으면 → 날짜 추정이 불확실\n\n"
                "━━━ 추정 연산 상세 ━━━\n"
                "  [1단계] 각 참조 이미지에 대해 거리 계산\n"
                "          dist_i = √(0.45·(Δb*)² + 0.30·(ΔS)² + 0.25·(ΔYI)²)\n\n"
                "  [2단계] 전체 scores 정렬 → Top-K 추출\n"
                f"          전체 참조 수: {n_refs}개\n\n"
                "  [3단계] Top-3의 역거리(1/dist)를 가중치로 날짜 추정\n"
                "          est_day = Σ(day_i / dist_i) / Σ(1 / dist_i)\n"
                "          → 거리가 가까울수록 해당 날짜에 더 큰 가중치\n\n"
                "  [4단계] 신뢰도 계산\n"
                "          confidence = max(0,  100 - dist_min × 200)\n"
                "          → 최근접 거리가 작을수록 신뢰도 높음\n\n"
                "━━━ 추정 결과 ━━━\n"
                f"  추정 노출 일수: {fv(est_day,1)}일\n"
                f"  신뢰도       : {fv(conf,0)}% → {conf_interp}\n"
                f"  총 참조 수   : {n_refs}개\n"
                f"  {days_spread}\n\n"
                "━━━ Top-3 최근접 참조 (가중치 포함) ━━━\n"
                + "\n".join(top3_lines) + "\n\n"
                "━━━ 이 결과의 의미 ━━━\n"
                f"  추정 {fv(est_day,1)}일은 색상 지표 기준으로 대상 이미지가\n"
                f"  참조 DB에서 같은 조건으로 {fv(est_day,1)}일 노출된 샘플과\n"
                "  가장 유사함을 의미한다. 실제 라만 측정을 통해 교차 검증을 권장한다."
            )
        return (
            "━━━ CONCEPT & DEFINITION ━━━\n"
            "The Distance Timeline plots colorimetric similarity between the target "
            "image and every reference in the DB as a function of aging day. "
            "It shows which day range in the DB best matches the target's color state.\n\n"
            "━━━ HOW TO READ ━━━\n"
            "  x-axis: Aging day — actual measurement day of each reference image\n"
            "  y-axis: Weighted Euclidean distance — lower = more similar\n"
            "  Each line: one experimental condition (e.g. Native-70%, PMMA-70%)\n"
            "  Red dashed vertical: estimated day (inverse-distance weighted mean)\n"
            "  → Multiple condition curves forming a valley near the line → high confidence\n"
            "  → Shallow, wide valley → uncertain day estimate\n\n"
            "━━━ ESTIMATION SEQUENCE ━━━\n"
            "  [Step 1] Compute distance to each reference image\n"
            "           dist_i = √(0.45·(Δb*)² + 0.30·(ΔS)² + 0.25·(ΔYI)²)\n\n"
            "  [Step 2] Sort all scores → extract Top-K\n"
            f"           Total references: {n_refs}\n\n"
            "  [Step 3] Inverse-distance weighted mean of Top-3 days\n"
            "           est_day = Σ(day_i / dist_i) / Σ(1 / dist_i)\n"
            "           → Closer references receive proportionally more weight\n\n"
            "  [Step 4] Confidence calculation\n"
            "           confidence = max(0,  100 - dist_min × 200)\n"
            "           → Smaller minimum distance → higher confidence\n\n"
            "━━━ RESULT ━━━\n"
            f"  Estimated aging day : {fv(est_day,1)} days\n"
            f"  Confidence          : {fv(conf,0)}% → {conf_interp}\n"
            f"  Total references    : {n_refs}\n"
            f"  {days_spread}\n\n"
            "━━━ TOP-3 CLOSEST REFERENCES (with weights) ━━━\n"
            + "\n".join(top3_lines) + "\n\n"
            "━━━ INTERPRETATION ━━━\n"
            f"  The estimate of {fv(est_day,1)} days means the target image's color\n"
            "  state most closely matches reference samples that have been exposed\n"
            f"  for {fv(est_day,1)} days under similar conditions.\n"
            "  Cross-validation with physical Raman measurement is recommended."
        )

    def _chart_desc_regression(self, eval_ctx, ko=False) -> str:
        """Regression 차트 상세 설명 (R², pred 값 포함)"""
        if not eval_ctx:
            return ""
        def fv(v, d=4):
            try: return f"{float(v):.{d}f}"
            except: return "N/A"
        tgt      = eval_ctx.get("target", {})
        r2_map   = eval_ctx.get("r2_map", {})
        coef_map = eval_ctx.get("coef_map", {})
        pairs    = eval_ctx.get("pairs", [])

        def _safe(v):
            try:
                fv2 = float(v)
                return 0.0 if fv2 != fv2 else fv2
            except: return 0.0

        t_map = {
            "b":  _safe(tgt.get("lab",{}).get("b")),
            "s":  _safe(tgt.get("s_mean")),
            "yi": _safe(tgt.get("yellowness_idx")),
            "de": _safe(tgt.get("delta_e")),
        }
        label_map    = {"b":"Lab b*", "s":"S-channel", "yi":"YI", "de":"ΔE"}
        label_map_ko = {"b":"Lab b*(노란도)", "s":"S-채널(채도)", "yi":"황색지수(YI)", "de":"색차(ΔE)"}
        wb_map = {"b":0.45, "s":0.25, "yi":0.20, "de":0.10}

        reg_lines = []
        total_ew = 0.0
        preds = []
        for key2 in ["b","s","yi","de"]:
            r2   = r2_map.get(key2, float("nan"))
            coef = coef_map.get(key2)
            tv   = t_map.get(key2, 0.0)
            lbl  = label_map_ko[key2] if ko else label_map[key2]
            if coef is not None:
                import numpy as _np
                pred = float(_np.polyval(coef, tv))
                ew   = wb_map.get(key2, 0) * (r2 if r2 == r2 else 0)
                total_ew += ew
                preds.append((pred, ew))
                if ko:
                    reg_lines.append(
                        f"  {lbl:18s}: 대상값={tv:.2f}  예측피크={pred:.4f}  "
                        f"R²={r2:.3f}  기본가중치={wb_map[key2]:.2f}  "
                        f"유효가중치={ew:.4f}")
                else:
                    reg_lines.append(
                        f"  {lbl:12s}: target={tv:.2f}  pred_peak={pred:.4f}  "
                        f"R²={r2:.3f}  base_w={wb_map[key2]:.2f}  eff_w={ew:.4f}")
            else:
                reg_lines.append(
                    f"  {lbl}: " +
                    ("분산 부족 — 앙상블에서 제외됨" if ko
                     else "insufficient variance — excluded from ensemble"))

        ensemble_str = ""
        if preds and total_ew > 0:
            ens_peak = sum(p*w for p,w in preds) / total_ew
            ensemble_str = (
                f"\n  ➜ 앙상블 추정 라만 피크 = {ens_peak:.4f}\n"
                f"     (각 지표 예측값 × 유효가중치 합산 후 정규화)" if ko else
                f"\n  ➜ Ensemble estimated Raman peak = {ens_peak:.4f}\n"
                f"     (weighted average of per-metric predictions)")

        if ko:
            return (
                "━━━ 그래프 개념 및 정의 ━━━\n"
                "회귀 차트(Regression Chart)는 이미지 색상 지표(x축)와 "
                "실측 라만 A₁g 피크 강도(y축)의 관계를 산점도로 표현하고, "
                "선형회귀로 적합한 결과이다. 이 관계를 통해 "
                "라만 측정 없이 이미지만으로 라만 피크 강도를 예측(Pseudo-Raman)한다.\n\n"
                "━━━ 4개 서브패널 의미 ━━━\n"
                "  각 패널은 서로 다른 색상 지표 하나와 라만 피크의 관계:\n"
                "  • Lab b*   : 황색도 ↔ A₁g 피크. 가장 직접적 상관관계 (가중치 0.45)\n"
                "  • S-채널   : 채도 ↔ A₁g 피크. 산화 시 채도 급감 (가중치 0.25)\n"
                "  • YI       : 황색지수 ↔ A₁g 피크. 절대적 황색도 (가중치 0.20)\n"
                "  • ΔE       : 색차 ↔ A₁g 피크. 0일 기준 색변화량 (가중치 0.10)\n\n"
                "━━━ 그래프 읽는 방법 ━━━\n"
                "  • 파란 점: 매칭된 이미지-라만 참조 쌍\n"
                "  • 주황 점선: 선형회귀 적합선 y = ax + b\n"
                "  • 빨간 삼각형: 대상 이미지에 대한 예측 라만값\n"
                "  • R²이 높을수록 → 해당 지표가 라만 상태의 신뢰할 수 있는 예측자\n"
                "  • R² < 0.5이면 → 해당 지표는 앙상블에서 낮은 가중치 적용\n\n"
                "━━━ 추정 연산 순서 ━━━\n"
                f"  [1단계] 이미지-라만 매칭 쌍 구성: {len(pairs)}개\n\n"
                "  [2단계] 각 지표에 대해 선형회귀 피팅\n"
                "          y(라만 피크) = a × x(지표값) + b\n\n"
                "  [3단계] 대상 이미지의 지표값을 회귀식에 대입 → 예측 피크\n\n"
                "  [4단계] 유효가중치 = 기본가중치 × R²\n"
                "          R²이 높은 지표에 더 큰 가중치 부여\n\n"
                "  [5단계] 유효가중치로 정규화한 예측 피크 가중 평균\n"
                "          est_peak = Σ(pred_i × eff_w_i) / Σ(eff_w_i)\n\n"
                "━━━ 지표별 회귀 결과 ━━━\n"
                + "\n".join(reg_lines)
                + ensemble_str + "\n\n"
                "━━━ R² 해석 기준 ━━━\n"
                "  R² > 0.80 : 강한 색상-라만 상관관계 → 이 지표 예측 신뢰 가능\n"
                "  R² 0.50–0.80: 중간 → 참조 가능하나 가중치 낮음\n"
                "  R² < 0.50 : 약한 상관관계 → 앙상블에 최소 기여"
            )
        return (
            "━━━ CONCEPT & DEFINITION ━━━\n"
            "The Regression Chart is a scatter plot of image color metrics (x-axis) "
            "versus measured Raman A1g peak intensity (y-axis) for all matched "
            "image-Raman reference pairs, fitted with linear regression. "
            "This relationship enables Pseudo-Raman prediction: "
            "estimating the Raman peak from image metrics alone.\n\n"
            "━━━ THE 4 SUB-PANELS ━━━\n"
            "  Each panel shows one color metric vs. Raman peak:\n"
            "  • Lab b*   : yellowness ↔ A1g peak. Most direct correlation (weight 0.45)\n"
            "  • S-channel: saturation ↔ A1g peak. Drops sharply on oxidation (weight 0.25)\n"
            "  • YI       : yellowness index ↔ A1g peak. Absolute yellowness (weight 0.20)\n"
            "  • ΔE       : color difference from Day-0 ↔ A1g peak (weight 0.10)\n\n"
            "━━━ HOW TO READ ━━━\n"
            "  • Blue dots: matched image-Raman reference pairs\n"
            "  • Orange dashed line: linear regression fit y = ax + b\n"
            "  • Red triangle: predicted Raman value for the target image\n"
            "  • Higher R² → that metric is a more reliable predictor\n"
            "  • R² < 0.5 → metric receives low weight in ensemble\n\n"
            "━━━ ESTIMATION SEQUENCE ━━━\n"
            f"  [Step 1] Build image-Raman matched pairs: {len(pairs)} pairs\n\n"
            "  [Step 2] Fit linear regression per metric\n"
            "           y(Raman peak) = a × x(metric value) + b\n\n"
            "  [Step 3] Apply target image's metric value to regression → pred peak\n\n"
            "  [Step 4] Effective weight = base weight × R²\n"
            "           Metrics with higher R² get proportionally more weight\n\n"
            "  [Step 5] Weighted average of predictions\n"
            "           est_peak = Σ(pred_i × eff_w_i) / Σ(eff_w_i)\n\n"
            "━━━ REGRESSION RESULTS PER METRIC ━━━\n"
            + "\n".join(reg_lines)
            + ensemble_str + "\n\n"
            "━━━ R² INTERPRETATION ━━━\n"
            "  R² > 0.80  : Strong color-Raman correlation → sub-estimate reliable\n"
            "  R² 0.50–0.80: Moderate → usable but lower weight\n"
            "  R² < 0.50  : Weak → minimal contribution to ensemble"
        )

    def _chart_desc_spectrum(self, eval_ctx, ko=False) -> str:
        """Estimated Spectrum 상세 설명 (CI, peak 포함)"""
        if not eval_ctx:
            return ""
        def fv(v, d=4):
            try: return f"{float(v):.{d}f}"
            except: return "N/A"
        est_peak  = eval_ctx.get("est_peak")
        ci_lo     = eval_ctx.get("ci_lo")
        ci_hi     = eval_ctx.get("ci_hi")
        est_se    = eval_ctx.get("est_se")
        dists     = eval_ctx.get("spec_dists", [])

        if est_peak is None:
            return ("Pseudo-Raman 추정값 없음. 라만 데이터를 로드하고 평가를 실행해야 한다."
                    if ko else
                    "No Pseudo-Raman estimate. Load Raman data and run Evaluation first.")

        # 산화 단계 판정
        if est_peak >= 0.80:
            stage = "Stage I" if ko else "Stage I"
            interp = ("신선 HfS₂ — A₁g 피크가 완전 강도에 근접. 최소 산화. HfS₂ 지배적 표면."
                      if ko else
                      "Pristine HfS2 — A1g near full intensity. Minimal oxidation. HfS2-dominated surface.")
        elif est_peak >= 0.55:
            stage = "Stage II"
            interp = (f"부분 산화. A₁g 억제 {(1-est_peak)*100:.0f}%. HfS₂/HfO₂ 혼합 표면 조성."
                      if ko else
                      f"Partial oxidation. A1g suppressed {(1-est_peak)*100:.0f}%. Mixed HfS2/HfO2 surface.")
        elif est_peak >= 0.35:
            stage = "Stage III"
            interp = (f"진행성 산화. A₁g {(1-est_peak)*100:.0f}% 억제. HfO₂가 지배적 상."
                      if ko else
                      f"Advanced oxidation. A1g suppressed {(1-est_peak)*100:.0f}%. HfO2 dominant phase.")
        else:
            stage = "Stage IV"
            interp = (f"HfO₂ 거의 완전 전환. A₁g 잔존 강도 {est_peak:.1%}. 물리 라만 측정 강력 권장."
                      if ko else
                      f"Near-complete HfO2 conversion. A1g residual {est_peak:.1%}. Physical Raman strongly recommended.")

        ref_lines = []
        for i, p in enumerate(dists[:2], 1):
            lbl = "참조 스펙트럼" if ko else "Reference spectrum"
            ref_lines.append(
                f"  {lbl} {i}: {p.get('cond','')} Day {p.get('day','')}  "
                f"norm_peak={fv(p.get('norm_peak'))}")

        if ko:
            return (
                "━━━ 그래프 개념 및 정의 ━━━\n"
                "추정 스펙트럼(Pseudo-Raman Estimated Spectrum)은 이미지 색상 지표만으로 "
                "라만 스펙트럼을 추정한 결과이다. 실제 라만 분광기 없이도 "
                "A₁g 피크 강도와 스펙트럼 형태를 예측할 수 있어, "
                "현장에서 빠른 산화도 판정의 대안으로 사용할 수 있다.\n\n"
                "━━━ 그래프 구성요소 ━━━\n"
                "  • 파란 실선    : 이미지 지표 기반 보간 추정 스펙트럼\n"
                "  • 음영 밴드    : 회귀 앙상블 불확도 기반 95% 신뢰 구간\n"
                "  • 회색 얇은 선 : 보간에 사용된 2개 최근접 참조 스펙트럼\n"
                "  • ~337 cm⁻¹   : HfS₂ A₁g 모드 — 산화의 핵심 지표 피크\n\n"
                "━━━ 추정 연산 순서 ━━━\n"
                "  [1단계] 회귀 앙상블로 라만 피크 강도 추정\n"
                f"          est_peak = {fv(est_peak)} (정규화 A₁g 강도)\n\n"
                "  [2단계] 라만 DB에서 추정 피크 강도와 가장 가까운 2개 참조 스펙트럼 선택\n"
                + "\n".join(ref_lines) + "\n\n"
                "  [3단계] 두 참조 스펙트럼 사이를 선형 보간\n"
                "          alpha = (est_peak - peak_ref2) / (peak_ref1 - peak_ref2)  [0~1 클리핑]\n"
                "          V_est(x) = alpha × Spectrum_ref1(x) + (1-alpha) × Spectrum_ref2(x)\n\n"
                "  [4단계] 95% 신뢰 구간 산출\n"
                "          신뢰 구간 = 회귀 앙상블의 표준 오차(SE) × 1.96\n"
                f"          SE = {fv(est_se)}   CI = [{fv(ci_lo)}, {fv(ci_hi)}]\n\n"
                "━━━ 추정 결과 및 해석 ━━━\n"
                f"  산화 단계   : {stage}\n"
                f"  A₁g 정규화 강도  : {fv(est_peak)} (1.000 = 신선 HfS₂)\n"
                f"  95% 신뢰 구간    : [{fv(ci_lo)}, {fv(ci_hi)}]\n"
                f"  불확도 (1 SE)    : {fv(est_se)}\n\n"
                f"  해석: {interp}\n\n"
                "━━━ 주요 라만 피크 참조값 ━━━\n"
                "  ~337 cm⁻¹  : HfS₂ A₁g 모드 — 산화 시 감소 (핵심 지표)\n"
                "  ~260 cm⁻¹  : HfS₂ E₂g 모드 — 산화에 민감\n"
                "  ~500 cm⁻¹  : HfO₂ Ag 모드  — 산화 시 출현·증가\n"
                "  ~630 cm⁻¹  : HfO₂ Bg 모드  — 산화 시 출현·증가\n\n"
                "  ⚠ 주의: 이 스펙트럼은 이미지 기반 추정값이다.\n"
                "  정밀 분석에는 물리적 라만 측정이 필요하다."
            )
        return (
            "━━━ CONCEPT & DEFINITION ━━━\n"
            "The Pseudo-Raman Estimated Spectrum predicts the Raman spectrum of the "
            "target film using image color metrics alone — without a physical Raman "
            "spectrometer. It provides a rapid on-site oxidation assessment tool by "
            "estimating the A1g peak intensity and spectral shape.\n\n"
            "━━━ CHART COMPONENTS ━━━\n"
            "  • Blue solid line    : Interpolated spectrum estimated from image metrics\n"
            "  • Shaded band        : 95% confidence interval from regression uncertainty\n"
            "  • Thin gray lines    : Two nearest reference spectra used for interpolation\n"
            "  • ~337 cm⁻¹ marker  : HfS2 A1g mode — primary oxidation fingerprint\n\n"
            "━━━ ESTIMATION SEQUENCE ━━━\n"
            "  [Step 1] Regression ensemble → Raman peak intensity estimate\n"
            f"           est_peak = {fv(est_peak)} (normalized A1g intensity)\n\n"
            "  [Step 2] Select two nearest reference spectra from Raman DB\n"
            + "\n".join(ref_lines) + "\n\n"
            "  [Step 3] Linear interpolation between the two reference spectra\n"
            "           alpha = (est_peak - peak_ref2) / (peak_ref1 - peak_ref2)  [clipped 0-1]\n"
            "           V_est(x) = alpha × Spectrum_ref1(x) + (1-alpha) × Spectrum_ref2(x)\n\n"
            "  [Step 4] 95% confidence interval from regression standard error\n"
            f"           SE = {fv(est_se)}   CI = [{fv(ci_lo)}, {fv(ci_hi)}]\n\n"
            "━━━ RESULT & INTERPRETATION ━━━\n"
            f"  Oxidation stage          : {stage}\n"
            f"  A1g norm. intensity      : {fv(est_peak)} (1.000 = pristine HfS2)\n"
            f"  95% Confidence Interval  : [{fv(ci_lo)}, {fv(ci_hi)}]\n"
            f"  Uncertainty (1 SE)       : {fv(est_se)}\n\n"
            f"  Interpretation: {interp}\n\n"
            "━━━ KEY RAMAN REFERENCE VALUES ━━━\n"
            "  ~337 cm⁻¹ : HfS2 A1g mode — decreases with oxidation (primary marker)\n"
            "  ~260 cm⁻¹ : HfS2 E2g mode — also oxidation-sensitive\n"
            "  ~500 cm⁻¹ : HfO2 Ag mode  — emerges and grows with oxidation\n"
            "  ~630 cm⁻¹ : HfO2 Bg mode  — emerges and grows with oxidation\n\n"
            "  ⚠ Note: This is an image-based estimate.\n"
            "  Physical Raman measurement is required for precision analysis."
        )

    def _chart_desc_raman_trend(self, raman_data, ko=False) -> str:
        """Raman Peak Trend 차트 설명"""
        if not raman_data:
            return ""
        n = len(raman_data)
        conds = list(dict.fromkeys(r.get("cond","") for r in raman_data))
        days  = sorted(set(r.get("day","") for r in raman_data))
        try:
            peaks = [float(r.get("norm_peak",0)) for r in raman_data
                     if r.get("norm_peak") is not None]
            p_min = f"{min(peaks):.4f}"
            p_max = f"{max(peaks):.4f}"
        except Exception:
            p_min = p_max = "N/A"

        if ko:
            return (
                "━━━ 그래프 개념 및 정의 ━━━\n"
                "라만 피크 추세 차트(Raman Peak Trend Chart)는 노출 일수에 따른 "
                "HfS₂ A₁g 라만 피크 강도의 시간적 변화를 나타냅니다. "
                "A₁g 피크(~337 cm⁻¹)는 HfS₂의 고유 진동 모드로, "
                "산화가 진행될수록 강도가 감소한다. "
                "이 그래프는 산화 속도의 조건별 차이를 직접 비교하는 데 핵심적이다.\n\n"
                "━━━ 그래프 읽는 방법 ━━━\n"
                "  x축: 노출 일수 (Day 0 = 제조 직후)\n"
                "  y축: 정규화 A₁g 피크 강도 (1.0 = Day-0 기준, 0.0 = 완전 소멸)\n"
                "  각 선: 실험 조건 하나 (Native/PMMA/Al₂O₃ × 습도)\n\n"
                "━━━ 이 차트에서 볼 점 ━━━\n"
                "  ① 기울기(감쇠 속도): 가파를수록 해당 조건에서 산화가 빠름\n"
                "     → 고습도(70%RH) Native HfS₂: 1주일 내 A₁g 완전 소멸 (Hwang et al., 2025)\n"
                "     → Al₂O₃ 캡층: 감쇠 현저히 느림\n"
                "  ② 고원(Plateau): 곡선이 수평으로 유지되면 보호층 형성 가능\n"
                "  ③ 조건 간 교차: 초기엔 비슷하다가 나중에 갈리면\n"
                "     → 초기 산화는 동일한 메커니즘, 이후 조건별 차이 발생\n\n"
                "━━━ 물리적 해석 ━━━\n"
                "  산화 반응: HfS₂ + O₂/H₂O → HfS₂₋ₓOₓ → HfO₂\n"
                "  A₁g 감쇠 모델: I(t) = I₀ × exp(-k × t)\n"
                "  k(감쇠상수)가 클수록 산화가 빠른 조건\n\n"
                "━━━ 데이터 요약 ━━━\n"
                f"  총 데이터: {n}건  |  조건: {', '.join(conds)}\n"
                f"  날짜 범위: {days[0] if days else 'N/A'}–{days[-1] if days else 'N/A'}일\n"
                f"  피크 강도 범위: {p_min}–{p_max} (정규화)\n\n"
                "  이 추세 데이터가 본 프로그램의 날짜 추정 및\n"
                "  Pseudo-Raman 회귀 모델의 기준 참조 데이터로 사용된다."
            )
        return (
            "━━━ CONCEPT & DEFINITION ━━━\n"
            "The Raman Peak Trend Chart shows the time-evolution of HfS₂ A1g Raman "
            "peak intensity (~337 cm⁻¹) as a function of aging day under each experimental "
            "condition. The A1g mode is a characteristic vibrational signature of HfS₂ "
            "that decreases monotonically as oxidation proceeds. "
            "This chart is the key reference for comparing oxidation rates across conditions.\n\n"
            "━━━ HOW TO READ ━━━\n"
            "  x-axis: Aging day (Day 0 = as-fabricated)\n"
            "  y-axis: Normalized A1g peak intensity (1.0 = Day-0 baseline, 0.0 = fully gone)\n"
            "  Each line: one experimental condition (Native/PMMA/Al2O3 × humidity)\n\n"
            "━━━ WHAT TO LOOK FOR ━━━\n"
            "  ① Slope (decay rate): steeper → faster oxidation under that condition\n"
            "     → Native HfS2 at 70% RH: complete A1g suppression within 1 week\n"
            "     → Al2O3-capped samples: significantly slower decay\n"
            "  ② Plateau: curve levels off → protective oxide layer may be forming\n"
            "  ③ Crossover: conditions diverge after initial similarity\n"
            "     → Same early mechanism, condition-specific kinetics emerge later\n\n"
            "━━━ PHYSICAL INTERPRETATION ━━━\n"
            "  Oxidation pathway: HfS₂ + O₂/H₂O → HfS₂₋ₓOₓ → HfO₂\n"
            "  Decay model: I(t) = I₀ × exp(-k × t)\n"
            "  Larger k = faster oxidation kinetics\n\n"
            "━━━ DATA SUMMARY ━━━\n"
            f"  Total entries: {n}  |  Conditions: {', '.join(conds)}\n"
            f"  Day range: {days[0] if days else 'N/A'}–{days[-1] if days else 'N/A'}\n"
            f"  Peak range: {p_min}–{p_max} (normalized)\n\n"
            "  This trend data is used as the reference baseline for\n"
            "  day estimation and Pseudo-Raman regression in this program."
        )

    def _chart_desc_raman_spectrum(self, raman_data, ko=False) -> str:
        """Raman Spectra Overlay 차트 설명"""
        specs = [r for r in raman_data if r.get("spectrum")]
        n_spec = len(specs)
        n_total = len(raman_data)
        if ko:
            return (
                "━━━ 그래프 개념 및 정의 ━━━\n"
                "라만 스펙트럼 오버레이(Raman Spectra Overlay)는 "
                "서로 다른 노출 조건과 일수에서 측정된 라만 스펙트럼을 "
                "한 화면에 겹쳐 표시한다. 산화가 진행됨에 따라 스펙트럼이 "
                "어떻게 변하는지를 시각적으로 보여주는 핵심 참조 데이터이다.\n\n"
                "━━━ 그래프 읽는 방법 ━━━\n"
                "  x축: 라만 시프트 (Raman shift, cm⁻¹)\n"
                "  y축: 정규화 강도 (1.0 = 해당 조건 Day-0의 A₁g 피크)\n"
                "  색상: 노출 일수 또는 조건별로 구분 (범례 참조)\n\n"
                "━━━ 반드시 봐야 할 특징 피크 ━━━\n"
                "  ~337 cm⁻¹ : HfS₂ A₁g 모드\n"
                "    - 이 피크의 높이가 산화도의 핵심 지표\n"
                "    - Day-0: 최대 강도 / 시간이 지날수록 감소 / 완전 산화 시 소멸\n"
                "    - 이 변화가 색상 지표(b*, YI 등) 감소와 상관관계\n\n"
                "  ~260 cm⁻¹ : HfS₂ E₂g 모드\n"
                "    - A₁g와 함께 감소하나 보통 더 빠르게 소멸\n\n"
                "  ~500 cm⁻¹ : HfO₂ Ag 모드\n"
                "    - 산화가 진행될수록 이 피크가 출현·증가\n"
                "    - HfS₂ → HfO₂ 전환의 직접적 증거\n\n"
                "  ~630 cm⁻¹ : HfO₂ Bg 모드\n"
                "    - 고도 산화 시 명확히 출현\n\n"
                "━━━ 이 차트가 추정에 사용되는 방법 ━━━\n"
                "  Pseudo-Raman 추정 시 이 스펙트럼들 중 추정 피크값에\n"
                "  가장 가까운 2개를 선택하여 선형 보간한다.\n"
                "  스펙트럼 데이터 포함: "
                f"{n_spec}/{n_total}건\n\n"
                "━━━ 정규화 방법 ━━━\n"
                "  동일 조건의 Day-0 A₁g 피크 기준 정규화\n"
                "  → 조건 간 절대 강도 차이를 제거하고 산화 진행도만 비교"
            )
        return (
            "━━━ CONCEPT & DEFINITION ━━━\n"
            "The Raman Spectra Overlay superimposes all measured Raman spectra "
            "from different conditions and aging days in a single plot. "
            "It is the core reference dataset showing how spectra evolve with oxidation.\n\n"
            "━━━ HOW TO READ ━━━\n"
            "  x-axis: Raman shift (cm⁻¹)\n"
            "  y-axis: Normalized intensity (1.0 = Day-0 A1g peak of same condition)\n"
            "  Color: coded by aging day or condition (see legend)\n\n"
            "━━━ KEY SPECTRAL FEATURES TO IDENTIFY ━━━\n"
            "  ~337 cm⁻¹ : HfS₂ A1g mode  — PRIMARY OXIDATION MARKER\n"
            "    - Height of this peak is the key oxidation indicator\n"
            "    - Day-0: maximum / decreases with time / disappears at full oxidation\n"
            "    - This change correlates with color metric (b*, YI) decrease\n\n"
            "  ~260 cm⁻¹ : HfS₂ E2g mode\n"
            "    - Decreases with A1g but typically disappears faster\n\n"
            "  ~500 cm⁻¹ : HfO₂ Ag mode\n"
            "    - Appears and grows as oxidation proceeds\n"
            "    - Direct evidence of HfS₂ → HfO₂ conversion\n\n"
            "  ~630 cm⁻¹ : HfO₂ Bg mode\n"
            "    - Appears clearly at advanced oxidation stages\n\n"
            "━━━ HOW THIS CHART IS USED IN ESTIMATION ━━━\n"
            "  During Pseudo-Raman estimation, the two spectra with peak values\n"
            "  closest to the estimated peak are selected from this collection\n"
            "  and linearly interpolated to produce the estimated spectrum.\n"
            f"  Spectra with full data: {n_spec}/{n_total}\n\n"
            "━━━ NORMALIZATION ━━━\n"
            "  Normalized to Day-0 A1g peak of the same condition.\n"
            "  This removes absolute intensity differences between conditions,\n"
            "  enabling direct comparison of oxidation progression."
        )

    def _chart_desc_raman_decay(self, raman_data, ko=False) -> str:
        """Raman Decay Rate 차트 설명"""
        if not raman_data:
            return ""
        conds = list(dict.fromkeys(r.get("cond","") for r in raman_data))

        # 각 조건별 간단한 통계
        cond_stats = {}
        for r in raman_data:
            c = r.get("cond","")
            try:
                d = float(r.get("day", 0))
                p = float(r.get("norm_peak", 1))
                if c not in cond_stats:
                    cond_stats[c] = []
                cond_stats[c].append((d, p))
            except Exception:
                pass

        stat_lines = []
        for c, pts in cond_stats.items():
            pts.sort(key=lambda x: x[0])
            if len(pts) >= 2:
                d0_p = pts[0][1]
                d_last, p_last = pts[-1]
                decay = d0_p - p_last
                stat_lines.append(
                    f"  {c}: Day 0={d0_p:.3f} → Day {d_last:.0f}={p_last:.3f} "
                    f"(감쇠={decay:.3f})" if ko else
                    f"  {c}: Day 0={d0_p:.3f} → Day {d_last:.0f}={p_last:.3f} "
                    f"(decay={decay:.3f})")

        if ko:
            return (
                "━━━ 그래프 개념 및 정의 ━━━\n"
                "A₁g 피크 감쇠율 비교 차트(A1g Peak Decay Chart)는 "
                "실험 조건별 산화 속도(kinetics)를 직접 비교한다. "
                "어떤 보관/처리 조건이 HfS₂를 가장 잘 보호하는지, "
                "그리고 각 조건에서 박막이 얼마나 빨리 열화되는지를 한눈에 보여준다.\n\n"
                "━━━ 그래프 읽는 방법 ━━━\n"
                "  x축: 노출 일수\n"
                "  y축: 정규화 A₁g 피크 강도 (1.0 = Day-0)\n"
                "  각 선: 하나의 실험 조건\n"
                "  기울기(slope)가 가파를수록 → 산화가 빠른 조건\n"
                "  0.5 라인(반감기): 피크가 절반으로 감쇠한 시점\n\n"
                "━━━ 이 차트에서 볼 점 ━━━\n"
                "  ① 조건별 곡선의 기울기 비교\n"
                "     - Native + 고습도: 가장 가파름 (가장 빠른 산화)\n"
                "     - Al₂O₃ 캡층: 가장 완만 (가장 효과적인 보호)\n"
                "     - PMMA 캡층: 중간 수준의 보호\n\n"
                "  ② 지수 감쇠 모델 적합\n"
                "     I(t) = I₀ × exp(-k × t)\n"
                "     k값이 클수록 빠른 산화. Kinetic 모델 추정에 사용됨\n\n"
                "  ③ 반감기 (t½ = ln(2)/k)\n"
                "     피크가 0.5가 되는 날짜 → 조건 심각도 비교 기준\n\n"
                "━━━ 조건별 감쇠 현황 ━━━\n"
                + "\n".join(stat_lines) + "\n\n"
                "━━━ 실용적 활용 ━━━\n"
                "  감쇠 곡선으로 보관 조건별 유효 수명(shelf life) 추정 가능.\n"
                "  외삽 적합선의 x절편 ≈ A₁g가 검출 불가(≈0.05) 수준이 되는 예상 일수.\n"
                "  이 차트는 본 프로그램의 Kinetic 모델(지수 감쇠 피팅)의\n"
                "  이론적 근거이자 파라미터 추정 데이터로 사용된다."
            )
        return (
            "━━━ CONCEPT & DEFINITION ━━━\n"
            "The A1g Peak Decay Chart directly compares oxidation kinetics across "
            "experimental conditions. It reveals which storage/treatment conditions "
            "best protect HfS₂ and how rapidly the film degrades under each condition.\n\n"
            "━━━ HOW TO READ ━━━\n"
            "  x-axis: Aging day\n"
            "  y-axis: Normalized A1g peak intensity (1.0 = Day-0)\n"
            "  Each line: one experimental condition\n"
            "  Steeper slope → faster oxidation under that condition\n"
            "  0.5 line = half-life: when peak decays to half its initial value\n\n"
            "━━━ WHAT TO LOOK FOR ━━━\n"
            "  ① Compare slopes across conditions:\n"
            "     - Native + high humidity: steepest (fastest oxidation)\n"
            "     - Al2O3 capped: most gradual (most effective protection)\n"
            "     - PMMA capped: intermediate protection\n\n"
            "  ② Exponential decay model fit:\n"
            "     I(t) = I₀ × exp(-k × t)\n"
            "     Larger k = faster oxidation. Used in Kinetic Model estimation.\n\n"
            "  ③ Half-life (t½ = ln(2)/k):\n"
            "     Day at which peak reaches 0.5 → benchmark for condition severity\n\n"
            "━━━ DECAY STATISTICS BY CONDITION ━━━\n"
            + "\n".join(stat_lines) + "\n\n"
            "━━━ PRACTICAL USE ━━━\n"
            "  Use decay curves to estimate shelf life under different storage conditions.\n"
            "  X-intercept of an extrapolated fit ≈ expected day when A1g becomes\n"
            "  undetectable (noise floor ~0.05).\n"
            "  This chart provides the theoretical basis and fitting data\n"
            "  for the Kinetic exponential decay model in this program."
        )

    def _report_final_opinion(self, images, eval_ctx,
                               eval_comment, pseudo_result,
                               ko=False) -> str:
        if not eval_ctx or not eval_ctx.get("target"):
            return ("평가 대상 이미지가 분석되지 않았다." if ko
                    else "No evaluation target was analyzed in this session.")

        def fv(v, d=2):
            try: return f"{float(v):.{d}f}"
            except: return "N/A"
        def fs(v):
            try: return float(v)
            except: return float("nan")

        tgt      = eval_ctx.get("target", {})
        est_day  = eval_ctx.get("est_day")
        conf     = eval_ctx.get("confidence", 0)
        top      = eval_ctx.get("top", [])
        est_peak = eval_ctx.get("est_peak")
        ci_lo    = eval_ctx.get("ci_lo")
        ci_hi    = eval_ctx.get("ci_hi")

        t_b  = fs(tgt.get("lab",{}).get("b"))
        t_s  = fs(tgt.get("s_mean"))
        t_yi = fs(tgt.get("yellowness_idx"))
        t_de = fs(tgt.get("delta_e"))

        # 산화 단계 판정
        if t_b > 40 and t_s > 120:
            stage = ("신선 / 초기 단계 (Pristine / Stage I)" if ko
                     else "Pristine / Early-Stage (Stage I)")
            stage_desc = (
                "박막이 특유의 황금-호박색 외관을 유지하고 있다. "
                "채도(S-채널)와 b* 값이 높아 표면 산화가 거의 없거나 없는 상태이다. "
                "HfS₂ 격자 무결성이 보존되어 있을 가능성이 높다."
                if ko else
                "The film retains its characteristic yellow-amber appearance "
                "with high saturation and strong b* signal. Surface oxidation "
                "is minimal or absent. HfS2 lattice integrity is likely preserved.")
        elif t_b > 20 and t_s > 50:
            stage = ("중등도 산화 (Moderate Oxidation / Stage II)" if ko
                     else "Moderate Oxidation (Stage II)")
            stage_desc = (
                "황색이 눈에 띄게 감소하였다. b* 감소와 채도 손실은 "
                "표면 HfS₂의 HfO₂로의 부분 전환을 나타냅니다. "
                "A₁g 라만 모드 강도 감소(정규화 피크 약 0.5~0.8 예상)가 측정될 것으로 예상된다."
                if ko else
                "Noticeable yellowing has diminished. b* reduction and "
                "saturation loss indicate partial conversion of surface HfS2 to HfO2. "
                "A1g Raman mode intensity reduction (norm. peak ~0.5-0.8) is expected.")
        elif t_b > 8:
            stage = ("진행성 산화 (Advanced Oxidation / Stage III)" if ko
                     else "Advanced Oxidation (Stage III)")
            stage_desc = (
                "상당한 산화가 진행되었다. 낮은 b*와 채도로 인한 근백색 외관은 "
                "HfO₂ 특성이 지배적임을 반영한다. A₁g 라만 강도가 크게 억제될 것으로 "
                "예상된다(정규화 피크 < 0.5). 전기적 특성이 크게 저하되었을 수 있다."
                if ko else
                "Substantial oxidation. Near-white appearance (low b*, low S) "
                "reflects dominant HfO2 character. A1g Raman intensity expected "
                "to be significantly suppressed (norm. peak < 0.5). "
                "Electrical properties will be strongly affected.")
        else:
            stage = ("심각한 산화 / 거의 완전 전환 (Severe / Stage IV)" if ko
                     else "Severe Oxidation / Near-Complete Conversion (Stage IV)")
            stage_desc = (
                "광학적 특성이 HfO₂로의 거의 완전한 전환과 일치한다. "
                "소재가 TMD 특성을 상실하였다. "
                "A₁g 라만 피크가 노이즈 바닥 이상으로 검출되지 않을 수 있다."
                if ko else
                "Optical signature consistent with near-complete conversion "
                "to HfO2. Material has lost TMD characteristics. "
                "A1g Raman peak may no longer be detectable above noise floor.")

        # 신뢰도 해석
        if conf >= 80:
            conf_interp = (
                "높음 (HIGH) — 다수의 근접 참조 매칭으로 추정이 잘 제약되어 있다."
                if ko else
                "HIGH -- estimate well-constrained by multiple close reference matches.")
        elif conf >= 50:
            conf_interp = (
                "중간 (MEDIUM) — 추정은 합리적이나 불확도가 남아 있다; "
                "이 산화 상태 근방의 추가 참조 데이터가 정밀도를 향상시킬 것이다."
                if ko else
                "MEDIUM -- estimate reasonable but some uncertainty remains; "
                "additional reference data near this oxidation state would improve precision.")
        else:
            conf_interp = (
                "낮음 (LOW) — 대상이 참조 데이터가 희박한 영역에 위치한다. "
                "이 산화 수준 근방의 참조 이미지를 추가 수집해야 한다."
                if ko else
                "LOW -- target falls in sparsely sampled region. "
                "Collect more reference images near this oxidation level.")

        # Top-3 비교
        top3_lines = []
        for rank, (dist, img) in enumerate(top[:3], 1):
            r_b = fs(img.get("lab",{}).get("b"))
            r_s = fs(img.get("s_mean"))
            db  = t_b - r_b
            ds  = t_s - r_s
            if ko:
                ab = ("더 신선함" if db > 3 else "더 산화됨" if db < -3 else "유사 수준")
                sa = ("더 선명함" if ds > 8 else "더 색 바램" if ds < -8 else "유사 채도")
                top3_lines.append(
                    f"  #{rank}: {img['cond']} / Day {img['day']}  (거리(dist)={dist:.4f})\n"
                    f"       b*(노란도): 대상={fv(t_b)} vs 참조={fv(r_b)} --> 대상이 {ab}\n"
                    f"       S(채도):  대상={fv(t_s,0)} vs 참조={fv(r_s,0)} --> 대상이 {sa}")
            else:
                ab = ("more pristine" if db > 3 else "more oxidized" if db < -3 else "same level")
                sa = ("more vivid"    if ds > 8 else "more faded"    if ds < -8 else "same saturation")
                top3_lines.append(
                    f"  #{rank}: {img['cond']} / Day {img['day']}  (dist={dist:.4f})\n"
                    f"       b*: target={fv(t_b)} vs ref={fv(r_b)} --> target is {ab}\n"
                    f"       S:  target={fv(t_s,0)} vs ref={fv(r_s,0)} --> target is {sa}")

        # Pseudo-Raman 의견
        raman_op = ""
        if est_peak is not None:
            if ko:
                if est_peak >= 0.80:
                    ri = "주로 HfS₂ 표면에 해당 — A₁g 피크가 신선 강도에 가깝다."
                elif est_peak >= 0.55:
                    ri = "부분 산화에 해당 — A₁g 강도 감소(~20~45%)로 HfS₂/HfO₂ 혼합 표면이 시사된다."
                elif est_peak >= 0.35:
                    ri = "진행성 산화에 해당 — A₁g 피크가 크게 억제되었다; HfO₂ 상이 지배적이다."
                else:
                    ri = "HfO₂ 거의 완전 전환에 해당 — A₁g 피크 강도가 임계적으로 낮다."
                raman_op = (
                    f"\nPseudo-Raman 추정 결과:\n"
                    f"  A₁g 정규화 피크 (norm. peak) = {fv(est_peak,4)}\n"
                    f"  95% 신뢰 구간 (CI): [{fv(ci_lo,4)} -- {fv(ci_hi,4)}]\n"
                    f"  해석: {ri}")
            else:
                if est_peak >= 0.80:
                    ri = "consistent with predominantly HfS2 surface -- A1g peak near pristine intensity."
                elif est_peak >= 0.55:
                    ri = "consistent with partial oxidation -- A1g intensity reduction (~20-45%) suggests mixed HfS2/HfO2 surface."
                elif est_peak >= 0.35:
                    ri = "consistent with advanced oxidation -- A1g peak substantially suppressed; HfO2 phase is dominant."
                else:
                    ri = "consistent with near-complete HfO2 conversion -- A1g peak intensity critically low."
                raman_op = (
                    f"\nPseudo-Raman Estimate:\n"
                    f"  A1g norm. peak = {fv(est_peak,4)}\n"
                    f"  95% CI: [{fv(ci_lo,4)} -- {fv(ci_hi,4)}]\n"
                    f"  Interpretation: {ri}")

        if ko:
            header      = "최종 분석 의견"
            sec_stage   = "산화 단계 (Oxidation Stage)"
            sec_est     = f"추정 산화 일수 (Estimated Day): {fv(est_day,1)}일"
            sec_conf    = f"신뢰도 (Confidence Level): {fv(conf,0)}%  --  {conf_interp}"
            sec_metrics = "대상 이미지 색상 지표 (Color Metrics):"
            m1 = f"  CIE Lab b*(노란도)     = {fv(t_b,2)}   (신선: 45-65 | 산화: 5-20)"
            m2 = f"  S-채널(채도)           = {fv(t_s,1)}   (신선: 150-200 | 산화: <20)"
            m3 = f"  황색지수 YI            = {fv(t_yi,1)}   (신선: 60-120 | 산화: 20-35)"
            m4 = f"  색차 ΔE (Day-0 기준)   = {fv(t_de,2)}   (≥10 = 주요 색 변화)"
            sec_top3    = "Top-3 참조 비교 (Reference Comparisons):"
            sec_recs    = "권고사항 (Recommendations):"
        else:
            header      = "FINAL ANALYTICAL OPINION"
            sec_stage   = "Oxidation Stage"
            sec_est     = f"Estimated Aging Day: {fv(est_day,1)} days"
            sec_conf    = f"Confidence Level: {fv(conf,0)}%  --  {conf_interp}"
            sec_metrics = "Color Metrics of Target Image:"
            m1 = f"  CIE Lab b*         = {fv(t_b,2)}   (pristine: 45-65 | oxidized: 5-20)"
            m2 = f"  S-channel          = {fv(t_s,1)}   (pristine: 150-200 | oxidized: <20)"
            m3 = f"  Yellowness Index   = {fv(t_yi,1)}   (pristine: 60-120 | oxidized: 20-35)"
            m4 = f"  DeltaE from Day-0  = {fv(t_de,2)}   (>=10 = major color change)"
            sec_top3    = "Top-3 Reference Comparisons:"
            sec_recs    = "Recommendations:"

        out_lines = [
            "=" * 62,
            header,
            "=" * 62,
            "",
            f"{sec_stage}: {stage}",
            f"{stage_desc}",
            "",
            sec_est,
            sec_conf,
            "",
            sec_metrics,
            m1, m2, m3, m4,
            "",
            sec_top3,
        ] + top3_lines + [raman_op, "", sec_recs]

        # 권고사항
        if ko:
            if conf < 50:
                out_lines.append("  * 이 산화 상태 근방의 참조 데이터셋을 확장해야 한다.")
            if t_b < 15:
                out_lines.append("  * 물리적 라만 분석으로 교차 검증해야 한다 — "
                                  "광학 방법은 이 진행 단계에서 과소평가할 수 있다.")
            if t_de > 15:
                out_lines.append("  * 높은 ΔE는 초기 상태 대비 주요 구조 변화를 나타냅니다; "
                                  "XPS 및 홀 측정(Hall measurement)을 권장한다.")
            if est_peak is not None and est_peak < 0.4:
                out_lines.append("  * A₁g 임계 억제가 시사된다 — "
                                  "결론 도출 전 물리적 라만 측정으로 검증해야 한다.")
            out_lines.append("  * 논문 수준 분석을 위해 광학 결과를 XPS 및 라만 분광법과 결합해야 한다.")
        else:
            if conf < 50:
                out_lines.append("  * Expand reference dataset near this oxidation state.")
            if t_b < 15:
                out_lines.append("  * Cross-validate with physical Raman spectroscopy -- "
                                  "optical method may underestimate at this advanced stage.")
            if t_de > 15:
                out_lines.append("  * High DeltaE indicates major structural change; "
                                  "XPS and Hall measurement recommended.")
            if est_peak is not None and est_peak < 0.4:
                out_lines.append("  * Critical A1g suppression suggested -- "
                                  "verify with physical Raman before drawing conclusions.")
            out_lines.append("  * For publication-quality analysis, combine optical results "
                              "with XPS and Raman spectroscopy.")

        return "\n".join(str(l) for l in out_lines if l is not None)


    def _build_report_html(self, images, raman_data, eval_ctx,
                            chart_b64, eval_comment,
                            pseudo_result, gen_time, lang="en") -> str:
        KO = (lang == "ko")
        def fv(v, d=2):
            try: return f"{float(v):.{d}f}"
            except: return "-"

        def itag(b64, alt="chart"):
            if not b64: return ""
            return (
                "<figure>"
                "<img src=\"data:image/png;base64," + b64 + "\" "
                "alt=\"" + alt + "\" style=\"width:100%;max-width:860px;"
                "display:block;margin:10px auto;border-radius:4px;\">"
                "<figcaption>" + alt + "</figcaption>"
                "</figure>")

        def desc_box(txt):
            if not txt: return ""
            esc = txt.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
            return "<div class=\"desc-box\"><pre>" + esc + "</pre></div>"

        def tbl(headers, rows):
            ths = "".join("<th>" + h + "</th>" for h in headers)
            trs = ""
            for i, row in enumerate(rows):
                bg = "#f0f4fa" if i % 2 == 0 else "#ffffff"
                tds = "".join("<td>" + str(c) + "</td>" for c in row)
                trs += "<tr style=\"background:" + bg + "\">" + tds + "</tr>"
            return "<table><thead><tr>" + ths + "</tr></thead><tbody>" + trs + "</tbody></table>"

        def sec(title, body, level=2):
            tag = "h" + str(level)
            return "<" + tag + ">" + title + "</" + tag + ">\n" + body + "\n"

        def pre(txt):
            if not txt: return ""
            esc = txt.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
            return "<pre>" + esc + "</pre>"

        # ── 이미지 표 ────────────────────────────
        img_tbl = ""
        if images:
            if KO:
                hdrs = ["파일명(Name)","조건(Cond)","날짜(Day)",
                        "b*(노란도)","채도(S-ch)","황색지수(YI)",
                        "색차(ΔE)","황색비율(Yellow%)"]
            else:
                hdrs = ["Name","Cond","Day","b*","S-ch","YI","dE","Yellow%"]
            rows = [[img["name"][:24], img["cond"], img["day"],
                     img["b"], img["s"], img["yi"],
                     img["de"], img["yr"]] for img in images[:60]]
            img_tbl = tbl(hdrs, rows)

        thumbs = ""
        th_title = "대표 썸네일 (분석 이미지)" if KO else "Sample Thumbnails"
        for img in [i for i in images if i.get("tb64","")][:12]:
            thumbs += (
                "<div class=\"thumb-item\">"
                "<img src=\"data:image/png;base64," + img["tb64"] + "\" "
                "style=\"width:150px;height:110px;object-fit:contain;\">"
                "<div class=\"thumb-cap\">" + img["cond"] + " D" + img["day"] +
                "<br>" + img["name"][:20] + "</div></div>")
        if thumbs:
            thumbs = "<h3>" + th_title + "</h3><div class=\"thumb-grid\">" + thumbs + "</div>"

        # ── Raman 표 ─────────────────────────────
        raman_tbl = ""
        if raman_data:
            if KO:
                r_h = ["조건(Condition)","날짜(Day)","피크(Peak)",
                       "정규화피크(Norm Peak)","피크위치(Shift cm⁻¹)","범위(Range)"]
            else:
                r_h = ["Condition","Day","Peak","Norm Peak","Shift(cm-1)","Range"]
            r_r = []
            for r in raman_data[:60]:
                r_r.append([r.get("cond",""), r.get("day",""),
                    fv(r.get("peak"),4), fv(r.get("norm_peak"),4),
                    fv(r.get("peak_shift"),1), r.get("peak_range","-")])
            raman_tbl = tbl(r_h, r_r)

        # ── Raman 차트 + 개별 설명 ───────────────
        raman_charts = ""
        fig5 = "그림 5. " if KO else "Figure 5. "
        fig6 = "그림 6. " if KO else "Figure 6. "
        fig7 = "그림 7. " if KO else "Figure 7. "
        for k, t, desc_fn in [
            ("raman_trend",
             fig5 + ("라만 피크 강도 추세 (시간 경과)" if KO else "Raman Peak Trend Over Time"),
             lambda: self._chart_desc_raman_trend(raman_data, KO)),
            ("raman_spectrum",
             fig6 + ("라만 스펙트럼 오버레이" if KO else "Raman Spectra Overlay"),
             lambda: self._chart_desc_raman_spectrum(raman_data, KO)),
            ("raman_decay",
             fig7 + ("조건별 A₁g 피크 감쇠율" if KO else "A1g Peak Decay by Condition"),
             lambda: self._chart_desc_raman_decay(raman_data, KO)),
        ]:
            if chart_b64.get(k):
                raman_charts += "<h3>" + t + "</h3>" + itag(chart_b64[k], t) + desc_box(desc_fn())

        # ── Evaluation 차트 + 개별 설명 ──────────
        eval_part = ""
        if eval_ctx.get("target"):
            est  = eval_ctx.get("est_day")
            conf = eval_ctx.get("confidence",0)
            tgt2 = eval_ctx.get("target",{})
            if KO:
                eval_part = (
                    "<div class=\"eval-box\">"
                    "<div class=\"est-day\">추정 산화 일수 (Estimated Day): "
                    + fv(est,1) + "일</div>"
                    "<div class=\"conf\">신뢰도 (Confidence): " + fv(conf,0) + "% &nbsp;|&nbsp; "
                    "b*(노란도)=" + fv(tgt2.get("lab",{}).get("b")) + " &nbsp; "
                    "S-채널(채도)=" + fv(tgt2.get("s_mean"),0) + " &nbsp; "
                    "YI(황색지수)=" + fv(tgt2.get("yellowness_idx"),0) + "</div>"
                    "</div>")
                chart_labels = [
                    ("radar",         "그림 1. 지표 유사도 레이더 차트 (Metric Similarity Radar)"),
                    ("timeline",      "그림 2. 거리 타임라인 (Distance Timeline)"),
                    ("pseudo_reg_fig","그림 3. 회귀 분석: 이미지 지표 → 라만 피크 (Regression)"),
                    ("pseudo_spec_fig","그림 4. 추정 라만 스펙트럼 + 95% 신뢰 구간 (Estimated Spectrum + 95% CI)"),
                ]
            else:
                eval_part = (
                    "<div class=\"eval-box\">"
                    "<div class=\"est-day\">Estimated Aging Day: " + fv(est,1) + " days</div>"
                    "<div class=\"conf\">Confidence: " + fv(conf,0) + "% &nbsp;|&nbsp; "
                    "b*=" + fv(tgt2.get("lab",{}).get("b")) + " &nbsp; "
                    "S-ch=" + fv(tgt2.get("s_mean"),0) + " &nbsp; "
                    "YI=" + fv(tgt2.get("yellowness_idx"),0) + "</div>"
                    "</div>")
                chart_labels = [
                    ("radar",         "Figure 1. Metric Similarity Radar"),
                    ("timeline",      "Figure 2. Distance Timeline"),
                    ("pseudo_reg_fig","Figure 3. Regression: Image Metrics to Raman Peak"),
                    ("pseudo_spec_fig","Figure 4. Estimated Spectrum + 95% CI"),
                ]
            desc_fns = [
                lambda: self._chart_desc_radar(eval_ctx, KO),
                lambda: self._chart_desc_timeline(eval_ctx, KO),
                lambda: self._chart_desc_regression(eval_ctx, KO),
                lambda: self._chart_desc_spectrum(eval_ctx, KO),
            ]
            for (k, t), desc_fn in zip(chart_labels, desc_fns):
                if chart_b64.get(k):
                    eval_part += ("<h3>" + t + "</h3>"
                                  + itag(chart_b64[k], t)
                                  + desc_box(desc_fn()))
        else:
            eval_part = ("<p><em>" +
                         ("평가 대상 이미지가 분석되지 않았다." if KO
                          else "No evaluation target analyzed.") + "</em></p>")

        opinion = self._report_final_opinion(
            images, eval_ctx, eval_comment, pseudo_result, KO)

        css = """
*{box-sizing:border-box;margin:0;padding:0}
body{font-family:"Segoe UI",Arial,sans-serif;font-size:13px;color:#1a1a2e;background:#f5f7fa;line-height:1.75}
.cover{background:linear-gradient(135deg,#1f3864,#2e75b6);color:white;padding:52px 72px}
.cover h1{font-size:28px;margin-bottom:8px}
.cover p{font-size:13px;opacity:0.82}
.content{max-width:1060px;margin:0 auto;padding:36px 48px}
h2{font-size:19px;color:#1f3864;border-bottom:2.5px solid #2e75b6;padding-bottom:6px;margin:36px 0 14px}
h3{font-size:14px;color:#2e75b6;margin:22px 0 8px}
p{margin:8px 0 12px;white-space:pre-line}
table{width:100%;border-collapse:collapse;margin:14px 0 22px;font-size:12px}
th{background:#1f3864;color:white;padding:8px 10px;text-align:left}
td{padding:6px 10px;border-bottom:1px solid #dde3f0}
figure{margin:12px 0 8px;text-align:center}
figcaption{font-size:11px;color:#666;margin-top:4px;font-style:italic}
.desc-box{background:#1a1a2e;border-left:4px solid #2e75b6;border-radius:0 6px 6px 0;
          margin:0 0 24px;padding:0}
.desc-box pre{background:transparent;color:#cdd6f4;padding:16px 20px;
              font-size:11px;line-height:1.8;overflow-x:auto;white-space:pre-wrap;margin:0}
pre{background:#1e1e2e;color:#cdd6f4;padding:18px 22px;border-radius:6px;
    font-size:11px;line-height:1.75;overflow-x:auto;white-space:pre-wrap;margin:12px 0}
.eval-box{background:#fff8f0;border-left:5px solid #e65100;
          padding:16px 22px;border-radius:4px;margin:14px 0 20px}
.est-day{font-size:22px;font-weight:bold;color:#c55a11}
.conf{font-size:13px;color:#555;margin-top:4px}
.thumb-grid{display:flex;flex-wrap:wrap;gap:12px;margin:12px 0}
.thumb-item{background:#fff;border:1px solid #dde3f0;border-radius:6px;padding:8px;text-align:center}
.thumb-cap{font-size:10px;color:#666;margin-top:4px}
.opinion{background:#f0f8ff;border:1px solid #90cdf4;border-radius:6px;
         padding:20px 28px;margin:16px 0;white-space:pre-wrap;
         font-family:monospace;font-size:12px;line-height:1.85}
@media print{body{background:white}.cover{-webkit-print-color-adjust:exact;print-color-adjust:exact}}
"""
        # 한글 섹션 제목
        if KO:
            s_abs   = "초록 (Abstract)"
            s_flow  = "1. 분석 전체 흐름 (Analysis Workflow)"
            s_color = "2. 색상 모델 및 산화 지표 (Color Models & Oxidation Metrics)"
            s_meth  = "3. 추정 방법론 (Estimation Methodology)"
            s_img   = "4. 이미지 분석 데이터 (Image Analysis Data)"
            s_raman = "5. 라만 참조 데이터 (Raman Reference Data)"
            s_eval  = "6. 평가 결과 (Evaluation Results)"
            s_est_c = "6.1 날짜 추정 근거 (Date Estimation Comment)"
            s_psr   = "6.2 Pseudo-Raman 추정 상세"
            s_guide = "7. 차트 해석 가이드 (Chart Guide)"
            s_final = "8. 최종 분석 의견 (Final Analytical Opinion)"
            cover_title  = "HfS₂ 박막 산화도 분석 보고서"
            cover_sub    = "다채널 색상계 분석 + Pseudo-Raman 추정"
            cover_gen    = "생성일시: " + gen_time
            no_raman     = "<p><em>라만 데이터가 로드되지 않았다.</em></p>"
            img_count    = "분석 완료 이미지: <strong>" + str(len(images)) + "</strong>개"
            r_count      = "총 라만 데이터: <strong>" + str(len(raman_data)) + "</strong>건"
            intro_txt    = self._report_intro_ko()
            flow_txt     = self._report_flowchart_svg(ko=True)
            color_txt    = self._report_color_model_ko()
            meth_txt     = self._report_estimation_method_ko()
            guide_txt    = self._report_chart_guide_ko()
        else:
            s_abs   = "Abstract"
            s_flow  = "1. Analysis Workflow Overview"
            s_color = "2. Color Space Models and Oxidation Metrics"
            s_meth  = "3. Estimation Methodology"
            s_img   = "4. Image Analysis Data"
            s_raman = "5. Raman Reference Data"
            s_eval  = "6. Evaluation Results"
            s_est_c = "6.1 Date Estimation Comment"
            s_psr   = "6.2 Pseudo-Raman Detail"
            s_guide = "7. Chart Interpretation Guide"
            s_final = "8. Final Analytical Opinion"
            cover_title  = "HfS\u2082 Thin Film Oxidation Analysis Report"
            cover_sub    = "Multi-channel Colorimetric Analysis with Pseudo-Raman Estimation"
            cover_gen    = "Generated: " + gen_time
            no_raman     = "<p><em>No Raman data loaded.</em></p>"
            img_count    = "Total analyzed: <strong>" + str(len(images)) + "</strong> images"
            r_count      = "Total entries: <strong>" + str(len(raman_data)) + "</strong>"
            intro_txt    = self._report_intro()
            flow_txt     = self._report_flowchart_svg()
            color_txt    = self._report_color_model()
            meth_txt     = self._report_estimation_method()
            guide_txt    = self._report_chart_guide()

        # ── Advanced 탭 결과 ──────────────────────────────
        adv_ctx = getattr(self, "_last_adv_ctx", None)
        if adv_ctx:
            if KO:
                s_adv = "9. 고급 분석 결과 (Advanced Analysis)"
            else:
                s_adv = "9. Advanced Analysis Results"
            adv_part = self._report_adv_section(adv_ctx, ko=KO)
        else:
            s_adv = ""
            adv_part = ""

        # ── 참고문헌 ──────────────────────────────────────
        ref_section = self._report_references(ko=KO)
        if KO:
            s_ref = "10. 참고문헌 (References)"
        else:
            s_ref = "10. References"

        return (
            "<!DOCTYPE html>\n<html lang=\"" + ("ko" if KO else "en") + "\">\n<head>\n"
            "<meta charset=\"UTF-8\">\n"
            "<title>" + cover_title + "</title>\n"
            "<style>" + css + "</style>\n"
            "</head>\n<body>\n"
            "<div class=\"cover\">\n"
            "  <h1>" + cover_title + "</h1>\n"
            "  <p>" + cover_sub + "</p>\n"
            "  <p style=\"margin-top:12px\">" + cover_gen + "</p>\n"
            "</div>\n"
            "<div class=\"content\">\n"
            + sec(s_abs,   "<p>" + intro_txt + "</p>")
            + sec(s_flow,  flow_txt)
            + sec(s_color, "<p>" + color_txt + "</p>")
            + sec(s_meth,  "<p>" + meth_txt  + "</p>")
            + sec(s_img,   "<p>" + img_count + "</p>" + img_tbl + thumbs)
            + sec(s_raman,
                  ("<p>" + r_count + "</p>" + raman_tbl + raman_charts)
                  if raman_data else no_raman)
            + sec(s_eval, eval_part)
            + (sec(s_est_c, pre(eval_comment), 3) if eval_comment else "")
            + (sec(s_psr,   pre(pseudo_result), 3) if pseudo_result else "")
            + sec(s_guide, "<p>" + guide_txt + "</p>")
            + sec(s_final, "<div class=\"opinion\">" + opinion + "</div>")
            + (sec(s_adv,  adv_part) if adv_part else "")
            + sec(s_ref,   ref_section)
            + "\n</div>\n</body>\n</html>"
        )


    def _report_adv_section(self, adv_ctx: dict, ko: bool = False) -> str:
        """Advanced 분석 결과 HTML 섹션 - 모델 설명 + 연산 순서 + 결과 포함"""
        def fmt(day, conf):
            if day is None: return "N/A"
            return f"{day:.1f}d ({conf:.0f}%)"

        if ko:
            overview = (
                "<h3>9.1 분석 방법론 개요</h3>"
                "<p>기존 KNN(가중 유클리드 거리) 외 4가지 독립적 추정 방법을 추가하여 "
                "5가지 방법의 앙상블로 날짜를 추정한다. 각 방법은 이미지의 서로 다른 측면을 분석하므로, "
                "여러 방법이 동의할수록 추정의 신뢰도가 높아집니다.</p>"

                "<h3>9.2 각 방법의 개념과 연산 순서</h3>"

                "<h4>① KNN — 가중 유클리드 거리</h4>"
                "<p><strong>개념:</strong> 4개 색상 지표(b*, S-채널, YI, ΔE)를 정규화 후 "
                "가중 유클리드 거리를 계산하여 역거리 가중 평균으로 날짜를 추정한다.<br>"
                "<strong>연산:</strong> dist = √(0.45·Δb*² + 0.30·ΔS² + 0.25·ΔYI²)<br>"
                "est_day = Σ(day_i/dist_i) / Σ(1/dist_i)<br>"
                "confidence = max(0, 100 - dist_min × 200)</p>"

                "<h4>② Wasserstein EMD — b* 히스토그램 Earth Mover Distance</h4>"
                "<p><strong>개념:</strong> 단순 평균이 아닌 b* 확률 분포 전체를 비교한다. "
                "같은 평균 b*라도 단봉(신선) vs 쌍봉(부분 산화)은 완전히 다른 산화 상태를 의미한다. "
                "EMD는 한 분포를 다른 분포로 변환하는 데 필요한 최소 비용이다.<br>"
                "<strong>연산:</strong> b* 값을 64-bin 히스토그램으로 변환(-30~80 범위) 후<br>"
                "W(P,Q) = Σ|CDF_P(i) - CDF_Q(i)| / N (누적분포함수 L1-norm)<br>"
                "confidence = max(0, 100 - W_min × 500)</p>"

                "<h4>③ FFT Texture — 주파수 도메인 텍스처 분석</h4>"
                "<p><strong>개념:</strong> 색상 정보와 무관한 표면 거칠기(texture)를 2D FFT로 분석한다. "
                "산화될수록 표면이 거칠어져 고주파 성분이 증가한다.<br>"
                "<strong>연산:</strong> 그레이스케일 변환 → 2D FFT → 주파수 중심 이동<br>"
                "고주파 에너지 비율 = (반경 40% 초과 에너지) / (전체 에너지, DC 제외)<br>"
                "dist = 0.5·|ΔHF비율| + 0.3·|Δ엔트로피| + 0.2·(1-코사인유사도)<br>"
                "confidence = max(0, 100 - dist_min × 300)</p>"

                "<h4>④ Spatial Pattern — 공간 산화 패턴 분석</h4>"
                "<p><strong>개념:</strong> 세그먼트 그리드(예: 3×3=9개 구역)로 이미지를 분할하여 "
                "각 구역의 b* 값을 비교한다. 산화는 공간적으로 불균일하게 진행되므로 "
                "이 공간 패턴이 산화 단계 정보를 담고 있다.<br>"
                "<strong>추출 피처:</strong><br>"
                "• 공간 엔트로피: 세그먼트 간 b* 표준편차 → 산화 불균일도<br>"
                "• 경계-중심 기울기: 경계 세그먼트 - 중심 세그먼트 b* 평균 → 산화 전파 방향<br>"
                "• 이방성: 가로/세로 방향 분산 비율 → 방향성 산화 여부<br>"
                "<strong>연산:</strong> dist = 0.4·|Δ엔트로피| + 0.4·|Δ경계기울기| + 0.2·|Δ이방성|<br>"
                "confidence = max(0, 100 - dist_min × 400)</p>"

                "<h4>⑤ Kinetic Model — 지수 감쇠 물리 모델</h4>"
                "<p><strong>개념:</strong> 산화는 물리화학적 반응 법칙을 따릅니다. "
                "참조 DB에서 조건별로 b*(t) = b*∞ + (b*₀ - b*∞)·exp(-k·t) 모델을 피팅하고, "
                "대상 이미지의 b* 값을 역대입하여 t(날짜)를 역산한다. "
                "이 방법은 참조 DB 범위 밖도 추정 가능한 유일한 방법이다.<br>"
                "<strong>연산:</strong><br>"
                "1단계: 참조 DB를 조건별로 그룹화 → scipy.optimize.curve_fit으로 피팅<br>"
                "   파라미터: b*₀(초기값), b*∞(포화값), k(감쇠상수)<br>"
                "2단계: R²(결정계수)로 모델 적합도 평가 → 최적 조건 선택<br>"
                "3단계: 역산 — t = -ln((b* - b*∞) / (b*₀ - b*∞)) / k<br>"
                "4단계: confidence = R² × 100 × 역산 안정도 (ratio가 0~1 범위일 때 높음)</p>"

                "<h3>9.3 앙상블 통합 연산</h3>"
                "<p>5가지 방법의 추정치와 신뢰도를 통합한다.<br>"
                "1단계: 신뢰도 20% 미만인 추정치 제외<br>"
                "2단계: weight_i = confidence_i / Σ(confidence_i)<br>"
                "3단계: ensemble_day = Σ(est_day_i × weight_i)<br>"
                "4단계: 일치도 보너스 = max(0, 10 - 방법 간 최대 편차 × 2)<br>"
                "5단계: 앙상블 신뢰도 = 가중 평균 신뢰도 + 일치도 보너스<br>"
                "<em>→ 여러 방법이 같은 날짜를 가리킬수록 신뢰도가 높아집니다.</em></p>"
            )
        else:
            overview = (
                "<h3>9.1 Methodology Overview</h3>"
                "<p>Four additional estimation methods augment the baseline KNN approach. "
                "Each analyzes a different image aspect; agreement among methods increases confidence.</p>"

                "<h3>9.2 Method Concepts and Computation</h3>"

                "<h4>① KNN — Weighted Euclidean Distance</h4>"
                "<p><strong>Concept:</strong> Normalizes 4 color metrics (b*, S-channel, YI, ΔE), "
                "computes weighted Euclidean distance to each reference, estimates day by "
                "inverse-distance weighted mean of the closest matches.<br>"
                "<strong>Computation:</strong> dist = √(0.45·Δb*² + 0.30·ΔS² + 0.25·ΔYI²)<br>"
                "est_day = Σ(day_i/dist_i) / Σ(1/dist_i)<br>"
                "confidence = max(0, 100 - dist_min × 200)</p>"

                "<h4>② Wasserstein EMD — b* Histogram Earth Mover Distance</h4>"
                "<p><strong>Concept:</strong> Compares the full b* probability distribution "
                "rather than just the mean. The same mean b* can represent very different oxidation "
                "states: unimodal (fresh) vs bimodal (partially oxidized). "
                "EMD measures the minimum cost to transform one distribution into another.<br>"
                "<strong>Computation:</strong> Convert b* to 64-bin histogram (range -30 to +80), then<br>"
                "W(P,Q) = Σ|CDF_P(i) - CDF_Q(i)| / N (L1-norm of CDFs)<br>"
                "confidence = max(0, 100 - W_min × 500)</p>"

                "<h4>③ FFT Texture — Frequency Domain Analysis</h4>"
                "<p><strong>Concept:</strong> Analyzes surface roughness via 2D FFT, "
                "independent of color information. As HfS₂ oxidizes, surface roughness increases, "
                "boosting high-frequency spectral energy.<br>"
                "<strong>Computation:</strong> Grayscale → 2D FFT → center shift<br>"
                "HF energy ratio = (energy at radius > 40%) / (total energy, DC excluded)<br>"
                "dist = 0.5·|ΔHF_ratio| + 0.3·|Δentropy| + 0.2·(1-cosine_similarity)<br>"
                "confidence = max(0, 100 - dist_min × 300)</p>"

                "<h4>④ Spatial Pattern — Spatial Oxidation Pattern Analysis</h4>"
                "<p><strong>Concept:</strong> Divides the image into a segment grid (e.g. 3×3=9 zones) "
                "and compares b* across zones. Oxidation is spatially non-uniform, "
                "and the spatial pattern carries oxidation-stage information.<br>"
                "<strong>Extracted features:</strong><br>"
                "• Spatial entropy: std of b* across segments → oxidation uniformity<br>"
                "• Boundary-center gradient: boundary mean - center mean → propagation direction<br>"
                "• Anisotropy: row/column variance ratio → directional oxidation<br>"
                "<strong>Computation:</strong> dist = 0.4·|Δentropy| + 0.4·|Δboundary_grad| + 0.2·|Δanisotropy|<br>"
                "confidence = max(0, 100 - dist_min × 400)</p>"

                "<h4>⑤ Kinetic Model — Exponential Decay Physical Model</h4>"
                "<p><strong>Concept:</strong> Oxidation follows physical chemistry laws. "
                "The model b*(t) = b*∞ + (b*₀ - b*∞)·exp(-k·t) is fitted from the reference DB, "
                "and the target b* is back-calculated to estimate t (day). "
                "This is the only method capable of extrapolating beyond the reference DB range.<br>"
                "<strong>Computation:</strong><br>"
                "Step 1: Group reference DB by condition → fit with scipy.optimize.curve_fit<br>"
                "   Parameters: b*₀ (initial), b*∞ (saturation), k (decay constant)<br>"
                "Step 2: Evaluate R² → select best-matching condition<br>"
                "Step 3: Back-calculate: t = -ln((b* - b*∞) / (b*₀ - b*∞)) / k<br>"
                "Step 4: confidence = R² × 100 × stability (higher when ratio is in 0-1 range)</p>"

                "<h3>9.3 Ensemble Integration</h3>"
                "<p>Integrates estimates from up to 5 methods.<br>"
                "Step 1: Exclude methods with confidence < 20%<br>"
                "Step 2: weight_i = confidence_i / Σ(confidence_i)<br>"
                "Step 3: ensemble_day = Σ(est_day_i × weight_i)<br>"
                "Step 4: consistency_bonus = max(0, 10 - max_spread × 2)<br>"
                "Step 5: ensemble_confidence = weighted_mean_confidence + consistency_bonus<br>"
                "<em>→ The more methods agree, the higher the ensemble confidence.</em></p>"
            )

        methods = [
            ("knn",     "① KNN (Weighted Euclidean)",  "🔵"),
            ("wass",    "② Wasserstein (b* EMD)",       "🟢"),
            ("fft",     "③ FFT Texture",                "🟡"),
            ("spatial", "④ Spatial Pattern",            "🟠"),
            ("kinetic", "⑤ Kinetic Model",              "🟣"),
            ("ens",     "⑥ Ensemble",                   "🔴"),
        ]
        rows_html = ""
        for key, label, icon in methods:
            day  = adv_ctx.get(f"{key}_day")
            conf = adv_ctx.get(f"{key}_conf", 0)
            w    = adv_ctx.get("weights", {}).get(key, 0)
            bold_style = ' style="font-weight:bold;background:#fff8e1"' if key == "ens" else ""
            rows_html += (
                f"<tr{bold_style}><td>{icon} {label}</td>"
                f'<td style="text-align:center">{fmt(day, conf)}</td>'
                f'<td style="text-align:center">{w*100:.0f}%</td></tr>\n'
            )
        tbl_lbl = "9.4 방법별 추정 결과" if ko else "9.4 Method Comparison Results"
        tbl_html = (
            f"<h3>{tbl_lbl}</h3>"
            '<table border="1" style="border-collapse:collapse;width:100%;margin:8px 0">'
            f'<tr style="background:#e8eaf6"><th>{"방법" if ko else "Method"}</th>'
            f'<th>{"추정일(신뢰도)" if ko else "Est. Day (Confidence)"}</th>'
            f'<th>{"앙상블 가중치" if ko else "Ensemble Weight"}</th></tr>'
            + rows_html + "</table>"
        )

        k_detail = adv_ctx.get("kinetic_detail", {})
        p = k_detail.get("model_params", {})
        kp_html = ""
        if p:
            kp_lbl = "9.5 Kinetic 모델 파라미터" if ko else "9.5 Kinetic Model Parameters"
            kp_html = (
                f"<h3>{kp_lbl}</h3><p>"
                f'{"적용 조건" if ko else "Applied condition"}: '
                f"<strong>{k_detail.get('cond_used', '')}</strong><br>"
                "b*(t) = b*&infin; + (b*&#8320; &minus; b*&infin;)&middot;exp(&minus;k&middot;t)<br>"
                f"b*&#8320; = {p.get('b0', 0):.2f} &nbsp;|&nbsp; "
                f"b*&infin; = {p.get('b_inf', 0):.2f} &nbsp;|&nbsp; "
                f"k = {p.get('k', 0):.5f} &nbsp;|&nbsp; "
                f"R&sup2; = {p.get('r2', 0):.4f}</p>"
            )

        ts = adv_ctx.get("t_spatial", {})
        tf = adv_ctx.get("t_fft", {})
        feat_html = ""
        if ts or tf:
            feat_lbl = "9.6 피처 상세" if ko else "9.6 Feature Details"
            feat_html = f"<h3>{feat_lbl}</h3><p>"
            if ts:
                ent = ts.get("entropy", 0)
                bg  = ts.get("boundary_grad", 0)
                ani = ts.get("anisotropy", 1)
                if ko:
                    ent_i = "불균일 산화" if ent > 15 else ("중간" if ent > 8 else "균일")
                    bg_i  = "경계부 산화 우세" if bg > 3 else ("중심부 산화" if bg < -3 else "균일 분포")
                else:
                    ent_i = "non-uniform" if ent > 15 else ("moderate" if ent > 8 else "uniform")
                    bg_i  = "boundary-dominated" if bg > 3 else ("center-dominated" if bg < -3 else "uniform")
                feat_html += (
                    f'{"공간 엔트로피" if ko else "Spatial Entropy"}: {ent:.3f} ({ent_i}) | '
                    f'{"경계 기울기" if ko else "Boundary Gradient"}: {bg:.3f} ({bg_i}) | '
                    f'{"이방성" if ko else "Anisotropy"}: {ani:.3f}<br>'
                )
            if tf:
                hf   = tf.get("hf_ratio", 0)
                ent2 = tf.get("entropy", 0)
                if ko:
                    hf_i = "표면 거칠음 높음" if hf > 0.35 else ("중간" if hf > 0.20 else "균일 표면")
                else:
                    hf_i = "rough surface" if hf > 0.35 else ("moderate" if hf > 0.20 else "uniform surface")
                feat_html += (
                    f'{"고주파 에너지 비율" if ko else "HF Energy Ratio"}: {hf*100:.1f}% ({hf_i}) | '
                    f'{"스펙트럼 엔트로피" if ko else "Spectral Entropy"}: {ent2:.4f}'
                )
            feat_html += "</p>"

        return overview + tbl_html + kp_html + feat_html

    def _report_references(self, ko: bool = False) -> str:
        """참고문헌 섹션 HTML"""
        if ko:
            return """
<ol style="line-height:1.9">
<li>
  Hwang, J., Mun, J., Lee, K.-T., Lee, T., Kim, J., Min, J., &amp; <strong>Park, K.</strong> (2025).
  Impact of humidity on long-term stability of HfS<sub>2</sub> grown on sapphire substrate
  by chemical vapor deposition and strategies to prevent native oxidation.
  <em>Materials Science in Semiconductor Processing</em>, <strong>192</strong>, 109471.
  https://doi.org/10.1016/j.mssp.2025.109471
</li>
<li>
  CIE (1976). Colorimetry (2nd ed.). CIE Publication 15.2.
  Commission Internationale de l'Éclairage.
</li>
<li>
  ASTM International. (2015). ASTM E313: Standard Practice for Calculating
  Yellowness and Whiteness Indices from Instrumentally Measured Color Coordinates.
  West Conshohocken, PA.
</li>
<li>
  Villanueva-Luna, A., et al. (2019). Wasserstein distance as a similarity metric
  for spectral analysis of thin film oxidation. <em>Applied Surface Science</em>.
</li>
</ol>"""
        return """
<ol style="line-height:1.9">
<li>
  Hwang, J., Mun, J., Lee, K.-T., Lee, T., Kim, J., Min, J., &amp; <strong>Park, K.</strong> (2025).
  Impact of humidity on long-term stability of HfS<sub>2</sub> grown on sapphire substrate
  by chemical vapor deposition and strategies to prevent native oxidation.
  <em>Materials Science in Semiconductor Processing</em>, <strong>192</strong>, 109471.
  https://doi.org/10.1016/j.mssp.2025.109471
</li>
<li>
  CIE (1976). Colorimetry (2nd ed.). CIE Publication 15.2.
  Commission Internationale de l'Éclairage.
</li>
<li>
  ASTM International. (2015). ASTM E313: Standard Practice for Calculating
  Yellowness and Whiteness Indices from Instrumentally Measured Color Coordinates.
  West Conshohocken, PA.
</li>
</ol>"""

    def _build_report_docx(self, path, images, raman_data, eval_ctx,
                             chart_b64, eval_comment,
                             pseudo_result, gen_time, lang="en"):
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()
        sec = doc.sections[0]
        sec.page_width = Cm(21); sec.page_height = Cm(29.7)
        sec.left_margin = sec.right_margin = Cm(2.5)
        sec.top_margin  = sec.bottom_margin = Cm(2)

        def ap(text="", bold=False, size=10, color=None, align=None):
            p = doc.add_paragraph()
            if align: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            r.font.name = "Arial"; r.font.size = Pt(size); r.font.bold = bold
            if color: r.font.color.rgb = RGBColor(*color)
            return p

        def ah(text, level=1):
            p = doc.add_heading(text, level=level)
            for r in p.runs:
                r.font.name = "Arial"
            return p

        def add_img(b64, width=5.5):
            if not b64: return
            try:
                doc.add_picture(io.BytesIO(base64.b64decode(b64)),
                                width=Inches(width))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception: pass

        def add_tbl(headers, rows):
            t = doc.add_table(rows=1, cols=len(headers))
            t.style = "Table Grid"
            for j, h2 in enumerate(headers):
                c = t.rows[0].cells[j]; c.text = ""
                r2 = c.paragraphs[0].add_run(h2)
                r2.font.bold = True; r2.font.size = Pt(8); r2.font.name = "Arial"
                r2.font.color.rgb = RGBColor(255,255,255)
                pr = c._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto")
                shd.set(qn("w:fill"),"1F3864"); pr.append(shd)
            for row in rows:
                tr = t.add_row()
                for j, val in enumerate(row):
                    c = tr.cells[j]; c.text = ""
                    r2 = c.paragraphs[0].add_run(str(val))
                    r2.font.size = Pt(8); r2.font.name = "Arial"

        KO = (lang == "ko")
        def fv(v, d=2):
            try: return f"{float(v):.{d}f}"
            except: return "N/A"

        # 표지
        doc.add_paragraph()
        cover_title = ("HfS₂ 박막 산화도 분석 보고서" if KO
                       else "HfS2 Thin Film Oxidation Analysis Report")
        ap(cover_title, bold=True, size=20, color=(31,56,100), align=True)
        gen_label = ("생성일시: " if KO else "Generated: ")
        ap(f"{gen_label}{gen_time}", size=11, color=(100,100,100), align=True)
        doc.add_paragraph()

        workflow_ko = (
            "전체 분석 흐름:\n\n"
            "1단계. 입력 (INPUT) — 대상 이미지와 참조 이미지 DB(알려진 산화 일수로 분석된 이미지들)를 로드한다.\n\n"
            "2단계. ROI 선택 — 대상 이미지에서 사용자 정의 관심 영역(ROI)을 크롭한다. "
            "ROI 미설정 시 전체 이미지를 사용한다.\n\n"
            "3단계. 색상 특성 추출 (Colorimetric Feature Extraction) — 4개 지표 계산: "
            "CIE Lab b*(노란-파란 축), HSI S-채널(채도), ASTM E313 황색지수(YI), CIE ΔE(Day-0 대비 색차).\n\n"
            "4단계. 가중 유클리드 거리 — d = sqrt(w_b*Δb*^2 + w_S*ΔS^2 + w_YI*ΔYI^2) "
            "기본 가중치: w_b=0.45, w_S=0.30, w_YI=0.25.\n\n"
            "5단계A. 날짜 추정 (Date Estimation) — Top-k 최근접 참조를 선택하고, "
            "역거리 가중 평균으로 노출 일수 추정. 신뢰도 = max(0, 100 - dist_min * 200).\n\n"
            "5단계B. Pseudo-Raman 회귀 — 병렬로 4개 단변량 선형 회귀(b*→라만, S→라만, YI→라만, ΔE→라만)를 "
            "이미지-라만 쌍에 학습, R²-가중 앙상블로 정규화 A₁g 피크 강도 + 95% CI 추정.\n\n"
            "6단계. 통합 판정 — 날짜 추정, Pseudo-Raman 피크, 지표값을 결합하여 Stage I~IV 산화 단계 분류.\n\n"
            "7단계. 출력 — HTML 또는 Word 보고서 (차트, 표, 최종 분석 의견 포함)."
        )
        workflow_en = (
            "The complete analysis pipeline proceeds as follows:\n\n"
            "Step 1. INPUT — A target photographic image is loaded along with a "
            "reference image database (images analyzed at known oxidation days).\n\n"
            "Step 2. ROI SELECTION — A user-defined Region of Interest is cropped "
            "from the target image. If no ROI is defined, the full image is used.\n\n"
            "Step 3. COLORIMETRIC FEATURE EXTRACTION — Four complementary metrics "
            "are computed: CIE Lab b* (yellow-blue axis), HSI S-channel (chromatic "
            "purity), ASTM E313 Yellowness Index (YI), and CIE DeltaE (cumulative "
            "color change from Day-0 reference).\n\n"
            "Step 4. WEIGHTED EUCLIDEAN DISTANCE — Target metrics are compared to "
            "each reference image using: d = sqrt(w_b*Db*^2 + w_S*DS^2 + w_YI*DYI^2) "
            "where default weights are w_b=0.45, w_S=0.30, w_YI=0.25.\n\n"
            "Step 5A. DATE ESTIMATION — Top-k nearest references are selected; "
            "the aging day is estimated as their inverse-distance weighted average.\n\n"
            "Step 5B. PSEUDO-RAMAN REGRESSION — In parallel, four univariate linear "
            "regressions combined via R2-weighted ensemble to estimate the normalized "
            "A1g Raman peak intensity with 95% CI.\n\n"
            "Step 6. INTEGRATED ASSESSMENT — Stage I-IV oxidation classification.\n\n"
            "Step 7. OUTPUT — HTML or Word report with all charts, tables, and opinion."
        )

        if KO:
            sections = [
                ("초록 (Abstract)", self._report_intro_ko()),
                ("1. 분석 전체 흐름 (Analysis Workflow)", workflow_ko),
                ("2. 색상 모델 및 산화 지표 (Color Models)", self._report_color_model_ko()),
                ("3. 추정 방법론 (Estimation Methodology)", self._report_estimation_method_ko()),
            ]
        else:
            sections = [
                ("Abstract", self._report_intro()),
                ("1. Analysis Workflow Overview", workflow_en),
                ("2. Color Space Models and Oxidation Metrics", self._report_color_model()),
                ("3. Estimation Methodology", self._report_estimation_method()),
            ]

        for title, body in sections:
            ah(title)
            for line in body.split("\n"):
                if line.strip(): ap(line, size=10)

        ah("4. 이미지 분석 데이터" if KO else "4. Image Analysis Data")
        ap(f"Total analyzed images: {len(images)}", bold=True)
        if images:
            if KO:
                hdrs = ["파일명(Name)","조건(Cond)","날짜(Day)",
                        "b*(노란도)","채도(S-ch)","황색지수(YI)",
                        "색차(ΔE)","황색비율(Yellow%)"]
            else:
                hdrs = ["Name","Cond","Day","b*","S-ch","YI","dE","Yellow%"]
            rows = [[i["name"][:20], i["cond"], i["day"], i["b"], i["s"],
                     i["yi"], i["de"], i["yr"]] for i in images[:40]]
            add_tbl(hdrs, rows)

        ah("5. 라만 참조 데이터 (Raman Reference Data)" if KO
           else "5. Raman Reference Data")
        if raman_data:
            ap(f"{'총 라만 데이터' if KO else 'Total entries'}: {len(raman_data)}",
               bold=True)
            if KO:
                r_h = ["조건(Condition)","날짜(Day)","피크(Peak)",
                       "정규화피크(Norm Peak)","위치(Shift)","범위(Range)"]
            else:
                r_h = ["Condition","Day","Peak","Norm Peak","Shift","Range"]
            r_r = []
            for r in raman_data[:40]:
                def fv2(v, d=4):
                    try: return f"{float(v):.{d}f}"
                    except: return "-"
                r_r.append([r.get("cond",""), r.get("day",""),
                    fv2(r.get("peak")), fv2(r.get("norm_peak")),
                    fv2(r.get("peak_shift"),1), r.get("peak_range","-")])
            add_tbl(r_h, r_r)
            if KO:
                raman_chart_list = [
                    ("raman_trend",    "그림 5. 라만 피크 강도 추세 (Raman Peak Trend)",
                     lambda: self._chart_desc_raman_trend(raman_data, ko=True)),
                    ("raman_spectrum", "그림 6. 라만 스펙트럼 오버레이 (Spectra Overlay)",
                     lambda: self._chart_desc_raman_spectrum(raman_data, ko=True)),
                    ("raman_decay",    "그림 7. 조건별 A₁g 피크 감쇠율 (Decay Rate)",
                     lambda: self._chart_desc_raman_decay(raman_data, ko=True)),
                ]
            else:
                raman_chart_list = [
                    ("raman_trend",    "Figure 5. Raman Peak Trend",
                     lambda: self._chart_desc_raman_trend(raman_data)),
                    ("raman_spectrum", "Figure 6. Raman Spectra Overlay",
                     lambda: self._chart_desc_raman_spectrum(raman_data)),
                    ("raman_decay",    "Figure 7. A1g Peak Decay by Condition",
                     lambda: self._chart_desc_raman_decay(raman_data)),
                ]
            for k, t, desc_fn in raman_chart_list:
                if chart_b64.get(k):
                    ah(t, level=2)
                    add_img(chart_b64[k])
                    for line in desc_fn().split("\n"):
                        if line.strip(): ap(line, size=9)
                    doc.add_paragraph()

        ah("6. 평가 결과 (Evaluation Results)" if KO
           else "6. Evaluation Results")
        if eval_ctx.get("target"):
            est  = eval_ctx.get("est_day"); conf = eval_ctx.get("confidence",0)
            tgt2 = eval_ctx.get("target",{})
            p2 = doc.add_paragraph()
            if KO:
                est_txt = f"추정 산화 일수 (Estimated Day): {fv(est,1)}일  (신뢰도: {fv(conf,0)}%)"
            else:
                est_txt = f"Estimated Day: {fv(est,1)}d  (Confidence: {fv(conf,0)}%)"
            r2 = p2.add_run(est_txt)
            r2.font.bold = True; r2.font.size = Pt(14); r2.font.name = "Arial"
            r2.font.color.rgb = RGBColor(197,90,17)
            if KO:
                ap(f"b*(노란도)={fv(tgt2.get('lab',{}).get('b'))}  "
                   f"S-채널(채도)={fv(tgt2.get('s_mean'),0)}  "
                   f"황색지수(YI)={fv(tgt2.get('yellowness_idx'),0)}")
            else:
                ap(f"b*={fv(tgt2.get('lab',{}).get('b'))}  "
                   f"S={fv(tgt2.get('s_mean'),0)}  YI={fv(tgt2.get('yellowness_idx'),0)}")
            if KO:
                eval_chart_list = [
                    ("radar",          "그림 1. 지표 유사도 레이더 (Metric Similarity Radar)",
                     lambda: self._chart_desc_radar(eval_ctx, ko=True)),
                    ("timeline",       "그림 2. 거리 타임라인 (Distance Timeline)",
                     lambda: self._chart_desc_timeline(eval_ctx, ko=True)),
                    ("pseudo_reg_fig", "그림 3. 회귀: 이미지 지표 → 라만 피크 (Regression)",
                     lambda: self._chart_desc_regression(eval_ctx, ko=True)),
                    ("pseudo_spec_fig","그림 4. 추정 스펙트럼 + 95% CI (Estimated Spectrum)",
                     lambda: self._chart_desc_spectrum(eval_ctx, ko=True)),
                ]
            else:
                eval_chart_list = [
                    ("radar",          "Figure 1. Metric Similarity Radar",
                     lambda: self._chart_desc_radar(eval_ctx)),
                    ("timeline",       "Figure 2. Distance Timeline",
                     lambda: self._chart_desc_timeline(eval_ctx)),
                    ("pseudo_reg_fig", "Figure 3. Regression: Image Metrics to Raman Peak",
                     lambda: self._chart_desc_regression(eval_ctx)),
                    ("pseudo_spec_fig","Figure 4. Estimated Spectrum + 95% CI",
                     lambda: self._chart_desc_spectrum(eval_ctx)),
                ]
            for k, t, desc_fn in eval_chart_list:
                if chart_b64.get(k):
                    ah(t, level=2)
                    add_img(chart_b64[k])
                    for line in desc_fn().split("\n"):
                        if line.strip(): ap(line, size=9)
                    doc.add_paragraph()
            if eval_comment:
                ah("날짜 추정 근거 (Date Estimation Comment)" if KO
                   else "Date Estimation Comment", level=2)
                for l in eval_comment.split("\n"):
                    if l.strip(): ap(l, size=9)
            if pseudo_result:
                ah("Pseudo-Raman 추정 상세" if KO
                   else "Pseudo-Raman Detail", level=2)
                for l in pseudo_result.split("\n"):
                    if l.strip(): ap(l, size=9)

        ah("7. 차트 해석 가이드 (Chart Interpretation Guide)" if KO
           else "7. Chart Interpretation Guide")
        guide = self._report_chart_guide_ko() if KO else self._report_chart_guide()
        for line in guide.split("\n"):
            if line.strip(): ap(line, size=10)

        ah("8. 최종 분석 의견 (Final Analytical Opinion)" if KO
           else "8. Final Analytical Opinion")
        for line in self._report_final_opinion(
                images, eval_ctx, eval_comment, pseudo_result,
                ko=KO).split("\n"):
            ap(line, size=10)

        # ── 9. Advanced 분석 결과 ──────────────────────────
        adv_ctx = getattr(self, "_last_adv_ctx", None)
        if adv_ctx:
            ah("9. 고급 분석 결과 (Advanced Analysis)" if KO
               else "9. Advanced Analysis Results")
            def fmt_d(day, conf):
                return f"{day:.1f}일 ({conf:.0f}%)" if day is not None else "N/A"

            methods_w = [
                ("knn",     "KNN (Weighted Euclidean)"),
                ("wass",    "Wasserstein EMD"),
                ("fft",     "FFT Texture"),
                ("spatial", "Spatial Pattern"),
                ("kinetic", "Kinetic Model"),
                ("ens",     "Ensemble"),
            ]
            hdr_row = (["방법(Method)", "추정일(Day)", "가중치(Weight)"]
                       if KO else ["Method", "Est. Day", "Weight"])
            rows_w = []
            for key, lbl in methods_w:
                day  = adv_ctx.get(f"{key}_day")
                conf = adv_ctx.get(f"{key}_conf", 0)
                w    = adv_ctx.get("weights", {}).get(key, 0)
                rows_w.append([lbl, fmt_d(day, conf), f"{w*100:.0f}%"])
            add_tbl(hdr_row, rows_w)

            # Kinetic 모델 상세
            k_detail = adv_ctx.get("kinetic_detail", {})
            p = k_detail.get("model_params", {})
            if p:
                ap(f"Kinetic: b*(t)=b*∞+(b*₀−b*∞)·exp(−k·t)  "
                   f"[{k_detail.get('cond_used','')}]  "
                   f"k={p.get('k',0):.5f}  R²={p.get('r2',0):.4f}", size=9)

            # Spatial 상세
            ts = adv_ctx.get("t_spatial", {})
            if ts:
                ap(f"Spatial: entropy={ts.get('entropy',0):.3f}  "
                   f"boundary_grad={ts.get('boundary_grad',0):.3f}  "
                   f"anisotropy={ts.get('anisotropy',1):.3f}", size=9)

        # ── 10. 참고문헌 ───────────────────────────────────
        ah("10. 참고문헌 (References)" if KO else "10. References")
        refs = [
            ("Hwang, J., Mun, J., Lee, K.-T., Lee, T., Kim, J., Min, J., "
             "& Park, K. (2025). Impact of humidity on long-term stability "
             "of HfS₂ grown on sapphire substrate by chemical vapor deposition "
             "and strategies to prevent native oxidation. "
             "Materials Science in Semiconductor Processing, 192, 109471. "
             "https://doi.org/10.1016/j.mssp.2025.109471"),
            ("CIE (1976). Colorimetry (2nd ed.). CIE Publication 15.2."),
            ("ASTM International. (2015). ASTM E313: Standard Practice for "
             "Calculating Yellowness and Whiteness Indices from Instrumentally "
             "Measured Color Coordinates."),
        ]
        for i, ref in enumerate(refs, 1):
            p_ref = doc.add_paragraph(style="List Number")
            r_ref = p_ref.add_run(ref)
            r_ref.font.size = Pt(9)
            r_ref.font.name = "Arial"

        doc.save(path)


    def _load_all_db(self):
        """
        여러 DB 파일을 한꺼번에 선택 → 종류 자동 감지 후 일괄 로드.
        지원 파일:
          *.db         → 이미지 분석 데이터 (images 테이블)
          *.raman.db   → Raman 데이터 (raman_data 테이블)
          *.target.db  → 평가 대상 이미지 (eval_target 테이블)
          이름에 raman / target 포함하지 않는 .db → 이미지 우선 시도
        """
        paths = filedialog.askopenfilenames(
            title="Load All — Select DB files (multiple OK)",
            filetypes=[
                ("All DB files", "*.db"),
                ("Image DB",     "*.db"),
                ("Raman DB",     "*.raman.db"),
                ("Target DB",    "*.target.db"),
                ("All",          "*.*"),
            ])
        if not paths:
            return

        import sqlite3 as _sq

        n_img = 0; n_raman = 0; loaded_target = False
        errors = []

        for path in paths:
            fname = os.path.basename(path).lower()
            try:
                # read-only 로 테이블 조회 — 9p 마운트 lock 회피
                con = _db_open_read(path)
                tables = {r[0] for r in
                          con.execute("SELECT name FROM sqlite_master"
                                      " WHERE type='table'").fetchall()}
                con.close()
            except Exception as ex:
                errors.append(f"{os.path.basename(path)}: {ex}")
                continue

            # ── 이미지 데이터 ──────────────────────────
            if "images" in tables:
                try:
                    records = db_load_all(path)
                    existing = {(img["name"], img["cond"], img["day"])
                                for img in self.images}
                    added = 0
                    for rec in records:
                        key_ = (rec["name"], rec["cond"], rec["day"])
                        if key_ in existing:
                            continue
                        rec["day_var"]  = tk.StringVar(value=rec["day"])
                        rec["cond_var"] = tk.StringVar(value=rec["cond"])
                        # DB 의 ROI 는 그대로 보존 + 품질만 평가 (자동 추정 X)
                        if rec.get("roi") is not None and rec.get("rgb") is not None:
                            flg, why = evaluate_roi_quality(rec["rgb"], rec["roi"])
                            rec["roi_flag"] = flg
                            rec["roi_reason"] = f"DB 로드: {why}"
                            rec["roi_source"] = "db"
                        else:
                            rec["roi_flag"] = None
                            rec["roi_reason"] = ""
                            rec["roi_source"] = "db"
                        self.images.append(rec)
                        existing.add(key_)
                        added += 1
                    n_img += added
                except Exception as ex:
                    errors.append(f"{os.path.basename(path)} (images): {ex}")

            # ── Raman 데이터 (id 보존 — 매칭 복원에 필수) ──
            if "raman_data" in tables:
                try:
                    new_raman = db_load_raman_all(path)
                    if new_raman:
                        # 기존 _raman_data 와 _id 충돌 회피 (다른 통합 DB 합칠 때)
                        existing_ids = set(r.get("_id")
                                           for r in self._raman_data
                                           if r.get("_id") is not None)
                        max_id = max(existing_ids, default=0)
                        for r in new_raman:
                            rid = r.get("_id")
                            if rid in existing_ids:
                                # id 충돌 — 새 id 로 재할당, 매칭은 깨짐
                                max_id += 1
                                r["_id"] = max_id
                                existing_ids.add(max_id)
                            else:
                                existing_ids.add(rid)
                                if rid is not None:
                                    max_id = max(max_id, rid)
                            self._raman_data.append(r)
                            n_raman += 1
                except Exception as ex:
                    errors.append(f"{os.path.basename(path)} (raman): {ex}")

            # ── 평가 대상 이미지 ───────────────────────
            if "eval_target" in tables:
                try:
                    ok = self._db_load_target(path)
                    if ok:
                        loaded_target = True
                except Exception as ex:
                    errors.append(f"{os.path.basename(path)} (target): {ex}")

        # UI 갱신
        if n_img > 0:
            if self.sel_idx < 0 and self.images:
                self.sel_idx = 0
            self._sort_images_by_cond_day()
            self._rebuild_list()
            self._refresh_orig()
            self._refresh_hsi()

        if n_raman > 0:
            self._normalize_raman()
            # 매칭 보강: 이미 raman_id 가 복원된 이미지는 그대로,
            # 새로 매칭 가능한 (cond+day 일치 + raman_id None) 이미지에는 자동 매칭
            self._auto_link_raman_by_cond_day()
            self._rebuild_raman_tree()
            self._refresh_raman_tab()

        if loaded_target and hasattr(self, "_pred_draw_preview"):
            self.after(200, self._pred_draw_preview)

        # 결과 요약
        parts = []
        if n_img    > 0: parts.append(f"{n_img} image records")
        if n_raman  > 0: parts.append(f"{n_raman} Raman entries")
        if loaded_target:  parts.append("evaluation target")
        if not parts:
            summary = "Nothing new loaded (all duplicates or empty files)."
        else:
            summary = "Loaded: " + ",  ".join(parts)

        if errors:
            summary += f"\n\nWarnings ({len(errors)}):\n" + "\n".join(errors[:5])

        msg_title = "Load All — Done"
        messagebox.showinfo(msg_title, summary)
        self._set_status("📂 " + summary.split("\n")[0])

    def _db_save(self):
        if not self.images:
            messagebox.showwarning(_L("주의", "Warning"),
                _L("저장할 이미지가 없습니다.\n먼저 이미지를 추가하세요.",
                   "No images to save.\nAdd images first."))
            return
        analyzed_n = sum(1 for img in self.images
                         if not np.isnan(img.get("s_mean", np.nan)))
        # 분석 결과가 없어도 ROI 보존용 checkpoint 저장 가능 — 안내만 표시
        if analyzed_n == 0:
            if not messagebox.askyesno(
                    _L("미분석 저장 확인", "Save without analysis?"),
                    _L(f"분석된 이미지가 없습니다.\n"
                       f"이미지/ROI/조건만 {len(self.images)}장 저장할까요?\n"
                       f"(분석 결과는 모두 NULL 로 저장 — 다음에 ▶ Analyze All)",
                       f"No analyzed images.\n"
                       f"Save {len(self.images)} image(s) with ROI/cond only?\n"
                       f"(Metrics saved as NULL — run ▶ Analyze All later)")):
                return

        path = filedialog.asksaveasfilename(
            title="Save DB File",
            defaultextension=".db",
            filetypes=[("HfS2 DB","*.db"),("All","*.*")])
        if not path: return

        # SQLite 작업은 OS local temp 디렉토리에서 수행 → \\wsl$ 9p 마운트의
        # SQLite lock 충돌 회피. 완성된 DB 파일만 마지막에 dst 로 shutil.move.
        tmp_dir = tempfile.mkdtemp(prefix="hfs2_save_")
        tmp_path = os.path.join(tmp_dir, "save.db")
        dst_pre_size = os.path.getsize(path) if os.path.exists(path) else -1
        dst_leftovers = [ext for ext in ("-wal", "-shm", "-journal")
                         if os.path.exists(path + ext)]
        print(f"[_db_save] === START ===")
        print(f"[_db_save] dst path={path} pre_size={dst_pre_size} "
              f"leftovers={dst_leftovers}")
        print(f"[_db_save] tmp path={tmp_path} (local temp dir)")
        print(f"[_db_save] images={len(self.images)} analyzed={analyzed_n} "
              f"eval_targets={len(self._pred_targets) if hasattr(self, '_pred_targets') and self._pred_targets else 0}")

        try:
            # step 0: 라만 _id 미부여 항목에 id 부여 (매칭 보존)
            #         self._raman_data 인메모리 갱신 — 다음 단계의 img.raman_id 가 정확히 참조
            if hasattr(self, "_raman_data") and self._raman_data:
                self._ensure_raman_ids()
            n_raman_local = (len(self._raman_data)
                             if hasattr(self, "_raman_data") else 0)
            n_img_linked = sum(1 for img in self.images
                               if img.get("raman_id") is not None)
            print(f"[_db_save] step 0: raman _id ensured "
                  f"(raman_entries={n_raman_local} "
                  f"images_linked={n_img_linked})")

            print(f"[_db_save] step 1: db_save_all -> tmp")
            n = db_save_all(tmp_path, self.images)
            tmp_size_after_imgs = os.path.getsize(tmp_path) if os.path.exists(tmp_path) else -1
            print(f"[_db_save] step 1 ok: n={n} tmp_size={tmp_size_after_imgs}")

            # 평가 대상 + backup 정리를 단일 connection 으로 통합
            eval_note = ""
            import pickle as _pk, json as _js
            print(f"[_db_save] step 2: open tmp for eval_target")
            con = _db_open_safe(tmp_path)
            try:
                if self._pred_targets:
                    _migrate_eval_target_schema(con)
                    con.execute("DELETE FROM eval_target")
                    n_targets = 0
                    for t in self._pred_targets:
                        rgb = t.get("rgb")
                        if rgb is None:
                            continue
                        roi_str = _js.dumps(
                            list(t["roi"]) if t.get("roi") else None)
                        blob = _pk.dumps(rgb)
                        con.execute(
                            "INSERT INTO eval_target VALUES "
                            "(?,?,?,?,?,?,?,datetime('now','localtime'))",
                            (int(t.get("tid", 0)),
                             t.get("name", ""),
                             blob, roi_str,
                             t.get("color", ""),
                             t.get("cond_hint", ""),
                             None))
                        n_targets += 1
                    eval_note = (f" + {len(self._pred_targets)} "
                                 "evaluation targets")
                    print(f"[_db_save] step 2: inserted {n_targets} eval_targets")
                else:
                    print(f"[_db_save] step 2: no pred_targets, skipping inserts")
                # step 3: 라만 데이터 저장 (같은 connection)
                if hasattr(self, "_raman_data") and self._raman_data:
                    n_r_saved = db_save_raman_all(con, self._raman_data)
                    print(f"[_db_save] step 3: saved {n_r_saved} raman entries")
                else:
                    con.execute("DELETE FROM raman_data")
                    print(f"[_db_save] step 3: no raman entries (cleared table)")
                con.execute("DROP TABLE IF EXISTS eval_target_v1_backup")
                con.commit()
                print(f"[_db_save] step 2+3 ok: commit done")
            finally:
                con.close()
                print(f"[_db_save] step 2+3: connection closed")

            tmp_size_final = os.path.getsize(tmp_path) if os.path.exists(tmp_path) else -1
            print(f"[_db_save] step 3: prepare to move tmp -> dst, tmp_final_size={tmp_size_final}")
            # 기존 path + lock 잔재 정리 (shutil.move 는 dst 가 있으면 실패할 수 있음)
            for ext in ("", "-wal", "-shm", "-journal"):
                target = path + ext
                if os.path.exists(target):
                    try:
                        os.remove(target)
                        print(f"[_db_save] cleaned dst: {target}")
                    except OSError as e:
                        print(f"[_db_save] WARN cannot remove dst {target}: {e}")
            # Cross-fs safe move (local temp -> 9p / Windows / Linux 모두 OK)
            print(f"[_db_save] step 3: shutil.move({tmp_path} -> {path})")
            try:
                shutil.move(tmp_path, path)
            except Exception as e:
                dst_now = os.path.getsize(path) if os.path.exists(path) else -1
                tmp_now = os.path.getsize(tmp_path) if os.path.exists(tmp_path) else -1
                print(f"[_db_save] FAIL shutil.move: {type(e).__name__}: {e} "
                      f"(tmp_size_now={tmp_now} dst_size_now={dst_now})")
                raise
            final_dst_size = os.path.getsize(path) if os.path.exists(path) else -1
            print(f"[_db_save] === DONE === final dst_size={final_dst_size} records={n}")
            # tmp_dir 정리
            try:
                shutil.rmtree(tmp_dir, ignore_errors=True)
            except Exception: pass

            unanalyzed = max(0, n - analyzed_n)
            ana_note = (f"  ({analyzed_n} analyzed, {unanalyzed} ROI-only)"
                        if unanalyzed > 0 else "")
            messagebox.showinfo(_L("저장 완료", "Saved"),
                f"Saved {n} images{ana_note}{eval_note} to DB.\n{path}")
            self._set_status(
                f"🗄 DB saved — {n} records{ana_note}  ({os.path.basename(path)})")
        except Exception as ex:
            tmp_size_now = os.path.getsize(tmp_path) if os.path.exists(tmp_path) else -1
            dst_size_now = os.path.getsize(path) if os.path.exists(path) else -1
            print(f"[_db_save] === FAILED === {type(ex).__name__}: {ex}")
            print(f"[_db_save] state: tmp_size={tmp_size_now} dst_size={dst_size_now}")
            traceback.print_exc()
            # 실패 시 tmp_dir 통째로 정리
            try:
                shutil.rmtree(tmp_dir, ignore_errors=True)
                print(f"[_db_save] cleanup: removed tmp_dir {tmp_dir}")
            except Exception as ce:
                print(f"[_db_save] cleanup WARN: {ce}")
            err_msg = str(ex)
            if "locked" in err_msg.lower():
                err_msg += _L(
                    "\n\n[해결법] 다른 프로그램(DB Browser 등)이 같은 파일을\n"
                    "열고 있는지 확인하세요. 또는 다른 경로(C:\\Users\\..)에 저장.",
                    "\n\n[Fix] Check if DB Browser or other tool has the file open.\n"
                    "Or save to a different path (e.g. C:\\Users\\..).")
            messagebox.showerror(_L("저장 실패", "Save Error"), err_msg)

    def _db_load(self):
        """이미지 DB 파일 1개 또는 복수 로드 (다중 선택 지원)."""
        paths = filedialog.askopenfilenames(
            title=_L("이미지 DB 파일 선택 (복수 가능)",
                     "Open Image DB Files (multi-select)"),
            filetypes=[("HfS2 DB","*.db"),("All","*.*")])
        if not paths:
            return

        existing_keys = {
            (img["name"], img["cond"], img["day"])
            for img in self.images}
        total_added = 0
        total_dup = 0
        total_eval = 0
        per_file = []
        errors = []

        for path in paths:
            try:
                records = db_load_all(path)
                if not records:
                    per_file.append((os.path.basename(path), 0, 0))
                    continue

                added = 0
                dup = 0
                for rec in records:
                    key = (rec["name"], rec["cond"], rec["day"])
                    if key in existing_keys:
                        dup += 1
                        continue
                    rec["day_var"]  = tk.StringVar(value=rec["day"])
                    rec["cond_var"] = tk.StringVar(value=rec["cond"])
                    if rec.get("roi") is not None and rec.get("rgb") is not None:
                        flg, why = evaluate_roi_quality(rec["rgb"], rec["roi"])
                        rec["roi_flag"] = flg
                        rec["roi_reason"] = f"DB 로드: {why}"
                        rec["roi_source"] = "db"
                    else:
                        rec["roi_flag"] = None
                        rec["roi_reason"] = ""
                        rec["roi_source"] = "db"
                    self.images.append(rec)
                    existing_keys.add(key)
                    added += 1
                total_added += added
                total_dup += dup
                per_file.append((os.path.basename(path), added, dup))

                # 평가 대상 (있으면 첫 파일에서만 — 누적 시 충돌 가능)
                try:
                    import sqlite3 as _sq
                    # read-only 로 열어 9p lock 회피. 마이그레이션은 SAVE 시점에만.
                    con = _db_open_read(path)
                    cur = con.execute(
                        "SELECT name FROM sqlite_master "
                        "WHERE type='table' AND name='eval_target'")
                    if cur.fetchone():
                        rows = list(con.execute(
                            "SELECT target_id, name, rgb_blob, roi, "
                            "color, cond_hint "
                            "FROM eval_target ORDER BY target_id"))
                        n_loaded = self._pred_load_rows(rows)
                        total_eval += n_loaded
                    con.close()
                except Exception:
                    pass
            except Exception as ex:
                errors.append((os.path.basename(path), str(ex)))

        if self.sel_idx < 0 and self.images:
            self.sel_idx = 0
        self._sort_images_by_cond_day()
        self._rebuild_list()
        if total_eval > 0:
            self.after(150, self._pred_rebuild_cards)

        # 결과 메시지
        lines = [_L(f"DB 파일 {len(paths)}개 로드 완료",
                    f"Loaded {len(paths)} DB file(s)"),
                 ""]
        for name, a, d in per_file:
            lines.append(f"  • {name}: +{a}장 (중복 {d})")
        lines.append("")
        lines.append(_L(f"총 추가: {total_added}장 | 중복 제외: {total_dup}장",
                        f"Total: +{total_added} | Duplicates: {total_dup}"))
        if total_eval > 0:
            lines.append(_L(f"평가대상 복원: {total_eval}개",
                            f"Evaluation targets restored: {total_eval}"))
        if errors:
            lines.append("")
            lines.append(_L("⚠ 오류:", "⚠ Errors:"))
            for name, err in errors:
                lines.append(f"  • {name}: {err}")
        messagebox.showinfo(_L("로드 결과", "Load Result"), "\n".join(lines))
        self._set_status(
            _L(f"📂 DB 로드 — {total_added}장  ({len(paths)}개 파일)",
               f"📂 DB loaded — {total_added} records  ({len(paths)} files)"))



    def _pred_load_target_dialog(self):
        """파일 다이얼로그로 평가 대상 이미지 로드"""
        path = filedialog.askopenfilename(
            title="Load Evaluation Target",
            filetypes=[("Target DB","*.target.db *.db"),("All","*.*")])
        if not path: return
        ok = self._db_load_target(path)
        if ok:
            self._set_status(f"🎯 Target loaded from: {os.path.basename(path)}")
        else:
            messagebox.showwarning("Load Target",
                "No evaluation target found.\n"
                "Save a target first using [💾 Save Target].")

    def _db_save_target(self):
        """평가 대상 이미지 별도 저장 (다중 target)"""
        if not self._pred_targets:
            messagebox.showwarning("Warning",
                "No evaluation target loaded.\nAdd targets in the [🎯 Evaluation] tab first.")
            return
        path = filedialog.asksaveasfilename(
            title="Save Evaluation Target",
            defaultextension=".target.db",
            filetypes=[("Target DB","*.target.db *.db"),("All","*.*")])
        if not path: return
        try:
            import sqlite3 as _sq, pickle as _pk, json as _js
            con = _db_open_safe(path)
            n_saved = 0
            try:
                con.execute("PRAGMA busy_timeout=15000")
                _migrate_eval_target_schema(con)
                con.execute("DELETE FROM eval_target")
                for t in self._pred_targets:
                    rgb = t.get("rgb")
                    if rgb is None:
                        continue
                    roi_str = _js.dumps(
                        list(t["roi"]) if t.get("roi") else None)
                    blob = _pk.dumps(rgb)
                    con.execute(
                        "INSERT INTO eval_target VALUES "
                        "(?,?,?,?,?,?,?,datetime('now','localtime'))",
                        (int(t.get("tid", 0)),
                         t.get("name", ""),
                         blob, roi_str,
                         t.get("color", ""),
                         t.get("cond_hint", ""),
                         None))
                    n_saved += 1
                # 마이그레이션 백업 정리
                con.execute("DROP TABLE IF EXISTS eval_target_v1_backup")
                con.commit()
            finally:
                con.close()
            messagebox.showinfo("Saved",
                f"{n_saved} evaluation target(s) saved.\n{path}")
            self._set_status(
                f"🎯 {n_saved} target(s) saved: {os.path.basename(path)}")
        except Exception as ex:
            messagebox.showerror("Save Error", str(ex))

    def _pred_load_rows(self, rows) -> int:
        """eval_target 행 리스트 → self._pred_targets 적재.

        rows 의 각 row 는 (target_id, name, rgb_blob, roi, color, cond_hint).
        Returns: 적재된 개수 (사용자 취소 시 0)
        """
        import pickle as _pk, json as _js
        n_loaded = 0
        # 메모리에 미저장 target 이 있으면 사용자에게 확인
        if self._pred_targets:
            try:
                ok = messagebox.askyesno(
                    _L("기존 target 교체",
                       "Replace existing targets"),
                    _L(
                        f"메모리에 {len(self._pred_targets)}개의 평가대상이 있습니다.\n"
                        f"DB 로드로 모두 교체할까요?\n(취소 시 로드 안 함)",
                        f"{len(self._pred_targets)} targets in memory.\n"
                        f"Replace all with DB load?\n(Cancel skips load)"))
            except Exception:
                ok = True
            if not ok:
                return 0
        # 기존 target 제거
        self._pred_targets.clear()
        self._pred_sel_tid = None
        for row in rows[:PRED_MAX_TARGETS]:
            try:
                tid = int(row[0]) if row[0] is not None else (n_loaded+1)
                name = row[1] or f"Target #{tid}"
                blob = row[2]
                roi_data = _js.loads(row[3]) if row[3] else None
                roi = tuple(roi_data) if roi_data else None
                color = row[4] or self._pred_assign_color()
                cond_hint = row[5] or ""
                rgb = _pk.loads(blob)
                # ROI 품질 평가
                if roi is not None:
                    try:
                        flag, reason = evaluate_roi_quality(rgb, roi)
                    except Exception:
                        flag, reason = "manual", "DB 로드"
                else:
                    # ROI 없으면 자동 추정
                    try:
                        roi, flag, reason = auto_detect_roi(rgb, cond=cond_hint)
                    except Exception:
                        h, w = rgb.shape[:2]
                        roi = (w//4, h//4, 3*w//4, 3*h//4)
                        flag = "failed"
                        reason = "auto_detect_roi 실패"
                try:
                    thumb = make_thumb(rgb, 90, 70, roi)
                except Exception:
                    thumb = None
                self._pred_targets.append({
                    "tid":   tid,
                    "name":  name,
                    "rgb":   rgb,
                    "roi":   roi,
                    "roi_flag":   flag,
                    "roi_reason": f"DB 로드: {reason}",
                    "roi_source": "db",
                    "color": color,
                    "thumb": thumb,
                    "cond_hint": cond_hint,
                    "result": None,
                })
                n_loaded += 1
            except Exception:
                continue
        if self._pred_targets:
            self._pred_sel_tid = self._pred_targets[0]["tid"]
        return n_loaded

    def _db_load_target(self, path: str):
        """평가 대상 이미지 로드 (내부 공용 — 다중 target). read-only."""
        try:
            import sqlite3 as _sq
            # read-only — 9p lock 회피. 마이그레이션은 SAVE 시점에만.
            con = _db_open_read(path)
            cur = con.execute(
                "SELECT name FROM sqlite_master "
                "WHERE type='table' AND name='eval_target'")
            if not cur.fetchone():
                con.close()
                return False
            rows = list(con.execute(
                "SELECT target_id, name, rgb_blob, roi, "
                "color, cond_hint "
                "FROM eval_target ORDER BY target_id"))
            con.close()
            if not rows:
                return False
            n = self._pred_load_rows(rows)
            if n > 0:
                self.after(150, self._pred_rebuild_cards)
                return True
        except Exception:
            pass
        return False

    def _db_save_raman(self):
        """Raman 데이터 별도 저장"""
        if not self._raman_data:
            messagebox.showwarning("Warning",
                "No Raman data loaded.\n"
                "Import Excel files in the [📡 Raman Analysis] tab first.")
            return
        path = filedialog.asksaveasfilename(
            title="Save Raman Data",
            defaultextension=".raman.db",
            filetypes=[("Raman DB","*.raman.db *.db"),("All","*.*")])
        if not path: return
        try:
            import sqlite3 as _sq, json as _js
            con = _db_open_safe(path)
            con.execute("""CREATE TABLE IF NOT EXISTS raman_data
                (id INTEGER PRIMARY KEY,
                 cond TEXT, day TEXT, peak REAL,
                 norm_peak REAL, peak_shift REAL,
                 peak_range TEXT, spectrum_json TEXT,
                 saved_at TEXT)""")
            con.execute("DELETE FROM raman_data")
            for r in self._raman_data:
                spec_j = (_js.dumps(r["spectrum"])
                          if r.get("spectrum") else None)
                con.execute(
                    "INSERT INTO raman_data"
                    " (cond,day,peak,norm_peak,peak_shift,"
                    "  peak_range,spectrum_json,saved_at)"
                    " VALUES (?,?,?,?,?,?,?,datetime('now','localtime'))",
                    (r.get("cond",""), r.get("day",""),
                     float(r.get("peak",0)),
                     float(r.get("norm_peak") or 0),
                     float(r.get("peak_shift") or 0),
                     str(r.get("peak_range","")),
                     spec_j))
            con.commit(); con.close()
            n = len(self._raman_data)
            conds = list(dict.fromkeys(r["cond"] for r in self._raman_data))
            messagebox.showinfo("Saved",
                f"Saved {n} Raman entries.\nConditions: {', '.join(conds[:5])}\n{path}")
            self._set_status(
                f"📡 Raman saved: {n} entries  ({os.path.basename(path)})")
        except Exception as ex:
            messagebox.showerror("Save Error", str(ex))

    def _raman_load_db_dialog(self):
        """파일 다이얼로그로 Raman DB 로드"""
        path = filedialog.askopenfilename(
            title="Load Raman Data",
            filetypes=[("Raman DB","*.raman.db *.db"),("All","*.*")])
        if not path: return
        n = self._db_load_raman(path)
        if n > 0:
            self._set_status(
                f"📡 Raman loaded: {n} entries from {os.path.basename(path)}")
            messagebox.showinfo("Loaded",
                f"Loaded {n} Raman entries from:\n{path}")
        else:
            messagebox.showwarning("Load Raman DB",
                "No Raman data found in the selected file.\n"
                "Save Raman data first using [Save Raman DB].")
    def _db_load_raman(self, path: str) -> int:
        """Raman 데이터 로드 (내부 공용). 추가된 항목 수 반환.

        read-only 모드 — 9p 마운트 lock 회피.
        """
        try:
            import sqlite3 as _sq, json as _js
            con = _db_open_read(path)
            rows = con.execute(
                "SELECT cond,day,peak,norm_peak,peak_shift,"
                "peak_range,spectrum_json FROM raman_data"
            ).fetchall()
            con.close()
            added = 0
            for row in rows:
                (cond,day,peak,norm_peak,
                 peak_shift,peak_range,spec_j) = row
                spec = _js.loads(spec_j) if spec_j else None
                self._raman_data.append({
                    "cond":      cond or "",
                    "day":       day  or "",
                    "peak":      float(peak or 0),
                    "norm_peak": float(norm_peak or 0),
                    "peak_shift":float(peak_shift or 0),
                    "peak_range":str(peak_range or ""),
                    "spectrum":  spec,
                })
                added += 1
            if added:
                self._normalize_raman()
                self._auto_link_raman_by_cond_day()
                self._rebuild_raman_tree()
                self._refresh_raman_tab()
            return added
        except Exception:
            return 0

    # ─────────────────────────────────────────
    #  2. 날짜 추정 탭
    # ─────────────────────────────────────────
    def _build_predict_tab(self):
        """
        Evaluation 탭 — 평가 대상 이미지 입력 후
        ① Date Estimation  ② Pseudo-Raman  동시 출력
        """
        f = self._tfs["predict"]

        # ══════════════════════════════════
        # 좌: 평가 대상 입력 패널 (다중 target 카드 리스트)
        # ══════════════════════════════════
        left = tk.Frame(f, bg=PANEL,
                        highlightbackground=BORDER, highlightthickness=1,
                        width=310)
        left.pack(side="left", fill="y", padx=(0,4))
        left.pack_propagate(False)

        # 헤더 (N/8 표시)
        hdr_row = tk.Frame(left, bg=PANEL2,
                           highlightbackground=BORDER,
                           highlightthickness=1)
        hdr_row.pack(fill="x")
        tk.Label(hdr_row,
                 text=_L("  🎯 평가대상", "  🎯 Evaluation Targets"),
                 bg=PANEL2, fg=TXT, font=MFB).pack(side="left", pady=4)
        self._pred_count_var = tk.StringVar(value="0/8")
        tk.Label(hdr_row, textvariable=self._pred_count_var,
                 bg=PANEL2, fg=ACCENT, font=MFB).pack(side="right", padx=8)

        # 안내 박스
        guide = tk.Frame(left, bg=CARD,
                         highlightbackground=BORDER,
                         highlightthickness=1)
        guide.pack(fill="x", padx=6, pady=(4,2))
        tk.Label(guide,
                 text=_L(
                    "1. 평가대상 이미지 추가 (드래그앤드롭 OK,\n"
                    "   최대 8개)\n"
                    "2. 카드 테두리 색으로 ROI 품질 확인\n"
                    "   🟢 OK   🟠 검토   🔴 실패\n"
                    "3. ▶ 모두 분석 → 모든 target 동시 분석\n"
                    "4. 카드 클릭 = 선택 (Top-3 후보 갱신)",
                    "1. Add target images (drag&drop OK,\n"
                    "   max 8)\n"
                    "2. Border color = ROI quality\n"
                    "   🟢 OK   🟠 review   🔴 failed\n"
                    "3. ▶ Run All → analyze all targets\n"
                    "4. Click card = select (Top-3 update)"),
                 bg=CARD, fg=SUB, font=("Segoe UI",7),
                 justify="left").pack(anchor="w", padx=6, pady=4)

        # 액션 버튼 행
        btn_row = tk.Frame(left, bg=PANEL)
        btn_row.pack(fill="x", padx=6, pady=(2,2))
        self._pred_add_btn = tk.Button(
            btn_row, text=_L("➕ 추가", "➕ Add"),
            command=self._pred_load_image,
            bg=ACCENT, fg="white", font=MF,
            relief="flat", padx=8, pady=4,
            cursor="hand2")
        self._pred_add_btn.pack(side="left", expand=True,
                                fill="x", padx=(0,2))
        tk.Button(btn_row, text=_L("📋 붙여넣기", "📋 Paste"),
                  command=self._pred_paste_image,
                  bg=BTN, fg=TXT, font=MF,
                  relief="flat", padx=6, pady=4,
                  cursor="hand2").pack(side="left", expand=True,
                                       fill="x", padx=(2,0))

        # ▶ 모두 분석
        tk.Button(left, text=_L("▶  모두 분석", "▶  Run All"),
                  command=self._pred_run_all,
                  bg=ACCENT, fg="white",
                  font=("Segoe UI",10,"bold"),
                  relief="flat", pady=8, cursor="hand2").pack(
                  fill="x", padx=6, pady=(2,2))

        # 🗑 전체 삭제 + Save/Load
        sub_row = tk.Frame(left, bg=PANEL)
        sub_row.pack(fill="x", padx=6, pady=(0,2))
        tk.Button(sub_row, text=_L("🗑 전체삭제","🗑 Clear All"),
                  command=self._pred_clear_all_targets,
                  bg=BTN, fg=RED, font=LF,
                  relief="flat", padx=4, pady=2,
                  cursor="hand2").pack(side="left", expand=True,
                                       fill="x", padx=(0,1))
        tk.Button(sub_row, text="💾",
                  command=self._db_save_target,
                  bg=BTN, fg=TXT, font=LF,
                  relief="flat", padx=4, pady=2,
                  cursor="hand2").pack(side="left", padx=1)
        tk.Button(sub_row, text="📂",
                  command=self._pred_load_target_dialog,
                  bg=BTN, fg=TXT, font=LF,
                  relief="flat", padx=4, pady=2,
                  cursor="hand2").pack(side="left", padx=1)

        tk.Frame(left, bg=BORDER, height=1).pack(fill="x", padx=8, pady=2)

        # 조건 힌트 (입력 시 모든 신규 target 의 cond_hint 기본값)
        cond_lbl = tk.Frame(left, bg=PANEL)
        cond_lbl.pack(fill="x", padx=6, pady=(0,1))
        tk.Label(cond_lbl, text=_L("Cond hint:","Cond hint:"),
                 bg=PANEL, fg=SUB, font=("Segoe UI",7)).pack(side="left")
        self._pred_cond_var = tk.StringVar()
        tk.Entry(cond_lbl, textvariable=self._pred_cond_var,
                 bg=PANEL2, fg=TXT, font=("Segoe UI",7),
                 insertbackground=TXT, relief="flat",
                 highlightbackground=BORDER,
                 highlightthickness=1, width=20).pack(
            side="left", fill="x", expand=True, padx=2)

        pf4 = tk.Frame(left, bg=PANEL)
        pf4.pack(fill="x", padx=6, pady=(0,2))
        for lbl2, col in self._presets:
            tk.Button(pf4, text=lbl2.replace("Native-","N-"),
                      command=lambda l=lbl2: self._pred_cond_var.set(l),
                      bg=BTN, fg=col,
                      font=("Segoe UI",7,"bold"),
                      relief="flat", cursor="hand2",
                      padx=2, pady=1).pack(side="left", padx=1)

        tk.Frame(left, bg=BORDER, height=1).pack(fill="x", padx=8, pady=2)

        # ── Raman 피크 처리 조건 (기존 유지) ─────────
        tk.Label(left, text=_L("  Raman 피크 추출","  Raman Peak Extraction"),
                 bg=PANEL, fg=SUB, font=MFB).pack(anchor="w", padx=8, pady=(0,2))

        pk_f = tk.Frame(left, bg=PANEL)
        pk_f.pack(fill="x", padx=8)

        self._pred_peak_mode = tk.StringVar(value="auto_max")
        for val, lbl in [
            ("auto_max",     "Global max peak"),
            ("range_single", "Max in range"),
            ("range_multi",  "Multiple ranges"),
            ("area",         "Area (integral)"),
        ]:
            tk.Radiobutton(pk_f, text=lbl,
                           variable=self._pred_peak_mode, value=val,
                           bg=PANEL, fg=TXT,
                           selectcolor=PANEL2,
                           activebackground=PANEL,
                           font=("Segoe UI",7)).pack(anchor="w")

        range_f = tk.Frame(left, bg=PANEL)
        range_f.pack(fill="x", padx=8, pady=(2,0))
        tk.Label(range_f, text="Range (cm⁻¹, comma-sep):",
                 bg=PANEL, fg=SUB,
                 font=("Segoe UI",7)).pack(anchor="w")
        self._pred_peak_ranges = tk.StringVar(value="300-360")
        tk.Entry(range_f, textvariable=self._pred_peak_ranges,
                 bg=PANEL2, fg=TXT, font=("Segoe UI",7),
                 relief="flat",
                 highlightbackground=BORDER,
                 highlightthickness=1).pack(fill="x", pady=2)
        tk.Label(range_f,
                 text='e.g. "300-360" or "300-360, 370-400"',
                 bg=PANEL, fg=SUB,
                 font=("Segoe UI",6)).pack(anchor="w")

        tk.Frame(left, bg=BORDER, height=1).pack(fill="x", padx=8, pady=4)

        # ── 카드 스크롤 영역 ─────────────────────────
        sc = tk.Frame(left, bg=PANEL)
        sc.pack(fill="both", expand=True, padx=4, pady=2)
        self._pred_lc = tk.Canvas(sc, bg=PANEL, highlightthickness=0)
        vsb = tk.Scrollbar(sc, orient="vertical",
                           command=self._pred_lc.yview)
        self._pred_lc.configure(yscrollcommand=vsb.set)
        vsb.pack(side="right", fill="y")
        self._pred_lc.pack(fill="both", expand=True)
        self._pred_lf = tk.Frame(self._pred_lc, bg=PANEL)
        _pw = self._pred_lc.create_window((0,0),
                                          window=self._pred_lf,
                                          anchor="nw")
        self._pred_lf.bind("<Configure>",
            lambda e: self._pred_lc.configure(
                scrollregion=self._pred_lc.bbox("all")))
        self._pred_lc.bind("<Configure>",
            lambda e: self._pred_lc.itemconfig(_pw, width=e.width))

        # 카드 패널 자체 DnD 등록
        if _DND:
            try:
                self._pred_lc.drop_target_register(DND_FILES)
                self._pred_lc.dnd_bind("<<Drop>>", self._pred_on_drop)
                self._pred_lf.drop_target_register(DND_FILES)
                self._pred_lf.dnd_bind("<<Drop>>", self._pred_on_drop)
            except Exception:
                pass

        # 호환용: 기존 코드가 _pred_cv / _pred_roi_info 를 참조해도 안 깨지도록
        # 더미 객체 유지 (실제로는 카드 패널이 미리보기 역할)
        self._pred_cv = self._pred_lc
        self._pred_roi_info = tk.StringVar(value="No ROI selected")

        # ══════════════════════════════════
        # 우: PanedWindow — 결과 영역 (수직)
        # ══════════════════════════════════
        right = tk.Frame(f, bg=BG)
        right.pack(side="left", fill="both", expand=True)

        right_pw = ttk.PanedWindow(right, orient="vertical")
        right_pw.pack(fill="both", expand=True)

        # ── 패널A: Date Estimation ─────────────
        pA = tk.Frame(right_pw, bg=BG)
        right_pw.add(pA, weight=3)

        # A-헤더
        a_hdr = tk.Frame(pA, bg=PANEL2,
                          highlightbackground=BORDER, highlightthickness=1)
        a_hdr.pack(fill="x")
        tk.Label(a_hdr, text="  📅 Date Estimation",
                 bg=PANEL2, fg=TXT, font=MFB).pack(side="left", pady=6, padx=8)
        tk.Label(a_hdr,
                 text="Compare with reference images by b* · S · YI",
                 bg=PANEL2, fg=SUB, font=("Segoe UI",7)).pack(
                 side="left", padx=4)

        # A-결과 요약 행
        res_row = tk.Frame(pA, bg=PANEL,
                           highlightbackground=BORDER, highlightthickness=1)
        res_row.pack(fill="x", padx=0, pady=(0,2))
        self._pred_result_var = tk.StringVar(
            value="Run Evaluation to see results.")
        tk.Label(res_row, textvariable=self._pred_result_var,
                 bg=PANEL, fg=GOLD, font=MFB,
                 anchor="w").pack(side="left", padx=12, pady=6)

        # A-차트 (레이더 + 타임라인)
        a_charts = tk.Frame(pA, bg=BG)
        a_charts.pack(fill="both", expand=True)
        a_charts.columnconfigure(0, weight=1)
        a_charts.columnconfigure(1, weight=1)
        a_charts.rowconfigure(0, weight=1)

        self._pred_figs = {}
        for ci2, (key2, title2) in enumerate([
                ("radar",    "Metric Similarity"),
                ("timeline", "Distance Timeline"),
        ]):
            cell = tk.Frame(a_charts, bg=PANEL,
                            highlightbackground=BORDER,
                            highlightthickness=1)
            cell.grid(row=0, column=ci2, padx=3, pady=3, sticky="nsew")
            hdr2 = tk.Frame(cell, bg=PANEL2,
                            highlightbackground=BORDER,
                            highlightthickness=1)
            hdr2.pack(fill="x")
            tk.Label(hdr2, text=f"  {title2}",
                     bg=PANEL2, fg=TXT, font=MFB).pack(side="left",
                     pady=3, padx=6)
            tk.Label(hdr2, text="⤢ dbl-click: enlarge",
                     bg=PANEL2, fg=SUB,
                     font=("Segoe UI",7)).pack(side="right", padx=6)
            fig2 = plt.Figure(figsize=(3.5, 2.8), facecolor=PANEL)
            cv2  = FigureCanvasTkAgg(fig2, master=cell)
            cv2.get_tk_widget().pack(fill="both", expand=True,
                                     padx=2, pady=2)
            # 초기 빈 화면 draw (Windows에서 blank 방지)
            ax0 = fig2.add_subplot(111)
            ax0.set_facecolor(PANEL2)
            ax0.text(0.5, 0.5, "Run Evaluation",
                     transform=ax0.transAxes,
                     ha="center", va="center",
                     color=SUB, fontsize=7)
            ax0.axis("off")
            cv2.draw()
            cv2.get_tk_widget().bind(
                "<Double-Button-1>",
                lambda e, k=key2, t=title2:
                    self._pred_chart_popup(k, t))
            self._pred_figs[key2] = {"fig": fig2, "cv": cv2}

        # ── 패널B: 후보 이미지 (Top-3) ─────────
        pB = tk.Frame(right_pw, bg=BG)
        right_pw.add(pB, weight=2)

        b_hdr = tk.Frame(pB, bg=PANEL2,
                          highlightbackground=BORDER, highlightthickness=1)
        b_hdr.pack(fill="x")
        tk.Label(b_hdr, text="  🖼 Top-3 Similar Candidates",
                 bg=PANEL2, fg=TXT, font=MFB).pack(
                 side="left", pady=4, padx=8)
        tk.Label(b_hdr,
                 text="Ranked by weighted distance (b*×0.45 + S×0.30 + YI×0.25)",
                 bg=PANEL2, fg=SUB,
                 font=("Segoe UI",7)).pack(side="left", padx=4)

        cand_body = tk.Frame(pB, bg=PANEL)
        cand_body.pack(fill="both", expand=True, padx=4, pady=4)
        for ci3 in range(3):
            cand_body.columnconfigure(ci3, weight=1)
        cand_body.rowconfigure(0, weight=1)

        self._pred_cand_frames = []
        rank_colors = [ACCENT, TEAL, PURPLE]
        for ci3 in range(3):
            cf2 = tk.Frame(cand_body, bg=CARD,
                           highlightbackground=BORDER, highlightthickness=1)
            cf2.grid(row=0, column=ci3, padx=3, pady=2, sticky="nsew")
            cf2.rowconfigure(1, weight=1)  # 썸네일 행 확장

            # 순위 헤더
            cf2_hdr = tk.Frame(cf2, bg=rank_colors[ci3])
            cf2_hdr.pack(fill="x")
            tk.Label(cf2_hdr, text=f"  #{ci3+1}",
                     bg=rank_colors[ci3], fg="white",
                     font=MFB).pack(side="left", pady=3, padx=6)
            info_var = tk.StringVar(value="(none)")
            tk.Label(cf2_hdr, textvariable=info_var,
                     bg=rank_colors[ci3], fg="white",
                     font=("Segoe UI",7),
                     wraplength=150).pack(
                     side="left", padx=4, pady=2)

            # 썸네일 캔버스 (fill=both로 공간 최대 활용)
            th_cv = tk.Canvas(cf2, bg=CARD2,
                              highlightthickness=0)
            th_cv.pack(fill="both", expand=True,
                       padx=2, pady=(2,0))

            # 지표 수치
            metric_var = tk.StringVar(value="")
            tk.Label(cf2, textvariable=metric_var,
                     bg=CARD, fg=SUB,
                     font=("Segoe UI",7),
                     justify="left",
                     anchor="w").pack(fill="x", padx=4, pady=(0,1))

            # 유사도 코멘트
            cmt_txt = tk.Text(cf2, height=4, wrap="word",
                              bg=CARD2, fg=TXT,
                              font=("Segoe UI",7),
                              relief="flat", padx=4, pady=3,
                              highlightthickness=0,
                              cursor="xterm",
                              state="disabled")
            cmt_txt.pack(fill="x", padx=2, pady=(0,3))
            cmt_txt.bind("<Control-c>",
                lambda e, w=cmt_txt: self._copy_text(w))

            self._pred_cand_frames.append({
                "frame":      cf2,
                "info_var":   info_var,
                "th_cv":      th_cv,
                "metric_var": metric_var,
                "cmt_txt":    cmt_txt,
            })

        # ── 패널C: Pseudo-Raman ─────────────────
        pC = tk.Frame(right_pw, bg=BG)
        right_pw.add(pC, weight=3)

        c_hdr = tk.Frame(pC, bg=PANEL2,
                          highlightbackground=BORDER, highlightthickness=1)
        c_hdr.pack(fill="x")
        tk.Label(c_hdr,
                 text="  🔮 Pseudo-Raman Estimation",
                 bg=PANEL2, fg=TXT, font=MFB).pack(
                 side="left", pady=6, padx=8)
        tk.Label(c_hdr,
                 text="Estimate Raman spectrum from image metrics"
                      "  (requires Raman reference data)",
                 bg=PANEL2, fg=SUB,
                 font=("Segoe UI",7)).pack(side="left", padx=4)

        c_body = tk.Frame(pC, bg=BG)
        c_body.pack(fill="both", expand=True)
        c_body.columnconfigure(0, weight=1)
        c_body.columnconfigure(1, weight=1)
        c_body.rowconfigure(0, weight=1)

        # 회귀 차트
        reg_f = tk.Frame(c_body, bg=PANEL,
                          highlightbackground=BORDER, highlightthickness=1)
        reg_f.grid(row=0, column=0, sticky="nsew", padx=3, pady=3)
        reg_hdr = tk.Frame(reg_f, bg=PANEL2,
                           highlightbackground=BORDER, highlightthickness=1)
        reg_hdr.pack(fill="x")
        tk.Label(reg_hdr,
                 text="  📈 Regression (Image → Raman Peak)",
                 bg=PANEL2, fg=TXT, font=MFB).pack(side="left",
                 pady=3, padx=6)
        tk.Label(reg_hdr, text="⤢ dbl-click",
                 bg=PANEL2, fg=SUB,
                 font=("Segoe UI",7)).pack(side="right", padx=6)
        self._pseudo_reg_fig = plt.Figure(figsize=(4, 3), facecolor=PANEL)
        reg_cv = FigureCanvasTkAgg(self._pseudo_reg_fig, master=reg_f)
        reg_cv.get_tk_widget().pack(fill="both", expand=True, padx=2, pady=2)
        _ax = self._pseudo_reg_fig.add_subplot(111)
        _ax.set_facecolor(PANEL2); _ax.axis("off")
        _ax.text(0.5,0.5,"Run Evaluation",transform=_ax.transAxes,
                 ha="center",va="center",color=SUB,fontsize=7)
        reg_cv.draw()
        reg_cv.get_tk_widget().bind(
            "<Double-Button-1>",
            lambda e: self._pred_chart_popup("pseudo_reg",
                                             "Regression (Image → Raman Peak)"))
        self._pseudo_reg_cv = reg_cv

        # 추정 스펙트럼 차트
        spec_f = tk.Frame(c_body, bg=PANEL,
                           highlightbackground=BORDER, highlightthickness=1)
        spec_f.grid(row=0, column=1, sticky="nsew", padx=3, pady=3)
        spec_hdr = tk.Frame(spec_f, bg=PANEL2,
                            highlightbackground=BORDER, highlightthickness=1)
        spec_hdr.pack(fill="x")
        tk.Label(spec_hdr,
                 text="  🔮 Estimated Spectrum + 95% CI",
                 bg=PANEL2, fg=TXT, font=MFB).pack(side="left",
                 pady=3, padx=6)
        tk.Label(spec_hdr, text="⤢ dbl-click",
                 bg=PANEL2, fg=SUB,
                 font=("Segoe UI",7)).pack(side="right", padx=6)
        self._pseudo_spec_fig = plt.Figure(figsize=(4, 3), facecolor=PANEL)
        spec_cv2 = FigureCanvasTkAgg(self._pseudo_spec_fig, master=spec_f)
        spec_cv2.get_tk_widget().pack(fill="both", expand=True, padx=2, pady=2)
        _ax2 = self._pseudo_spec_fig.add_subplot(111)
        _ax2.set_facecolor(PANEL2); _ax2.axis("off")
        _ax2.text(0.5,0.5,"Run Evaluation",transform=_ax2.transAxes,
                  ha="center",va="center",color=SUB,fontsize=7)
        spec_cv2.draw()
        spec_cv2.get_tk_widget().bind(
            "<Double-Button-1>",
            lambda e: self._pred_chart_popup("pseudo_spec",
                                             "Estimated Spectrum + 95% CI"))
        self._pseudo_spec_cv = spec_cv2

        # 추정 결과 텍스트
        c_txt_hdr = tk.Frame(pC, bg=PANEL2,
                              highlightbackground=BORDER, highlightthickness=1)
        c_txt_hdr.pack(fill="x")
        tk.Label(c_txt_hdr, text="  📋 Estimation Result",
                 bg=PANEL2, fg=TXT, font=MFB).pack(
                 side="left", pady=4, padx=8)
        tk.Button(c_txt_hdr, text="📋",
                  command=lambda: self._copy_text(self._pseudo_res_txt),
                  bg=BTN, fg=TXT, font=LF,
                  relief="flat", padx=6, pady=2,
                  cursor="hand2").pack(side="right", padx=6, pady=4)

        self._pseudo_res_txt = tk.Text(
            pC, height=5, wrap="word",
            bg=CARD2, fg=TXT,
            font=("Segoe UI",8),
            relief="flat", padx=10, pady=6,
            highlightthickness=0,
            cursor="xterm",
            state="disabled")
        self._pseudo_res_txt.pack(fill="x", padx=4, pady=(0,4))

        # 분석 근거 코멘트 (Date Estimation)
        cmt_hdr2 = tk.Frame(pC, bg=PANEL2,
                             highlightbackground=BORDER, highlightthickness=1)
        cmt_hdr2.pack(fill="x")
        tk.Label(cmt_hdr2, text="  💬 Date Estimation Comment",
                 bg=PANEL2, fg=TXT, font=MFB).pack(
                 side="left", pady=4, padx=8)
        self._pred_comment = tk.Text(
            pC, wrap="word", height=4,
            bg=CARD2, fg=TXT,
            font=("Segoe UI",8),
            relief="flat", padx=10, pady=6,
            highlightthickness=0,
            cursor="xterm",
            state="disabled")
        self._pred_comment.pack(fill="x", padx=4, pady=(0,4))
        self._pred_comment.bind("<Control-c>",
            lambda e: self._copy_text(self._pred_comment))

        # 내부 상태
        self._last_eval_ctx = None   # 차트 팝업 재렌더용

        # 초기 카드 빌드
        self._pred_rebuild_cards()


    # ─────────────────────────────────────────
    #  다중 평가 대상 — 핵심 메서드
    # ─────────────────────────────────────────
    def _pred_add_target(self, pil_img, name: str = "Target"):
        """평가대상 이미지 추가. 8개 도달 시 경고."""
        if len(self._pred_targets) >= PRED_MAX_TARGETS:
            messagebox.showwarning(
                _L("주의","Warning"),
                _L(f"평가대상은 최대 {PRED_MAX_TARGETS}개까지만 추가 가능합니다.",
                   f"Maximum {PRED_MAX_TARGETS} targets allowed."))
            return None
        try:
            pil_rgb = pil_img.convert("RGB")
            rgb = np.array(pil_rgb)
        except Exception as ex:
            messagebox.showerror(_L("오류","Error"), str(ex))
            return None

        # 자동 ROI 추정
        try:
            auto_roi, roi_flag, roi_reason = auto_detect_roi(rgb)
        except Exception as ex:
            h, w = rgb.shape[:2]
            auto_roi = (w//4, h//4, 3*w//4, 3*h//4)
            roi_flag = "failed"
            roi_reason = f"auto_detect_roi 실패: {ex}"

        tid = self._pred_max_tid() + 1
        color = self._pred_assign_color()
        cond_hint = (self._pred_cond_var.get().strip()
                     if hasattr(self, "_pred_cond_var") else "")
        try:
            thumb = make_thumb(rgb, 90, 70, auto_roi)
        except Exception:
            thumb = None

        target = {
            "tid":         tid,
            "name":        (name or f"Target #{tid}")[:64],
            "rgb":         rgb,
            "roi":         auto_roi,
            "roi_flag":    roi_flag,
            "roi_reason":  roi_reason,
            "roi_source":  "auto",
            "color":       color,
            "thumb":       thumb,
            "cond_hint":   cond_hint,
            "result":      None,
        }
        self._pred_targets.append(target)

        # 첫 추가 시 자동 선택
        if self._pred_sel_tid is None:
            self._pred_sel_tid = tid

        self._pred_rebuild_cards()
        self._set_status(
            _L(f"✓ 평가대상 추가: {target['name']}  ({len(self._pred_targets)}/{PRED_MAX_TARGETS})",
               f"✓ Target added: {target['name']}  ({len(self._pred_targets)}/{PRED_MAX_TARGETS})"))
        return target

    def _pred_remove_target(self, tid):
        """tid 의 target 삭제."""
        before = len(self._pred_targets)
        self._pred_targets = [t for t in self._pred_targets
                              if t.get("tid") != tid]
        if before == len(self._pred_targets):
            return
        if self._pred_sel_tid == tid:
            self._pred_sel_tid = (self._pred_targets[0]["tid"]
                                  if self._pred_targets else None)
        # Advanced 탭 stale 방지 (이전 target 의 _last_eval_ctx 무효화)
        self._last_eval_ctx = None
        self._pred_rebuild_cards()
        # 차트도 갱신 시도
        try:
            self._pred_draw_all_charts()
        except Exception:
            pass

    def _pred_clear_all_targets(self):
        """모든 target 삭제 (확인 후)."""
        if not self._pred_targets:
            return
        if not messagebox.askyesno(
                _L("확인","Confirm"),
                _L("모든 평가대상을 삭제할까요?",
                   "Delete all evaluation targets?")):
            return
        self._pred_targets.clear()
        self._pred_sel_tid = None
        # Advanced 탭 stale 방지
        self._last_eval_ctx = None
        self._pred_rebuild_cards()
        try:
            self._pred_draw_all_charts()
        except Exception:
            pass
        self._set_status(
            _L("🗑 모든 평가대상 삭제됨", "🗑 All targets cleared"))

    def _pred_replace_first_target(self, rgb: np.ndarray, name: str):
        """호환용 — 기존 단일 target 흐름 마이그레이션. 첫 target 만 교체."""
        try:
            auto_roi, roi_flag, roi_reason = auto_detect_roi(rgb)
        except Exception:
            h, w = rgb.shape[:2]
            auto_roi = (w//4, h//4, 3*w//4, 3*h//4)
            roi_flag = "failed"
            roi_reason = "auto_detect_roi 실패"
        try:
            thumb = make_thumb(rgb, 90, 70, auto_roi)
        except Exception:
            thumb = None
        if self._pred_targets:
            t = self._pred_targets[0]
            t["rgb"] = rgb
            t["name"] = (name or t.get("name","Target"))[:64]
            t["roi"] = auto_roi
            t["roi_flag"] = roi_flag
            t["roi_reason"] = roi_reason
            t["roi_source"] = "auto"
            t["thumb"] = thumb
            t["result"] = None
        else:
            tid = 1
            t = {
                "tid": tid, "name": (name or "Target #1")[:64],
                "rgb": rgb, "roi": auto_roi,
                "roi_flag": roi_flag, "roi_reason": roi_reason,
                "roi_source": "auto",
                "color": self._pred_assign_color(),
                "thumb": thumb,
                "cond_hint": (self._pred_cond_var.get().strip()
                              if hasattr(self, "_pred_cond_var") else ""),
                "result": None,
            }
            self._pred_targets.append(t)
            self._pred_sel_tid = tid
        self._pred_rebuild_cards()
        return t

    def _pred_select(self, tid):
        """카드 클릭 = 선택. Top-3 후보 + 메인 차트 + Pseudo-Raman + 결과 라벨/코멘트 갱신."""
        if not any(t.get("tid")==tid for t in self._pred_targets):
            return
        self._pred_sel_tid = tid
        t = self._pred_get_target_by_tid(tid)
        if t and t.get("result"):
            try:
                res = t["result"]
                self._update_pred_candidates(
                    res.get("target_metrics",{}),
                    res.get("top",[]))
            except Exception:
                pass
            # ── 선택된 target 의 결과를 모든 의존 표시에 반영 ──
            res = t["result"]
            tm = res.get("target_metrics") or res.get("target")
            # 메인 차트 (radar + timeline) 갱신
            try:
                self._pred_draw_all_charts()
            except Exception:
                pass
            # Pseudo-Raman 회귀/스펙트럼 + ctx 갱신
            if tm is not None and getattr(self, "_raman_data", None):
                try:
                    self._run_pseudo_raman_inline(tm)
                except Exception:
                    pass
            # 결과 요약 라벨 (_pred_run_all 과 동일 포맷)
            try:
                if hasattr(self, "_pred_result_var"):
                    ed   = res.get("est_day")
                    conf = res.get("confidence", 0)
                    best = res.get("best_match", "—")
                    nm   = (t.get("name","?") or "?")[:18]
                    if ed is not None:
                        txt = (f"[{nm}] "
                               f"Day {ed:.1f} (Conf {conf:.0f}%)  "
                               f"Best={best}")
                    else:
                        txt = _L(f"[{nm}] 추정 불가",
                                 f"[{nm}] cannot estimate")
                    self._pred_result_var.set(txt)
            except Exception:
                pass
            # 분석 근거 코멘트 (재생성)
            try:
                if hasattr(self, "_pred_comment") and tm is not None:
                    comment = self._generate_comment(
                        tm, res.get("top", []),
                        res.get("est_day"),
                        res.get("confidence", 0),
                        res.get("cond_input",""),
                        res.get("pool_cond", []))
                    self._pred_comment.configure(state="normal")
                    self._pred_comment.delete("1.0","end")
                    self._pred_comment.insert("end", comment)
                    self._pred_comment.configure(state="disabled")
            except Exception:
                pass
        self._pred_rebuild_cards()

    def _pred_on_drop(self, event):
        """카드 패널 DnD 핸들러."""
        paths = parse_drop_paths(event.data)
        img_paths = [p for p in paths
                     if os.path.splitext(p)[1].lower() in _IMG_EXTS]
        if not img_paths:
            self._set_status(_L("⚠ 이미지 파일이 아닙니다.",
                                "⚠ Not an image file."))
            return
        for p in img_paths:
            try:
                self._pred_add_target(Image.open(p),
                                      os.path.basename(p))
            except Exception as ex:
                messagebox.showerror(_L("오류","Error"),
                                     f"{p}\n{ex}")

    def _pred_load_image(self):
        """[➕ 추가] 핸들러 — 다중 파일 선택 가능"""
        if len(self._pred_targets) >= PRED_MAX_TARGETS:
            messagebox.showwarning(
                _L("주의","Warning"),
                _L(f"평가대상은 최대 {PRED_MAX_TARGETS}개까지만 가능합니다.",
                   f"Maximum {PRED_MAX_TARGETS} targets allowed."))
            return
        paths = filedialog.askopenfilenames(
            title=_L("평가대상 이미지","Target Images"),
            filetypes=[(_L("이미지","Image"),"*.png *.jpg *.jpeg *.bmp *.tiff"),
                       (_L("전체","All"),"*.*")])
        if not paths:
            return
        for path in paths:
            try:
                self._pred_add_target(Image.open(path),
                                      os.path.basename(path))
            except Exception as ex:
                messagebox.showerror(_L("오류","Error"), str(ex))

    def _pred_paste_image(self):
        try:
            from PIL import ImageGrab
            pil = ImageGrab.grabclipboard()
            if pil is None:
                messagebox.showwarning(_L("클립보드","Clipboard"),
                                       _L("이미지가 없다.","No image found.")); return
            if isinstance(pil, list):
                pil = Image.open(pil[0]) if pil else None
            if pil is None: return
            ts = datetime.datetime.now().strftime("%H%M%S")
            self._pred_add_target(pil, f"clipboard_{ts}.png")
        except Exception as ex:
            messagebox.showerror(_L("오류","Error"), str(ex))

    def _pred_draw_preview(self):
        """호환용 stub — 카드 리스트가 미리보기 역할."""
        try:
            self._pred_rebuild_cards()
        except Exception:
            pass

    def _pred_open_roi(self, tid=None):
        """[🎯 ROI 보정] 핸들러. tid 가 None 이면 첫 target."""
        if tid is None:
            if not self._pred_targets:
                messagebox.showinfo(_L("알림","Info"),
                                    _L("평가대상을 먼저 추가한다.",
                                       "Add a target first."))
                return
            t = self._pred_targets[0]
        else:
            t = self._pred_get_target_by_tid(tid)
            if t is None:
                return
        if t.get("rgb") is None:
            return
        fake_entry = {"name": t.get("name","target"),
                      "rgb":  t["rgb"],
                      "roi":  t.get("roi")}
        def on_confirm(roi, _img=None):
            t["roi"]        = roi
            t["roi_source"] = "manual"
            try:
                flag, reason = evaluate_roi_quality(t["rgb"], roi)
                t["roi_flag"]   = flag
                t["roi_reason"] = reason
            except Exception:
                t["roi_flag"]   = "manual"
                t["roi_reason"] = "manual"
            try:
                t["thumb"] = make_thumb(t["rgb"], 90, 70, roi)
            except Exception:
                pass
            # ROI 변경 시 결과 무효화 (재계산 필요)
            t["result"] = None
            # Advanced 탭 stale 방지
            self._last_eval_ctx = None
            self._pred_rebuild_cards()
        ROISelector(self, fake_entry, on_confirm)

    def _pred_open_color(self, tid):
        """[🎨 색상] 핸들러 — 다음 사용 가능한 색상으로 cycling."""
        t = self._pred_get_target_by_tid(tid)
        if t is None:
            return
        cur = t.get("color")
        try:
            idx = TARGET_COLOR_PALETTE.index(cur)
        except ValueError:
            idx = -1
        used = {x.get("color") for x in self._pred_targets if x is not t}
        changed = False
        for j in range(1, len(TARGET_COLOR_PALETTE)+1):
            cand = TARGET_COLOR_PALETTE[(idx+j) % len(TARGET_COLOR_PALETTE)]
            if cand not in used:
                t["color"] = cand
                changed = True
                break
        if not changed:
            # 모든 색이 사용 중 → 강제 다음 색 (중복 허용) + 사용자 안내
            next_idx = (idx + 1) % len(TARGET_COLOR_PALETTE)
            t["color"] = TARGET_COLOR_PALETTE[next_idx]
            try:
                self._set_status(_L(
                    "⚠ 모든 색이 사용 중 — 중복 색으로 변경",
                    "⚠ All colors in use — duplicated"))
            except Exception:
                pass
        self._pred_rebuild_cards()

    def _pred_clear_roi(self, tid=None):
        """호환용 — tid 의 ROI 를 자동 재추정."""
        if tid is None:
            if not self._pred_targets:
                return
            t = self._pred_targets[0]
        else:
            t = self._pred_get_target_by_tid(tid)
            if t is None:
                return
        try:
            roi, flag, reason = auto_detect_roi(t["rgb"], cond=t.get("cond_hint"))
        except Exception:
            h, w = t["rgb"].shape[:2]
            roi = (w//4, h//4, 3*w//4, 3*h//4)
            flag = "failed"
            reason = "auto_detect_roi 실패"
        t["roi"]        = roi
        t["roi_flag"]   = flag
        t["roi_reason"] = reason
        t["roi_source"] = "auto"
        try:
            t["thumb"] = make_thumb(t["rgb"], 90, 70, roi)
        except Exception:
            pass
        t["result"] = None
        # Advanced 탭 stale 방지
        self._last_eval_ctx = None
        self._pred_rebuild_cards()

    # ─────────────────────────────────────────
    #  카드 빌드
    # ─────────────────────────────────────────
    def _pred_rebuild_cards(self):
        """카드 리스트 다시 그리기."""
        if not hasattr(self, "_pred_lf"):
            return
        for w in self._pred_lf.winfo_children():
            w.destroy()

        n = len(self._pred_targets)
        if hasattr(self, "_pred_count_var"):
            self._pred_count_var.set(f"{n}/{PRED_MAX_TARGETS}")
        # ➕ 추가 버튼 활성/비활성
        if hasattr(self, "_pred_add_btn"):
            self._pred_add_btn.configure(
                state="disabled" if n >= PRED_MAX_TARGETS else "normal",
                bg=BTN if n >= PRED_MAX_TARGETS else ACCENT,
                fg=SUB if n >= PRED_MAX_TARGETS else "white")

        if n == 0:
            tk.Label(self._pred_lf,
                     text=_L("평가대상을 추가하세요.\n드래그앤드롭 OK",
                             "Add evaluation targets.\n(Drag&drop supported)"),
                     bg=PANEL, fg=SUB, font=("Segoe UI",8),
                     justify="center").pack(padx=4, pady=20)
            return

        for t in self._pred_targets:
            tid = t.get("tid")
            is_sel = (tid == self._pred_sel_tid)
            roi_flag = t.get("roi_flag")
            has_roi = t.get("roi") is not None
            # 카드 테두리 색 (PURPLE 제외 — 그룹 일관성 적용 X)
            brd = _border_color_for_roi(
                roi_flag, has_roi,
                {"green": GREEN, "amber": AMBER, "red": RED,
                 "purple": PURPLE, "border": BORDER},
                inconsistent=False)
            if is_sel:
                brd = ACCENT
            thick = 2 if (is_sel or
                          roi_flag in ("warn_small","warn_off",
                                       "warn_paper","failed")) else 1

            card = tk.Frame(self._pred_lf,
                            bg=CARD2 if is_sel else CARD,
                            highlightbackground=brd,
                            highlightthickness=thick,
                            cursor="hand2")
            card.pack(fill="x", padx=4, pady=3)
            card.bind("<Button-1>",
                      lambda e, x=tid: self._pred_select(x))

            # 1행: 썸네일 + 색상 + 이름 + ❌
            r1 = tk.Frame(card, bg=card["bg"])
            r1.pack(fill="x", padx=5, pady=(5,2))

            th = t.get("thumb")
            if th is not None:
                try:
                    pil_copy = th.copy()
                    pil_copy.thumbnail((90,70), Image.LANCZOS)
                    tk_th = ImageTk.PhotoImage(pil_copy)
                    self._refs[f"pred_card_th_{tid}"] = tk_th
                    th_lbl = tk.Label(r1, image=tk_th,
                                       bg=card["bg"], cursor="hand2")
                    th_lbl.pack(side="left", padx=(0,4))
                    th_lbl.bind("<Button-1>",
                                lambda e, x=tid: self._pred_select(x))
                except Exception:
                    tk.Label(r1, text="📷", bg=card["bg"],
                             fg=SUB, font=("Segoe UI",16),
                             width=4).pack(side="left", padx=(0,4))
            else:
                tk.Label(r1, text="📷", bg=card["bg"],
                         fg=SUB, font=("Segoe UI",16),
                         width=4).pack(side="left", padx=(0,4))

            # 색상 사각형
            color_sq = tk.Frame(r1, bg=t.get("color",ACCENT),
                                width=12, height=12,
                                highlightbackground=BORDER,
                                highlightthickness=1)
            color_sq.pack(side="left", padx=(0,4))
            color_sq.pack_propagate(False)

            name = t.get("name","")
            if len(name) > 18:
                name_disp = name[:17] + "…"
            else:
                name_disp = name
            tk.Label(r1, text=name_disp,
                     bg=card["bg"], fg=TXT, font=MFB,
                     anchor="w").pack(side="left", fill="x",
                                      expand=True)

            tk.Button(r1, text="❌",
                      command=lambda x=tid:
                          self._pred_remove_target(x),
                      bg=card["bg"], fg=RED, font=LF,
                      relief="flat", cursor="hand2",
                      padx=3).pack(side="right")

            # 2행: ROI 상태
            r2 = tk.Frame(card, bg=card["bg"])
            r2.pack(fill="x", padx=8, pady=1)
            if has_roi:
                roi = t["roi"]
                rw = roi[2]-roi[0]; rh = roi[3]-roi[1]
                sym = "✔" if roi_flag in ("good","manual") else \
                      ("⚠" if roi_flag in ("warn_small","warn_off",
                                            "warn_paper") else
                       ("✗" if roi_flag=="failed" else "•"))
                col = (GREEN if roi_flag in ("good","manual")
                       else AMBER if roi_flag in
                            ("warn_small","warn_off","warn_paper")
                       else RED if roi_flag=="failed"
                       else SUB)
                src = t.get("roi_source","auto")
                tk.Label(r2,
                         text=f"ROI: {sym} {rw}×{rh} ({src})",
                         bg=card["bg"], fg=col,
                         font=("Segoe UI",7,"bold")
                         ).pack(side="left")
            else:
                tk.Label(r2, text=_L("ROI: 없음","ROI: none"),
                         bg=card["bg"], fg=AMBER,
                         font=("Segoe UI",7,"bold")
                         ).pack(side="left")

            # 3행: 결과 요약
            r3 = tk.Frame(card, bg=card["bg"])
            r3.pack(fill="x", padx=8, pady=1)
            res = t.get("result")
            if res:
                ed = res.get("est_day")
                conf = res.get("confidence", 0)
                bm = res.get("best_match", "—")
                ed_str = f"{ed:.1f}" if ed is not None else "—"
                tk.Label(r3,
                         text=f"Day={ed_str}  "
                              f"Conf={conf:.0f}%  "
                              f"Best={bm}",
                         bg=card["bg"], fg=GOLD,
                         font=("Segoe UI",7),
                         anchor="w").pack(side="left", fill="x",
                                          expand=True)
            else:
                tk.Label(r3, text=_L("(미분석)","(not analyzed)"),
                         bg=card["bg"], fg=SUB,
                         font=("Segoe UI",7,"italic")
                         ).pack(side="left")

            # 4행: 버튼들
            r4 = tk.Frame(card, bg=card["bg"])
            r4.pack(fill="x", padx=5, pady=(2,5))
            tk.Button(r4, text=_L("🎯 ROI","🎯 ROI"),
                      command=lambda x=tid: self._pred_open_roi(x),
                      bg=BTN, fg=TEAL,
                      font=("Segoe UI",7,"bold"),
                      relief="flat", cursor="hand2",
                      padx=4, pady=1).pack(side="left", padx=(0,3))
            tk.Button(r4, text=_L("🎨 색상","🎨 Color"),
                      command=lambda x=tid: self._pred_open_color(x),
                      bg=BTN, fg=t.get("color",ACCENT),
                      font=("Segoe UI",7,"bold"),
                      relief="flat", cursor="hand2",
                      padx=4, pady=1).pack(side="left", padx=3)

    # ─────────────────────────────────────────
    #  분석 — 단일 target / 전체 실행
    # ─────────────────────────────────────────
    def _pred_compute_one(self, t: dict):
        """단일 target 분석. 결과 dict 를 반환 (또는 None)."""
        rgb = t.get("rgb")
        if rgb is None:
            return None

        # 참조 풀
        ref_pool = [img for img in self.images
                    if not np.isnan(img.get("lab",{}).get("b", np.nan))]
        if not ref_pool:
            return None

        roi = t.get("roi")
        if roi:
            mask = roi_to_mask(rgb.shape, roi)
        else:
            mask = np.ones(rgb.shape[:2], bool)
            roi  = (0,0,rgb.shape[1],rgb.shape[0])

        target = {
            "s_mean":       compute_s_mean(rgb, mask),
            "yellow_ratio": compute_yellow_ratio(rgb, mask),
            "yellowness_idx": compute_yellowness_index(rgb, mask),
            "lab":          compute_lab_metrics(rgb, mask),
            "glcm":         compute_glcm_metrics(rgb, mask),
            "rgb":          rgb,
            "mask":         mask,
            "roi":          roi,
        }

        cond_input = (t.get("cond_hint") or "").strip()
        if cond_input:
            pool_cond = [img for img in ref_pool
                         if img["cond"] == cond_input]
            if not pool_cond:
                pool_cond = [img for img in ref_pool
                             if cond_input.lower() in img["cond"].lower()
                             or img["cond"].lower() in cond_input.lower()]
            if not pool_cond:
                pool_cond = ref_pool
        else:
            pool_cond = ref_pool

        def _safe(v): return 0.0 if np.isnan(v) else float(v)

        all_b  = [_safe(img["lab"]["b"])       for img in pool_cond]
        all_s  = [_safe(img["s_mean"])         for img in pool_cond]
        all_yi = [_safe(img["yellowness_idx"]) for img in pool_cond]

        def norm_range(vals):
            mn,mx = min(vals),max(vals)
            r = mx-mn if mx!=mn else 1.0
            return mn, r

        b_mn,b_r   = norm_range(all_b)
        s_mn,s_r   = norm_range(all_s)
        yi_mn,yi_r = norm_range(all_yi)

        t_b  = (_safe(target["lab"]["b"])       - b_mn)  / b_r
        t_s  = (_safe(target["s_mean"])         - s_mn)  / s_r
        t_yi = (_safe(target["yellowness_idx"]) - yi_mn) / yi_r

        scores = []
        for img in pool_cond:
            i_b  = (_safe(img["lab"]["b"])       - b_mn)  / b_r
            i_s  = (_safe(img["s_mean"])         - s_mn)  / s_r
            i_yi = (_safe(img["yellowness_idx"]) - yi_mn) / yi_r
            wb  = self.cfg_w_b.get()
            ws  = self.cfg_w_s.get()
            wyi = self.cfg_w_yi.get()
            wt  = wb + ws + wyi
            if wt > 0: wb,ws,wyi = wb/wt, ws/wt, wyi/wt
            dist = (wb *(t_b-i_b)**2 +
                    ws *(t_s-i_s)**2 +
                    wyi*(t_yi-i_yi)**2) ** 0.5
            scores.append((dist, img))
        scores.sort(key=lambda x: x[0])
        top = scores[:min(5, len(scores))]

        def df(d):
            try: return float(d)
            except: return None

        day_weights = []
        for dist, img in top[:3]:
            d = df(img["day"])
            if d is not None:
                w = 1/(dist+1e-6)
                day_weights.append((d, w))

        if day_weights:
            total_w = sum(w for _,w in day_weights)
            est_day = sum(d*w for d,w in day_weights) / total_w
            confidence = max(0, 100 - top[0][0]*200)
        else:
            est_day = None
            confidence = 0

        # delta_e
        ref0 = next((img for img in pool_cond
                     if str(img.get("day","")) == "0"), None)
        if ref0 and "lab" in ref0:
            target["delta_e"] = compute_delta_e(target["lab"], ref0["lab"])
        else:
            target["delta_e"] = float("nan")

        best_dist, best_img = top[0] if top else (0.0, None)
        best_match = (f"{best_img.get('cond','?')[:12]}-"
                      f"{best_img.get('day','?')}d"
                      if best_img else "—")

        return {
            "target_metrics": target,
            "scores":         scores,
            "top":            top,
            "est_day":        est_day,
            "confidence":     confidence,
            "best_match":     best_match,
            "best_dist":      best_dist,
            "pool_cond":      pool_cond,
            "cond_input":     cond_input,
        }

    def _pred_run_all(self):
        """모든 평가대상을 동기 실행."""
        if not self._pred_targets:
            messagebox.showwarning(
                _L("주의","Warning"),
                _L("평가대상을 먼저 추가하세요.",
                   "Add evaluation targets first."))
            return
        ref_pool = [img for img in self.images
                    if not np.isnan(img.get("lab",{}).get("b", np.nan))]
        if not ref_pool:
            messagebox.showwarning(
                _L("주의","Warning"),
                _L("분석된 참조 이미지가 없습니다.\n"
                   "[▶ Analyze All] 또는 DB 로드 먼저 실행하세요.",
                   "No analyzed reference images.\n"
                   "Run [▶ Analyze All] or load DB first."))
            return

        n = len(self._pred_targets)
        # 분석 시작 — 이전 결과 클리어 (차트 stale 방지)
        for _t in self._pred_targets:
            _t["result"] = None
        try:
            self._pred_rebuild_cards()
            self.update_idletasks()
        except Exception:
            pass
        for i, t in enumerate(self._pred_targets, start=1):
            self._set_status(
                _L(f"분석 중 ({i}/{n}): {t.get('name','')[:24]}",
                   f"Analyzing ({i}/{n}): {t.get('name','')[:24]}"))
            try:
                self.update_idletasks()
            except Exception:
                pass
            try:
                res = self._pred_compute_one(t)
                t["result"] = res
            except Exception as ex:
                t["result"] = None
                self._set_status(
                    _L(f"⚠ {t.get('name','?')} 분석 실패: {ex}",
                       f"⚠ {t.get('name','?')} failed: {ex}"))

        # 차트 + 카드 갱신
        try:
            self._pred_draw_all_charts()
        except Exception as ex:
            self._set_status(f"chart err: {ex}")
        self._pred_rebuild_cards()

        # 선택된 target 의 후보/Pseudo-Raman 갱신
        sel = self._pred_get_target_by_tid(self._pred_sel_tid)
        if sel and sel.get("result"):
            res = sel["result"]
            try:
                self._update_pred_candidates(
                    res.get("target_metrics",{}),
                    res.get("top",[]))
            except Exception:
                pass
            # 결과 요약 텍스트
            ed = res.get("est_day")
            conf = res.get("confidence",0)
            best = res.get("best_match","—")
            if ed is not None:
                txt = (f"[{sel['name'][:18]}] "
                       f"Day {ed:.1f} (Conf {conf:.0f}%)  "
                       f"Best={best}")
            else:
                txt = _L(f"[{sel['name'][:18]}] 추정 불가",
                         f"[{sel['name'][:18]}] cannot estimate")
            try:
                self._pred_result_var.set(txt)
            except Exception:
                pass
            # 분석 근거 코멘트
            try:
                tm = res.get("target_metrics",{})
                comment = self._generate_comment(
                    tm, res.get("top",[]),
                    ed, conf,
                    res.get("cond_input",""),
                    res.get("pool_cond",[]))
                self._pred_comment.configure(state="normal")
                self._pred_comment.delete("1.0","end")
                self._pred_comment.insert("end", comment)
                self._pred_comment.configure(state="disabled")
            except Exception:
                pass
            # Pseudo-Raman
            if self._raman_data:
                try:
                    self._run_pseudo_raman_inline(tm)
                except Exception as ex:
                    # 차트 명시 clear + 에러 표시 (이전 결과 표시 방지)
                    try:
                        self._pseudo_reg_fig.clear()
                        ax = self._pseudo_reg_fig.add_subplot(111)
                        ax.text(0.5, 0.5,
                                _L(f"Pseudo-Raman 오류:\n{ex}",
                                   f"Pseudo-Raman error:\n{ex}"),
                                ha="center", va="center",
                                color=RED, fontsize=8,
                                transform=ax.transAxes)
                        ax.axis("off")
                        self._pseudo_reg_cv.draw()
                    except Exception:
                        pass
                    try:
                        self._pseudo_spec_fig.clear()
                        ax = self._pseudo_spec_fig.add_subplot(111)
                        ax.text(0.5, 0.5,
                                _L(f"Pseudo-Raman 오류:\n{ex}",
                                   f"Pseudo-Raman error:\n{ex}"),
                                ha="center", va="center",
                                color=RED, fontsize=8,
                                transform=ax.transAxes)
                        ax.axis("off")
                        self._pseudo_spec_cv.draw()
                    except Exception:
                        pass
                    self._set_status(
                        _L(f"⚠ Pseudo-Raman 분석 오류: {ex}",
                           f"⚠ Pseudo-Raman error: {ex}"))

        ok = sum(1 for t in self._pred_targets if t.get("result"))
        self._set_status(
            _L(f"✓ 분석 완료: {ok}/{n}",
               f"✓ Done: {ok}/{n}"))

    def _run_predict(self):
        """호환용 — 단일 흐름 호출 시 _pred_run_all 로 위임"""
        self._pred_run_all()

    def _pred_draw_all_charts(self):
        """모든 target 의 결과를 차트에 동시 표시 (선택된 target 컨텍스트 기반).
        Pseudo-Raman 차트 도 함께 갱신 (m9: stale 방지)."""
        # 선택된 target 의 결과를 메인 ctx 로 사용
        sel = self._pred_get_target_by_tid(self._pred_sel_tid)
        sel_res = sel.get("result") if sel else None
        if sel_res:
            target = sel_res.get("target_metrics",{})
            top    = sel_res.get("top",[])
            est_day= sel_res.get("est_day")
            scores = sel_res.get("scores",[])
            pool   = sel_res.get("pool_cond",[])
            self._draw_pred_charts(target, top, est_day, scores, pool)
        else:
            # 결과 없음 — 빈 화면
            for key in ("radar","timeline"):
                if key in self._pred_figs:
                    fig = self._pred_figs[key]["fig"]
                    fig.clear()
                    ax = fig.add_subplot(111)
                    ax.set_facecolor(PANEL2)
                    ax.text(0.5,0.5,"No results",
                            transform=ax.transAxes,
                            ha="center",va="center",
                            color=SUB, fontsize=8)
                    ax.axis("off")
                    self._pred_figs[key]["cv"].draw()

        # Pseudo-Raman 차트 도 함께 갱신 (선택된 target 기준)
        try:
            if getattr(self, "_raman_data", None):
                # 선택된 target 우선, 없으면 첫 번째 결과 있는 target
                sel_t = sel
                if sel_t is None or not sel_t.get("result"):
                    sel_t = next(
                        (t for t in self._pred_targets if t.get("result")),
                        None)
                if sel_t and sel_t.get("result"):
                    tm = (sel_t["result"].get("target_metrics")
                          or sel_t["result"].get("target"))
                    if tm is not None:
                        try:
                            self._run_pseudo_raman_inline(tm)
                        except Exception:
                            pass
                else:
                    # 분석 결과 없음 — Pseudo 차트 비우기
                    if hasattr(self, "_pseudo_reg_fig"):
                        self._pseudo_reg_fig.clear()
                        if hasattr(self, "_pseudo_reg_cv"):
                            self._pseudo_reg_cv.draw()
                    if hasattr(self, "_pseudo_spec_fig"):
                        self._pseudo_spec_fig.clear()
                        if hasattr(self, "_pseudo_spec_cv"):
                            self._pseudo_spec_cv.draw()
        except Exception:
            pass

    def _run_pseudo_raman_inline(self, target: dict):
        """Pseudo-Raman: 회귀 앙상블 → 추정 스펙트럼. 컨텍스트 _last_eval_ctx에 저장."""
        def _safe(v):
            try:
                fv = float(v)
                return 0.0 if fv != fv else fv
            except Exception:
                return 0.0

        # ── Evaluation 탭에서 지정한 피크 처리 조건 읽기 ──
        mode = getattr(self, "_pred_peak_mode",
                       tk.StringVar(value="auto_max")).get()
        ranges_str = getattr(self, "_pred_peak_ranges",
                             tk.StringVar(value="300-360")).get()
        peak_ranges = []
        for seg in ranges_str.split(","):
            seg = seg.strip()
            try:
                lo, hi = map(float, seg.split("-"))
                peak_ranges.append((lo, hi))
            except Exception:
                pass
        if not peak_ranges:
            peak_ranges = [(300, 360)]

        def _extract_peak(r_m) -> float:
            """선택된 모드에 따라 Raman 레코드에서 피크 강도 추출"""
            spec = r_m.get("spectrum")
            if spec and mode != "auto_max":
                sh = np.array(spec.get("shifts", []))
                iv = np.array(spec.get("intensities", []))
                if len(sh) == 0:
                    return _safe(r_m.get("peak", 0))
                if mode == "area":
                    vals = []
                    for lo, hi in peak_ranges:
                        mask = (sh >= lo) & (sh <= hi)
                        if mask.any():
                            vals.append(float(np.trapz(iv[mask], sh[mask])))
                    return float(np.mean(vals)) if vals else _safe(r_m.get("peak", 0))
                else:  # range_single / range_multi
                    vals = []
                    for lo, hi in peak_ranges:
                        mask = (sh >= lo) & (sh <= hi)
                        if mask.any():
                            vals.append(float(np.max(iv[mask])))
                    return float(np.mean(vals)) if vals else _safe(r_m.get("peak", 0))
            # auto_max 또는 스펙트럼 없을 때
            return _safe(r_m.get("norm_peak") or r_m.get("peak", 0))

        an = [img for img in self.images
              if not np.isnan(img.get("lab",{}).get("b", np.nan))]
        rd = self._raman_data

        pairs = []
        for img in an:
            # 명시적 매칭(raman_id) 우선, 없으면 cond+day 자연 매칭 fallback
            r_m = None
            rid = img.get("raman_id")
            if rid is not None:
                r_m = self._raman_by_id(rid)
            if r_m is None:
                r_m = next((r for r in rd
                            if r["cond"] == img["cond"]
                            and r["day"]  == img["day"]), None)
            if r_m:
                peak_val = _extract_peak(r_m)
                pairs.append({
                    "cond": img["cond"], "day": img["day"],
                    "b":    _safe(img["lab"]["b"]),
                    "s":    _safe(img["s_mean"]),
                    "yi":   _safe(img.get("yellowness_idx", float("nan"))),
                    "de":   _safe(img.get("delta_e", float("nan"))),
                    "peak": peak_val,
                    "norm_peak": peak_val,   # 이미 추출된 값 사용
                    "spectrum":  r_m.get("spectrum"),
                })
        # 추출 후 조건별 0일차 대비 정규화
        if pairs:
            from collections import defaultdict
            cond_max: dict = {}
            def df(d):
                try: return float(d)
                except: return 9999
            for cond in set(p["cond"] for p in pairs):
                pts = sorted([p for p in pairs if p["cond"]==cond],
                             key=lambda x: df(x["day"]))
                if pts:
                    cond_max[cond] = pts[0]["peak"] or 1.0
            for p in pairs:
                ref = cond_max.get(p["cond"], 1.0)
                p["norm_peak"] = p["peak"] / ref if ref != 0 else 0.0

        t_b  = _safe(target.get("lab",{}).get("b",  float("nan")))
        t_s  = _safe(target.get("s_mean",            float("nan")))
        t_yi = _safe(target.get("yellowness_idx",    float("nan")))
        t_de = _safe(target.get("delta_e",           float("nan")))

        # 컨텍스트 갱신
        ctx = getattr(self, "_last_eval_ctx", {}) or {}
        ctx["pairs"] = pairs

        # 부족한 경우 처리
        if len(pairs) < 3:
            msg = (f"Need ≥3 matched pairs (currently {len(pairs)}).\n"
                   "Ensure cond+day match between images and Raman data.")
            for fig_, cv_ in [(self._pseudo_reg_fig, self._pseudo_reg_cv),
                               (self._pseudo_spec_fig, self._pseudo_spec_cv)]:
                fig_.clear()
                fig_.patch.set_facecolor(PANEL)
                ax_ = fig_.add_subplot(111)
                ax_.set_facecolor(PANEL2)
                ax_.text(0.5, 0.5, msg, transform=ax_.transAxes,
                         ha="center", va="center", color=SUB, fontsize=8,
                         wrap=True)
                ax_.axis("off")
                cv_.draw()
            self._set_cmt(self._pseudo_res_txt,
                f"Insufficient matched pairs: {len(pairs)} (need ≥3).\n"
                "Ensure cond+day match between image analysis and Raman data.")
            return

        # ── 회귀 앙상블 ────────────────────────────────
        wb_map = {"b":0.45,"s":0.25,"yi":0.20,"de":0.10}
        t_map  = {"b":t_b, "s":t_s, "yi":t_yi, "de":t_de}

        r2_map = {}; coef_map = {}; estimates = []
        wt_total = 0.0
        for key2, wt in wb_map.items():
            xs = np.array([p[key2] for p in pairs])
            ys = np.array([p["norm_peak"] for p in pairs])
            if xs.std() < 1e-6:
                continue
            coef = np.polyfit(xs, ys, 1)
            pred = float(np.polyval(coef, t_map[key2]))
            res  = ys - np.polyval(coef, xs)
            se   = float(np.std(res))
            r2   = max(0.0, 1 - np.sum(res**2) /
                       (np.sum((ys-ys.mean())**2)+1e-9))
            r2_map[key2]   = r2
            coef_map[key2] = coef
            ew = wt * r2
            estimates.append((pred, se, ew))
            wt_total += ew

        ctx["r2_map"]   = r2_map
        ctx["coef_map"] = coef_map

        if not estimates or wt_total < 1e-9:
            self._set_cmt(self._pseudo_res_txt,
                "Regression failed (all R²≈0). "
                "Check that data covers a range of oxidation states.")
            return

        est_peak = sum(p*w for p,se,w in estimates) / wt_total
        est_se   = (sum(se**2*w for p,se,w in estimates)/wt_total)**0.5
        ci_lo    = max(0.0, est_peak - 1.96*est_se)
        ci_hi    = min(1.5, est_peak + 1.96*est_se)

        dists_sorted = sorted(pairs, key=lambda p: abs(p["norm_peak"]-est_peak))

        # 컨텍스트 저장
        ctx["est_peak"]   = est_peak
        ctx["ci_lo"]      = ci_lo
        ctx["ci_hi"]      = ci_hi
        ctx["est_se"]     = est_se
        ctx["spec_dists"] = dists_sorted
        self._last_eval_ctx = ctx

        # ── 회귀 차트 ──────────────────────────────────
        self._pseudo_reg_fig.clear()
        self._pseudo_reg_fig.patch.set_facecolor(PANEL)
        self._draw_pseudo_reg_fig(self._pseudo_reg_fig, target, large=False)
        self._pseudo_reg_cv.draw()

        # ── 추정 스펙트럼 차트 ─────────────────────────
        self._pseudo_spec_fig.clear()
        self._pseudo_spec_fig.patch.set_facecolor(PANEL)
        self._draw_pseudo_spec_fig(self._pseudo_spec_fig, target, large=False)
        self._pseudo_spec_cv.draw()

        # ── 결과 텍스트 ────────────────────────────────
        def _oxid(peak):
            if peak >= 0.85:   return "Pristine (unoxidized)"
            elif peak >= 0.65: return "Early oxidation"
            elif peak >= 0.40: return "Significant oxidation"
            else:              return "Severe oxidation (HfO₂ dominant)"

        best_ref = dists_sorted[0]
        r2_lines = [f"  {k:4s}: R²={v:.3f}  (weight={wb_map.get(k,0)*v:.4f})"
                    for k,v in r2_map.items()]
        result_lines = [
            "━━ Pseudo-Raman Estimation Result ━━━━━━━━━━━━━━━━━━━",
            f"Input metrics:  b*={t_b:.2f}  S={t_s:.1f}  YI={t_yi:.1f}  ΔE={t_de:.2f}",
            f"Matched pairs used for regression: {len(pairs)}",
            "",
            f"Estimated A₁g peak:  {est_peak:.4f}  (normalized, 1.0=pristine)",
            f"95% Confidence Interval:  [{ci_lo:.4f},  {ci_hi:.4f}]",
            f"Uncertainty (σ):  ±{est_se:.4f}",
            "",
            f"Closest reference: {best_ref['cond']}  Day {best_ref['day']}",
            f"  (ref peak={best_ref['norm_peak']:.4f}, diff={abs(best_ref['norm_peak']-est_peak):.4f})",
            f"Oxidation stage:  {_oxid(est_peak)}",
            "",
            "Regression R² per metric:",
            *r2_lines,
            "",
            "Ensemble formula:",
            f"  est = Σ(pred_m × base_wt_m × R²_m) / Σ(base_wt_m × R²_m)",
        ]
        self._set_cmt(self._pseudo_res_txt, "\n".join(result_lines))
        self._set_status(
            f"✓ Pseudo-Raman: est. peak={est_peak:.4f}  [{ci_lo:.4f}, {ci_hi:.4f}]")

    def _draw_pseudo_reg_fig(self, fig, target, large: bool):
        """회귀 차트 — 팝업/인라인 공용. _last_eval_ctx에서 데이터 읽음."""
        fs_t = 11 if large else 7
        fs_a = 9  if large else 6
        fs_k = 8  if large else 5
        ms   = 8  if large else 4

        ctx     = getattr(self, "_last_eval_ctx", {}) or {}
        pairs   = ctx.get("pairs", [])
        r2_map  = ctx.get("r2_map", {})
        coef_map= ctx.get("coef_map", {})
        est_peak= ctx.get("est_peak")

        def _safe(v):
            try:
                fv = float(v)
                return 0.0 if fv != fv else fv
            except Exception:
                return 0.0

        target = ctx.get("target", target)
        t_map  = {
            "b":  _safe(target.get("lab",{}).get("b",  float("nan"))),
            "s":  _safe(target.get("s_mean",            float("nan"))),
            "yi": _safe(target.get("yellowness_idx",    float("nan"))),
            "de": _safe(target.get("delta_e",           float("nan"))),
        }

        metrics = [("b","Lab b*"),("s","S-ch"),("yi","YI"),("de","ΔE")]
        # 인라인 vs 팝업 gridspec 간격 분리
        if large:
            gs = fig.add_gridspec(2, 2, hspace=0.52, wspace=0.40,
                                   top=0.88, bottom=0.10,
                                   left=0.10, right=0.97)
        else:
            gs = fig.add_gridspec(2, 2, hspace=0.70, wspace=0.50,
                                   top=0.85, bottom=0.12,
                                   left=0.13, right=0.97)
        fig.patch.set_facecolor(PANEL)

        for ax_i, (key2, lbl) in enumerate(metrics):
            ax = fig.add_subplot(gs[ax_i//2, ax_i%2])
            ax.set_facecolor(PANEL2)
            for sp in ax.spines.values(): sp.set_color(BORDER)
            ax.tick_params(colors=SUB, labelsize=fs_k,
                           pad=1 if not large else 3)

            if not pairs:
                ax.text(0.5, 0.5, "No data",
                        transform=ax.transAxes,
                        ha="center", va="center",
                        color=SUB, fontsize=fs_k)
                ax.set_xlabel(lbl, fontsize=fs_k, color=SUB)
                continue

            xs = np.array([p[key2] for p in pairs])
            ys = np.array([p["norm_peak"] for p in pairs])
            ax.scatter(xs, ys, color=ACCENT, s=ms*8, zorder=5, alpha=0.85)

            if key2 in coef_map:
                coef = coef_map[key2]
                xl   = np.linspace(xs.min(), xs.max(), 60)
                ax.plot(xl, np.polyval(coef, xl),
                        "--", color=AMBER,
                        lw=1.5 if large else 0.8, alpha=0.85)

                # ── 모든 평가대상의 ★ 마커 표시 ──
                targets_with_result = [
                    tt for tt in getattr(self, "_pred_targets", [])
                    if tt.get("result")]
                for tt in targets_with_result:
                    tm_t = tt["result"].get("target_metrics", {})
                    tv_t = {
                        "b":  _safe(tm_t.get("lab",{}).get("b",0)),
                        "s":  _safe(tm_t.get("s_mean",0)),
                        "yi": _safe(tm_t.get("yellowness_idx",0)),
                        "de": _safe(tm_t.get("delta_e",0)),
                    }.get(key2, 0)
                    pred_t = float(np.polyval(coef, tv_t))
                    col_t = tt.get("color", RED)
                    ax.plot(tv_t, pred_t, "*",
                            color=col_t, ms=ms+4, zorder=7,
                            markeredgecolor="white",
                            markeredgewidth=0.6)

                # 호환용 — 단일 target (선택된) 도 v 마커 (for 루프 밖)
                if not targets_with_result:
                    tv = t_map.get(key2, 0)
                    pred = float(np.polyval(coef, tv))
                    ax.axvline(tv, color=RED,
                               lw=1.5 if large else 0.8,
                               ls="--", alpha=0.8)
                    ax.plot(tv, pred, "v", color=RED,
                            ms=ms+2, zorder=6)

                r2 = r2_map.get(key2, 0)
                ann = f"R²={r2:.2f}"
                ax.text(0.04, 0.96, ann,
                        transform=ax.transAxes, fontsize=fs_k,
                        color=AMBER, va="top",
                        bbox=dict(boxstyle="round,pad=0.2",
                                  fc=PANEL, ec=BORDER, alpha=0.85))

            ax.set_xlabel(lbl, fontsize=fs_a, color=SUB, labelpad=1)
            ax.set_ylabel("Raman", fontsize=max(4,fs_k-1),
                          color=SUB, labelpad=1)

        fig.suptitle("Regression: Metrics → Raman Peak",
                     fontsize=fs_t, color=TXT,
                     y=0.98 if large else 0.99)

    def _draw_pseudo_spec_fig(self, fig, target, large: bool):
        """추정 스펙트럼 차트 — 팝업/인라인 공용."""
        fs_t = 10 if large else 8
        fs_a = 9  if large else 7
        fs_k = 8  if large else 6
        lw   = 2.0 if large else 1.5

        ctx       = getattr(self, "_last_eval_ctx", {}) or {}
        est_peak  = ctx.get("est_peak")
        ci_lo     = ctx.get("ci_lo")
        ci_hi     = ctx.get("ci_hi")
        est_se    = ctx.get("est_se", 0)
        dists     = ctx.get("spec_dists", [])

        ax = fig.add_subplot(111)
        ax.set_facecolor(PANEL2)
        fig.patch.set_facecolor(PANEL)
        for sp in ax.spines.values(): sp.set_color(BORDER)
        ax.tick_params(colors=SUB, labelsize=fs_k)

        if est_peak is None:
            ax.text(0.5, 0.5, "No estimation yet.\nRun Evaluation first.",
                    transform=ax.transAxes, ha="center", va="center",
                    color=SUB, fontsize=fs_k+2)
            ax.axis("off"); return

        ref_spectra = []
        for p in dists[:3]:
            spec = p.get("spectrum")
            if spec:
                sh = np.array(spec["shifts"])
                iv = np.array(spec["intensities"])
                ref_spectra.append((sh, iv, p["day"], p.get("norm_peak",0)))
                # 참조 — 알파 낮춰
                ax.plot(sh, iv, lw=0.8, alpha=0.15,
                        color=SUB, zorder=1)

        if len(ref_spectra) >= 2:
            s1,v1,d1,p1 = ref_spectra[0]
            s2,v2,d2,p2 = ref_spectra[1]
            v2i = np.interp(s1, s2, v2)
            alpha = ((est_peak-p2)/(p1-p2+1e-9)
                     if abs(p1-p2)>1e-6 else 0.5)
            alpha = max(0.0, min(1.0, alpha))
            v_est = alpha*v1 + (1-alpha)*v2i
            band  = abs(ci_hi-ci_lo)/2 * 0.5
            # 신뢰구간 밴드 — 선택된 target 색상
            sel_t = self._pred_get_target_by_tid(
                getattr(self, "_pred_sel_tid", None))
            sel_color = sel_t.get("color", ACCENT) if sel_t else ACCENT
            ax.fill_between(s1,
                            v_est*(1-band), v_est*(1+band),
                            alpha=0.18, color=sel_color, zorder=2)
            # 선택된 target 의 곡선만 (다른 target 은 자체 ctx 가 없어
            # 동일 v_est 표시 시 거짓 표현 — 곡선 overlay 제거)
            sel_label = (f"T{sel_t.get('tid','?')}"
                         if sel_t else "Selected")
            ax.plot(s1, v_est, lw=lw, color=sel_color,
                    alpha=0.85, zorder=3, label=sel_label)

            # ── 다른 target 의 est_peak 위치를 ★ 마커로 표시 ──
            # (각 target 의 스칼라 est_peak 가 ctx 에 없으므로,
            #  result 가 있는 target 만 A₁g(337) 위치에 색 마커로 위치 표시)
            try:
                a1g_x = 337.0
                # 선택 target 의 v_est 피크 강도를 기준으로 보조 마커 배치
                v_max = float(np.max(v_est)) if len(v_est) else 1.0
                for tt in getattr(self, "_pred_targets", []):
                    if tt is sel_t or not tt.get("result"):
                        continue
                    col_t = tt.get("color", ACCENT)
                    # 자체 est_peak 가 없으므로 v_max 위치에 ★ 만 표시
                    # (시각적 위치 = "다른 target 도 결과 있음" 표식)
                    ax.plot(a1g_x, v_max, marker="*",
                            markersize=12,
                            color=col_t, linestyle="None",
                            markeredgecolor=BORDER,
                            markeredgewidth=0.5,
                            zorder=4)
            except Exception:
                pass

            # 참조선
            ax.annotate(
                f"A₁g~337cm⁻¹",
                xy=(337, max(v_est)), xytext=(310, max(v_est)*1.08),
                fontsize=fs_k, color=RED,
                arrowprops=dict(arrowstyle="->", color=RED, lw=0.8))

        ax.set_xlabel("Raman Shift (cm⁻¹)", color=SUB, fontsize=fs_a)
        ax.set_ylabel("Intensity (normalized)", color=SUB, fontsize=fs_a)
        ax.set_title(
            f"Pseudo-Raman Spectrum (selected)  "
            f"(A₁g≈{est_peak:.3f}±{est_se:.3f})",
            color=TXT, fontsize=fs_t)
        # 차트 내 범례 생략 (카드 색상 표시가 범례)
        fig.tight_layout(pad=0.6)


    def _pred_chart_popup(self, key: str, title: str):
        """
        차트 확대 팝업 — FigureCanvasTkAgg 직접 사용.
        다른 탭(Raman 등)과 완전히 동일한 방식.
        """
        win = tk.Toplevel(self)
        win.title(f"🔍 {title}")
        win.configure(bg=PANEL)
        sw, sh = self.winfo_screenwidth(), self.winfo_screenheight()
        ww, wh = int(sw * 0.85), int(sh * 0.88)
        win.geometry(f"{ww}x{wh}+{(sw-ww)//2}+{(sh-wh)//2}")
        # ★ Toplevel 크기 확정 후 child 배치 (Windows tkinter 레이아웃 안정화)
        win.update_idletasks()

        # ── 헤더 ─────────────────────────────────
        hdr = tk.Frame(win, bg=PANEL2)
        hdr.pack(side="top", fill="x")
        tk.Label(hdr, text=f"  🔍 {title}",
                 bg=PANEL2, fg=TXT,
                 font=("Segoe UI",11,"bold")).pack(
                 side="left", pady=8, padx=10)
        tk.Button(hdr, text="✕ Close",
                  command=win.destroy,
                  bg=BTN, fg=TXT, font=MF,
                  relief="flat", padx=10, pady=4,
                  cursor="hand2").pack(side="right", padx=8, pady=6)

        # ── 우측: 설명 패널 ───────────────────────
        ctx = getattr(self, "_last_eval_ctx", None)
        # ★ _EVAL_CHART_HELP 속성이 없을 때 방어
        _help_map = getattr(self, "_EVAL_CHART_HELP", {}) or {}
        hi = _help_map.get(key, {}) if isinstance(_help_map, dict) else {}
        ht = hi.get("text", "Chart detail view.")
        try:
            dyn = self._build_chart_explanation(key, ctx)
        except Exception as _ex:
            dyn = f"(explanation unavailable: {type(_ex).__name__})"
        full = ht + ("\n\n" + dyn if dyn else "")

        help_f = tk.Frame(win, bg=PANEL,
                          highlightbackground=BORDER, highlightthickness=1,
                          width=310)
        help_f.pack(side="right", fill="y", padx=(4,4), pady=4)
        help_f.pack_propagate(False)

        tk.Label(help_f, text=f"  📖 {hi.get('title', title)}",
                 bg=PANEL2, fg=TXT, font=MFB,
                 highlightbackground=BORDER,
                 highlightthickness=1).pack(fill="x")

        txt_f = tk.Frame(help_f, bg=PANEL)
        txt_f.pack(fill="both", expand=True, padx=4, pady=4)
        vsb = tk.Scrollbar(txt_f)
        vsb.pack(side="right", fill="y")
        tw = tk.Text(txt_f, wrap="word", bg=CARD2, fg=TXT,
                     font=("Segoe UI",9), relief="flat",
                     padx=10, pady=8, highlightthickness=0,
                     cursor="xterm", yscrollcommand=vsb.set,
                     state="disabled")
        tw.pack(fill="both", expand=True)
        vsb.configure(command=tw.yview)
        self._set_cmt(tw, full)
        tk.Button(help_f, text="📋 Copy",
                  command=lambda: (self.clipboard_clear(),
                                   self.clipboard_append(full)),
                  bg=BTN, fg=TXT, font=LF,
                  relief="flat", padx=8, pady=3,
                  cursor="hand2").pack(fill="x", padx=6, pady=4)

        # ── 좌측: matplotlib canvas ───────────────
        chart_f = tk.Frame(win, bg=PANEL,
                           highlightbackground=BORDER, highlightthickness=1)
        chart_f.pack(side="left", fill="both", expand=True,
                     padx=(4,0), pady=4)

        tb_holder = tk.Frame(chart_f, bg=PANEL2)
        tb_holder.pack(side="bottom", fill="x")

        # 창 크기로 Figure 인치 계산 (100dpi 기준)
        fig_w_px = ww - 310 - 24
        fig_h_px = wh - 80
        dpi = 100
        pop_fig = plt.Figure(
            figsize=(max(4, fig_w_px/dpi),
                     max(3, fig_h_px/dpi)),
            dpi=dpi, facecolor=PANEL)

        pop_cv = FigureCanvasTkAgg(pop_fig, master=chart_f)
        pop_cv.get_tk_widget().pack(side="top", fill="both",
                                    expand=True, padx=2, pady=2)
        NavigationToolbar2Tk(pop_cv, tb_holder)

        # 창을 화면에 먼저 표시
        win.update()
        win.update_idletasks()

        # ★ 핵심 수정: canvas widget의 실제 크기로 Figure resize
        #   FigureCanvasTkAgg는 초기 figsize로 만들어지지만,
        #   pack(fill='both', expand=True) 이후 widget 크기가 바뀌어도
        #   Figure는 자동으로 따라가지 않음. 수동으로 맞춰야 함.
        cw_actual = pop_cv.get_tk_widget().winfo_width()
        ch_actual = pop_cv.get_tk_widget().winfo_height()
        if cw_actual > 50 and ch_actual > 50:
            pop_fig.set_size_inches(cw_actual / dpi, ch_actual / dpi,
                                    forward=False)

        # ── 차트 렌더링 ───────────────────────────
        live = getattr(self, "_last_eval_ctx", None)
        pop_fig.clear()
        pop_fig.patch.set_facecolor(PANEL)

        try:
            if live is None:
                ax = pop_fig.add_subplot(111)
                ax.set_facecolor(PANEL2)
                ax.text(0.5, 0.5,
                        "Run [▶ Run Evaluation] first.",
                        transform=ax.transAxes,
                        ha="center", va="center",
                        color=SUB, fontsize=14)
                ax.axis("off")
            else:
                t_  = live["target"]
                tp_ = live["top"]
                ed_ = live["est_day"]
                sc_ = live["scores"]
                if   key == "radar":
                    self._draw_radar_fig(pop_fig, t_, tp_, large=True)
                elif key == "timeline":
                    self._draw_timeline_fig(pop_fig, t_, ed_,
                                            sc_, large=True)
                elif key == "pseudo_reg":
                    self._draw_pseudo_reg_fig(pop_fig, t_, large=True)
                elif key == "pseudo_spec":
                    self._draw_pseudo_spec_fig(pop_fig, t_, large=True)
        except Exception as ex:
            pop_fig.clear()
            ax = pop_fig.add_subplot(111)
            ax.set_facecolor(PANEL2)
            ax.text(0.5, 0.5, f"Render error:\n{type(ex).__name__}: {ex}",
                    transform=ax.transAxes, ha="center", va="center",
                    color=RED, fontsize=10)
            ax.axis("off")

        try:
            pop_fig.tight_layout(pad=0.8)
        except Exception:
            pass
        pop_cv.draw()
        pop_cv.flush_events()

        # ★ 추가: 윈도우 리사이즈 시 Figure도 따라가도록 바인딩
        def _on_resize(event, cv=pop_cv, fg=pop_fig, w=pop_cv.get_tk_widget()):
            try:
                nw = w.winfo_width()
                nh = w.winfo_height()
                if nw > 50 and nh > 50:
                    fg.set_size_inches(nw / dpi, nh / dpi, forward=False)
                    try:
                        fg.tight_layout(pad=0.8)
                    except Exception:
                        pass
                    cv.draw()
                    cv.flush_events()
            except Exception:
                pass
        win.bind("<Configure>", _on_resize)


    def _update_pred_candidates(self, target: dict, top: list):
        """날짜 추정 후보 이미지 3개 패널 갱신 — 썸네일 + 지표 + 유사도 코멘트"""
        def _safe(v):
            try:
                f = float(v)
                return 0.0 if (f != f) else f  # NaN check
            except Exception:
                return 0.0

        t_b  = _safe(target.get("lab",{}).get("b",  float("nan")))
        t_s  = _safe(target.get("s_mean",           float("nan")))
        t_yi = _safe(target.get("yellowness_idx",   float("nan")))

        for ci, cand_f in enumerate(self._pred_cand_frames):
            if ci >= len(top):
                cand_f["info_var"].set("(none)")
                self._set_cmt(cand_f["cmt_txt"], "")
                cand_f["metric_var"].set("")
                th_cv = cand_f["th_cv"]
                th_cv.delete("all")
                w = th_cv.winfo_width() or 120
                h = th_cv.winfo_height() or 90
                th_cv.create_text(w//2, h//2, text="—",
                                  fill=BORDER, font=MFB)
                continue

            dist, img = top[ci]

            # ── 헤더 정보 ───────────────────────────
            cand_f["info_var"].set(
                f"{img['cond']}\nDay {img['day']}  dist={dist:.3f}")

            # ── 썸네일 렌더링 ──────────────────────
            th_cv = cand_f["th_cv"]
            th_cv.delete("all")
            th = img.get("thumb")   # PIL Image
            if th is not None:
                try:
                    tw = th_cv.winfo_width()  or 120
                    th_h = th_cv.winfo_height() or 90
                    # PIL Image → 리사이즈 → PhotoImage
                    pil_copy = th.copy()
                    pil_copy.thumbnail((tw, th_h), Image.LANCZOS)
                    tk_th = ImageTk.PhotoImage(pil_copy)
                    ref_key = f"pred_cand_th_{ci}"
                    self._refs[ref_key] = tk_th   # GC 방지
                    tw2 = th_cv.winfo_width() or 120
                    th_h2 = th_cv.winfo_height() or 90
                    th_cv.create_image(tw2//2, th_h2//2,
                                       anchor="center",
                                       image=tk_th)
                except Exception as e:
                    th_cv.create_text(60, 45,
                                      text=f"thumb\nerr:\n{e}",
                                      fill=SUB, font=("Segoe UI",6),
                                      justify="center")
            else:
                w = th_cv.winfo_width() or 120
                h = th_cv.winfo_height() or 90
                th_cv.create_text(w//2, h//2,
                                  text="No\nthumb",
                                  fill=BORDER, font=LF,
                                  justify="center")

            # ── 지표 요약 ──────────────────────────
            r_b  = _safe(img.get("lab",{}).get("b",  float("nan")))
            r_s  = _safe(img.get("s_mean",           float("nan")))
            r_yi = _safe(img.get("yellowness_idx",   float("nan")))
            r_de = _safe(img.get("delta_e",          float("nan")))
            cand_f["metric_var"].set(
                f"b*={r_b:.1f}  S={r_s:.0f}  YI={r_yi:.0f}  ΔE={r_de:.1f}")

            # ── 유사도 코멘트 ──────────────────────
            def _diff_str(tv, rv, label, thresh_eq=3, higher_good=True):
                diff = tv - rv
                if abs(diff) < thresh_eq:
                    return f"{label}: ≈identical ({tv:.1f}≈{rv:.1f})"
                direction = (("higher→more pristine" if diff>0 else "lower→more oxidized")
                             if higher_good
                             else ("higher→more oxidized" if diff>0 else "lower→less change"))
                return f"{label}: target={tv:.1f} ref={rv:.1f} → {direction}"

            diffs_abs = {
                "b*":   abs(t_b  - r_b),
                "S-ch": abs(t_s  - r_s) / 255 * 60,
                "YI":   abs(t_yi - r_yi) / 120 * 60,
            }
            most_sim = min(diffs_abs, key=diffs_abs.get)

            lines = [
                f"[Rank #{ci+1}  dist={dist:.3f}]",
                _diff_str(t_b,  r_b,  "b*",   thresh_eq=3),
                _diff_str(t_s,  r_s,  "S-ch", thresh_eq=10),
                _diff_str(t_yi, r_yi, "YI",   thresh_eq=8),
                f"→ {most_sim} is most similar",
            ]
            self._set_cmt(cand_f["cmt_txt"], "\n".join(lines))

    def _generate_comment(self, target, top, est_day,
                           confidence, cond_input, pool):
        """날짜 추정 분석 근거 — 구체적 수치 포함 상세 설명"""
        def _safe(v):
            try:
                f = float(v)
                return 0.0 if f != f else f
            except Exception:
                return 0.0

        t_b  = _safe(target["lab"]["b"])
        t_s  = _safe(target["s_mean"])
        t_yi = _safe(target["yellowness_idx"])
        t_de = _safe(target.get("delta_e", float("nan")))

        if not top:
            return "Insufficient reference data for analysis."

        best_dist, best_img = top[0]
        lines = []

        # ── 추정 결과 요약 ────────────────────────────────
        if est_day is not None:
            conf_desc = ("High" if confidence >= 70
                         else "Medium" if confidence >= 40 else "Low")
            lines.append(
                f"━━ Date Estimation Result ━━━━━━━━━━━━━━━━━━━━━━━━\n"
                f"Estimated Day: {est_day:.1f}d  "
                f"(Confidence: {confidence:.0f}%  [{conf_desc}])\n"
                f"Formula: Σ(dayᵢ × 1/distᵢ) / Σ(1/distᵢ)\n"
            )
            # 가중 평균 분해
            wb  = self.cfg_w_b.get()
            ws  = self.cfg_w_s.get()
            wyi = self.cfg_w_yi.get()
            wt  = wb + ws + wyi
            if wt > 0: wb, ws, wyi = wb/wt, ws/wt, wyi/wt
            lines.append(
                f"Weights: b*×{wb:.2f} + S×{ws:.2f} + YI×{wyi:.2f}"
            )
            day_weights = []
            for dist, img in top[:3]:
                try:
                    d = float(img["day"])
                    day_weights.append((d, 1/(dist+1e-6), dist))
                except Exception:
                    pass
            if day_weights:
                total_w = sum(w for _,w,_ in day_weights)
                parts   = " + ".join(
                    f"({d:.0f}d × {w:.1f})" for d,w,_ in day_weights)
                lines.append(
                    f"= ({parts}) / {total_w:.1f}\n"
                    f"= {est_day:.2f} days\n"
                )
        else:
            lines.append("Cannot estimate: no reference images found.\n")

        # ── 타겟 이미지 지표 ──────────────────────────────
        def _grade(val, good, warn, higher_good=True):
            if higher_good:
                return ("✅ pristine" if val >= good
                        else "⚠️ boundary" if val >= warn
                        else "🔴 oxidized")
            else:
                return ("✅ minimal" if val <= warn
                        else "⚠️ moderate" if val <= good
                        else "🔴 large")

        lines.append(
            f"━━ Target Image Metrics ━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            f"  Lab b*  = {t_b:6.2f}  "
            f"{_grade(t_b, self.cfg_b_good.get(), self.cfg_b_warn.get())}\n"
            f"            (pristine≈50–60 / boundary≈20–40 / oxidized<15)\n"
            f"  S-ch    = {t_s:6.1f}  "
            f"{_grade(t_s, self.cfg_s_good.get(), self.cfg_s_warn.get())}\n"
            f"            (pristine≈150–200 / boundary≈40–80 / oxidized<20)\n"
            f"  YI      = {t_yi:6.1f}  "
            f"{_grade(t_yi, self.cfg_yi_good.get(), self.cfg_yi_warn.get())}\n"
            f"            (pristine≈50–110 / boundary≈35–50 / oxidized<35)\n"
            f"  ΔE      = {t_de:6.2f}  "
            f"{_grade(t_de, 3, 10, higher_good=False)}\n"
            f"            (ΔE<3 minor / 3–10 clear / >10 major change)\n"
        )

        # ── Top-3 비교 ────────────────────────────────────
        lines.append("━━ Top-3 Reference Comparison ━━━━━━━━━━━━━━━━━━━")
        for rank, (dist, img) in enumerate(top[:3], 1):
            r_b  = _safe(img["lab"]["b"])
            r_s  = _safe(img["s_mean"])
            r_yi = _safe(img.get("yellowness_idx", float("nan")))

            db   = t_b  - r_b
            ds   = t_s  - r_s
            dyi  = t_yi - r_yi

            def _arrow(diff, thresh=3):
                if abs(diff) < thresh: return "≈same"
                return f"{'↑' if diff>0 else '↓'}{abs(diff):.1f}"

            lines.append(
                f"\n  #{rank} {img['cond']}  Day {img['day']}  "
                f"dist={dist:.4f}\n"
                f"    b*: target={t_b:.1f} ref={r_b:.1f} diff={_arrow(db)}\n"
                f"    S:  target={t_s:.0f} ref={r_s:.0f} diff={_arrow(ds,5)}\n"
                f"    YI: target={t_yi:.0f} ref={r_yi:.0f} diff={_arrow(dyi,5)}"
            )

        # ── 신뢰도 해석 ───────────────────────────────────
        lines.append(f"\n\n━━ Confidence Interpretation ━━━━━━━━━━━━━━━━━━━━")
        lines.append(
            f"  Confidence = max(0, 100 − dist_min × 200)\n"
            f"             = max(0, 100 − {best_dist:.4f} × 200) = {confidence:.0f}%"
        )
        if len(top) >= 3:
            days_top3 = []
            for _, img in top[:3]:
                try: days_top3.append(float(img["day"]))
                except Exception: pass
            if days_top3:
                spread = max(days_top3) - min(days_top3)
                if spread <= 3:
                    note = f"Narrow spread ({spread:.0f}d) → reliable estimate"
                elif spread <= 7:
                    note = f"Moderate spread ({spread:.0f}d) → use with caution"
                else:
                    note = (f"Wide spread ({spread:.0f}d) → low confidence\n"
                            "  → Collect more reference data across days")
                lines.append(
                    f"  Top-3 day range: {min(days_top3):.0f}–{max(days_top3):.0f}d  {note}"
                )

        if cond_input:
            matched = list(set(img["cond"] for _,img in top))
            if all(c==cond_input for c in matched):
                lines.append(f"\n  Filtered by condition: '{cond_input}' (exact match)")
            else:
                lines.append(
                    f"\n  Condition '{cond_input}' not exact; "
                    f"compared with: {', '.join(matched[:3])}")
        else:
            lines.append(f"\n  No condition filter — compared against all {len(pool)} references")

        return "\n".join(lines)


    def _build_chart_explanation(self, key: str, ctx) -> str:
        """실제 추정 수치를 포함한 동적 설명 생성"""
        if ctx is None:
            return ""
        target  = ctx.get("target", {})
        top     = ctx.get("top", [])
        est_day = ctx.get("est_day")
        conf    = ctx.get("confidence", 0)
        pool    = ctx.get("pool", [])
        pairs   = ctx.get("pairs", [])
        r2_map  = ctx.get("r2_map", {})
        coef_map= ctx.get("coef_map", {})

        def _s(v):
            try:
                f = float(v)
                return 0.0 if f != f else f
            except Exception:
                return 0.0

        t_b  = _s(target.get("lab",{}).get("b", float("nan")))
        t_s  = _s(target.get("s_mean", float("nan")))
        t_yi = _s(target.get("yellowness_idx", float("nan")))
        t_de = _s(target.get("delta_e", float("nan")))

        lines = []

        if key in ("radar", "timeline"):
            lines += [
                "━━ Estimation Basis (numeric) ━━━━━━━━━━━━━━━━━━",
                f"Target image metrics:",
                f"  Lab b*  = {t_b:.2f}  (typical pristine ≈50–60, oxidized ≈3–15)",
                f"  S-ch    = {t_s:.1f}  (pristine ≈150–200, oxidized <20)",
                f"  YI      = {t_yi:.1f}  (pristine ≈50–110, oxidized ≈20–35)",
                f"  ΔE      = {t_de:.2f}  (vs Day-0 of best-match condition)",
                "",
            ]
            if top:
                wb  = self.cfg_w_b.get()
                ws  = self.cfg_w_s.get()
                wyi = self.cfg_w_yi.get()
                wt  = wb + ws + wyi
                if wt > 0: wb, ws, wyi = wb/wt, ws/wt, wyi/wt
                lines.append(f"Weights used: b*={wb:.2f}, S={ws:.2f}, YI={wyi:.2f}")
                lines.append("")
                lines.append("Top-3 distances:")
                day_weights = []
                for rank, (dist, img) in enumerate(top[:3], 1):
                    r_b  = _s(img["lab"]["b"])
                    r_s  = _s(img["s_mean"])
                    r_yi = _s(img.get("yellowness_idx", float("nan")))
                    try:
                        d_day = float(img["day"])
                    except Exception:
                        d_day = None
                    w = 1/(dist+1e-6) if d_day is not None else 0
                    if d_day is not None:
                        day_weights.append((d_day, w))
                    lines.append(
                        f"  #{rank}: {img['cond']} D{img['day']}  "
                        f"dist={dist:.4f}  "
                        f"(b*={r_b:.1f}, S={r_s:.0f}, YI={r_yi:.0f})"
                    )
                if day_weights and est_day is not None:
                    total_w = sum(w for _,w in day_weights)
                    lines.append("")
                    lines.append("Weighted day calculation:")
                    parts = " + ".join(
                        f"({d:.0f}d×{w:.1f})" for d,w in day_weights)
                    lines.append(f"  = ({parts}) / {total_w:.1f}")
                    lines.append(f"  = {est_day:.2f} days")
                    lines.append(f"  Confidence = max(0, 100 - {top[0][0]:.4f}×200)")
                    lines.append(f"             = {conf:.0f}%")

        elif key == "pseudo_reg":
            lines += [
                "━━ Regression Basis (numeric) ━━━━━━━━━━━━━━━━━━",
                f"Target values used as X-axis intercepts:",
                f"  b* = {t_b:.2f},  S = {t_s:.1f},  YI = {t_yi:.1f},  ΔE = {t_de:.2f}",
                f"  Number of matched pairs: {len(pairs)}",
                "",
            ]
            wb_map = {"b":0.45,"s":0.25,"yi":0.20,"de":0.10}
            t_map  = {"b":t_b,"s":t_s,"yi":t_yi,"de":t_de}
            estimates = []
            wt_total  = 0.0
            for key2, wt in wb_map.items():
                if key2 not in coef_map or key2 not in r2_map:
                    continue
                pred = float(np.polyval(coef_map[key2], t_map[key2]))
                r2   = r2_map[key2]
                ew   = wt * r2
                estimates.append((key2, pred, r2, wt, ew))
                wt_total += ew
                a, b = coef_map[key2]
                lines.append(
                    f"  {key2:4s}: coef=({a:.4f}, {b:.4f})  "
                    f"pred={pred:.4f}  R²={r2:.3f}  eff_wt={ew:.4f}")
            if estimates and wt_total > 0:
                est = sum(p*w for _,p,_,_,w in estimates) / wt_total
                lines.append("")
                lines.append(f"Ensemble estimate = Σ(pred×eff_wt) / Σ(eff_wt)")
                lines.append(f"  = {' + '.join(f'{p:.4f}×{w:.4f}' for _,p,_,_,w in estimates)}")
                lines.append(f"  / {wt_total:.4f} = {est:.4f}")

        elif key == "pseudo_spec":
            est_peak = ctx.get("est_peak")
            ci_lo    = ctx.get("ci_lo")
            ci_hi    = ctx.get("ci_hi")
            est_se   = ctx.get("est_se")
            dists    = ctx.get("spec_dists", [])
            if est_peak is not None:
                lines += [
                    "━━ Spectrum Estimation Basis ━━━━━━━━━━━━━━━━━━",
                    f"Estimated A₁g peak:  {est_peak:.4f}  (normalized)",
                    f"Uncertainty (σ):      {est_se:.4f}",
                    f"95% CI:  [{ci_lo:.4f}, {ci_hi:.4f}]",
                    "",
                    "Interpolation sources:",
                ]
                for i, p in enumerate(dists[:2]):
                    lines.append(
                        f"  Ref {i+1}: {p['cond']} D{p['day']}  "
                        f"norm_peak={p.get('norm_peak',0):.4f}"
                    )
                if len(dists) >= 2:
                    p1 = dists[0].get("norm_peak",1) or 1
                    p2 = dists[1].get("norm_peak",0.5) or 0.5
                    if abs(p1-p2) > 1e-6:
                        alpha = (est_peak-p2)/(p1-p2)
                        alpha = max(0.0, min(1.0, alpha))
                        lines.append("")
                        lines.append(f"Interpolation weight α:")
                        lines.append(f"  α = ({est_peak:.4f} - {p2:.4f})")
                        lines.append(f"      / ({p1:.4f} - {p2:.4f}) = {alpha:.3f}")
                        lines.append(f"  V_est(x) = {alpha:.3f}×Spec1 + {1-alpha:.3f}×Spec2")

        return "\n".join(lines)

    # ─────────────────────────────────────────
    #  PNG 캐시 기반 차트 표시 시스템
    # ─────────────────────────────────────────
    def _render_chart_to_label(self, draw_fn, lbl: tk.Label,
                                cache_key: str, large_dpi=180):
        """
        draw_fn(fig) 으로 고해상도 Figure를 PNG 버퍼에 저장(_chart_png_cache)
        → lbl이 배치된 후(after) 크기에 맞게 축소 표시.
        """
        import io
        if not hasattr(self, "_chart_png_cache"):
            self._chart_png_cache = {}
        if not hasattr(self, "_chart_photo_refs"):
            self._chart_photo_refs = {}

        try:
            fig = plt.Figure(facecolor=PANEL, dpi=large_dpi)
            draw_fn(fig)
            buf = io.BytesIO()
            fig.savefig(buf, format="png", dpi=large_dpi,
                        facecolor=PANEL, bbox_inches="tight")
            plt.close(fig)
            buf.seek(0)
            self._chart_png_cache[cache_key] = buf.getvalue()
        except Exception as e:
            self._chart_png_cache[cache_key] = None
            lbl.configure(image="", text=f"Error: {e}",
                          font=("Segoe UI", 7))
            return

        # 레이아웃 확정 후 표시 (즉시 + 500ms 후 재시도)
        self.after(10,  lambda: self._update_chart_label(lbl, cache_key))
        self.after(500, lambda: self._update_chart_label(lbl, cache_key))

    def _update_chart_label(self, lbl: tk.Label, cache_key: str):
        """PNG 캐시 → lbl 실제 크기에 맞게 축소 후 PhotoImage 표시"""
        import io
        cache = getattr(self, "_chart_png_cache", {})
        png_data = cache.get(cache_key)
        if not png_data:
            return
        try:
            # 위젯이 파괴됐으면 skip
            if not lbl.winfo_exists():
                return
            w = lbl.winfo_width()
            h = lbl.winfo_height()
            # 아직 배치 안 됨 → 부모에서 추정
            if w < 10:
                w = lbl.winfo_reqwidth() or 280
            if h < 10:
                h = lbl.winfo_reqheight() or 180

            pil_img = Image.open(io.BytesIO(png_data)).convert("RGB")
            pil_img.thumbnail((max(10,w), max(10,h)), Image.LANCZOS)

            photo = ImageTk.PhotoImage(pil_img)
            if not hasattr(self, "_chart_photo_refs"):
                self._chart_photo_refs = {}
            self._chart_photo_refs[cache_key] = photo
            lbl.configure(image=photo, text="", compound="center")
        except Exception as e:
            try:
                lbl.configure(image="", text=f"Display error: {e}",
                              font=("Segoe UI", 7))
            except Exception:
                pass

    def _draw_pred_charts(self, target, top, est_day, scores, pool):
        """메인 Evaluation 차트 렌더링 + 컨텍스트 저장"""
        ctx: dict = {
            "target": target, "top": top,
            "est_day": est_day, "scores": scores,
            "pool": pool,
            "confidence": max(0, 100 - top[0][0]*200) if top else 0,
            "pairs": [], "r2_map": {}, "coef_map": {},
            "est_peak": None, "ci_lo": None,
            "ci_hi": None, "est_se": None, "spec_dists": [],
        }
        self._last_eval_ctx = ctx

        fig_r = self._pred_figs["radar"]["fig"]
        fig_r.clear()
        fig_r.patch.set_facecolor(PANEL)
        self._draw_radar_fig(fig_r, target, top, large=False)
        self._pred_figs["radar"]["cv"].draw()

        fig_t = self._pred_figs["timeline"]["fig"]
        fig_t.clear()
        fig_t.patch.set_facecolor(PANEL)
        self._draw_timeline_fig(fig_t, target, est_day, scores, large=False)
        self._pred_figs["timeline"]["cv"].draw()

    def _draw_radar_fig(self, fig, target, top, large: bool):
        fs_t = 11 if large else 7
        fs_k = 9  if large else 5
        lw   = 2.5 if large else 1.5

        metrics = ["b*","S-ch","YI","YR%"]
        n = len(metrics)
        angles = [i/n*2*3.14159+3.14159/n for i in range(n)]
        angles += angles[:1]

        def _safe(v):
            try:
                f = float(v)
                return 0.0 if f!=f else f
            except Exception: return 0.0

        ax = fig.add_subplot(111, polar=True)
        ax.set_facecolor(PANEL2)
        fig.patch.set_facecolor(PANEL)

        # ── 모든 평가대상 target 을 동시 표시 ──
        targets_with_result = [
            t for t in getattr(self, "_pred_targets", [])
            if t.get("result")]

        if targets_with_result:
            for ti, t in enumerate(targets_with_result):
                tm = t["result"].get("target_metrics", {})
                v = [_safe(tm.get("lab",{}).get("b",0))/80,
                     _safe(tm.get("s_mean",0))/200,
                     _safe(tm.get("yellowness_idx",0))/120,
                     _safe(tm.get("yellow_ratio",0))*3]
                v += v[:1]
                col = t.get("color", RED)
                ax.plot(angles, v, "o-", color=col, lw=lw, alpha=0.9)
                ax.fill(angles, v, alpha=0.10, color=col)
        else:
            # 호환용 — 단일 target 표시
            target_v = [_safe(target.get("lab",{}).get("b",0))/80,
                        _safe(target.get("s_mean",0))/200,
                        _safe(target.get("yellowness_idx",0))/120,
                        _safe(target.get("yellow_ratio",0))*3]
            target_v += target_v[:1]
            ax.plot(angles, target_v, "o-", color=RED, lw=lw)
            ax.fill(angles, target_v, alpha=0.15, color=RED)

        # Top-3 후보 (선택된 target 의 top)
        for ci2, (dist, img) in enumerate(top[:3]):
            iv = [_safe(img["lab"]["b"])/80,
                  _safe(img["s_mean"])/200,
                  _safe(img.get("yellowness_idx",0))/120,
                  _safe(img.get("yellow_ratio",0))*3]
            iv += iv[:1]
            col = [COND_COLORS[0], COND_COLORS[1], COND_COLORS[2]][ci2]
            ax.plot(angles, iv, "--", color=col,
                    lw=max(0.8,lw*0.7), alpha=0.6)
            ax.fill(angles, iv, alpha=0.04, color=col)

        ax.set_xticks(angles[:-1])
        ax.set_xticklabels(metrics, fontsize=fs_k, color=SUB)
        ax.tick_params(colors=SUB, labelsize=max(4,fs_k-1))
        ax.set_title("Metric Similarity", color=TXT, fontsize=fs_t,
                     pad=8 if large else 4)
        # 차트 내 범례 생략 (카드 색상 표시가 범례)
        fig.tight_layout(pad=0.3)

    def _draw_timeline_fig(self, fig, target, est_day, scores, large: bool):
        fs_t = 11 if large else 7
        fs_a = 9  if large else 6
        fs_k = 8  if large else 5
        lw   = 2.0 if large else 1.2
        ms   = 8   if large else 4

        ax = fig.add_subplot(111)
        ax.set_facecolor(PANEL2)
        fig.patch.set_facecolor(PANEL)
        for sp in ax.spines.values(): sp.set_color(BORDER)
        ax.tick_params(colors=SUB, labelsize=fs_k)

        day_dist = []
        for dist, img in scores:
            try:
                d = float(img["day"])
                day_dist.append((d, dist, img["cond"]))
            except Exception:
                pass

        if day_dist:
            days  = [x[0] for x in day_dist]
            dists = [x[1] for x in day_dist]
            conds = [x[2] for x in day_dist]
            cond_set = list(dict.fromkeys(conds))
            for ci2, cond in enumerate(cond_set):
                pts = [(d,v) for d,v,c in zip(days,dists,conds) if c==cond]
                pts.sort(key=lambda x: x[0])
                xs = [p[0] for p in pts]
                ys = [p[1] for p in pts]
                col = COND_COLORS[ci2 % len(COND_COLORS)]
                ax.plot(xs, ys, "o-", color=col, lw=lw, ms=ms,
                        alpha=0.8)
                for x,y in zip(xs, ys):
                    ax.annotate(f"D{x:.0f}",
                                xy=(x,y), xytext=(0,5),
                                textcoords="offset points",
                                fontsize=max(4,fs_k-2),
                                color=col, ha="center")

        # ── 모든 평가대상의 est_day 를 ★ 마커로 표시 ──
        targets_with_result = [
            t for t in getattr(self, "_pred_targets", [])
            if t.get("result") and t["result"].get("est_day") is not None]
        # y 위치: 화면 하단 5% 지점
        if targets_with_result:
            ymin, ymax = ax.get_ylim() if day_dist else (0, 1)
            y_marker = ymin + (ymax-ymin)*0.05
            for ti, t in enumerate(targets_with_result):
                ed_t = t["result"]["est_day"]
                col = t.get("color", RED)
                # tid 인덱스 — 카드 표시 순서대로 T1, T2 ...
                idx = next((j for j, x in
                            enumerate(self._pred_targets,start=1)
                            if x is t), ti+1)
                ax.plot(ed_t, y_marker, "*",
                        color=col, ms=ms*2, zorder=10,
                        markeredgecolor="white", markeredgewidth=0.8)
                ax.annotate(f"T{idx}",
                            xy=(ed_t, y_marker),
                            xytext=(0, ms*1.5+2),
                            textcoords="offset points",
                            fontsize=max(5,fs_k-1),
                            color=col, ha="center",
                            fontweight="bold")
        elif est_day is not None:
            # 호환용
            ax.axvline(est_day, color=RED, lw=lw, ls="--")

        ax.set_xlabel("Day", color=SUB, fontsize=fs_a)
        ax.set_ylabel("Distance (lower=more similar)",
                      color=SUB, fontsize=fs_a)
        ax.set_title("Distance Timeline", color=TXT, fontsize=fs_t)
        # 차트 내 범례 생략 (카드 색상 표시가 범례)
        fig.tight_layout(pad=0.5)


    # ─────────────────────────────────────────────────────────
    #  🧪 Advanced Analysis 탭
    # ─────────────────────────────────────────────────────────
    def _build_advanced_tab(self):
        f = self._tfs["advanced"]
        self._last_adv_ctx = None        # 마지막 분석 결과 캐시

        # ── 헤더 바 ──────────────────────────────────────────
        hdr = tk.Frame(f, bg=PANEL2, highlightbackground=BORDER,
                       highlightthickness=1)
        hdr.pack(fill="x", padx=6, pady=(6, 0))

        tk.Label(hdr, text="🧪 Advanced Analysis — Histogram · FFT · Ensemble",
                 bg=PANEL2, fg=TXT, font=MFB).pack(side="left", padx=12, pady=6)

        self._adv_run_btn = tk.Button(
            hdr, text="▶ Run Advanced Estimation",
            command=self._adv_run,
            bg=ACCENT, fg="white", font=MFB,
            relief="flat", padx=14, pady=4, cursor="hand2")
        self._adv_run_btn.pack(side="right", padx=8, pady=6)

        # ── 메인 영역 (좌: 차트 3개 / 우: 결과 패널) ────────
        body = tk.Frame(f, bg=BG)
        body.pack(fill="both", expand=True, padx=6, pady=6)

        # 좌측: 차트 3개 (세로 배치)
        charts_f = tk.Frame(body, bg=BG)
        charts_f.pack(side="left", fill="both", expand=True)

        # 차트1: Histogram Distribution
        hist_lbl = tk.Frame(charts_f, bg=PANEL2,
                            highlightbackground=BORDER, highlightthickness=1)
        hist_lbl.pack(fill="x", pady=(0, 3))
        tk.Label(hist_lbl,
                 text="📊 b* Histogram — Target vs Reference (Wasserstein EMD)",
                 bg=PANEL2, fg=TXT, font=MF).pack(side="left", padx=10, pady=4)

        self._adv_hist_frame = tk.Frame(charts_f, bg=CARD,
                                        highlightbackground=BORDER,
                                        highlightthickness=1, height=220)
        self._adv_hist_frame.pack(fill="x", pady=(0, 6))
        self._adv_hist_frame.pack_propagate(False)
        self._adv_hist_canvas = None

        # 차트2: FFT Power Spectrum
        fft_lbl = tk.Frame(charts_f, bg=PANEL2,
                           highlightbackground=BORDER, highlightthickness=1)
        fft_lbl.pack(fill="x", pady=(0, 3))
        tk.Label(fft_lbl,
                 text="📡 FFT Analysis — Frequency Domain & High-Freq Energy",
                 bg=PANEL2, fg=TXT, font=MF).pack(side="left", padx=10, pady=4)

        self._adv_fft_frame = tk.Frame(charts_f, bg=CARD,
                                       highlightbackground=BORDER,
                                       highlightthickness=1, height=220)
        self._adv_fft_frame.pack(fill="x", pady=(0, 6))
        self._adv_fft_frame.pack_propagate(False)
        self._adv_fft_canvas = None

        # 차트3: Radial Spectrum + HF Ratio 비교
        rad_lbl = tk.Frame(charts_f, bg=PANEL2,
                           highlightbackground=BORDER, highlightthickness=1)
        rad_lbl.pack(fill="x", pady=(0, 3))
        tk.Label(rad_lbl,
                 text="🔁 Radial Spectrum — Reference Comparison",
                 bg=PANEL2, fg=TXT, font=MF).pack(side="left", padx=10, pady=4)

        self._adv_rad_frame = tk.Frame(charts_f, bg=CARD,
                                       highlightbackground=BORDER,
                                       highlightthickness=1, height=200)
        self._adv_rad_frame.pack(fill="x")
        self._adv_rad_frame.pack_propagate(False)
        self._adv_rad_canvas = None

        # ── 더블클릭 팝업 바인딩 ─────────────────────────────
        for frame, key in [(self._adv_hist_frame, "hist"),
                           (self._adv_fft_frame,  "fft"),
                           (self._adv_rad_frame,  "radial")]:
            frame.bind("<Double-Button-1>",
                       lambda e, k=key: self._adv_chart_popup(k))

        # 우측: 결과 패널
        res_f = tk.Frame(body, bg=PANEL,
                         highlightbackground=BORDER, highlightthickness=1,
                         width=280)
        res_f.pack(side="right", fill="y", padx=(6, 0))
        res_f.pack_propagate(False)

        tk.Label(res_f, text="  📋 Method Comparison",
                 bg=PANEL2, fg=TXT, font=MFB,
                 highlightbackground=BORDER,
                 highlightthickness=1).pack(fill="x")

        # 각 방법의 추정치 표시
        self._adv_result_vars = {}
        methods = [
            ("knn",     "🔵 KNN (기존)",       "#4f8ef7"),
            ("wass",    "🟢 Wasserstein",       "#22c55e"),
            ("fft",     "🟡 FFT Texture",       "#fbbf24"),
            ("spatial", "🟠 Spatial Pattern",   "#f97316"),
            ("kinetic", "🟣 Kinetic Model",     "#a78bfa"),
            ("ens",     "🔴 Ensemble",          "#ef4444"),
        ]
        for key, label, color in methods:
            mf = tk.Frame(res_f, bg=PANEL)
            mf.pack(fill="x", padx=8, pady=4)
            tk.Label(mf, text=label, bg=PANEL, fg=color,
                     font=MFB, width=16, anchor="w").pack(side="left")
            var = tk.StringVar(value="—")
            tk.Label(mf, textvariable=var, bg=PANEL, fg=TXT,
                     font=MF, anchor="e").pack(side="right")
            self._adv_result_vars[key] = var

        tk.Frame(res_f, bg=BORDER, height=1).pack(fill="x", padx=8, pady=6)

        # 가중치 바
        tk.Label(res_f, text="  Ensemble Weights",
                 bg=PANEL, fg=SUB, font=LF).pack(anchor="w", padx=8)
        self._adv_weight_frame = tk.Frame(res_f, bg=PANEL)
        self._adv_weight_frame.pack(fill="x", padx=8, pady=4)

        tk.Frame(res_f, bg=BORDER, height=1).pack(fill="x", padx=8, pady=6)

        # 해석 텍스트
        tk.Label(res_f, text="  Interpretation",
                 bg=PANEL, fg=SUB, font=LF).pack(anchor="w", padx=8)
        interp_f = tk.Frame(res_f, bg=CARD)
        interp_f.pack(fill="both", expand=True, padx=8, pady=(0, 8))
        vsb = tk.Scrollbar(interp_f)
        vsb.pack(side="right", fill="y")
        self._adv_interp_txt = tk.Text(
            interp_f, wrap="word", bg=CARD, fg=TXT,
            font=("Segoe UI", 8), relief="flat",
            padx=8, pady=6, yscrollcommand=vsb.set,
            state="disabled", highlightthickness=0)
        self._adv_interp_txt.pack(fill="both", expand=True)
        vsb.configure(command=self._adv_interp_txt.yview)

        # 초기 안내 메시지
        self._adv_set_interp(
            "🧪 Advanced Analysis\n\n"
            "Evaluation 탭에서 시편을 로드하고\n"
            "▶ Run Advanced Estimation 을\n"
            "클릭한다.\n\n"
            "방법:\n"
            "① KNN — 기존 가중 유클리드 거리\n"
            "② Wasserstein — b* 히스토그램\n"
            "   Earth Mover Distance\n"
            "③ FFT — 주파수 도메인 텍스처\n"
            "④ Ensemble — ①②③ 신뢰도 가중\n"
            "   앙상블\n\n"
            "차트를 더블클릭하면 확대 팝업이\n"
            "표시된다."
        )

    def _adv_set_interp(self, txt: str):
        """해석 텍스트 박스 내용 업데이트"""
        self._adv_interp_txt.configure(state="normal")
        self._adv_interp_txt.delete("1.0", "end")
        self._adv_interp_txt.insert("end", txt)
        self._adv_interp_txt.configure(state="disabled")

    def _adv_on_switch(self):
        """탭 전환 시 호출 — 이전 결과 있으면 유지, 없으면 안내만"""
        pass   # 자동 재실행하지 않음 (Run 버튼 클릭 시만 실행)

    def _adv_run(self):
        """▶ Run Advanced Estimation 버튼 핸들러"""
        # ── Evaluation 탭 결과 재사용 ────────────────────────
        ctx = getattr(self, "_last_eval_ctx", None)
        if ctx is None or ctx.get("target") is None:
            self._adv_set_interp(
                "⚠ 먼저 Evaluation 탭에서\n"
                "▶ Run Evaluation을 실행한다.")
            return

        target   = ctx["target"]
        pool     = ctx.get("pool", [])
        if not pool:
            self._adv_set_interp("⚠ 참조 DB가 비어 있다.")
            return

        self._adv_run_btn.configure(state="disabled", text="⏳ 계산 중...")
        self.update_idletasks()

        try:
            self._adv_run_inner(target, pool, ctx)
        except Exception as ex:
            import traceback
            self._adv_set_interp(f"❌ 오류:\n{traceback.format_exc()}")
        finally:
            self._adv_run_btn.configure(state="normal",
                                         text="▶ Run Advanced Estimation")

    def _adv_run_inner(self, target: dict, pool: list, ctx: dict):
        """실제 연산 — 5가지 방법 + 앙상블"""
        # ── 0. RGB/mask 확인 ─────────────────────────────────
        t_rgb  = target.get("rgb")
        t_mask = target.get("mask")
        if t_rgb is None or t_mask is None:
            self._adv_set_interp("⚠ 타겟 이미지 RGB/Mask가 없다.")
            return

        t_roi = target.get("roi",
                           (0, 0, t_rgb.shape[1], t_rgb.shape[0]))
        rows = self.cfg_rows.get() if hasattr(self, "cfg_rows") else 3
        cols = self.cfg_cols.get() if hasattr(self, "cfg_cols") else 3

        # ── 1. 참조 DB 사전계산 (캐시) ───────────────────────
        adv_precompute_pool(pool, rows=rows, cols=cols)

        # ── 2. 타겟 피처 계산 ────────────────────────────────
        t_hist    = adv_hist_signature(t_rgb, t_mask)
        t_fft     = adv_fft_features(t_rgb, t_mask)
        t_spatial = adv_spatial_features(t_rgb, t_mask, t_roi, rows, cols)

        # ── 3. Kinetic 모델 피팅 (참조 DB 전체) ─────────────
        try:
            kinetic_params = adv_kinetic_fit(pool)
        except Exception:
            kinetic_params = {}

        # ── 4. 각 방법으로 추정 ──────────────────────────────
        knn_day  = ctx.get("est_day")
        knn_conf = ctx.get("confidence", 0)

        w_res = adv_wasserstein_estimate(t_hist, pool)
        wass_day, wass_conf = w_res["est_day"], w_res["confidence"]
        wass_scores = w_res["scores"]

        f_res = adv_fft_estimate(t_fft, pool)
        fft_day, fft_conf = f_res["est_day"], f_res["confidence"]
        fft_scores = f_res["scores"]

        s_res = adv_spatial_estimate(t_spatial, pool, rows, cols)
        spatial_day, spatial_conf = s_res["est_day"], s_res["confidence"]
        spatial_scores = s_res["scores"]

        target_cond = target.get("cond", "")
        t_b = target.get("lab", {}).get("b", np.nan)
        k_res = adv_kinetic_estimate(t_b, kinetic_params, target_cond)
        kinetic_day, kinetic_conf = k_res["est_day"], k_res["confidence"]

        # ── 5. 앙상블 ────────────────────────────────────────
        ens = adv_ensemble(knn_day,     knn_conf,
                           wass_day,    wass_conf,
                           fft_day,     fft_conf,
                           spatial_day, spatial_conf,
                           kinetic_day, kinetic_conf)
        ens_day  = ens["est_day"]
        ens_conf = ens["confidence"]
        weights  = ens["weights"]

        # ── 6. 결과 저장 ─────────────────────────────────────
        self._last_adv_ctx = {
            "target":         target,
            "pool":           pool,
            "t_hist":         t_hist,
            "t_fft":          t_fft,
            "t_spatial":      t_spatial,
            "kinetic_params": kinetic_params,
            "knn_day":        knn_day,     "knn_conf":     knn_conf,
            "wass_day":       wass_day,    "wass_conf":    wass_conf,
            "wass_scores":    wass_scores,
            "fft_day":        fft_day,     "fft_conf":     fft_conf,
            "fft_scores":     fft_scores,
            "spatial_day":    spatial_day, "spatial_conf": spatial_conf,
            "spatial_scores": spatial_scores,
            "kinetic_day":    kinetic_day, "kinetic_conf": kinetic_conf,
            "kinetic_detail": k_res,
            "ens_day":        ens_day,     "ens_conf":     ens_conf,
            "weights":        weights,
        }

        # ── 7. UI 업데이트 ───────────────────────────────────
        def fmt(day, conf):
            if day is None:
                return "추정 불가"
            return f"{day:.1f}일  (신뢰도 {conf:.0f}%)"

        self._adv_result_vars["knn"    ].set(fmt(knn_day,     knn_conf))
        self._adv_result_vars["wass"   ].set(fmt(wass_day,    wass_conf))
        self._adv_result_vars["fft"    ].set(fmt(fft_day,     fft_conf))
        self._adv_result_vars["spatial"].set(fmt(spatial_day, spatial_conf))
        self._adv_result_vars["kinetic"].set(fmt(kinetic_day, kinetic_conf))
        self._adv_result_vars["ens"    ].set(fmt(ens_day,     ens_conf))

        # 가중치 바 업데이트
        for w in self._adv_weight_frame.winfo_children():
            w.destroy()
        for name, color in [("knn",     "#4f8ef7"),
                             ("wass",    "#22c55e"),
                             ("fft",     "#fbbf24"),
                             ("spatial", "#f97316"),
                             ("kinetic", "#a78bfa")]:
            w = weights.get(name, 0)
            rf = tk.Frame(self._adv_weight_frame, bg=PANEL)
            rf.pack(fill="x", pady=1)
            tk.Label(rf, text=f"{name[:7]:7s}", bg=PANEL,
                     fg=color, font=LF, width=8).pack(side="left")
            bar_bg = tk.Frame(rf, bg=CARD2, height=8)
            bar_bg.pack(side="left", fill="x", expand=True)
            bar_bg.update_idletasks()
            bar_w = max(2, int(bar_bg.winfo_width() * w))
            tk.Frame(bar_bg, bg=color, height=8,
                     width=bar_w).place(x=0, y=0)
            tk.Label(rf, text=f"{w*100:.0f}%", bg=PANEL,
                     fg=SUB, font=LF, width=4).pack(side="right")

        # 해석 텍스트
        interp = self._adv_build_interp(
            knn_day, knn_conf, wass_day, wass_conf,
            fft_day, fft_conf, spatial_day, spatial_conf,
            kinetic_day, kinetic_conf, k_res,
            ens_day, ens_conf, weights, t_fft, t_spatial)
        self._adv_set_interp(interp)

        # ── 8. 차트 렌더링 ───────────────────────────────────
        self.after(50,  self._adv_draw_hist)
        self.after(150, self._adv_draw_fft)
        self.after(250, self._adv_draw_radial)

    def _adv_build_interp(self,
                           knn_day, knn_conf,
                           wass_day, wass_conf,
                           fft_day, fft_conf,
                           spatial_day, spatial_conf,
                           kinetic_day, kinetic_conf, k_res,
                           ens_day, ens_conf,
                           weights, t_fft, t_spatial) -> str:
        """결과 해석 텍스트 생성"""
        lines = ["=== 분석 결과 ===\n"]

        if ens_day is not None:
            lines.append(f"🔴 앙상블 추정: {ens_day:.1f}일\n"
                         f"   신뢰도: {ens_conf:.0f}%\n")

        # 방법 간 일치도
        days = [d for d in [knn_day, wass_day, fft_day,
                             spatial_day, kinetic_day] if d is not None]
        if len(days) >= 2:
            spread = max(days) - min(days)
            if spread < 2:
                lines.append(f"✅ {len(days)}가지 방법 잘 일치\n"
                              "   (편차 <2일)\n")
            elif spread < 5:
                lines.append(f"⚠ 방법 간 편차: {spread:.1f}일\n"
                              "   결과 해석 시 주의\n")
            else:
                lines.append(f"❌ 방법 간 큰 불일치\n"
                              f"   편차: {spread:.1f}일\n"
                              "   데이터 재확인 권장\n")

        # Kinetic 모델 해석
        lines.append("\n=== Kinetic 모델 ===\n")
        if kinetic_day is not None:
            p    = k_res.get("model_params", {})
            cond = k_res.get("cond_used", "")
            fail = k_res.get("fail_reason", "")
            lines.append(f"추정: {kinetic_day:.1f}일\n"
                         f"신뢰도: {kinetic_conf:.0f}%\n"
                         f"조건: {cond}\n"
                         f"k(감쇠율): {p.get('k', 0):.4f}\n"
                         f"b*₀: {p.get('b0', 0):.1f}\n"
                         f"b*∞: {p.get('b_inf', 0):.1f}\n"
                         f"R²: {p.get('r2', 0):.3f}\n"
                         f"피팅방법: {p.get('fit_method','exponential')}\n")
            if fail:
                lines.append(f"참고: {fail}\n")
            if p.get("r2", 0) > 0.9:
                lines.append("→ 모델 적합도 우수\n")
            elif p.get("r2", 0) > 0.7:
                lines.append("→ 모델 적합도 양호\n")
            else:
                lines.append("→ 선형 근사 사용 (데이터 부족)\n")
        else:
            fail = k_res.get("fail_reason", "참조 DB 데이터 부족")
            lines.append(f"추정 불가\n이유: {fail}\n"
                         "해결: 같은 조건의 참조 이미지를\n"
                         "3개 이상 로드하면 된다.\n")

        # Spatial 패턴 해석
        lines.append("\n=== Spatial 패턴 ===\n")
        ent = t_spatial.get("entropy", 0)
        bg  = t_spatial.get("boundary_grad", 0)
        ani = t_spatial.get("anisotropy", 1)
        lines.append(f"공간 엔트로피: {ent:.2f}\n")
        if ent > 15:
            lines.append("→ 불균일 산화\n  (부분적 진행)\n")
        elif ent > 8:
            lines.append("→ 중간 불균일\n")
        else:
            lines.append("→ 균일한 산화\n  (전면적 진행)\n")
        lines.append(f"경계-중심 기울기: {bg:.2f}\n")
        if bg < -3:
            lines.append("→ 중심부 산화 우세\n")
        elif bg > 3:
            lines.append("→ 경계부 산화 우세\n")
        else:
            lines.append("→ 균일 산화 분포\n")
        if spatial_day is not None:
            lines.append(f"공간패턴 추정: {spatial_day:.1f}일\n"
                         f"신뢰도: {spatial_conf:.0f}%\n")

        # FFT 해석
        hf  = t_fft.get("hf_ratio", 0)
        ent2 = t_fft.get("entropy", 0)
        lines.append(f"\n=== FFT 분석 ===\n"
                     f"고주파 에너지: {hf*100:.1f}%\n")
        if hf > 0.35:
            lines.append("→ 표면 거칠음 높음\n  산화 진행 ↑\n")
        elif hf > 0.20:
            lines.append("→ 표면 거칠음 중간\n")
        else:
            lines.append("→ 표면 균일 (초기)\n")

        lines.append("\n💡 차트 더블클릭 → 확대 팝업")
        return "".join(lines)

    # ── 차트 렌더링 ─────────────────────────────────────────

    def _adv_embed_canvas(self, fig, frame, attr: str):
        """figure를 frame에 embed + 기존 canvas 교체"""
        old = getattr(self, attr, None)
        if old is not None:
            try:
                old.get_tk_widget().destroy()
            except Exception:
                pass
        cv = FigureCanvasTkAgg(fig, master=frame)
        cv.get_tk_widget().pack(fill="both", expand=True)
        setattr(self, attr, cv)
        return cv

    def _adv_draw_hist(self):
        """b* Histogram 차트"""
        ctx = self._last_adv_ctx
        if ctx is None:
            return

        t_hist  = ctx["t_hist"]
        pool    = ctx["pool"]
        wass_s  = ctx.get("wass_scores", [])
        bins    = len(t_hist)
        x_edges = np.linspace(-30, 80, bins + 1)
        x_mid   = (x_edges[:-1] + x_edges[1:]) / 2

        fig = plt.Figure(figsize=(7, 2.0), dpi=96, facecolor=PANEL)
        ax  = fig.add_subplot(111)
        styled_ax(ax)

        # 타겟
        ax.bar(x_mid, t_hist, width=(x_edges[1]-x_edges[0])*0.85,
               color=RED, alpha=0.85, label="Target", zorder=3)

        # Top-3 참조 (Wasserstein 기준)
        colors = ["#4f8ef7", "#22c55e", "#fbbf24"]
        for i, (dist, img) in enumerate(wass_s[:3]):
            ref_hist = img.get("_adv_hist")
            if ref_hist is None:
                continue
            day_lbl = img.get("day", "?")
            cond    = img.get("cond", "")
            ax.step(x_mid, ref_hist, where="mid",
                    color=colors[i], alpha=0.75, linewidth=1.4,
                    label=f"Ref [{day_lbl}d/{cond}] W={dist:.3f}")

        ax.set_xlabel("b* (CIE Lab)", color=SUB, fontsize=7)
        ax.set_ylabel("Probability", color=SUB, fontsize=7)
        ax.set_title("b* Histogram Distribution (Wasserstein EMD)",
                     color=TXT, fontsize=8)
        ax.legend(fontsize=6, framealpha=0.3, edgecolor=BORDER,
                  loc="upper left")
        ax.tick_params(labelsize=6)

        # Wasserstein day 추정 표시
        wass_day = ctx.get("wass_day")
        if wass_day is not None:
            ax.axvline(x=0, color=SUB, lw=0.5, linestyle=":")
            ax.text(0.98, 0.95,
                    f"EMD est: {wass_day:.1f}d",
                    transform=ax.transAxes,
                    ha="right", va="top",
                    color="#22c55e", fontsize=7,
                    bbox=dict(boxstyle="round,pad=0.3",
                              facecolor=PANEL2, alpha=0.8))

        fig.tight_layout(pad=0.4)
        cv = self._adv_embed_canvas(fig, self._adv_hist_frame,
                                     "_adv_hist_canvas")
        cv.draw()
        cv.flush_events()

    def _adv_draw_fft(self):
        """FFT 파워스펙트럼 차트 (1×2)"""
        ctx = self._last_adv_ctx
        if ctx is None:
            return

        t_fft = ctx["t_fft"]
        power_map = t_fft["power_map"]
        hf_ratio  = t_fft["hf_ratio"]

        fig = plt.Figure(figsize=(7, 2.0), dpi=96, facecolor=PANEL)
        fig.subplots_adjust(left=0.08, right=0.97,
                             top=0.88, bottom=0.18, wspace=0.35)

        # 좌: 2D 로그 파워스펙트럼
        ax1 = fig.add_subplot(1, 2, 1)
        styled_ax(ax1)
        H, W = power_map.shape
        ax1.imshow(power_map, cmap="inferno", aspect="auto",
                   extent=[-W//2, W//2, -H//2, H//2])
        # 40% 반경 원 (고주파 경계)
        r40 = min(H, W) / 2 * 0.40
        theta = np.linspace(0, 2*np.pi, 200)
        ax1.plot(r40*np.cos(theta), r40*np.sin(theta),
                 color="#22c55e", lw=1.0, linestyle="--",
                 alpha=0.8, label="HF boundary")
        ax1.set_title("2D FFT (log power)", color=TXT, fontsize=7, pad=3)
        ax1.tick_params(labelsize=5)
        ax1.legend(fontsize=5, loc="upper right", framealpha=0.3,
                   edgecolor=BORDER)

        # 우: 참조 HF ratio 비교 막대
        ax2 = fig.add_subplot(1, 2, 2)
        styled_ax(ax2)

        pool = ctx["pool"]
        fft_scores = ctx.get("fft_scores", [])
        top5 = fft_scores[:5] if fft_scores else []

        ref_labels = []
        ref_hf     = []
        for dist, img in top5:
            ff = img.get("_adv_fft")
            if ff:
                ref_labels.append(f"{img.get('day','?')}d")
                ref_hf.append(ff["hf_ratio"] * 100)

        # 타겟 추가
        all_labels = ["Target"] + ref_labels
        all_hf     = [hf_ratio * 100] + ref_hf
        bar_colors = [RED] + ["#4f8ef7"]*len(ref_labels)

        bars = ax2.bar(range(len(all_labels)), all_hf,
                       color=bar_colors, alpha=0.85, width=0.6)
        ax2.set_xticks(range(len(all_labels)))
        ax2.set_xticklabels(all_labels, fontsize=5, rotation=30)
        ax2.set_ylabel("HF Energy %", color=SUB, fontsize=6)
        ax2.set_title("High-Freq Ratio vs Refs", color=TXT,
                      fontsize=7, pad=3)
        ax2.tick_params(labelsize=5)
        ax2.axhline(y=hf_ratio*100, color=RED, lw=0.8,
                    linestyle="--", alpha=0.6)

        # 값 레이블
        for bar, val in zip(bars, all_hf):
            ax2.text(bar.get_x() + bar.get_width()/2, val + 0.3,
                     f"{val:.1f}", ha="center", va="bottom",
                     color=TXT, fontsize=4)

        # FFT 추정치 표시
        fft_day = ctx.get("fft_day")
        if fft_day is not None:
            ax2.text(0.98, 0.95,
                     f"FFT est: {fft_day:.1f}d",
                     transform=ax2.transAxes,
                     ha="right", va="top",
                     color="#fbbf24", fontsize=7,
                     bbox=dict(boxstyle="round,pad=0.3",
                               facecolor=PANEL2, alpha=0.8))

        cv = self._adv_embed_canvas(fig, self._adv_fft_frame,
                                     "_adv_fft_canvas")
        cv.draw()
        cv.flush_events()

    def _adv_draw_radial(self):
        """방사형 스펙트럼 + 앙상블 요약 차트 (1×2)"""
        ctx = self._last_adv_ctx
        if ctx is None:
            return

        t_fft = ctx["t_fft"]
        t_rad = t_fft["radial_mean"]
        bins  = len(t_rad)
        x     = np.arange(bins)

        fig = plt.Figure(figsize=(7, 1.9), dpi=96, facecolor=PANEL)
        fig.subplots_adjust(left=0.08, right=0.97,
                             top=0.87, bottom=0.20, wspace=0.38)

        # 좌: 방사형 스펙트럼 비교
        ax1 = fig.add_subplot(1, 2, 1)
        styled_ax(ax1)

        ax1.fill_between(x, t_rad, alpha=0.5, color=RED, label="Target")
        ax1.plot(x, t_rad, color=RED, lw=1.2)

        fft_scores = ctx.get("fft_scores", [])
        cols = ["#4f8ef7", "#22c55e", "#fbbf24"]
        for i, (dist, img) in enumerate(fft_scores[:3]):
            ff = img.get("_adv_fft")
            if ff is None:
                continue
            rad = ff["radial_mean"]
            day_lbl = img.get("day", "?")
            ax1.plot(x, rad, color=cols[i], lw=0.9, alpha=0.75,
                     label=f"[{day_lbl}d]")

        ax1.set_xlabel("Spatial Frequency (bins)", color=SUB, fontsize=6)
        ax1.set_ylabel("Norm. Power", color=SUB, fontsize=6)
        ax1.set_title("Radial Spectrum (Target vs Refs)",
                      color=TXT, fontsize=7, pad=3)
        ax1.legend(fontsize=5, framealpha=0.3, edgecolor=BORDER)
        ax1.tick_params(labelsize=5)

        # 우: 앙상블 요약 (4개 추정치 + 오차 막대)
        ax2 = fig.add_subplot(1, 2, 2)
        styled_ax(ax2)

        knn_d  = ctx.get("knn_day")
        wass_d = ctx.get("wass_day")
        fft_d  = ctx.get("fft_day")
        ens_d  = ctx.get("ens_day")

        method_days  = [(d, lab, col) for d, lab, col in [
            (knn_d,  "KNN",   "#4f8ef7"),
            (wass_d, "W.EMD", "#22c55e"),
            (fft_d,  "FFT",   "#fbbf24"),
            (ens_d,  "Ens",   RED),
        ] if d is not None]

        if method_days:
            days = [d for d, _, _ in method_days]
            labs = [l for _, l, _ in method_days]
            cols2 = [c for _, _, c in method_days]
            yp = range(len(days))
            ax2.barh(list(yp), days, color=cols2, alpha=0.85, height=0.5)
            ax2.set_yticks(list(yp))
            ax2.set_yticklabels(labs, fontsize=6)
            ax2.set_xlabel("Estimated Day", color=SUB, fontsize=6)
            ax2.set_title("Method Comparison", color=TXT, fontsize=7, pad=3)
            ax2.tick_params(labelsize=5)
            # 값 레이블
            for y, d in zip(yp, days):
                ax2.text(d + 0.1, y, f"{d:.1f}d",
                         va="center", color=TXT, fontsize=6)

        cv = self._adv_embed_canvas(fig, self._adv_rad_frame,
                                     "_adv_rad_canvas")
        cv.draw()
        cv.flush_events()

    # ── 팝업 ────────────────────────────────────────────────

    def _adv_chart_popup(self, key: str):
        """차트 더블클릭 → 확대 팝업"""
        ctx = self._last_adv_ctx
        if ctx is None:
            return

        titles = {
            "hist":   "📊 b* Histogram — Wasserstein EMD",
            "fft":    "📡 FFT Analysis — Frequency Domain",
            "radial": "🔁 Radial Spectrum & Method Comparison",
        }
        win = tk.Toplevel(self)
        win.title(f"🔍 {titles.get(key, key)}")
        win.configure(bg=PANEL)
        sw, sh = self.winfo_screenwidth(), self.winfo_screenheight()
        ww, wh = int(sw * 0.80), int(sh * 0.80)
        win.geometry(f"{ww}x{wh}+{(sw-ww)//2}+{(sh-wh)//2}")
        win.update_idletasks()

        hdr = tk.Frame(win, bg=PANEL2)
        hdr.pack(side="top", fill="x")
        tk.Label(hdr, text=f"  🔍 {titles.get(key, key)}",
                 bg=PANEL2, fg=TXT, font=MFB).pack(side="left", pady=7, padx=10)
        tk.Button(hdr, text="✕ Close", command=win.destroy,
                  bg=BTN, fg=TXT, font=MF, relief="flat",
                  padx=10, pady=4, cursor="hand2").pack(side="right", padx=8, pady=6)

        chart_f = tk.Frame(win, bg=PANEL)
        chart_f.pack(fill="both", expand=True, padx=8, pady=8)
        tb_holder = tk.Frame(chart_f, bg=PANEL2)
        tb_holder.pack(side="bottom", fill="x")

        dpi = 100
        fw = (ww - 20) / dpi
        fh = (wh - 80) / dpi
        pop_fig = plt.Figure(figsize=(max(6,fw), max(4,fh)),
                              dpi=dpi, facecolor=PANEL)

        pop_cv = FigureCanvasTkAgg(pop_fig, master=chart_f)
        pop_cv.get_tk_widget().pack(fill="both", expand=True)
        NavigationToolbar2Tk(pop_cv, tb_holder)
        win.update()
        win.update_idletasks()

        # ★ canvas widget 실제 크기로 Figure 동기화
        cw_actual = pop_cv.get_tk_widget().winfo_width()
        ch_actual = pop_cv.get_tk_widget().winfo_height()
        if cw_actual > 50 and ch_actual > 50:
            pop_fig.set_size_inches(cw_actual / dpi, ch_actual / dpi,
                                    forward=False)

        try:
            if key == "hist":
                self._adv_draw_hist_fig(pop_fig, ctx, large=True)
            elif key == "fft":
                self._adv_draw_fft_fig(pop_fig, ctx, large=True)
            elif key == "radial":
                self._adv_draw_radial_fig(pop_fig, ctx, large=True)
        except Exception as ex:
            ax = pop_fig.add_subplot(111)
            ax.set_facecolor(PANEL2)
            ax.text(0.5, 0.5, f"Render error:\n{ex}",
                    transform=ax.transAxes, ha="center", va="center",
                    color=RED, fontsize=10)
            ax.axis("off")

        try:
            pop_fig.tight_layout(pad=0.8)
        except Exception:
            pass
        pop_cv.draw()
        pop_cv.flush_events()

        # 윈도우 리사이즈 대응
        def _on_resize(event, cv=pop_cv, fg=pop_fig, w=pop_cv.get_tk_widget()):
            try:
                nw = w.winfo_width()
                nh = w.winfo_height()
                if nw > 50 and nh > 50:
                    fg.set_size_inches(nw / dpi, nh / dpi, forward=False)
                    try:
                        fg.tight_layout(pad=0.8)
                    except Exception:
                        pass
                    cv.draw()
                    cv.flush_events()
            except Exception:
                pass
        win.bind("<Configure>", _on_resize)

    # ── 팝업용 고해상도 figure 렌더 함수 ─────────────────────

    def _adv_draw_hist_fig(self, fig, ctx, large=False):
        t_hist = ctx["t_hist"]
        wass_s = ctx.get("wass_scores", [])
        bins   = len(t_hist)
        x_e    = np.linspace(-30, 80, bins + 1)
        x_mid  = (x_e[:-1] + x_e[1:]) / 2

        ax = fig.add_subplot(111)
        styled_ax(ax)

        ax.bar(x_mid, t_hist, width=(x_e[1]-x_e[0])*0.85,
               color=RED, alpha=0.85, label="Target", zorder=3)

        colors = ["#4f8ef7", "#22c55e", "#fbbf24"]
        for i, (dist, img) in enumerate(wass_s[:3]):
            ref_h = img.get("_adv_hist")
            if ref_h is None:
                continue
            d_lbl = img.get("day", "?")
            cond  = img.get("cond", "")
            ax.step(x_mid, ref_h, where="mid", color=colors[i],
                    alpha=0.80, linewidth=1.8,
                    label=f"Ref [{d_lbl}d/{cond}]  W={dist:.4f}")

        fs = 11 if large else 8
        ax.set_xlabel("b* (CIE Lab)", color=SUB, fontsize=fs)
        ax.set_ylabel("Probability", color=SUB, fontsize=fs)
        ax.set_title("b* Histogram Distribution — Wasserstein (EMD) Comparison",
                     color=TXT, fontsize=fs+1)
        ax.legend(fontsize=fs-2, framealpha=0.35, edgecolor=BORDER)
        ax.tick_params(labelsize=fs-2)

        wass_day  = ctx.get("wass_day")
        wass_conf = ctx.get("wass_conf", 0)
        if wass_day is not None:
            ax.text(0.98, 0.95,
                    f"Wasserstein estimate: {wass_day:.1f}d  "
                    f"(conf {wass_conf:.0f}%)",
                    transform=ax.transAxes, ha="right", va="top",
                    color="#22c55e", fontsize=fs-1,
                    bbox=dict(boxstyle="round,pad=0.4",
                              facecolor=PANEL2, alpha=0.85))

    def _adv_draw_fft_fig(self, fig, ctx, large=False):
        t_fft    = ctx["t_fft"]
        power    = t_fft["power_map"]
        hf_ratio = t_fft["hf_ratio"]
        fs = 11 if large else 8

        fig.subplots_adjust(left=0.07, right=0.97,
                             top=0.92, bottom=0.10, wspace=0.30)
        ax1 = fig.add_subplot(1, 2, 1)
        styled_ax(ax1)
        H, W = power.shape
        ax1.imshow(power, cmap="inferno", aspect="equal",
                   extent=[-W//2, W//2, -H//2, H//2])
        r40 = min(H, W) / 2 * 0.40
        theta = np.linspace(0, 2*np.pi, 300)
        ax1.plot(r40*np.cos(theta), r40*np.sin(theta),
                 color="#22c55e", lw=1.5, linestyle="--",
                 alpha=0.9, label=f"HF boundary (40%)")
        ax1.set_title("2D Log Power Spectrum", color=TXT, fontsize=fs+1)
        ax1.set_xlabel("Freq X", color=SUB, fontsize=fs-1)
        ax1.set_ylabel("Freq Y", color=SUB, fontsize=fs-1)
        ax1.tick_params(labelsize=fs-2)
        ax1.legend(fontsize=fs-2, framealpha=0.3, edgecolor=BORDER)

        ax2 = fig.add_subplot(1, 2, 2)
        styled_ax(ax2)

        fft_scores = ctx.get("fft_scores", [])
        all_labels = ["Target"]
        all_hf     = [hf_ratio * 100]
        all_ent    = [t_fft["entropy"]]
        colors_b   = [RED]
        for dist, img in fft_scores[:6]:
            ff = img.get("_adv_fft")
            if ff:
                all_labels.append(f"{img.get('day','?')}d/{img.get('cond','')[:4]}")
                all_hf.append(ff["hf_ratio"] * 100)
                all_ent.append(ff["entropy"])
                colors_b.append("#4f8ef7")

        xp = np.arange(len(all_labels))
        ax2.bar(xp - 0.2, all_hf, width=0.35,
                color=colors_b, alpha=0.85, label="HF%")
        ax2_b = ax2.twinx()
        ax2_b.set_facecolor("none")
        ax2_b.plot(xp, all_ent, "o--", color="#a78bfa",
                   lw=1.2, ms=4, label="Entropy")
        ax2_b.tick_params(colors=SUB, labelsize=fs-3)
        ax2_b.set_ylabel("Entropy", color="#a78bfa", fontsize=fs-1)

        ax2.set_xticks(xp)
        ax2.set_xticklabels(all_labels, fontsize=fs-3, rotation=35, ha="right")
        ax2.set_ylabel("HF Energy %", color=SUB, fontsize=fs-1)
        ax2.set_title("HF Ratio & Entropy (Target vs Refs)",
                      color=TXT, fontsize=fs+1)
        ax2.tick_params(labelsize=fs-2)

        lines1, lbs1 = ax2.get_legend_handles_labels()
        lines2, lbs2 = ax2_b.get_legend_handles_labels()
        ax2.legend(lines1+lines2, lbs1+lbs2,
                   fontsize=fs-2, framealpha=0.3, edgecolor=BORDER)

        fft_day  = ctx.get("fft_day")
        fft_conf = ctx.get("fft_conf", 0)
        if fft_day is not None:
            ax2.text(0.98, 0.95,
                     f"FFT estimate: {fft_day:.1f}d  "
                     f"(conf {fft_conf:.0f}%)",
                     transform=ax2.transAxes, ha="right", va="top",
                     color="#fbbf24", fontsize=fs-1,
                     bbox=dict(boxstyle="round,pad=0.4",
                               facecolor=PANEL2, alpha=0.85))

    def _adv_draw_radial_fig(self, fig, ctx, large=False):
        t_rad = ctx["t_fft"]["radial_mean"]
        fs    = 11 if large else 8
        bins  = len(t_rad)
        x     = np.arange(bins)

        fig.subplots_adjust(left=0.08, right=0.97,
                             top=0.92, bottom=0.12, wspace=0.35)
        ax1 = fig.add_subplot(1, 2, 1)
        styled_ax(ax1)

        ax1.fill_between(x, t_rad, alpha=0.45, color=RED)
        ax1.plot(x, t_rad, color=RED, lw=1.8, label="Target")

        fft_scores = ctx.get("fft_scores", [])
        cols = ["#4f8ef7", "#22c55e", "#fbbf24"]
        for i, (dist, img) in enumerate(fft_scores[:3]):
            ff = img.get("_adv_fft")
            if ff is None:
                continue
            ax1.plot(x, ff["radial_mean"], color=cols[i],
                     lw=1.2, alpha=0.80,
                     label=f"[{img.get('day','?')}d] d={dist:.3f}")

        ax1.set_xlabel("Spatial Frequency (bin)", color=SUB, fontsize=fs-1)
        ax1.set_ylabel("Normalised Power", color=SUB, fontsize=fs-1)
        ax1.set_title("Radial Spectrum Comparison", color=TXT, fontsize=fs+1)
        ax1.legend(fontsize=fs-2, framealpha=0.3, edgecolor=BORDER)
        ax1.tick_params(labelsize=fs-2)

        ax2 = fig.add_subplot(1, 2, 2)
        styled_ax(ax2)

        knn_d  = ctx.get("knn_day")
        wass_d = ctx.get("wass_day")
        fft_d  = ctx.get("fft_day")
        ens_d  = ctx.get("ens_day")

        knn_c  = ctx.get("knn_conf",  0)
        wass_c = ctx.get("wass_conf", 0)
        fft_c  = ctx.get("fft_conf",  0)
        ens_c  = ctx.get("ens_conf",  0)

        entries = [(d, c, lab, col) for d, c, lab, col in [
            (knn_d,  knn_c,  "KNN",         "#4f8ef7"),
            (wass_d, wass_c, "Wasserstein",  "#22c55e"),
            (fft_d,  fft_c,  "FFT",          "#fbbf24"),
            (ens_d,  ens_c,  "Ensemble",     RED),
        ] if d is not None]

        if entries:
            days  = [d for d,_,_,_ in entries]
            confs = [c for _,c,_,_ in entries]
            labs  = [l for _,_,l,_ in entries]
            cols2 = [c for _,_,_,c in entries]
            yp    = range(len(days))
            bars  = ax2.barh(list(yp), days, color=cols2,
                              alpha=0.85, height=0.5)
            ax2.set_yticks(list(yp))
            ax2.set_yticklabels(labs, fontsize=fs-1)
            ax2.set_xlabel("Estimated Day", color=SUB, fontsize=fs-1)
            ax2.set_title("Method Comparison (Day Estimate)",
                          color=TXT, fontsize=fs+1)
            ax2.tick_params(labelsize=fs-2)

            for y, d, c in zip(yp, days, confs):
                ax2.text(d + 0.15, y, f"{d:.1f}d  ({c:.0f}%)",
                         va="center", color=TXT, fontsize=fs-2)

    def _build_settings_tab(self):
        f = self._tfs["settings"]

        # 스크롤 가능 영역
        outer = tk.Frame(f, bg=BG)
        outer.pack(fill="both", expand=True)
        canvas = tk.Canvas(outer, bg=BG, highlightthickness=0)
        vsb = tk.Scrollbar(outer, orient="vertical",
                           command=canvas.yview)
        canvas.configure(yscrollcommand=vsb.set)
        vsb.pack(side="right", fill="y")
        canvas.pack(fill="both", expand=True)
        inner = tk.Frame(canvas, bg=BG)
        _cw = canvas.create_window((0,0), window=inner, anchor="nw")
        inner.bind("<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.bind("<Configure>",
            lambda e: canvas.itemconfig(_cw, width=e.width))

        # ── 헤더 ──────────────────────────────────
        tk.Label(inner,
                 text=_L("  ⚙  임계값 설정","  ⚙  Threshold Settings"),
                 bg=BG, fg=ACCENT,
                 font=("Segoe UI",12,"bold")).pack(
                 anchor="w", padx=12, pady=(12,4))
        tk.Label(inner,
                 text=_L("  설정값은 [적용] 버튼을 누르면 즉시 분석에 반영된다."
                         "  분석을 다시 실행하면 새 값으로 계산된다.",
                         "  Changes take effect on [Apply]."
                         "  Re-run analysis to recalculate."),
                 bg=BG, fg=SUB, font=LF).pack(anchor="w", padx=16, pady=(0,10))

        # 카테고리별 카드 생성 헬퍼
        def section(title, desc):
            card = tk.Frame(inner, bg=PANEL,
                            highlightbackground=BORDER, highlightthickness=1)
            card.pack(fill="x", padx=12, pady=6)
            hf = tk.Frame(card, bg=PANEL2,
                          highlightbackground=BORDER, highlightthickness=1)
            hf.pack(fill="x")
            tk.Label(hf, text=f"  {title}",
                     bg=PANEL2, fg=TXT,
                     font=("Segoe UI",9,"bold")).pack(
                     side="left", padx=10, pady=7)
            tk.Label(hf, text=desc,
                     bg=PANEL2, fg=SUB, font=LF).pack(
                     side="left", padx=4)
            body = tk.Frame(card, bg=PANEL)
            body.pack(fill="x", padx=10, pady=8)
            return body

        # 슬라이더 행 헬퍼
        def row(parent, label, var, lo, hi,
                unit="", digits=2, desc="", width=260):
            rf = tk.Frame(parent, bg=PANEL)
            rf.pack(fill="x", pady=4)

            # 라벨
            tk.Label(rf, text=label, bg=PANEL, fg=TXT,
                     font=MF, width=22, anchor="w").pack(side="left")

            # 슬라이더
            is_int = isinstance(var, tk.IntVar)
            res = 1 if is_int else 0.01
            sl = tk.Scale(rf, variable=var,
                          from_=lo, to=hi,
                          resolution=res,
                          orient="horizontal",
                          length=width,
                          bg=PANEL, fg=TXT,
                          troughcolor=PANEL2,
                          highlightthickness=0,
                          sliderlength=16,
                          font=LF,
                          showvalue=False)
            sl.pack(side="left", padx=6)

            # 현재값 표시
            val_lbl = tk.Label(rf, bg=PANEL, fg=ACCENT,
                               font=MFB, width=7, anchor="e")
            val_lbl.pack(side="left")

            def _update(*_):
                v = var.get()
                val_lbl.configure(
                    text=f"{v:.{digits}f}{unit}" if not is_int
                    else f"{int(v)}{unit}")
            var.trace_add("write", _update)
            _update()

            # 설명
            if desc:
                tk.Label(rf, text=desc, bg=PANEL, fg=SUB,
                         font=("Segoe UI",7),
                         wraplength=300,
                         justify="left").pack(side="left", padx=8)

        # ══ 1. 황색 판정 (HSI) ════════════════════
        b1 = section(
            _L("🟡  황색 판정 기준  (HSI)","🟡  Yellow Classification (HSI)"),
            _L("ROI 내 픽셀을 '황색'으로 분류하는 조건","Criteria to classify pixels as yellow in ROI"))

        row(b1, _L("H 하한 각도 (h_lo)","H lower angle (h_lo)"),
            self.cfg_h_lo, 0, 90, unit="°", digits=1,
            desc="이 각도 이상인 픽셀만 황색으로 인식\n"
                 "기본 35° — 낮추면 더 많은 황색 계열 포함")
        row(b1, _L("H 상한 각도 (h_hi)","H upper angle (h_hi)"),
            self.cfg_h_hi, 30, 120, unit="°", digits=1,
            desc="이 각도 이하인 픽셀만 황색으로 인식\n"
                 "기본 75° — 높이면 연두색 계열까지 포함")
        row(b1, _L("S 최소 채도 (s_thresh)","S min saturation (s_thresh)"),
            self.cfg_s_thresh, 0.01, 0.50, unit="", digits=2,
            desc="Below this = achromatic (excludes white/beige). Default 0.10.")

        tk.Frame(b1, bg=BORDER, height=1).pack(fill="x", pady=(4,0))
        tk.Label(b1,
                 text=_L("  💡 미산화(pristine) HfS₂: H≈45~55°, S≈0.5  |  "
                         "Oxidized HfO₂: white S≈0.02~0.04",
                         "  💡 Pristine HfS₂: H≈45~55°, S≈0.5  |  "
                         "Oxidized HfO₂: white S≈0.02~0.04"),
                 bg=PANEL, fg=TEAL, font=("Segoe UI",7)).pack(
                 anchor="w", pady=(4,0))

        # ══ 2. S채널 색상 판정 ════════════════════
        b2 = section(
            _L("🔵  S채널 판정 색상 기준","🔵  S-Channel Color Thresholds"),
            _L("S채널 평균값에 따라 카드·비교 그리드의 색상이 바뀜","Color coding in cards and grid based on S-ch mean"))

        row(b2, _L("양호(녹색) 기준 ≥","Good(green) threshold ≥"),
            self.cfg_s_good, 20, 200, unit="", digits=0,
            desc="이 값 이상이면 녹색 표시\n기본 80")
        row(b2, _L("경고(주황) 기준 ≥","Warn(orange) threshold ≥"),
            self.cfg_s_warn, 5, 150, unit="", digits=0,
            desc="이 값 이상이면 주황, 미만이면 빨강\n기본 40")

        # ══ 3. YI 색상 판정 ══════════════════════
        b3 = section(
            _L("🟤  Yellowness Index 판정 기준","🟤  Yellowness Index Thresholds"),
            _L("YI(ASTM E313) 값에 따른 색상 표시 기준","Color thresholds based on YI (ASTM E313)"))

        row(b3, _L("양호(녹색) 기준 ≥","Good(green) threshold ≥"),
            self.cfg_yi_good, 10, 150, unit="", digits=0,
            desc="미산화 시편: YI ≈ 50~110\n기본 60")
        row(b3, _L("경고(주황) 기준 ≥","Warn(orange) threshold ≥"),
            self.cfg_yi_warn, 5, 100, unit="", digits=0,
            desc="산화 시편: YI ≈ 20~35\n기본 35")

        # ══ 4. Lab b* 판정 ═══════════════════════
        b4 = section(
            _L("🔬  Lab b* 판정 기준","🔬  Lab b* Thresholds"),
            "컬러 분석 탭의 b* 색상 표시 기준")

        row(b4, _L("양호(녹색) 기준 ≥","Good(green) threshold ≥"),
            self.cfg_b_good, 5, 80, unit="", digits=0,
            desc="미산화 시편: b* ≈ 50~60\n기본 40")
        row(b4, _L("경고(주황) 기준 ≥","Warn(orange) threshold ≥"),
            self.cfg_b_warn, 1, 50, unit="", digits=0,
            desc="산화 시편: b* ≈ 4~15\n기본 20")

        # ══ 5. 날짜 추정 가중치 ══════════════════
        b5 = section(
            _L("🔎  날짜 추정 가중치","🔎  Date Estimation Weights"),
            _L("세 지표의 합이 1.0이 되도록 설정 — 합계가 맞지 않으면 자동 정규화","Sum of 3 weights = 1.0; auto-normalized on apply"))

        row(b5, _L("Lab b* 가중치","Lab b* weight"),
            self.cfg_w_b, 0.0, 1.0, unit="", digits=2,
            desc="황색도 지표 (기본 0.45)\n"
                 "Raman과 가장 높은 상관관계")
        row(b5, _L("S채널 가중치","S-ch weight"),
            self.cfg_w_s, 0.0, 1.0, unit="", digits=2,
            desc=_L("HSI 채도 (기본 0.30)","HSI Saturation (default 0.30)"))
        row(b5, _L("YI 가중치","YI weight"),
            self.cfg_w_yi, 0.0, 1.0, unit="", digits=2,
            desc=_L("ASTM E313 황색도 지수 (기본 0.25)","ASTM E313 YI (default 0.25)"))

        # 가중치 합계 표시
        wsum_var = tk.StringVar()
        def _update_wsum(*_):
            total = (self.cfg_w_b.get() +
                     self.cfg_w_s.get() +
                     self.cfg_w_yi.get())
            ok = abs(total-1.0) < 0.01
            wsum_var.set(f'{_L("합계","Sum")}: {total:.2f}'
                         + (" ✔" if ok else _L("  ← 적용 시 자동 정규화","  ← auto-normalized on apply")))
            wsum_lbl.configure(fg=GREEN if ok else AMBER)
        for v in (self.cfg_w_b, self.cfg_w_s, self.cfg_w_yi):
            v.trace_add("write", _update_wsum)
        wsum_lbl = tk.Label(b5, textvariable=wsum_var,
                            bg=PANEL, fg=GREEN, font=MFB)
        wsum_lbl.pack(anchor="w", pady=(2,0))
        _update_wsum()

        # ══ 6. 세그먼트 최소 픽셀 ════════════════
        b6 = section(
            _L("🔲  세그먼트 유효 픽셀","🔲  Min Segment Pixels"),
            _L("이 수 미만의 픽셀을 가진 세그먼트는 통계에서 제외","Segments below this pixel count are excluded from stats"))

        row(b6, _L("최소 유효 픽셀","Min valid pixels"),
            self.cfg_min_pix, 1, 200, unit="px", digits=0,
            desc="기본 10px — ROI가 작을 때 줄이고,\n"
                 "노이즈가 많을 때 늘리세요")

        # ── 버튼 영역 ──────────────────────────────
        btn_f = tk.Frame(inner, bg=BG)
        btn_f.pack(fill="x", padx=12, pady=12)

        tk.Button(btn_f, text=_L("✔  현재 설정 적용","✔  Apply Settings"),
                  command=self._apply_settings,
                  bg=ACCENT, fg="white",
                  font=("Segoe UI",10,"bold"),
                  relief="flat", padx=20, pady=8,
                  cursor="hand2").pack(side="left", padx=(0,8))

        tk.Button(btn_f, text=_L("↺  기본값으로 초기화","↺  Reset to Default"),
                  command=self._reset_settings,
                  bg=BTN, fg=TXT, font=MF,
                  relief="flat", padx=14, pady=8,
                  cursor="hand2").pack(side="left")

        self._settings_status = tk.Label(
            btn_f, text="",
            bg=BG, fg=GREEN, font=MFB)
        self._settings_status.pack(side="left", padx=14)

    def _apply_settings(self):
        """설정값을 검증하고 내부 변수에 반영"""
        # 가중치 자동 정규화
        total = self.cfg_w_b.get() + self.cfg_w_s.get() + self.cfg_w_yi.get()
        if total > 0 and abs(total-1.0) > 0.01:
            self.cfg_w_b.set(round(self.cfg_w_b.get()/total, 3))
            self.cfg_w_s.set(round(self.cfg_w_s.get()/total, 3))
            self.cfg_w_yi.set(round(self.cfg_w_yi.get()/total, 3))

        # h_lo < h_hi 보정
        if self.cfg_h_lo.get() >= self.cfg_h_hi.get():
            self.cfg_h_lo.set(self.cfg_h_hi.get() - 5)

        # 판정 기준 good > warn 보정
        for good, warn in [(self.cfg_s_good, self.cfg_s_warn),
                           (self.cfg_yi_good, self.cfg_yi_warn),
                           (self.cfg_b_good,  self.cfg_b_warn)]:
            if good.get() <= warn.get():
                warn.set(good.get() - 1)

        self._settings_status.configure(
            text=_L("✔ 설정 적용 완료  —  [▶ 전체 분석]을 다시 실행한다","✔ Applied. Re-run analysis."),
            fg=GREEN)
        self._set_status(_L("⚙ 설정 적용 완료","⚙ Settings applied"))

    def _reset_settings(self):
        """모든 설정값을 기본값으로 초기화"""
        self.cfg_h_lo.set(35.0)
        self.cfg_h_hi.set(75.0)
        self.cfg_s_thresh.set(0.10)
        self.cfg_s_good.set(80)
        self.cfg_s_warn.set(40)
        self.cfg_yi_good.set(60)
        self.cfg_yi_warn.set(35)
        self.cfg_b_good.set(40)
        self.cfg_b_warn.set(20)
        self.cfg_w_b.set(0.45)
        self.cfg_w_s.set(0.30)
        self.cfg_w_yi.set(0.25)
        self.cfg_min_pix.set(10)
        self._settings_status.configure(
            text=_L("↺ 기본값으로 초기화됨","↺ Reset to defaults"), fg=AMBER)
        self._set_status(_L("⚙ 설정 초기화","⚙ Settings reset"))

    def _s_color(self, v):
        """S채널 값 → 판정 색상"""
        if np.isnan(v): return SUB
        return GREEN if v >= self.cfg_s_good.get() \
          else AMBER if v >= self.cfg_s_warn.get() \
          else RED

    def _yi_color(self, v):
        """YI 값 → 판정 색상"""
        if np.isnan(v): return SUB
        return GREEN if v >= self.cfg_yi_good.get() \
          else AMBER if v >= self.cfg_yi_warn.get() \
          else RED

    def _b_color(self, v):
        """Lab b* 값 → 판정 색상"""
        if np.isnan(v): return SUB
        return GREEN if v >= self.cfg_b_good.get() \
          else AMBER if v >= self.cfg_b_warn.get() \
          else RED

    # ─────────────────────────────────────────
    #  7. Raman 비교 탭
    # ─────────────────────────────────────────
    def _build_raman_tab(self):
        f = self._tfs["raman"]

        # Raman 데이터 저장소: [{cond, day, peak, norm_peak}]
        self._raman_data: list[dict] = []

        # ── 좌: 데이터 입력 패널 ──────────────────────
        left = tk.Frame(f, bg=PANEL,
                        highlightbackground=BORDER, highlightthickness=1,
                        width=360)
        left.pack(side="left", fill="y", padx=(0,4))
        left.pack_propagate(False)

        # 헤더
        tk.Label(left, text=_L("  📡 Raman 데이터 입력","  📡 Raman Data Input"),
                 bg=PANEL2, fg=TXT, font=MFB,
                 highlightbackground=BORDER,
                 highlightthickness=1).pack(fill="x", pady=0)

        # Excel 로드 — DnD 드롭존 + 버튼
        ef = tk.Frame(left, bg=PANEL)
        ef.pack(fill="x", padx=8, pady=(8,4))

        # DnD 드롭존 (여러 파일 동시 드롭 지원)
        self._raman_drop_cv = tk.Canvas(
            ef, bg=CARD2,
            highlightbackground=TEAL if _DND else BORDER,
            highlightthickness=2,
            height=70, cursor="hand2")
        self._raman_drop_cv.pack(fill="x", pady=(0,4))

        self._raman_drop_lbl_id = self._raman_drop_cv.create_text(
            170, 35, anchor="center",
            text=_L("📂 Excel 파일을 여기에 드래그 (여러 파일 동시 가능)\n"
                    "Supported: .xlsx  .xls  .csv",
                    "📂 Drag Excel files here  (multiple at once)\n"
                    "Supported: .xlsx  .xls  .csv"),
            fill=SUB, font=("Segoe UI", 8),
            justify="center")

        def _reanchor(event=None):
            w = self._raman_drop_cv.winfo_width() or 340
            self._raman_drop_cv.coords(self._raman_drop_lbl_id, w//2, 35)
        self._raman_drop_cv.bind("<Configure>", _reanchor)
        # 클릭으로도 파일 선택 가능
        self._raman_drop_cv.bind(
            "<Button-1>",
            lambda e: self._load_raman_smart())

        if _DND:
            self._raman_drop_cv.drop_target_register(DND_FILES)
            self._raman_drop_cv.dnd_bind(
                "<<Drop>>",
                lambda e: self._on_raman_excel_drop(e))
            def _raman_enter(e):
                self._raman_drop_cv.configure(
                    highlightbackground=ACCENT, bg=PANEL2)
                self._raman_drop_cv.itemconfigure(
                    self._raman_drop_lbl_id, fill=TXT)
            def _raman_leave(e):
                self._raman_drop_cv.configure(
                    highlightbackground=TEAL, bg=CARD2)
                self._raman_drop_cv.itemconfigure(
                    self._raman_drop_lbl_id, fill=SUB)
            self._raman_drop_cv.dnd_bind("<<DragEnter>>", _raman_enter)
            self._raman_drop_cv.dnd_bind("<<DragLeave>>", _raman_leave)

        tk.Button(ef,
                  text=_L("📂 파일 선택으로 로드","📂 Load via File Dialog"),
                  command=self._load_raman_smart,
                  bg=BTN, fg=TXT, font=LF,
                  relief="flat", padx=8, pady=3,
                  cursor="hand2").pack(fill="x", pady=(0,2))
        tk.Label(ef,
                 text=_L("구조 자동 감지 → 조건 매칭 → 피크 추출 방식 선택",
                          "Auto-detect → Match condition → Select peak extraction"),
                 bg=PANEL, fg=SUB, font=("Segoe UI",7),
                 justify="left").pack(anchor="w", pady=(0,4))

        # Raman DB Save / Load 버튼
        rdb_row = tk.Frame(ef, bg=PANEL)
        rdb_row.pack(fill="x", pady=(0,2))
        tk.Button(rdb_row, text="💾 Save Raman DB",
                  command=self._db_save_raman,
                  bg=BTN, fg=TXT, font=LF,
                  relief="flat", padx=6, pady=3,
                  cursor="hand2").pack(side="left", expand=True,
                                       fill="x", padx=(0,2))
        tk.Button(rdb_row, text="📂 Load Raman DB",
                  command=self._raman_load_db_dialog,
                  bg=BTN, fg=TXT, font=LF,
                  relief="flat", padx=6, pady=3,
                  cursor="hand2").pack(side="left", expand=True,
                                       fill="x", padx=(2,0))

        tk.Frame(left, bg=BORDER, height=1).pack(fill="x", padx=8)

        # 수동 입력
        tk.Label(left, text=_L("  수동 입력","  Manual Input"),
                 bg=PANEL, fg=SUB, font=MFB).pack(anchor="w", padx=8, pady=(8,2))

        mf2 = tk.Frame(left, bg=PANEL)
        mf2.pack(fill="x", padx=8)

        self._raman_cond_var = tk.StringVar()
        self._raman_day_var  = tk.StringVar()
        self._raman_peak_var = tk.StringVar()

        for lbl_txt, var in [("Condition (cond):", self._raman_cond_var),
                              ("Day (day):",  self._raman_day_var),
                              ("Peak Intensity:",   self._raman_peak_var)]:
            row = tk.Frame(mf2, bg=PANEL)
            row.pack(fill="x", pady=2)
            tk.Label(row, text=lbl_txt, bg=PANEL, fg=SUB,
                     font=LF, width=12, anchor="w").pack(side="left")
            tk.Entry(row, textvariable=var,
                     bg=PANEL2, fg=TXT, font=MF,
                     insertbackground=TXT, relief="flat",
                     highlightbackground=BORDER,
                     highlightthickness=1, width=18).pack(
                     side="left", padx=4)

        # 조건 프리셋 버튼
        pf3 = tk.Frame(mf2, bg=PANEL)
        pf3.pack(fill="x", pady=(2,4))
        tk.Label(pf3, text=_L("조건 프리셋:","Condition Presets:"), bg=PANEL,
                 fg=SUB, font=LF).pack(side="left")
        for lbl2,col in self._presets:
            tk.Button(pf3, text=lbl2.replace("Native-","N-"),
                      command=lambda l=lbl2: self._raman_cond_var.set(l),
                      bg=BTN, fg=col,
                      font=("Segoe UI",7,"bold"),
                      relief="flat", cursor="hand2",
                      padx=2, pady=1).pack(side="left", padx=1)

        tk.Button(mf2, text=_L("➕ 추가","➕ Add"),
                  command=self._add_raman_row,
                  bg=GREEN, fg="white", font=MFB,
                  relief="flat", padx=8, pady=4,
                  cursor="hand2").pack(fill="x", pady=4)

        tk.Frame(left, bg=BORDER, height=1).pack(fill="x", padx=8)

        # Raman 데이터 테이블
        tk.Label(left, text=_L("  입력된 Raman 데이터","  Raman Data"),
                 bg=PANEL, fg=SUB, font=MFB).pack(anchor="w", padx=8, pady=(6,2))

        rtf = tk.Frame(left, bg=PANEL)
        rtf.pack(fill="both", expand=True, padx=4, pady=4)

        rcols = ("cond","day","peak","norm")
        self._raman_tree = ttk.Treeview(
            rtf, columns=rcols, show="headings", height=12)
        for c,h,w in [("cond","Cond",100),("day","Day",40),
                      ("peak","Peak",70),("norm","Norm",60)]:
            self._raman_tree.heading(c, text=h)
            self._raman_tree.column(c, width=w, anchor="center")
        rsb = tk.Scrollbar(rtf, orient="vertical",
                           command=self._raman_tree.yview)
        self._raman_tree.configure(yscrollcommand=rsb.set)
        rsb.pack(side="right", fill="y")
        self._raman_tree.pack(fill="both", expand=True)
        self._raman_tree.bind("<Delete>", self._delete_raman_row)

        tk.Label(left, text=_L("  행 선택 후 Delete키로 삭제","  Select row + Delete to remove"),
                 bg=PANEL, fg=SUB, font=LF).pack(anchor="w", padx=8, pady=2)

        # ── 우: 비교 차트 + AI 분석 ───────────────────
        right = tk.Frame(f, bg=BG)
        right.pack(side="left", fill="both", expand=True)

        # 상단: 3개 차트
        chart_area = tk.Frame(right, bg=BG)
        chart_area.pack(fill="both", expand=True)
        for c in range(3): chart_area.columnconfigure(c, weight=1)
        chart_area.rowconfigure(0, weight=1)

        self._raman_charts = {}
        for col_i, (key, title) in enumerate([
            ("raman_trend",    "Raman Peak Trend  (normalized)"),
            ("raman_spectrum", "Raman Spectra Overlay"),
            ("raman_decay",    "Decay Rate by Condition"),
        ]):
            cell = tk.Frame(chart_area, bg=PANEL,
                            highlightbackground=BORDER, highlightthickness=1)
            cell.grid(row=0, column=col_i, padx=3, pady=3, sticky="nsew")
            h2 = tk.Frame(cell, bg=PANEL2,
                          highlightbackground=BORDER, highlightthickness=1)
            h2.pack(fill="x")
            tk.Label(h2, text=f"  {title}",
                     bg=PANEL2, fg=TXT, font=LF).pack(
                     side="left", pady=3, padx=6)
            tk.Label(h2, text="⤢", bg=PANEL2, fg=SUB,
                     font=("Segoe UI",9)).pack(side="right", padx=4)
            fig = plt.Figure(figsize=(4.5, 3.0), facecolor=PANEL)
            cv  = FigureCanvasTkAgg(fig, master=cell)
            cv.get_tk_widget().pack(
                fill="both", expand=True, padx=2, pady=(0,2))
            cv.get_tk_widget().bind(
                "<Double-Button-1>",
                lambda e, k=key, t=title: self._popup_raman(k, t))
            self._raman_charts[key] = {"fig":fig,"cv":cv,"title":title}

        # Evaluation 탭 안내 (Pseudo-Raman은 Evaluation 탭에서 자동 실행)
        eval_note = tk.Frame(right, bg=PANEL2,
                             highlightbackground=BORDER, highlightthickness=1)
        eval_note.pack(fill="x", pady=(4,0))
        tk.Label(eval_note,
                 text="  🎯 Pseudo-Raman",
                 bg=PANEL2, fg=TXT, font=MFB).pack(side="left", pady=6, padx=8)
        tk.Label(eval_note,
                 text="Auto-runs in [🎯 Evaluation] tab after date estimation",
                 bg=PANEL2, fg=SUB, font=("Segoe UI",7)).pack(side="left", padx=4)

        ai_frame = tk.Frame(right, bg=PANEL,
                            highlightbackground=BORDER, highlightthickness=1,
                            height=220)
        ai_frame.pack(fill="x", padx=0, pady=(4,0))
        ai_frame.pack_propagate(False)

        ai_hdr = tk.Frame(ai_frame, bg=PANEL2,
                          highlightbackground=BORDER, highlightthickness=1)
        ai_hdr.pack(fill="x")

        # 타이틀
        tk.Label(ai_hdr, text=_L("  🤖 AI 비교 분석","  🤖 AI Analysis"),
                 bg=PANEL2, fg=TXT, font=MFB).pack(side="left", pady=6, padx=8)

        # 분석 모드 선택 라디오버튼
        self._ai_mode = tk.StringVar(value="local")

        mode_frame = tk.Frame(ai_hdr, bg=PANEL2)
        mode_frame.pack(side="left", padx=12)

        tk.Radiobutton(
            mode_frame, text=_L("📊 로컬 통계","📊 Local Stats"),
            variable=self._ai_mode, value="local",
            bg=PANEL2, fg=TXT, selectcolor=PANEL2,
            activebackground=PANEL2, font=MF,
            command=self._on_ai_mode_change
        ).pack(side="left", padx=4)

        tk.Radiobutton(
            mode_frame, text="🤖 Claude API",
            variable=self._ai_mode, value="api",
            bg=PANEL2, fg=ACCENT, selectcolor=PANEL2,
            activebackground=PANEL2, font=MF,
            command=self._on_ai_mode_change
        ).pack(side="left", padx=4)

        # API 키 입력 (Claude API 선택 시 표시)
        self._api_key_frame = tk.Frame(ai_hdr, bg=PANEL2)
        self._api_key_frame.pack(side="left", padx=4)
        tk.Label(self._api_key_frame, text="API Key:",
                 bg=PANEL2, fg=SUB, font=LF).pack(side="left")
        _saved_key = self._settings.get("claude_api_key", "")
        self._api_key_var = tk.StringVar(value=_saved_key)
        self._api_key_entry = tk.Entry(
            self._api_key_frame,
            textvariable=self._api_key_var,
            width=32, bg=PANEL, fg=TXT,
            insertbackground=TXT, font=LF,
            relief="flat", show="*",
            highlightbackground=BORDER, highlightthickness=1)
        self._api_key_entry.pack(side="left", padx=4, pady=5)
        # 키 저장 체크박스 — ON 이면 settings.json 에 평문 저장
        self._api_key_save_var = tk.BooleanVar(value=bool(_saved_key))
        tk.Checkbutton(
            self._api_key_frame, text=_L("저장","Save"),
            variable=self._api_key_save_var,
            bg=PANEL2, fg=SUB, selectcolor=PANEL2,
            activebackground=PANEL2, font=LF,
            command=self._on_api_key_save_toggle,
        ).pack(side="left", padx=2)
        # 초기에는 숨김 (로컬 모드 기본)
        self._api_key_frame.pack_forget()

        # 버튼
        tk.Button(ai_hdr, text=_L("▶ 분석 실행","▶ Run Analysis"),
                  command=self._run_ai_analysis,
                  bg=ACCENT, fg="white", font=MFB,
                  relief="flat", padx=12, pady=4,
                  cursor="hand2").pack(side="right", padx=8, pady=5)
        tk.Button(ai_hdr, text=_L("📋 복사","📋 Copy"),
                  command=self._copy_ai_text,
                  bg=BTN, fg=TXT, font=MF,
                  relief="flat", padx=8, pady=4,
                  cursor="hand2").pack(side="right", padx=2)

        # AI 결과 텍스트
        ai_body = tk.Frame(ai_frame, bg=PANEL)
        ai_body.pack(fill="both", expand=True, padx=4, pady=4)
        self._ai_text = tk.Text(
            ai_body, wrap="word", height=8,
            bg=CARD2, fg=TXT, font=("Segoe UI",9),
            relief="flat", padx=8, pady=6,
            highlightbackground=BORDER, highlightthickness=1,
            state="disabled")
        ai_vsb = tk.Scrollbar(ai_body, orient="vertical",
                               command=self._ai_text.yview)
        self._ai_text.configure(yscrollcommand=ai_vsb.set)
        ai_vsb.pack(side="right", fill="y")
        self._ai_text.pack(fill="both", expand=True)
        self._ai_set_text(
            _L("Raman 데이터와 이미지 분석 완료 후 [▶ 분석 실행] 버튼을 누르세요.\n\n"
               "📊 로컬 통계 — API 키 없이 즉시 사용 가능. R² 상관계수 기반 분석.\n"
               "🤖 Claude API — enter key to use. Generates expert interpretation.",
               "After Raman data + image analysis, click [▶ Run Analysis].\n\n"
               "📊 Local Stats — no API key needed. R² correlation based.\n"
               "🤖 Claude API — enter API key. Generates expert interpretation."))

    # ─────────────────────────────────────────
    #  Raman 데이터 관리
    # ─────────────────────────────────────────
    def _add_raman_row(self):
        cond = self._raman_cond_var.get().strip()
        day  = self._raman_day_var.get().strip()
        peak_s = self._raman_peak_var.get().strip()
        if not cond or not day or not peak_s:
            messagebox.showwarning(_L("입력 오류","Input Error"),_L("조건·날짜·피크강도를 모두 입력한다.","Enter condition, day, and peak intensity."))
            return
        try:
            peak = float(peak_s)
        except ValueError:
            messagebox.showerror(_L("오류","Error"),"피크 강도는 숫자로 입력한다.")
            return
        new_entry = {"cond":cond,"day":day,"peak":peak}
        self._raman_data.append(new_entry)
        self._ensure_raman_ids()
        # 라만 추가 직후 매칭할 이미지(들) 선택
        if self.images:
            picked = self._pick_images_for_raman_dialog(new_entry)
            if picked is not None and picked:
                self._attach_raman_to_images(new_entry["_id"], picked)
        self._normalize_raman()
        self._rebuild_raman_tree()
        self._raman_cond_var.set(""); self._raman_day_var.set("")
        self._raman_peak_var.set("")
        self._set_status(_L(f"Raman 데이터 추가: {cond} / {day}일 / {peak}",f"Raman data added: {cond} / Day {day} / {peak}"))

    def _delete_raman_row(self, event=None):
        sel = self._raman_tree.selection()
        if not sel: return
        for item in sel:
            vals = self._raman_tree.item(item,"values")
            # 삭제 대상 _id 모아서 매칭 일괄 해제
            doomed_ids = [r.get("_id") for r in self._raman_data
                          if r["cond"] == vals[0] and r["day"] == str(vals[1])
                          and r.get("_id") is not None]
            for rid in doomed_ids:
                self._detach_raman_from_all(rid)
            self._raman_data = [
                r for r in self._raman_data
                if not (r["cond"]==vals[0] and r["day"]==str(vals[1]))]
        self._normalize_raman()
        self._rebuild_raman_tree()

    def _auto_link_raman_by_cond_day(self, raman_entries=None):
        """주어진 라만 entry 들을 cond+day 동일 이미지에 자동 매칭.

        raman_entries=None 이면 전체. 이미 raman_id 가 있는 이미지는 덮지 않음.
        반환: 새로 매칭된 (image_idx, raman_id) 쌍 수.
        """
        if raman_entries is None:
            raman_entries = self._raman_data
        self._ensure_raman_ids()
        n_linked = 0
        for r in raman_entries:
            rid = r.get("_id")
            if rid is None:
                continue
            rcond = (r.get("cond") or "").strip()
            rday = str(r.get("day") or "").strip()
            for i, img in enumerate(self.images):
                if img.get("raman_id") is not None:
                    continue
                if (img.get("cond") == rcond and
                        str(img.get("day", "")) == rday):
                    img["raman_id"] = rid
                    self._refresh_card_raman_badge(i)
                    n_linked += 1
        return n_linked

    def _ensure_raman_ids(self):
        """self._raman_data 의 각 entry 에 _id 가 없으면 새로 부여 (1부터, 충돌 회피).

        통합 DB 저장 시 호출됨 — img['raman_id'] 가 가리킬 안정적 id 보장.
        """
        used = {r["_id"] for r in self._raman_data
                if r.get("_id") is not None}
        next_id = max(used, default=0) + 1
        for r in self._raman_data:
            if r.get("_id") is None:
                while next_id in used:
                    next_id += 1
                r["_id"] = next_id
                used.add(next_id)
                next_id += 1

    def _raman_by_id(self, rid):
        """raman _id → entry. 없으면 None."""
        if rid is None:
            return None
        for r in self._raman_data:
            if r.get("_id") == rid:
                return r
        return None

    def _images_with_raman(self, rid):
        """라만 _id 에 매칭된 이미지 인덱스 리스트."""
        return [i for i, img in enumerate(self.images)
                if img.get("raman_id") == rid]

    def _detach_raman_from_all(self, rid):
        """라만 _id 매칭을 모든 이미지에서 해제. 카드 뱃지 갱신."""
        if rid is None:
            return
        for i, img in enumerate(self.images):
            if img.get("raman_id") == rid:
                img["raman_id"] = None
                self._refresh_card_raman_badge(i)

    def _attach_raman_to_images(self, rid, idx_list):
        """라만 _id 를 이미지 idx 들에 매칭 (기존 매칭 유지). 카드 뱃지 갱신."""
        if rid is None:
            return
        for i in idx_list:
            if 0 <= i < len(self.images):
                self.images[i]["raman_id"] = rid
                self._refresh_card_raman_badge(i)

    def _pick_images_for_raman_dialog(self, raman_entry,
                                      preselect_same_cond_day=True):
        """라만 entry 에 매칭할 이미지(들)를 사용자가 선택하는 모달 다이얼로그.

        Returns: 선택된 이미지 idx 리스트 (취소 시 None).
        """
        if not self.images:
            messagebox.showinfo(
                _L("이미지 없음", "No images"),
                _L("먼저 이미지를 추가하세요.", "Add images first."))
            return None
        win = tk.Toplevel(self)
        win.title(_L("라만 매칭 이미지 선택",
                     "Pick images to link with Raman"))
        win.transient(self)
        win.grab_set()
        win.geometry("520x460")

        rcond = (raman_entry.get("cond") or "").strip()
        rday = str(raman_entry.get("day") or "").strip()
        peak = raman_entry.get("peak")
        tk.Label(
            win, anchor="w",
            text=(f"Raman:  cond={rcond}  day={rday}  peak={peak}"),
            font=("Segoe UI", 9, "bold")
        ).pack(fill="x", padx=8, pady=(8, 4))

        # 체크박스 Treeview
        from tkinter import ttk as _ttk
        tree = _ttk.Treeview(
            win, columns=("name", "cond", "day", "match"),
            show="headings", height=14, selectmode="extended")
        for col, label, w in (("name", "name", 220),
                              ("cond", "cond", 130),
                              ("day", "day", 50),
                              ("match", "current raman", 90)):
            tree.heading(col, text=label)
            tree.column(col, width=w, anchor="w")
        tree.pack(fill="both", expand=True, padx=8, pady=4)

        # 행 채우기 + 같은 cond+day 사전 선택
        for i, img in enumerate(self.images):
            cur_rid = img.get("raman_id")
            cur_label = (f"#{cur_rid}" if cur_rid is not None
                         else "")
            tree.insert("", "end", iid=str(i),
                        values=(img.get("name", ""),
                                img.get("cond", ""),
                                img.get("day", ""),
                                cur_label))
            if preselect_same_cond_day:
                if (img.get("cond") == rcond and
                        str(img.get("day", "")) == rday):
                    tree.selection_add(str(i))

        # 버튼 행
        btn_row = tk.Frame(win)
        btn_row.pack(fill="x", padx=8, pady=6)

        def _select_all():
            tree.selection_set([str(i) for i in range(len(self.images))])

        def _select_same_cond():
            tree.selection_set([str(i) for i, img in enumerate(self.images)
                                if img.get("cond") == rcond])

        def _clear():
            tree.selection_remove(tree.selection())

        tk.Button(btn_row, text=_L("전체", "All"),
                  command=_select_all).pack(side="left", padx=2)
        tk.Button(btn_row, text=_L("동일 cond", "Same cond"),
                  command=_select_same_cond).pack(side="left", padx=2)
        tk.Button(btn_row, text=_L("해제", "Clear"),
                  command=_clear).pack(side="left", padx=2)

        result = {"idx": None}

        def _ok():
            result["idx"] = [int(i) for i in tree.selection()]
            win.destroy()

        def _skip():
            result["idx"] = []
            win.destroy()

        def _cancel():
            result["idx"] = None
            win.destroy()

        tk.Button(btn_row, text=_L("취소", "Cancel"),
                  command=_cancel).pack(side="right", padx=2)
        tk.Button(btn_row, text=_L("매칭 안 함", "Skip"),
                  command=_skip).pack(side="right", padx=2)
        tk.Button(btn_row, text=_L("확인", "OK"),
                  command=_ok, bg="#4f46e5", fg="white"
                  ).pack(side="right", padx=2)

        win.wait_window()
        return result["idx"]

    def _refresh_card_raman_badge(self, idx):
        """단일 카드의 라만 매칭 뱃지(⚛) 갱신. 캐시된 카드 frame 의 라벨을 토글."""
        if not hasattr(self, "_cards_by_idx"):
            return
        card = self._cards_by_idx.get(idx)
        if card is None:
            return
        # 기존 뱃지 제거
        for w in card.winfo_children():
            if getattr(w, "_is_raman_badge", False):
                try: w.destroy()
                except Exception: pass
        # 매칭 있으면 새로 추가
        img = self.images[idx] if 0 <= idx < len(self.images) else None
        if img is None or img.get("raman_id") is None:
            return
        try:
            lbl = tk.Label(card, text="⚛", fg="#a78bfa",
                           bg=card.cget("bg"),
                           font=("Segoe UI", 11, "bold"))
            lbl._is_raman_badge = True
            lbl.place(relx=1.0, rely=0.0, anchor="ne", x=-6, y=2)
        except Exception as e:
            print(f"[raman-badge] WARN idx={idx}: {e}")

    def _normalize_raman(self):
        """조건별로 첫날 피크를 1.0으로 정규화"""
        from collections import defaultdict
        cond_first: dict = {}
        def df(d):
            try: return float(d)
            except: return 9999
        for cond in set(r["cond"] for r in self._raman_data):
            rows_c = sorted([r for r in self._raman_data if r["cond"]==cond],
                            key=lambda x: df(x["day"]))
            if rows_c:
                cond_first[cond] = rows_c[0]["peak"]
        for r in self._raman_data:
            ref = cond_first.get(r["cond"], 1.0)
            r["norm_peak"] = r["peak"] / ref if ref != 0 else np.nan

    def _rebuild_raman_tree(self):
        self._raman_tree.delete(*self._raman_tree.get_children())
        def df(d):
            try: return float(d)
            except: return 9999
        for r in sorted(self._raman_data,
                        key=lambda x:(x["cond"],df(x["day"]))):
            np_ = r.get("norm_peak", np.nan)
            self._raman_tree.insert("","end",values=(
                r["cond"][:18], r["day"],
                f"{r['peak']:.4f}",
                f"{np_:.3f}" if not np.isnan(np_) else "-"))

    # ─────────────────────────────────────────
    #  스마트 Raman Excel 로드
    # ─────────────────────────────────────────
    # ─────────────────────────────────────────
    #  Raman Excel DnD 드롭 핸들러
    # ─────────────────────────────────────────
    def _on_raman_excel_drop(self, event):
        """여러 Excel 파일 동시 드롭 → 파일별 스마트 팝업 순차 실행"""
        # 드롭존 원래 색상으로 복구
        self._raman_drop_cv.configure(
            highlightbackground=TEAL, bg=CARD2)
        self._raman_drop_cv.itemconfigure(
            self._raman_drop_lbl_id, fill=SUB)

        paths = parse_drop_paths(event.data)
        excel_paths = [p for p in paths
                       if p.lower().endswith(('.xlsx','.xls','.csv'))]

        if not excel_paths:
            messagebox.showwarning(
                _L("형식 오류","Format Error"),
                _L("Excel/CSV 파일만 지원한다. (.xlsx .xls .csv)",
                   "Only Excel/CSV files are supported. (.xlsx .xls .csv)"))
            return

        # 여러 파일이면 배치 처리 안내 팝업
        if len(excel_paths) > 1:
            self._raman_batch_popup(excel_paths)
        else:
            self._load_raman_from_path(excel_paths[0])

    def _load_raman_from_path(self, path: str):
        """단일 파일 경로로 스마트 로드 팝업 실행"""
        try:
            import openpyxl
            if path.lower().endswith('.csv'):
                import csv as _csv
                with open(path, encoding='utf-8-sig') as f:
                    rows_raw = list(_csv.reader(f))
                all_vals = rows_raw
            else:
                wb = openpyxl.load_workbook(path, data_only=True)
                ws = wb.active
                all_vals = [[ws.cell(r,c).value
                             for c in range(1, ws.max_column+1)]
                            for r in range(1, ws.max_row+1)]

            parsed = self._detect_raman_structure(all_vals, path)
            if parsed is None:
                messagebox.showerror(
                    _L("구조 감지 실패","Detection Failed"),
                    _L(f"'{os.path.basename(path)}'의 구조를 자동 인식하지 못했다.",
                       f"Could not auto-detect structure of '{os.path.basename(path)}'."))
                return
            self._raman_smart_popup(parsed, path)

        except Exception as ex:
            messagebox.showerror(_L("로드 오류","Load Error"),
                                 f"{os.path.basename(path)}: {ex}")

    def _raman_batch_popup(self, paths: list):
        """여러 Excel 파일 배치 처리 팝업 — 파일별 조건 매칭 후 일괄 임포트"""
        import openpyxl

        win = tk.Toplevel(self)
        win.title(_L("Raman 일괄 로드","Raman Batch Load"))
        win.configure(bg=PANEL)
        sw = self.winfo_screenwidth(); sh = self.winfo_screenheight()
        ww, wh = min(900, int(sw*0.75)), int(sh*0.75)
        win.geometry(f"{ww}x{wh}+{(sw-ww)//2}+{(sh-wh)//2}")
        win.grab_set()

        # 헤더
        hdr = tk.Frame(win, bg=PANEL2,
                       highlightbackground=BORDER, highlightthickness=1)
        hdr.pack(fill="x")
        tk.Label(hdr,
                 text=_L(f"  📂 Raman 일괄 로드  —  {len(paths)}개 파일",
                          f"  📂 Raman Batch Load  —  {len(paths)} files"),
                 bg=PANEL2, fg=TXT, font=MFB).pack(
                 side="left", pady=8, padx=10)
        tk.Label(hdr,
                 text=_L("각 파일에 대응 조건을 지정한다",
                          "Assign a condition to each file"),
                 bg=PANEL2, fg=SUB, font=LF).pack(
                 side="left", padx=6)

        # 스크롤 영역
        outer = tk.Frame(win, bg=BG)
        outer.pack(fill="both", expand=True, padx=6, pady=6)
        canvas = tk.Canvas(outer, bg=BG, highlightthickness=0)
        vsb = tk.Scrollbar(outer, orient="vertical", command=canvas.yview)
        canvas.configure(yscrollcommand=vsb.set)
        vsb.pack(side="right", fill="y")
        canvas.pack(fill="both", expand=True)
        inner = tk.Frame(canvas, bg=BG)
        canvas.create_window((0,0), window=inner, anchor="nw")
        inner.bind("<Configure>",
                   lambda e: canvas.configure(
                       scrollregion=canvas.bbox("all")))

        # 기존 조건 후보
        cond_candidates = list(dict.fromkeys(
            img.get("cond","") for img in self.images
            if img.get("cond","")))
        if not cond_candidates:
            cond_candidates = [""]

        # 파일별 행 생성
        file_rows = []   # [{path, parsed, cond_var, include_var, status_var}]

        # 컬럼 헤더
        hdr_f = tk.Frame(inner, bg=PANEL2,
                         highlightbackground=BORDER, highlightthickness=1)
        hdr_f.pack(fill="x", padx=2, pady=(0,2))
        for col_txt, col_w in [
            (_L("포함","Include"), 6),
            (_L("파일명","Filename"), 28),
            (_L("감지 패턴","Pattern"), 12),
            (_L("대응 조건","Condition"), 22),
            (_L("날짜 수","Days"), 6),
            (_L("상태","Status"), 12),
        ]:
            tk.Label(hdr_f, text=col_txt,
                     bg=PANEL2, fg=SUB, font=MFB,
                     width=col_w, anchor="w").pack(side="left", padx=4)

        for path in paths:
            row_f = tk.Frame(inner, bg=CARD,
                             highlightbackground=BORDER, highlightthickness=1)
            row_f.pack(fill="x", padx=2, pady=2)

            include_var = tk.BooleanVar(value=True)
            cond_var    = tk.StringVar()
            status_var  = tk.StringVar(value=_L("분석 중...","Parsing..."))

            # 파싱 시도
            parsed = None
            try:
                if path.lower().endswith('.csv'):
                    import csv as _csv
                    with open(path, encoding='utf-8-sig') as f:
                        rows_raw = [r for r in _csv.reader(f)]
                    parsed = self._detect_raman_structure(rows_raw, path)
                else:
                    wb2 = openpyxl.load_workbook(path, data_only=True)
                    ws2 = wb2.active
                    all_v = [[ws2.cell(r,c).value
                               for c in range(1,ws2.max_column+1)]
                              for r in range(1,ws2.max_row+1)]
                    parsed = self._detect_raman_structure(all_v, path)
            except Exception as ex:
                status_var.set(_L(f"오류: {ex}",f"Error: {ex}"))

            if parsed:
                n_days = len(parsed.get("days",[]))
                pat = {"wide":_L("가로","Wide"),
                       "long":_L("세로","Long")}.get(
                       parsed.get("pattern",""),"?")
                status_var.set(_L(f"OK ({n_days}일차)",f"OK ({n_days} days)"))
                # 파일명 힌트로 조건 추천
                hint = parsed["filename_hint"].lower()
                def _batch_hint_score(c):
                    cl = c.lower()
                    sc = 0
                    nums_h = _re.findall(r'\d+', hint)
                    nums_c = _re.findall(r'\d+', cl)
                    for n in nums_h:
                        sc += (50 if n in nums_c else -20)
                    for tok in _re.split(r'[-_%\s]', hint):
                        if len(tok) >= 2 and not tok.isdigit() and tok in cl:
                            sc += len(tok)
                    return sc
                best = max(cond_candidates,
                           key=_batch_hint_score,
                           default=cond_candidates[0])
                cond_var.set(best)
            else:
                pat = "?"
                n_days = 0
                if status_var.get() == _L("분석 중...","Parsing..."):
                    status_var.set(_L("감지 실패","Not detected"))

            # 체크박스
            tk.Checkbutton(row_f, variable=include_var,
                           bg=CARD, fg=TXT,
                           selectcolor=CARD2, relief="flat",
                           activebackground=CARD).pack(side="left", padx=8)
            # 파일명
            tk.Label(row_f,
                     text=os.path.basename(path),
                     bg=CARD, fg=TXT, font=LF,
                     width=28, anchor="w").pack(side="left", padx=2)
            # 패턴
            tk.Label(row_f, text=pat,
                     bg=CARD, fg=TEAL, font=LF,
                     width=12, anchor="w").pack(side="left", padx=2)
            # 조건 콤보
            cb = ttk.Combobox(row_f, textvariable=cond_var,
                               values=cond_candidates, width=20,
                               state="normal")
            cb.pack(side="left", padx=4)
            # 날짜 수
            tk.Label(row_f, text=str(n_days) if n_days else "-",
                     bg=CARD, fg=SUB, font=LF,
                     width=6).pack(side="left", padx=2)
            # 상태
            tk.Label(row_f, textvariable=status_var,
                     bg=CARD, fg=GREEN, font=LF,
                     width=12, anchor="w").pack(side="left", padx=4)
            # 상세 설정 버튼
            if parsed:
                def _open_detail(p=path, prs=parsed):
                    win.grab_release()
                    self._raman_smart_popup(prs, p)
                    # 스마트 팝업이 닫힌 후 배치 창 grab 복원
                    win.wait_window(win.winfo_children()[-1]
                                    if win.winfo_children() else win)
                    try: win.grab_set()
                    except Exception: pass
                tk.Button(row_f,
                          text=_L("⚙ 상세","⚙ Detail"),
                          command=_open_detail,
                          bg=BTN, fg=TXT, font=LF,
                          relief="flat", padx=6, pady=2,
                          cursor="hand2").pack(side="left", padx=4)

            file_rows.append({
                "path":        path,
                "parsed":      parsed,
                "cond_var":    cond_var,
                "include_var": include_var,
                "status_var":  status_var,
            })

        # 하단 버튼
        btn_f = tk.Frame(win, bg=PANEL2,
                         highlightbackground=BORDER, highlightthickness=1)
        btn_f.pack(fill="x")

        total_var = tk.StringVar(value="")
        tk.Label(btn_f, textvariable=total_var,
                 bg=PANEL2, fg=GREEN, font=LF).pack(
                 side="left", padx=10)

        def _batch_import():
            def _np_arr(lst):
                return np.array([float(x) for x in lst if x is not None])
            total_added = 0
            for fr in file_rows:
                if not fr["include_var"].get(): continue
                parsed_f = fr["parsed"]
                if not parsed_f: continue
                cond = fr["cond_var"].get().strip()
                if not cond: continue

                for day_lbl, spec in parsed_f["spectra"].items():
                    nums = _re.findall(r'\d+', str(day_lbl))
                    day_str = nums[0] if nums else day_lbl
                    shifts = _np_arr(spec.get("shifts",[]))
                    intens = _np_arr(spec.get("intensities",[]))
                    if len(shifts)==0: continue
                    # 기본: 전체 최대 피크
                    peak_val = float(intens[np.argmax(intens)]) \
                               if len(intens) else 0.0
                    peak_shift = float(shifts[np.argmax(intens)]) \
                                 if len(intens) else 0.0
                    self._raman_data.append({
                        "cond":       cond,
                        "day":        day_str,
                        "peak":       peak_val,
                        "peak_shift": peak_shift,
                        "peak_range": "auto_max",
                        "spectrum":   {
                            "shifts":      shifts.tolist(),
                            "intensities": intens.tolist()},
                    })
                    total_added += 1
                fr["status_var"].set(
                    _L(f"✓ 임포트됨","✓ Imported"))

            self._normalize_raman()
            # 추가된 라만 항목들을 cond+day 일치 이미지에 자동 매칭
            n_link = self._auto_link_raman_by_cond_day()
            self._rebuild_raman_tree()
            self._refresh_raman_tab()
            total_var.set(
                _L(f"✓ 총 {total_added}개 피크 임포트 완료 (이미지 매칭 {n_link}건)",
                   f"✓ {total_added} peaks imported total ({n_link} image links)"))
            self._set_status(
                _L(f"✓ Raman 일괄 로드 완료 ({total_added}개, 매칭 {n_link})",
                   f"✓ Raman batch load complete ({total_added}, links {n_link})"))

        tk.Button(btn_f,
                  text=_L("✔ 선택 파일 일괄 임포트","✔ Import Selected Files"),
                  command=_batch_import,
                  bg=ACCENT, fg="white", font=MFB,
                  relief="flat", padx=16, pady=6,
                  cursor="hand2").pack(side="right", padx=8, pady=6)
        tk.Button(btn_f,
                  text=_L("닫기","Close"),
                  command=win.destroy,
                  bg=BTN, fg=TXT, font=MF,
                  relief="flat", padx=12, pady=6,
                  cursor="hand2").pack(side="right", padx=2, pady=6)

    def _load_raman_smart(self):
        """파일 다이얼로그로 Excel 파일 선택 (다중 선택 지원)"""
        paths = filedialog.askopenfilenames(
            title=_L("Raman Excel 파일 선택","Select Raman Excel Files"),
            filetypes=[("Excel","*.xlsx *.xls"),
                       ("CSV","*.csv"),
                       (_L("전체","All"),"*.*")])
        if not paths: return

        if len(paths) == 1:
            self._load_raman_from_path(paths[0])
        else:
            self._raman_batch_popup(list(paths))

    def _detect_raman_structure(self, vals: list, path: str) -> dict | None:
        """
        엑셀 배치 패턴 자동 감지.
        반환: {
          "pattern": "wide"|"long"|"multi_peak",
          "days":    ["0 day","3 days",...],
          "spectra": {day_str: {"shifts":[...], "intensities":[...]}},
          "filename_hint": "native_RH35",
        }
        """
        import re as _re
        filename_hint = os.path.splitext(os.path.basename(path))[0]

        # ── 패턴 A: 가로 배치 (열쌍: shift, intensity 반복) ──
        # 행1~3이 헤더, 행4~ 숫자
        try:
            # 숫자 데이터 시작 행 찾기
            data_start = None
            for ri, row in enumerate(vals):
                numeric_count = sum(1 for v in row
                                    if v is not None and
                                    isinstance(v, (int, float)))
                if numeric_count >= 4:
                    data_start = ri
                    break

            if data_start is not None:
                # 열쌍 개수
                n_cols = len(vals[0])
                n_pairs = n_cols // 2

                # 날짜 레이블 찾기 (헤더 행에서)
                days = []
                for ri in range(data_start):
                    row = vals[ri]
                    candidates = []
                    for ci in range(0, n_cols, 2):
                        v = row[ci] if ci < len(row) else None
                        if v and isinstance(v, str) and any(
                                kw in str(v).lower()
                                for kw in ["day","days","d","days", "days", "days", "일"]):
                            candidates.append(str(v).strip())
                    if len(candidates) >= 2:
                        days = candidates
                        break

                if not days:
                    # 날짜 라벨 없으면 열쌍 순서로 생성
                    days = [f"pair_{i}" for i in range(n_pairs)]

                # 스펙트럼 추출
                spectra = {}
                for pi, day_lbl in enumerate(days):
                    sc = pi * 2      # shift column
                    ic = pi * 2 + 1  # intensity column
                    if ic >= n_cols:
                        break
                    shifts, intens = [], []
                    for row in vals[data_start:]:
                        s = row[sc] if sc < len(row) else None
                        v = row[ic] if ic < len(row) else None
                        if s is not None and v is not None:
                            try:
                                shifts.append(float(s))
                                intens.append(float(v))
                            except: pass
                    if shifts:
                        spectra[day_lbl] = {
                            "shifts":     shifts,
                            "intensities": intens}

                if spectra:
                    return {
                        "pattern":       "wide",
                        "days":          list(spectra.keys()),
                        "spectra":       spectra,
                        "filename_hint": filename_hint,
                    }
        except Exception:
            pass

        # ── 패턴 B: 세로 배치 (cond/day/shift/intensity 컬럼) ──
        try:
            header = [str(v).lower().strip() if v else ""
                      for v in vals[0]]
            shift_col = next((i for i,h in enumerate(header)
                              if any(k in h for k in
                                     ["shift","raman","cm","wavenumber", "wavenumber", "wavenumber", "파수"])), None)
            intens_col = next((i for i,h in enumerate(header)
                               if any(k in h for k in
                                      ["intensity","count","signal",
                                       "normalized","intensity", "signal", "signal", "강도"])), None)
            day_col = next((i for i,h in enumerate(header)
                            if any(k in h for k in
                                   ["day","date","day", "date", "date", "date", "날짜", "일"])), None)

            if shift_col is not None and intens_col is not None:
                spectra_b: dict = {}
                for row in vals[1:]:
                    s = row[shift_col] if shift_col < len(row) else None
                    v = row[intens_col] if intens_col < len(row) else None
                    d = (str(row[day_col]).strip()
                         if day_col is not None and
                         day_col < len(row) and row[day_col] else "all")
                    if s is not None and v is not None:
                        try:
                            spectra_b.setdefault(d, {"shifts":[],"intensities":[]})
                            spectra_b[d]["shifts"].append(float(s))
                            spectra_b[d]["intensities"].append(float(v))
                        except: pass

                if spectra_b:
                    return {
                        "pattern":       "long",
                        "days":          list(spectra_b.keys()),
                        "spectra":       spectra_b,
                        "filename_hint": filename_hint,
                    }
        except Exception:
            pass

        return None

    def _raman_smart_popup(self, parsed: dict, path: str):
        """파싱 미리보기 + 조건 매칭 + 피크 추출 방식 통합 팝업"""
        win = tk.Toplevel(self)
        win.title(_L("Raman 스마트 로드","Raman Smart Load"))
        win.configure(bg=PANEL)
        sw = self.winfo_screenwidth(); sh = self.winfo_screenheight()
        ww, wh = min(1100, int(sw*0.82)), int(sh*0.82)
        win.geometry(f"{ww}x{wh}+{(sw-ww)//2}+{(sh-wh)//2}")
        win.grab_set()

        # ── 헤더 ──────────────────────────────────
        hdr = tk.Frame(win, bg=PANEL2,
                       highlightbackground=BORDER, highlightthickness=1)
        hdr.pack(fill="x")
        tk.Label(hdr,
                 text=_L(f"  📡 Raman 스마트 로드  —  {os.path.basename(path)}",
                          f"  📡 Raman Smart Load  —  {os.path.basename(path)}"),
                 bg=PANEL2, fg=TXT, font=MFB).pack(
                 side="left", pady=8, padx=10)
        pattern_str = {"wide":_L("가로 배치(열쌍)","Wide/Column-pair"),
                       "long":_L("세로 배치(행)","Long/Row-based")}.get(
                       parsed["pattern"],"")
        tk.Label(hdr,
                 text=_L(f"감지 패턴: {pattern_str}",
                          f"Detected: {pattern_str}"),
                 bg=PANEL2, fg=TEAL, font=LF).pack(
                 side="left", padx=8)

        body = tk.Frame(win, bg=BG)
        body.pack(fill="both", expand=True, padx=6, pady=6)
        body.columnconfigure(0, weight=3)
        body.columnconfigure(1, weight=2)
        body.rowconfigure(0, weight=1)

        # ── 좌: 스펙트럼 미리보기 ─────────────────
        left_f = tk.Frame(body, bg=PANEL,
                          highlightbackground=BORDER, highlightthickness=1)
        left_f.grid(row=0, column=0, sticky="nsew", padx=(0,4))

        tk.Label(left_f,
                 text=_L("  📊 스펙트럼 미리보기 (더블클릭: 날짜 선택)",
                          "  📊 Spectrum Preview (dbl-click: select day)"),
                 bg=PANEL2, fg=TXT, font=MFB,
                 highlightbackground=BORDER,
                 highlightthickness=1).pack(fill="x")

        prev_fig = plt.Figure(figsize=(5, 3.5), facecolor=PANEL)
        prev_cv  = FigureCanvasTkAgg(prev_fig, master=left_f)
        prev_cv.get_tk_widget().pack(fill="both", expand=True,
                                     padx=4, pady=4)

        # ── 우: 설정 패널 ─────────────────────────
        right_f = tk.Frame(body, bg=PANEL,
                           highlightbackground=BORDER, highlightthickness=1)
        right_f.grid(row=0, column=1, sticky="nsew")

        # §1 조건 매칭
        tk.Label(right_f,
                 text=_L("  §1 조건 매칭","  §1 Condition Matching"),
                 bg=PANEL2, fg=TXT, font=MFB,
                 highlightbackground=BORDER,
                 highlightthickness=1).pack(fill="x")

        match_f = tk.Frame(right_f, bg=PANEL)
        match_f.pack(fill="x", padx=8, pady=6)

        # 파일명 → 조건 후보 추천
        hint = parsed["filename_hint"].lower()
        cond_candidates = []
        for img in self.images:
            c = img.get("cond","")
            if c and c not in cond_candidates:
                cond_candidates.append(c)
        # 유사도 점수 — 숫자 토큰에 높은 가중치 부여
        def hint_score(cond):
            c = cond.lower()
            score = 0
            toks = _re.split(r'[-_%\s]', hint)
            # 숫자 토큰 별도 추출 (예: "35", "70")
            nums = _re.findall(r'\d+', hint)
            cond_nums = _re.findall(r'\d+', c)
            # 숫자 정확 일치: 높은 가중치
            for n in nums:
                if n in cond_nums:
                    score += 50
                else:
                    score -= 20   # 숫자 불일치 패널티
            # 문자 토큰 부분 매칭
            for tok in toks:
                if len(tok) >= 2 and not tok.isdigit() and tok in c:
                    score += len(tok)
            return score
        import re as _re  # already imported globally
        cond_candidates.sort(key=hint_score, reverse=True)
        if not cond_candidates:
            cond_candidates = [_L("(조건 없음)","(no conditions)")]

        tk.Label(match_f,
                 text=_L("파일명에서 추천:","Recommended from filename:"),
                 bg=PANEL, fg=SUB, font=LF).pack(anchor="w")
        tk.Label(match_f,
                 text=f"  '{parsed['filename_hint']}'",
                 bg=PANEL, fg=TEAL, font=MFB).pack(anchor="w", pady=(0,4))

        cond_var = tk.StringVar(value=cond_candidates[0])
        tk.Label(match_f, text=_L("대응 조건 선택:","Match to condition:"),
                 bg=PANEL, fg=TXT, font=LF).pack(anchor="w")
        cond_cb = ttk.Combobox(match_f, textvariable=cond_var,
                               values=cond_candidates, width=22,
                               state="normal")
        cond_cb.pack(fill="x", pady=(2,4))
        tk.Label(match_f,
                 text=_L("(또는 직접 입력)","(or type directly)"),
                 bg=PANEL, fg=SUB, font=("Segoe UI",7)).pack(anchor="w")

        # §2 날짜 매핑
        tk.Frame(right_f, bg=BORDER, height=1).pack(fill="x", padx=8, pady=4)
        tk.Label(right_f,
                 text=_L("  §2 날짜 매핑","  §2 Day Mapping"),
                 bg=PANEL2, fg=TXT, font=MFB,
                 highlightbackground=BORDER,
                 highlightthickness=1).pack(fill="x")

        day_map_f = tk.Frame(right_f, bg=PANEL)
        day_map_f.pack(fill="x", padx=8, pady=4)

        day_vars = {}
        for day_lbl in parsed["days"]:
            row_f = tk.Frame(day_map_f, bg=PANEL)
            row_f.pack(fill="x", pady=1)
            tk.Label(row_f,
                     text=f"{day_lbl} →",
                     bg=PANEL, fg=SUB, font=LF, width=12,
                     anchor="e").pack(side="left")
            # 숫자 추출 시도
            nums = _re.findall(r'\d+', str(day_lbl))
            default_day = nums[0] if nums else ""
            dv = tk.StringVar(value=default_day)
            day_vars[day_lbl] = dv
            tk.Entry(row_f, textvariable=dv, width=6,
                     bg=PANEL2, fg=TXT, font=LF,
                     relief="flat",
                     highlightbackground=BORDER,
                     highlightthickness=1).pack(side="left", padx=4)
            tk.Label(row_f,
                     text=_L("일","day"),
                     bg=PANEL, fg=SUB, font=LF).pack(side="left")

        # §3 피크 추출 방식
        tk.Frame(right_f, bg=BORDER, height=1).pack(fill="x", padx=8, pady=4)
        tk.Label(right_f,
                 text=_L("  §3 피크 추출 방식","  §3 Peak Extraction Method"),
                 bg=PANEL2, fg=TXT, font=MFB,
                 highlightbackground=BORDER,
                 highlightthickness=1).pack(fill="x")

        peak_f = tk.Frame(right_f, bg=PANEL)
        peak_f.pack(fill="x", padx=8, pady=6)

        peak_mode = tk.StringVar(value="auto_max")

        modes = [
            ("auto_max",
             _L("전체 최대 피크","Global maximum peak")),
            ("range_single",
             _L("특정 범위 내 최대 피크","Peak in specific range")),
            ("range_multi",
             _L("복수 범위 피크 추출","Multiple range peaks")),
            ("area",
             _L("범위 내 면적(적분)","Area under curve in range")),
        ]
        for val, lbl in modes:
            tk.Radiobutton(peak_f, text=lbl,
                           variable=peak_mode, value=val,
                           bg=PANEL, fg=TXT,
                           selectcolor=PANEL2,
                           activebackground=PANEL,
                           font=LF).pack(anchor="w", pady=1)

        # 범위 설정 (range 모드일 때 활성화)
        range_f = tk.Frame(peak_f, bg=PANEL)
        range_f.pack(fill="x", pady=(4,0))

        tk.Label(range_f,
                 text=_L("범위 설정 (cm⁻¹):","Range (cm⁻¹):"),
                 bg=PANEL, fg=SUB, font=LF).pack(anchor="w")

        # 복수 범위: 쉼표로 구분 "300-360, 370-400"
        peak_ranges_var = tk.StringVar(value="300-360")
        tk.Label(range_f,
                 text=_L("예) 300-360  또는  300-360, 370-400",
                          "e.g. 300-360  or  300-360, 370-400"),
                 bg=PANEL, fg=SUB, font=("Segoe UI",7)).pack(anchor="w")
        tk.Entry(range_f, textvariable=peak_ranges_var,
                 bg=PANEL2, fg=TXT, font=LF,
                 relief="flat",
                 highlightbackground=BORDER,
                 highlightthickness=1).pack(fill="x", pady=2)

        # 스펙트럼 미리보기 그리기
        def _draw_preview(day_lbl=None):
            prev_fig.clear()
            ax = prev_fig.add_subplot(111)
            ax.set_facecolor(PANEL2)
            prev_fig.patch.set_facecolor(PANEL)
            for sp in ax.spines.values(): sp.set_color(BORDER)
            ax.tick_params(colors=SUB, labelsize=7)

            days_to_draw = ([day_lbl] if day_lbl
                            else parsed["days"])
            for i, dl in enumerate(days_to_draw):
                spec = parsed["spectra"].get(dl, {})
                shifts = spec.get("shifts", [])
                intens = spec.get("intensities", [])
                if shifts:
                    col = COND_COLORS[i % len(COND_COLORS)]
                    ax.plot(shifts, intens, lw=1.2, color=col,
                            label=dl, alpha=0.85)

            # 범위 표시
            try:
                ranges_str = peak_ranges_var.get()
                for seg in ranges_str.split(","):
                    seg = seg.strip()
                    lo, hi = map(float, seg.split("-"))
                    ax.axvspan(lo, hi, alpha=0.12, color=ACCENT)
                    ax.axvline(lo, color=ACCENT, lw=0.8, ls="--")
                    ax.axvline(hi, color=ACCENT, lw=0.8, ls="--")
            except Exception:
                pass

            ax.set_xlabel("Raman Shift (cm⁻¹)", color=SUB, fontsize=8)
            ax.set_ylabel(_L("강도 (정규화)","Intensity (norm.)"),
                          color=SUB, fontsize=8)
            ax.legend(fontsize=7, framealpha=0.8, edgecolor=BORDER)
            prev_fig.tight_layout(pad=0.8)
            prev_cv.draw()

        _draw_preview()
        peak_ranges_var.trace_add("write", lambda *_: _draw_preview())
        peak_mode.trace_add("write", lambda *_: _draw_preview())

        # ── 하단 버튼 ─────────────────────────────
        btn_f = tk.Frame(win, bg=PANEL2,
                         highlightbackground=BORDER, highlightthickness=1)
        btn_f.pack(fill="x", pady=0)

        status_var = tk.StringVar(value="")
        tk.Label(btn_f, textvariable=status_var,
                 bg=PANEL2, fg=GREEN, font=LF).pack(
                 side="left", padx=12)

        def _do_import():
            cond = cond_var.get().strip()
            if not cond:
                messagebox.showwarning(
                    _L("경고","Warning"),
                    _L("조건을 선택한다.","Select a condition."),
                    parent=win)
                return

            mode = peak_mode.get()
            ranges_str = peak_ranges_var.get()

            # 범위 파싱
            peak_ranges = []
            for seg in ranges_str.split(","):
                seg = seg.strip()
                try:
                    lo, hi = map(float, seg.split("-"))
                    peak_ranges.append((lo, hi))
                except Exception:
                    pass
            if not peak_ranges and mode != "auto_max":
                peak_ranges = [(300, 360)]

            added = 0
            for day_lbl, dv in day_vars.items():
                day_str = dv.get().strip()
                if not day_str: continue
                spec = parsed["spectra"].get(day_lbl, {})
                shifts = np.array(spec.get("shifts", []))
                intens = np.array(spec.get("intensities", []))
                if len(shifts) == 0: continue

                if mode == "auto_max":
                    peaks = [{"range":"all",
                               "shift": float(shifts[np.argmax(intens)]),
                               "intensity": float(np.max(intens))}]
                elif mode == "area":
                    peak_list = []
                    for lo, hi in peak_ranges:
                        mask = (shifts >= lo) & (shifts <= hi)
                        if mask.any():
                            area = float(np.trapz(intens[mask], shifts[mask]))
                            peak_list.append({"range":f"{lo:.0f}-{hi:.0f}",
                                              "shift": float(np.mean(shifts[mask])),
                                              "intensity": area})
                    peaks = peak_list if peak_list else []
                else:  # range_single / range_multi
                    peak_list = []
                    for lo, hi in peak_ranges:
                        mask = (shifts >= lo) & (shifts <= hi)
                        if mask.any():
                            idx = np.argmax(intens[mask])
                            peak_list.append({
                                "range":     f"{lo:.0f}-{hi:.0f}",
                                "shift":     float(shifts[mask][idx]),
                                "intensity": float(intens[mask][idx])})
                    peaks = peak_list if peak_list else []

                # 대표 피크 = 첫 번째 (수동 입력 호환)
                for pk in peaks:
                    self._raman_data.append({
                        "cond":  cond,
                        "day":   day_str,
                        "peak":  pk["intensity"],
                        "peak_shift": pk["shift"],
                        "peak_range": pk["range"],
                        # 전체 스펙트럼 저장 (Pseudo-Raman에 활용)
                        "spectrum": {
                            "shifts":     shifts.tolist(),
                            "intensities": intens.tolist()},
                    })
                    added += 1

            self._normalize_raman()
            n_link = self._auto_link_raman_by_cond_day()
            self._rebuild_raman_tree()
            self._refresh_raman_tab()
            status_var.set(_L(f"✓ {added}개 피크 임포트 (이미지 매칭 {n_link}건)",
                               f"✓ {added} peaks imported ({n_link} image links)"))
            self._set_status(_L(f"✓ Raman 스마트 로드 완료 ({added}개, 매칭 {n_link})",
                                 f"✓ Raman smart load complete ({added}, links {n_link})"))

        tk.Button(btn_f,
                  text=_L("✔ 임포트","✔ Import"),
                  command=_do_import,
                  bg=ACCENT, fg="white", font=MFB,
                  relief="flat", padx=16, pady=6,
                  cursor="hand2").pack(side="right", padx=8, pady=6)
        tk.Button(btn_f,
                  text=_L("닫기","Close"),
                  command=win.destroy,
                  bg=BTN, fg=TXT, font=MF,
                  relief="flat", padx=12, pady=6,
                  cursor="hand2").pack(side="right", padx=2, pady=6)

    # ─────────────────────────────────────────
    #  Pseudo-Raman 추정
    # ─────────────────────────────────────────
    def _pseudo_raman_popup(self):
        """
        이미지 분석 지표(b*, S, YI, ΔE) + 실측 Raman 쌍으로
        회귀 모델 수립 → 평가 대상 이미지의 Raman 스펙트럼 추정값·범위 출력
        """
        # 매칭 쌍 구성
        an = [img for img in self.images
              if not np.isnan(img.get("lab",{}).get("b", np.nan))]
        rd = self._raman_data

        if not an:
            messagebox.showwarning(
                _L("주의","Warning"),
                _L("이미지 분석을 먼저 실행한다.","Run image analysis first."))
            return
        if not rd:
            messagebox.showwarning(
                _L("주의","Warning"),
                _L("Raman 데이터가 없다.","No Raman data loaded."))
            return

        def _safe(v): return 0.0 if v is None or np.isnan(float(v)) else float(v)

        # 명시적 매칭(raman_id) 우선, 없으면 cond+day 자연 매칭
        pairs = []
        for img in an:
            r_match = None
            rid = img.get("raman_id")
            if rid is not None:
                r_match = self._raman_by_id(rid)
            if r_match is None:
                r_match = next((r for r in rd
                                if r["cond"] == img["cond"]
                                and r["day"] == img["day"]), None)
            if r_match:
                pairs.append({
                    "cond": img["cond"], "day": img["day"],
                    "b":    _safe(img["lab"]["b"]),
                    "s":    _safe(img["s_mean"]),
                    "yi":   _safe(img.get("yellowness_idx", np.nan)),
                    "de":   _safe(img.get("delta_e", np.nan)),
                    "peak": _safe(r_match["peak"]),
                    "norm_peak": _safe(r_match.get("norm_peak", np.nan)),
                    "spectrum": r_match.get("spectrum"),
                })

        win = tk.Toplevel(self)
        win.title(_L("🔮 Pseudo-Raman 추정","🔮 Pseudo-Raman Estimation"))
        win.configure(bg=PANEL)
        sw = self.winfo_screenwidth(); sh = self.winfo_screenheight()
        ww, wh = int(sw*0.88), int(sh*0.88)
        win.geometry(f"{ww}x{wh}+{(sw-ww)//2}+{(sh-wh)//2}")

        hdr = tk.Frame(win, bg=PANEL2,
                       highlightbackground=BORDER, highlightthickness=1)
        hdr.pack(fill="x")
        tk.Label(hdr,
                 text=_L("  🔮 Pseudo-Raman  —  이미지 분석만으로 라만 스펙트럼 추정",
                          "  🔮 Pseudo-Raman  —  Estimate Raman spectrum from image analysis only"),
                 bg=PANEL2, fg=TXT, font=MFB).pack(
                 side="left", pady=8, padx=10)

        body = tk.Frame(win, bg=BG)
        body.pack(fill="both", expand=True, padx=6, pady=6)
        body.columnconfigure(0, weight=1)
        body.columnconfigure(1, weight=1)
        body.rowconfigure(0, weight=2)
        body.rowconfigure(1, weight=1)

        # ── 상단 좌: 회귀 모델 차트 ─────────────
        reg_f = tk.Frame(body, bg=PANEL,
                         highlightbackground=BORDER, highlightthickness=1)
        reg_f.grid(row=0, column=0, sticky="nsew", padx=(0,3), pady=(0,3))
        tk.Label(reg_f,
                 text=_L("  📈 이미지 지표 vs Raman 피크 회귀",
                          "  📈 Image Metrics vs Raman Peak Regression"),
                 bg=PANEL2, fg=TXT, font=MFB,
                 highlightbackground=BORDER,
                 highlightthickness=1).pack(fill="x")
        reg_fig = plt.Figure(figsize=(5.5, 4), facecolor=PANEL)
        reg_cv  = FigureCanvasTkAgg(reg_fig, master=reg_f)
        reg_cv.get_tk_widget().pack(fill="both", expand=True, padx=2, pady=2)

        # ── 상단 우: 추정 스펙트럼 차트 ─────────
        spec_f = tk.Frame(body, bg=PANEL,
                          highlightbackground=BORDER, highlightthickness=1)
        spec_f.grid(row=0, column=1, sticky="nsew", pady=(0,3))
        tk.Label(spec_f,
                 text=_L("  🔮 추정 스펙트럼 + 신뢰 범위",
                          "  🔮 Estimated Spectrum + Confidence Band"),
                 bg=PANEL2, fg=TXT, font=MFB,
                 highlightbackground=BORDER,
                 highlightthickness=1).pack(fill="x")
        spec_fig = plt.Figure(figsize=(5.5, 4), facecolor=PANEL)
        spec_cv  = FigureCanvasTkAgg(spec_fig, master=spec_f)
        spec_cv.get_tk_widget().pack(fill="both", expand=True, padx=2, pady=2)

        # ── 하단: 평가 대상 + 결과 텍스트 ───────
        bot_f = tk.Frame(body, bg=BG)
        bot_f.grid(row=1, column=0, columnspan=2, sticky="nsew")
        bot_f.columnconfigure(0, weight=1)
        bot_f.columnconfigure(1, weight=2)

        # 평가 대상 입력
        inp_f = tk.Frame(bot_f, bg=PANEL,
                         highlightbackground=BORDER, highlightthickness=1)
        inp_f.grid(row=0, column=0, sticky="nsew", padx=(0,3))
        tk.Label(inp_f,
                 text=_L("  📸 평가 대상 이미지","  📸 Target Image"),
                 bg=PANEL2, fg=TXT, font=MFB,
                 highlightbackground=BORDER,
                 highlightthickness=1).pack(fill="x")

        inp_inner = tk.Frame(inp_f, bg=PANEL)
        inp_inner.pack(fill="x", padx=8, pady=6)

        target_var = tk.StringVar()
        analyzed_names = [f"{img['cond']} / {img['day']}일"
                          for img in an]
        tk.Label(inp_inner,
                 text=_L("분석 완료 이미지 선택:","Select analyzed image:"),
                 bg=PANEL, fg=SUB, font=LF).pack(anchor="w")
        target_cb = ttk.Combobox(inp_inner, textvariable=target_var,
                                  values=analyzed_names, width=28,
                                  state="readonly")
        target_cb.pack(fill="x", pady=2)
        if analyzed_names:
            target_cb.current(0)

        tk.Button(inp_inner,
                  text=_L("🔮 Pseudo-Raman 실행","🔮 Run Pseudo-Raman"),
                  command=lambda: _run_pseudo(),
                  bg=ACCENT, fg="white", font=MFB,
                  relief="flat", padx=12, pady=5,
                  cursor="hand2").pack(fill="x", pady=(6,0))

        # 결과 텍스트
        res_f = tk.Frame(bot_f, bg=PANEL,
                         highlightbackground=BORDER, highlightthickness=1)
        res_f.grid(row=0, column=1, sticky="nsew")
        tk.Label(res_f,
                 text=_L("  📋 추정 결과","  📋 Estimation Result"),
                 bg=PANEL2, fg=TXT, font=MFB,
                 highlightbackground=BORDER,
                 highlightthickness=1).pack(fill="x")
        res_txt = tk.Text(res_f, wrap="word",
                          bg=CARD2, fg=TXT,
                          font=("Segoe UI",9),
                          relief="flat", padx=10, pady=6,
                          highlightthickness=0,
                          cursor="xterm",
                          state="disabled")
        res_txt.pack(fill="both", expand=True, padx=4, pady=4)

        # ── 회귀 차트 그리기 ──────────────────────
        def _draw_regression():
            reg_fig.clear()
            if len(pairs) < 3:
                ax = reg_fig.add_subplot(111)
                ax.text(0.5, 0.5,
                        _L(f"매칭 쌍이 {len(pairs)}개\n(최소 3개 필요)",
                           f"Only {len(pairs)} matched pairs\n(need ≥3)"),
                        transform=ax.transAxes,
                        ha="center", va="center", color=SUB)
                reg_cv.draw(); return

            gs = reg_fig.add_gridspec(2, 2, hspace=0.45, wspace=0.4,
                                      top=0.92, bottom=0.10)
            metrics = [
                ("b",  "Lab b*",  "fresh"),
                ("s",  "S-ch",    "fresh"),
                ("yi", "YI",      "fresh"),
                ("de", "ΔE",      "oxidized"),
            ]
            for ax_i, (key, lbl, _) in enumerate(metrics):
                ax = reg_fig.add_subplot(gs[ax_i//2, ax_i%2])
                ax.set_facecolor(PANEL2)
                for sp in ax.spines.values(): sp.set_color(BORDER)
                ax.tick_params(colors=SUB, labelsize=7)

                xs = np.array([p[key] for p in pairs])
                ys = np.array([p["norm_peak"] for p in pairs])
                ax.scatter(xs, ys, color=ACCENT, s=30, zorder=5)

                if len(xs) >= 2:
                    coef = np.polyfit(xs, ys, 1)
                    xl = np.linspace(xs.min(), xs.max(), 50)
                    ax.plot(xl, np.polyval(coef, xl),
                            "--", color=RED, lw=1.2, alpha=0.8)
                    r2 = 1 - np.sum((ys - np.polyval(coef,xs))**2) / \
                             np.sum((ys - ys.mean())**2) if ys.std() > 0 else 0
                    ax.text(0.05, 0.92, f"R²={r2:.3f}",
                            transform=ax.transAxes,
                            fontsize=7, color=RED,
                            va="top",
                            bbox=dict(boxstyle="round",
                                      fc=PANEL,ec=BORDER,alpha=0.8))
                ax.set_xlabel(lbl, fontsize=7, color=SUB)
                ax.set_ylabel(_L("Raman 피크(정규화)","Raman Peak(norm)"),
                              fontsize=6, color=SUB)

            reg_fig.suptitle(
                _L("이미지 지표 vs Raman 피크 회귀",
                   "Image Metrics vs Raman Peak"),
                fontsize=9, color=TXT)
            reg_cv.draw()

        _draw_regression()

        # ── Pseudo-Raman 실행 ─────────────────────
        def _run_pseudo():
            sel = target_var.get()
            if not sel: return
            # 선택된 이미지 찾기
            target_img = None
            for img in an:
                label = f"{img['cond']} / {img['day']}일"
                if label == sel:
                    target_img = img; break
            if target_img is None: return

            t_b  = _safe(target_img["lab"]["b"])
            t_s  = _safe(target_img["s_mean"])
            t_yi = _safe(target_img.get("yellowness_idx", np.nan))
            t_de = _safe(target_img.get("delta_e", np.nan))

            if len(pairs) < 3:
                self._set_cmt(res_txt,
                    _L("매칭 쌍 부족 (최소 3개 필요).\n"
                       "이미지 분석과 Raman 데이터의 cond+day가 일치하는 쌍이 필요하다.",
                       "Insufficient matched pairs (need ≥3).\n"
                       "cond+day must match between image analysis and Raman data."))
                return

            # 가중 앙상블 회귀 (b*, S, YI, ΔE)
            metrics_target = {"b": t_b, "s": t_s, "yi": t_yi, "de": t_de}
            wb_map = {"b": 0.45, "s": 0.25, "yi": 0.20, "de": 0.10}

            # 각 지표별 회귀 추정
            estimates = []
            weights_total = 0.0
            r2_map = {}
            for key, wt in wb_map.items():
                xs = np.array([p[key] for p in pairs])
                ys = np.array([p["norm_peak"] for p in pairs])
                if xs.std() < 1e-6: continue
                coef = np.polyfit(xs, ys, 1)
                pred = np.polyval(coef, metrics_target[key])
                res  = ys - np.polyval(coef, xs)
                se   = np.std(res)
                r2   = max(0, 1 - np.sum(res**2)/np.sum((ys-ys.mean())**2))
                r2_map[key] = r2
                estimates.append((pred, se, wt * r2))
                weights_total += wt * r2

            if not estimates or weights_total < 1e-9:
                self._set_cmt(res_txt,
                    _L("회귀 모델 수립 실패.","Regression model failed."))
                return

            # 가중 평균 추정값 + 95% 신뢰 구간
            est_peak = sum(p * w for p,se,w in estimates) / weights_total
            est_se   = (sum((se**2) * w for p,se,w in estimates)
                        / weights_total) ** 0.5
            ci_lo    = max(0, est_peak - 1.96 * est_se)
            ci_hi    = min(1.5, est_peak + 1.96 * est_se)

            # 가장 유사한 날짜 추정 (참조 쌍에서)
            dists = []
            for p in pairs:
                d = abs(p["norm_peak"] - est_peak)
                dists.append((d, p))
            dists.sort(key=lambda x: x[0])
            closest_day = dists[0][1]["day"] if dists else "?"

            # 추정 스펙트럼 생성 (참조 스펙트럼 선형 보간)
            _draw_estimated_spectrum(est_peak, ci_lo, ci_hi, dists[:3])

            # 결과 텍스트
            result = (
                f"【Pseudo-Raman Estimation Result】\n"
                f"{'─'*40}\n"
                f"Target:  {target_img['cond']} / Day {target_img['day']}\n\n"
                f"📊 Image Metrics:\n"
                f"  b*={t_b:.1f}  S={t_s:.1f}  YI={t_yi:.0f}  ΔE={t_de:.1f}\n\n"
                f"🔮 Estimated Raman A₁g Peak:\n"
                f"  Estimate: {est_peak:.4f} (normalized)\n"
                f"  95% CI: [{ci_lo:.4f}, {ci_hi:.4f}]\n"
                f"  Uncertainty: ±{est_se:.4f}\n\n"
                f"📅 Oxidation Stage:\n"
                f"  Closest ref: {dists[0][1]['cond']} Day {closest_day}\n"
                f"  (peak diff: {dists[0][0]:.4f})\n\n"
                f"📈 R² per metric:\n"
                + "".join(f"  {k:4s}: R²={v:.3f}\n"
                          for k,v in r2_map.items()) +
                f"\n💡 Interpretation:\n"
                f"  A₁g peak {est_peak:.3f}: {_peak_to_oxidation_en(est_peak)}."
            )

            self._set_cmt(res_txt, result)
            self._set_status(
                _L(f"✓ Pseudo-Raman 완료: 추정 피크={est_peak:.4f}",
                   f"✓ Pseudo-Raman done: est. peak={est_peak:.4f}"))

        def _peak_to_oxidation(peak: float) -> str:
            if peak >= 0.85: return "pristine (unoxidized)"
            elif peak >= 0.65: return "early oxidation"
            elif peak >= 0.40: return "significant oxidation"
            else: return "severe oxidation (HfO₂ dominant)"

        def _peak_to_oxidation_en(peak: float) -> str:
            if peak >= 0.85: return "pristine (unoxidized)"
            elif peak >= 0.65: return "early oxidation"
            elif peak >= 0.40: return "significant oxidation"
            else: return "severe oxidation (HfO₂ dominant)"

        def _draw_estimated_spectrum(est_peak, ci_lo, ci_hi, closest_pairs):
            """참조 스펙트럼을 보간하여 추정 스펙트럼 + 신뢰 범위 그리기"""
            spec_fig.clear()
            ax = spec_fig.add_subplot(111)
            ax.set_facecolor(PANEL2)
            for sp in ax.spines.values(): sp.set_color(BORDER)
            ax.tick_params(colors=SUB, labelsize=7)

            # 참조 스펙트럼 그리기 (반투명)
            ref_spectra = []
            for dist, p in closest_pairs:
                spec = p.get("spectrum")
                if spec:
                    s = np.array(spec["shifts"])
                    v = np.array(spec["intensities"])
                    ref_spectra.append((s, v, p["day"]))
                    ax.plot(s, v, lw=1, alpha=0.3, color=SUB,
                            label=f"ref {p['day']}d")

            # 추정 스펙트럼 (가중 보간)
            if ref_spectra and closest_pairs:
                # 피크 비율로 스케일링된 추정 스펙트럼
                # 가장 가까운 두 참조의 선형 보간
                if len(ref_spectra) >= 2:
                    s1, v1, d1 = ref_spectra[0]
                    s2, v2, d2 = ref_spectra[1]
                    # 공통 shift 축
                    s_common = s1
                    v2_interp = np.interp(s_common, s2, v2)
                    # 비율
                    p1 = closest_pairs[0][1].get("norm_peak", 1.0) or 1.0
                    p2 = closest_pairs[1][1].get("norm_peak", 1.0) or 1.0
                    if abs(p1 - p2) > 1e-6:
                        alpha = (est_peak - p2) / (p1 - p2)
                        alpha = max(0, min(1, alpha))
                    else:
                        alpha = 0.5
                    v_est = alpha * v1 + (1 - alpha) * v2_interp
                    # 신뢰 범위
                    band = abs(ci_hi - ci_lo) / 2
                    v_lo = v_est * (1 - band * 0.5)
                    v_hi = v_est * (1 + band * 0.5)

                    ax.fill_between(s_common, v_lo, v_hi,
                                    alpha=0.25, color=ACCENT,
                                    label=_L("95% 신뢰 범위","95% CI band"))
                    ax.plot(s_common, v_est,
                            lw=2, color=ACCENT,
                            label=_L(f"추정 스펙트럼 (피크={est_peak:.3f})",
                                     f"Estimated (peak={est_peak:.3f})"))

            ax.set_xlabel("Raman Shift (cm⁻¹)", color=SUB, fontsize=8)
            ax.set_ylabel(_L("강도 (정규화)","Intensity (norm.)"),
                          color=SUB, fontsize=8)
            ax.set_title(
                _L(f"Pseudo-Raman 추정  (A₁g≈{est_peak:.3f}±{(ci_hi-ci_lo)/2:.3f})",
                   f"Pseudo-Raman Est.  (A₁g≈{est_peak:.3f}±{(ci_hi-ci_lo)/2:.3f})"),
                fontsize=8, color=TXT)
            ax.legend(fontsize=7, framealpha=0.8, edgecolor=BORDER)
            spec_fig.tight_layout(pad=0.8)
            spec_cv.draw()

    def _load_raman_excel(self):
        """Excel 파일에서 Raman 데이터 일괄 로드"""
        path = filedialog.askopenfilename(
            title=_L("Raman Excel 파일 선택","Select Raman Excel File"),
            filetypes=[("Excel","*.xlsx *.xls *.csv"),(_L("전체","All"),"*.*")])
        if not path: return
        try:
            if path.lower().endswith(".csv"):
                import csv as _csv
                with open(path, encoding="utf-8-sig") as f:
                    reader = _csv.DictReader(f)
                    rows = list(reader)
            else:
                import openpyxl
                wb = openpyxl.load_workbook(path, data_only=True)
                ws = wb.active
                headers = [str(c.value).strip().lower()
                           if c.value else "" for c in ws[1]]
                rows = []
                for row in ws.iter_rows(min_row=2, values_only=True):
                    rows.append({headers[i]: str(v) if v is not None else ""
                                 for i,v in enumerate(row)})

            # 컬럼 매핑 (유연하게)
            added = 0
            for row in rows:
                keys = {k.lower().strip() for k in row}
                # cond 컬럼 후보
                cond_key = next((k for k in row
                                 if k.lower() in ("cond","condition","Cond")), None)
                day_key  = next((k for k in row
                                 if k.lower() in ("day","day", "date", "day", "date", "날짜", "일")), None)
                peak_key = next((k for k in row
                                 if k.lower() in ("peak","intensity","intensity", "intensity", "intensity", "강도",
                                                  "a1g","a_1g","peak_intensity")), None)
                if not all([cond_key, day_key, peak_key]): continue
                try:
                    cond = str(row[cond_key]).strip()
                    day  = str(row[day_key]).strip()
                    peak = float(row[peak_key])
                    if cond and day and not np.isnan(peak):
                        self._raman_data.append(
                            {"cond":cond,"day":day,"peak":peak})
                        added += 1
                except (ValueError, TypeError):
                    continue

            self._normalize_raman()
            n_link = self._auto_link_raman_by_cond_day()
            self._rebuild_raman_tree()
            self._set_status(_L(f"✓ Raman 데이터 {added}행 로드 (매칭 {n_link}건): {os.path.basename(path)}",f"✓ Raman data loaded: {added} rows, {n_link} links  ({os.path.basename(path)})"))

        except Exception as ex:
            messagebox.showerror(_L("로드 오류","Load Error"), str(ex))

    # ─────────────────────────────────────────
    #  Raman 차트 갱신
    # ─────────────────────────────────────────
    def _refresh_raman_tab(self):
        for key, cell in self._raman_charts.items():
            cell["fig"].clear()
            self._draw_raman_chart(cell["fig"], key, large=False)
            cell["cv"].draw()

    def _popup_raman(self, key: str, title: str):
        win = tk.Toplevel(self)
        win.title(title); win.configure(bg=PANEL)
        sw=self.winfo_screenwidth(); sh=self.winfo_screenheight()
        ww,wh=int(sw*0.80),int(sh*0.80)
        win.geometry(f"{ww}x{wh}+{(sw-ww)//2}+{(sh-wh)//2}")
        tk.Label(win, text=f"  {title}", bg=PANEL2, fg=TXT,
                 font=("Segoe UI",11,"bold"),
                 highlightbackground=BORDER,
                 highlightthickness=1).pack(fill="x")
        fig=plt.Figure(facecolor=PANEL)
        cv=FigureCanvasTkAgg(fig,master=win)
        cv.get_tk_widget().pack(fill="both",expand=True,padx=4,pady=4)
        tb_f=tk.Frame(win,bg=PANEL2); tb_f.pack(fill="x")
        NavigationToolbar2Tk(cv,tb_f)
        self._draw_raman_chart(fig, key, large=True)
        cv.draw()
        tk.Button(win,text=_L("닫기","Close"),command=win.destroy,
                  bg=BTN,fg=TXT,font=MF,relief="flat",
                  padx=16,pady=4).pack(pady=6)

    def _draw_raman_chart(self, fig: plt.Figure, key: str, large: bool):
        """Raman 데이터 자체 분석 차트 — 이미지 분석과 무관"""
        fs_t = 11 if large else 9
        fs_a = 9  if large else 7
        fs_k = 8  if large else 6
        lw   = 2.0 if large else 1.5
        ms   = 7   if large else 4

        rd = self._raman_data
        fig.clear()
        fig.patch.set_facecolor(PANEL)

        def _no_data(msg="No Raman data.\nLoad Excel file."):
            ax = fig.add_subplot(111)
            ax.set_facecolor(PANEL2)
            for sp in ax.spines.values(): sp.set_color(BORDER)
            ax.text(0.5, 0.5, msg,
                    transform=ax.transAxes,
                    ha="center", va="center",
                    color=SUB, fontsize=fs_a, linespacing=1.6)
            ax.set_xticks([]); ax.set_yticks([])

        if not rd:
            _no_data(); return

        conds = list(dict.fromkeys(r["cond"] for r in rd))

        def df(d):
            try: return float(d)
            except: return 9999

        # ── 차트 1: Raman 피크 정규화 추이 ────────────────────────────
        if key == "raman_trend":
            ax = fig.add_subplot(111)
            ax.set_facecolor(PANEL2)
            fig.patch.set_facecolor(PANEL)
            for sp in ax.spines.values(): sp.set_color(BORDER)
            ax.tick_params(colors=SUB, labelsize=fs_k)

            has_data = False
            for ci, cond in enumerate(conds):
                pts = sorted(
                    [(df(r["day"]), r.get("norm_peak", r["peak"]))
                     for r in rd if r["cond"] == cond
                     and df(r["day"]) != 9999],
                    key=lambda x: x[0])
                if not pts: continue
                has_data = True
                days = [p[0] for p in pts]
                peaks = [p[1] for p in pts]
                col = COND_COLORS[ci % len(COND_COLORS)]
                ax.plot(days, peaks, "o-",
                        color=col, lw=lw, ms=ms,
                        label=cond[:14])
                for d, p in zip(days, peaks):
                    ax.annotate(f"{p:.2f}",
                                xy=(d, p),
                                xytext=(0, 6),
                                textcoords="offset points",
                                ha="center", fontsize=fs_k-1,
                                color=col)

            if not has_data:
                _no_data("No day data found.\nCheck day mapping."); return

            ax.set_xlabel("Day", color=SUB, fontsize=fs_a)
            ax.set_ylabel("Raman Peak (normalized)", color=SUB, fontsize=fs_a)
            ax.set_title("Raman A₁g Peak Trend", color=TXT, fontsize=fs_t)
            ax.legend(fontsize=fs_k, framealpha=0.8, edgecolor=BORDER)
            ax.set_ylim(bottom=0)
            fig.tight_layout(pad=0.8)

        # ── 차트 2: 스펙트럼 오버레이 ──────────────────────────────────
        elif key == "raman_spectrum":
            # 스펙트럼 있는 항목만
            spec_rows = [r for r in rd if r.get("spectrum")]
            if not spec_rows:
                _no_data("No spectrum data.\nUse Smart Excel Load\nto import full spectra.")
                return

            ax = fig.add_subplot(111)
            ax.set_facecolor(PANEL2)
            fig.patch.set_facecolor(PANEL)
            for sp in ax.spines.values(): sp.set_color(BORDER)
            ax.tick_params(colors=SUB, labelsize=fs_k)

            # 조건×날짜별 스펙트럼
            plotted = 0
            for ci, cond in enumerate(conds):
                rows_c = sorted(
                    [r for r in spec_rows if r["cond"] == cond],
                    key=lambda r: df(r["day"]))
                for ri, r in enumerate(rows_c):
                    spec = r["spectrum"]
                    sh = spec["shifts"]
                    iv = spec["intensities"]
                    col = COND_COLORS[ci % len(COND_COLORS)]
                    alpha = 0.9 - ri * 0.12
                    lw2 = lw * (1.0 - ri * 0.08)
                    ls = ["-","--","-.",":"," "][min(ri, 4)]
                    label = f"{cond[:10]} D{r['day']}" if ri == 0 else f"  D{r['day']}"
                    ax.plot(sh, iv,
                            color=col, lw=max(0.6, lw2),
                            ls=ls, alpha=max(0.3, alpha),
                            label=label if plotted < 12 else "")
                    plotted += 1

            ax.set_xlabel("Raman Shift (cm⁻¹)", color=SUB, fontsize=fs_a)
            ax.set_ylabel("Intensity (normalized)", color=SUB, fontsize=fs_a)
            ax.set_title("Raman Spectra Overlay", color=TXT, fontsize=fs_t)
            if plotted <= 12:
                ax.legend(fontsize=fs_k-1, framealpha=0.8,
                          edgecolor=BORDER, ncol=2)
            fig.tight_layout(pad=0.8)

        # ── 차트 3: 조건별 감소율 바차트 ─────────────────────────────
        elif key == "raman_decay":
            ax = fig.add_subplot(111)
            ax.set_facecolor(PANEL2)
            fig.patch.set_facecolor(PANEL)
            for sp in ax.spines.values(): sp.set_color(BORDER)
            ax.tick_params(colors=SUB, labelsize=fs_k)

            decay_info = []
            for cond in conds:
                rows_c = sorted(
                    [r for r in rd if r["cond"] == cond
                     and df(r["day"]) != 9999],
                    key=lambda r: df(r["day"]))
                if len(rows_c) < 2: continue
                p0 = rows_c[0].get("norm_peak", rows_c[0]["peak"])
                pN = rows_c[-1].get("norm_peak", rows_c[-1]["peak"])
                if p0 > 0:
                    decay_pct = (p0 - pN) / p0 * 100
                    d0 = df(rows_c[0]["day"])
                    dN = df(rows_c[-1]["day"])
                    decay_info.append((cond, decay_pct, d0, dN))

            if not decay_info:
                _no_data("Need ≥2 days per condition\nto compute decay rate.")
                return

            decay_info.sort(key=lambda x: x[1], reverse=True)
            x_pos = range(len(decay_info))
            bars = ax.bar(
                x_pos,
                [d[1] for d in decay_info],
                color=[COND_COLORS[i % len(COND_COLORS)]
                       for i in range(len(decay_info))],
                edgecolor=BORDER, linewidth=0.8, alpha=0.85)

            for bar, (cond, pct, d0, dN) in zip(bars, decay_info):
                ax.text(bar.get_x() + bar.get_width()/2,
                        bar.get_height() + 1.5,
                        f"{pct:.0f}%",
                        ha="center", va="bottom",
                        fontsize=fs_k, color=TXT, fontweight="bold")
                ax.text(bar.get_x() + bar.get_width()/2,
                        -4,
                        f"D{d0:.0f}→D{dN:.0f}",
                        ha="center", va="top",
                        fontsize=fs_k-1, color=SUB)

            ax.set_xticks(list(x_pos))
            ax.set_xticklabels(
                [d[0][:12] for d in decay_info],
                rotation=20, ha="right", fontsize=fs_k)
            ax.set_ylabel("Peak Decay (%)", color=SUB, fontsize=fs_a)
            ax.set_title("Raman Peak Decay by Condition",
                         color=TXT, fontsize=fs_t)
            ax.set_ylim(-8, max(d[1] for d in decay_info)*1.2 + 5)
            fig.tight_layout(pad=0.8)

        else:
            _no_data(f"Unknown chart key: {key}")


    def _ai_set_text(self, txt: str):
        self._ai_text.configure(state="normal")
        self._ai_text.delete("1.0","end")
        self._ai_text.insert("end", txt)
        self._ai_text.configure(state="disabled")

    def _copy_ai_text(self):
        txt = self._ai_text.get("1.0","end").strip()
        self.clipboard_clear()
        self.clipboard_append(txt)
        self._set_status("AI 분석 텍스트 복사 완료")

    def _on_ai_mode_change(self):
        """모드 라디오버튼 변경 시 API 키 입력란 표시/숨김"""
        if self._ai_mode.get() == "api":
            self._api_key_frame.pack(side="left", padx=4)
        else:
            self._api_key_frame.pack_forget()

    def _run_ai_analysis(self):
        an = [img for img in self.images
              if img.get("roi") and
              not np.isnan(img.get("lab",{}).get("b",np.nan))]
        rd = self._raman_data

        if not an:
            messagebox.showwarning(_L("주의","Warning"),"이미지 분석 먼저 실행한다.")
            return
        if not rd:
            messagebox.showwarning(_L("주의","Warning"),_L("Raman 데이터를 먼저 입력한다.","Enter Raman data first."))
            return

        mode = self._ai_mode.get()

        if mode == "api":
            api_key = self._api_key_var.get().strip()
            if not api_key:
                messagebox.showwarning(
                    _L("API 키 필요","API Key Required"),
                    "Enter your Claude API key.\n\nWithout an API key, use [📊 Local Stats] mode.")
                return
            # 저장 체크박스 ON 이면 현재 키를 settings 에 갱신
            if self._api_key_save_var.get():
                if self._settings.get("claude_api_key") != api_key:
                    self._settings["claude_api_key"] = api_key
                    save_settings(self._settings)
            self._ai_set_text("⏳ Calling Claude API...")
            self.update_idletasks()
            import threading
            threading.Thread(
                target=self._call_claude_api,
                args=(an, rd, api_key), daemon=True).start()
        else:
            self._set_status("⏳ Running local analysis...")
            self.update_idletasks()
            import threading
            threading.Thread(
                target=self._local_analysis,
                args=(an, rd), daemon=True).start()

    def _on_api_key_save_toggle(self):
        """API 키 저장 체크박스 토글 — ON 이면 즉시 저장, OFF 이면 저장된 키 제거"""
        if self._api_key_save_var.get():
            key = self._api_key_var.get().strip()
            if key:
                self._settings["claude_api_key"] = key
                if save_settings(self._settings):
                    self._set_status(_L("✓ API 키 저장됨", "✓ API key saved"))
                else:
                    self._set_status(_L("⚠ API 키 저장 실패", "⚠ API key save failed"))
        else:
            if "claude_api_key" in self._settings:
                self._settings.pop("claude_api_key", None)
                save_settings(self._settings)
                self._set_status(_L("✓ 저장된 API 키 제거됨", "✓ Saved API key removed"))

    def _build_data_summary(self, an: list, rd: list) -> str:
        """이미지 + Raman 데이터를 LLM 입력용 텍스트로 요약"""
        def df(d):
            try: return float(d)
            except: return 9999

        conds = list(dict.fromkeys(
            [img["cond"] for img in an] + [r["cond"] for r in rd]))

        summary_lines = []
        for cond in conds:
            imgs_c = sorted([img for img in an if img["cond"] == cond],
                            key=lambda x: df(x["day"]))
            rams_c = sorted([r for r in rd if r["cond"] == cond],
                            key=lambda x: df(x["day"]))

            img_rows = []
            for img in imgs_c:
                b = img["lab"]["b"]; s = img["s_mean"]
                yi = img.get("yellowness_idx", np.nan)
                de = img.get("delta_e", np.nan)
                img_rows.append(
                    f"  day={img['day']}: b*={b:.1f}, S={s:.1f}, "
                    f"YI={yi:.0f}, ΔE={de:.1f}"
                    if not np.isnan(de) else
                    f"  day={img['day']}: b*={b:.1f}, S={s:.1f}, YI={yi:.0f}")
            ram_rows = [
                f"  day={r['day']}: 피크={r['peak']:.4f} "
                f"(norm={r.get('norm_peak', np.nan):.3f})"
                for r in rams_c]

            summary_lines.append(
                f"[Condition: {cond}]\n"
                f"Image metrics:\n" + "\n".join(img_rows) + "\n"
                f"Raman peaks:\n" + "\n".join(ram_rows))

        return "\n\n".join(summary_lines)

    # 캐싱되는 정적 system prompt — 매 요청마다 동일
    _CLAUDE_SYSTEM_PROMPT = """당신은 HfS₂ 박막 산화도 분석 전문가이다.
사용자가 제공하는 이미지 분석 지표와 Raman 분광 데이터를 바탕으로 산화도를 평가한다.

=== 지표 설명 ===
- b* (Lab 황색도): 미산화 시편 +40~60, 완전 산화 시 +5~15. 감소할수록 산화 진행.
- S채널: HSI 채도. 노란색 강도. 산화 시 감소.
- YI (Yellowness Index): 산업 표준 황색도. 미산화 50~110, 산화 20~35.
- ΔE: 0일 기준 색차. 값이 클수록 색 변화 큼.
- Raman A₁g 피크: HfS₂ 고유 진동 모드. 산화 시 소멸.

=== 분석 요청 형식 ===
사용자 데이터를 보고 다음 내용을 한국어로 3~5문장으로 분석한다:
1. 조건별 산화 진행 속도 비교 (빠른 순서 명시)
2. 이미지 지표(b*, S, YI)와 Raman 피크 감소 간의 일치/불일치 여부
3. 어떤 이미지 지표가 Raman과 가장 높은 상관관계를 보이는지
4. 실용적 결론: 이미지 분석만으로 Raman 없이 산화도 판단이 가능한지 여부

간결하고 핵심만 담아 연구자가 바로 보고서에 활용할 수 있는 수준으로 작성한다."""

    def _call_claude_api(self, an: list, rd: list, api_key: str = ""):
        """Claude API 호출 (별도 스레드). anthropic SDK + prompt caching 사용."""
        try:
            try:
                import anthropic
            except ImportError:
                self.after(0, lambda: self._ai_set_text(
                    "anthropic SDK 가 설치되지 않았습니다.\n\n"
                    "터미널에서 다음 명령으로 설치:\n"
                    "    pip install anthropic\n\n"
                    "또는 [📊 Local Stats] 모드를 사용하세요."))
                self.after(0, lambda: self._set_status(_L("⚠ anthropic 미설치", "⚠ anthropic not installed")))
                return

            data_summary = self._build_data_summary(an, rd)
            user_msg = f"=== 측정 데이터 ===\n{data_summary}"

            client = anthropic.Anthropic(api_key=api_key)
            response = client.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=1024,
                system=[{
                    "type": "text",
                    "text": self._CLAUDE_SYSTEM_PROMPT,
                    "cache_control": {"type": "ephemeral"},
                }],
                messages=[{"role": "user", "content": user_msg}],
            )

            text = next((b.text for b in response.content if b.type == "text"), "")
            if not text:
                text = "(응답에 텍스트 블록이 없습니다.)"

            # 캐시 사용량 표시 (디버깅 / 비용 가시화)
            usage = getattr(response, "usage", None)
            if usage is not None:
                cr = getattr(usage, "cache_read_input_tokens", 0) or 0
                cw = getattr(usage, "cache_creation_input_tokens", 0) or 0
                if cr or cw:
                    text = text + f"\n\n[cache: read={cr} write={cw} tokens]"

            self.after(0, lambda: self._ai_set_text(text))
            self.after(0, lambda: self._set_status(_L("✓ AI 분석 완료", "✓ AI analysis complete")))

        except anthropic.AuthenticationError:
            self.after(0, lambda: self._ai_set_text(
                "API 키가 유효하지 않습니다.\n\n"
                "console.anthropic.com 에서 키를 확인하세요."))
            self.after(0, lambda: self._set_status(_L("⚠ 인증 실패", "⚠ Auth failed")))
        except anthropic.RateLimitError:
            self.after(0, lambda: self._ai_set_text(
                "API 호출 한도 초과 (Rate limit).\n\n"
                "잠시 후 재시도하거나 [📊 Local Stats] 모드를 사용하세요."))
            self.after(0, lambda: self._set_status(_L("⚠ 한도 초과", "⚠ Rate limit")))
        except anthropic.NotFoundError:
            self.after(0, lambda: self._ai_set_text(
                "모델을 찾을 수 없습니다 (404). 모델 ID 가 변경되었을 수 있습니다.\n\n"
                "코드의 model='claude-sonnet-4-6' 부분을 최신 모델 ID 로 업데이트하세요."))
            self.after(0, lambda: self._set_status(_L("⚠ 모델 없음", "⚠ Model not found")))
        except anthropic.APIConnectionError:
            self.after(0, lambda: self._ai_set_text(
                "네트워크 연결 오류. 인터넷 연결을 확인하세요."))
            self.after(0, lambda: self._set_status(_L("⚠ 연결 실패", "⚠ Connection failed")))
        except anthropic.APIStatusError as e:
            err = f"API 오류 (HTTP {e.status_code}): {e.message}"
            self.after(0, lambda m=err: self._ai_set_text(m))
            self.after(0, lambda: self._set_status(_L("⚠ API 오류", "⚠ API error")))
        except Exception as ex:
            self.after(0, lambda m=str(ex): self._ai_set_text(
                f"예상치 못한 오류: {m}\n\n[📊 Local Stats] 모드로 전환 후 재시도하세요."))
            self.after(0, lambda: self._set_status(_L("⚠ API 호출 실패", "⚠ API call failed")))

    def _local_analysis(self, an: list, rd: list):
        """로컬 통계 기반 텍스트 분석 — 한/영 지원"""
        try:
            def df(d):
                try: return float(d)
                except: return 9999

            lines = [_L("📊 이미지-Raman 비교 분석 결과 (로컬 통계)",
                         "📊 Image-Raman Comparison (Local Stats)") + "\n",
                     "─" * 50]

            conds = list(dict.fromkeys(
                [img["cond"] for img in an]+[r["cond"] for r in rd]))

            decay_info = []
            for cond in conds:
                imgs_c = sorted([img for img in an if img["cond"]==cond], key=lambda x: df(x["day"]))
                rams_c = sorted([r  for r  in rd  if r["cond"]==cond],   key=lambda x: df(x["day"]))
                b_d = np.nan; r_d = np.nan
                if len(imgs_c)>=2 and imgs_c[0]["lab"]["b"]>0:
                    b_d=(imgs_c[0]["lab"]["b"]-imgs_c[-1]["lab"]["b"])/imgs_c[0]["lab"]["b"]*100
                if len(rams_c)>=2 and rams_c[0]["peak"]>0:
                    r_d=(rams_c[0]["peak"]-rams_c[-1]["peak"])/rams_c[0]["peak"]*100
                decay_info.append((cond,b_d,r_d))

            sorted_d = sorted([d for d in decay_info if not np.isnan(d[1])], key=lambda x:x[1], reverse=True)
            if sorted_d:
                lines.append("\n"+_L("■ 산화 진행 속도 (b* 기준, 빠른 순)","■ Oxidation Rate (b* basis, fastest first)"))
                for i,(cond,bd,rd_) in enumerate(sorted_d):
                    rstr = (f", Raman {_L('감소율','decay')} {rd_:.1f}%" if not np.isnan(rd_) else "")
                    lines.append(f"  {i+1}. {cond}: b* {_L('감소율','decay')} {bd:.1f}%{rstr}")

            matched = []
            for img in an:
                r_m = next((r for r in rd if r["cond"]==img["cond"] and r["day"]==img["day"]),None)
                if r_m:
                    matched.append((img["lab"]["b"],img["s_mean"],r_m.get("norm_peak",np.nan)))
            if len(matched)>=3:
                b_vals=[m[0] for m in matched]; s_vals=[m[1] for m in matched]
                r_vals=[m[2] for m in matched if not np.isnan(m[2])]
                if len(r_vals)==len(b_vals):
                    corr_b=np.corrcoef(b_vals,r_vals)[0,1]
                    corr_s=np.corrcoef(s_vals,r_vals)[0,1]
                    lines.append("\n"+_L("■ 이미지 지표 vs Raman 상관계수 (R)","■ Image Metric vs Raman Correlation (R)"))
                    lines.append(f"  • Lab b*   : R = {corr_b:.3f}")
                    lines.append(f"  • {_L('S채널','S-ch'):6s}  : R = {corr_s:.3f}")
                    best = "Lab b*" if abs(corr_b)>abs(corr_s) else _L("S채널","S-ch")
                    lines.append(_L(f"  → {best}가 Raman과 더 높은 상관관계를 보이다.",
                                    f"  → {best} shows higher correlation with Raman."))

            lines.append("\n"+_L("■ 종합 판단","■ Overall Assessment"))
            valid = [d[1] for d in decay_info if not np.isnan(d[1])]
            if valid:
                max_b = max(valid)
                if max_b > 50:
                    lines.append(_L("  일부 조건에서 b* 감소율 50% 이상 — 뚜렷한 산화 확인",
                                    "  Some conditions >50% b* decay — significant oxidation confirmed."))
                elif max_b > 20:
                    lines.append(_L("  중간 수준의 산화가 관찰된다.","  Moderate oxidation observed."))
                else:
                    lines.append(_L("  산화 진행이 아직 초기 단계이다.","  Oxidation is still at an early stage."))

            result_text = "\n".join(lines)
            self.after(0, lambda t=result_text: self._ai_set_text(t))
            self.after(0, lambda: self._set_status(
                _L("✓ 로컬 통계 분석 완료","✓ Local stats analysis complete")))

        except Exception as ex:
            self.after(0, lambda m=str(ex): self._ai_set_text(
                _L(f"분석 오류: {m}",f"Analysis error: {m}")))


    def _build_color_tab(self):
        f = self._tfs["color"]

        # 안내 헤더
        hdr = tk.Frame(f, bg=PANEL2,
                       highlightbackground=BORDER, highlightthickness=1)
        hdr.pack(fill="x")
        tk.Label(hdr,
                 text=_L("  🔬 컬러 분석  —  Lab b* · ΔE · GLCM 텍스처  |  더블클릭: 확대","  🔬 Color Analysis  —  Lab b* · ΔE · GLCM  |  dbl-click: enlarge"),
                 bg=PANEL2, fg=TXT, font=MFB).pack(side="left", pady=6, padx=8)

        # 2행×3열 그리드
        grid = tk.Frame(f, bg=BG)
        grid.pack(fill="both", expand=True, padx=4, pady=4)
        for c in range(3): grid.columnconfigure(c, weight=1)
        for r in range(2): grid.rowconfigure(r, weight=1)

        self._color_charts = {}
        defs = [
            (0,0,"lab_b",     _L("Lab b* 추이  (황색도 핵심 지표)",     "Lab b* Trend (key yellowness)")),
            (0,1,"delta_e",   _L("ΔE 색차  (0일 기준 변화량)",          "ΔE Color Difference")),
            (0,2,"lab_L",     _L("Lab L* 밝기  (산화→밝아짐)",           "Lab L* Lightness")),
            (1,0,"glcm_con",  _L("GLCM 대비  (텍스처 대비)",             "GLCM Contrast")),
            (1,1,"glcm_eng",  _L("GLCM 에너지/균일도",                   "GLCM Energy")),
            (1,2,"all_trend", _L("지표 종합 비교  (정규화)",              "Normalized Multi-metric")),
        ]
        for row,col,key,title in defs:
            cell = tk.Frame(grid, bg=PANEL,
                            highlightbackground=BORDER, highlightthickness=1)
            cell.grid(row=row, column=col,
                      padx=3, pady=3, sticky="nsew")
            h2 = tk.Frame(cell, bg=PANEL2,
                          highlightbackground=BORDER, highlightthickness=1)
            h2.pack(fill="x")
            tk.Label(h2, text=f"  {title}",
                     bg=PANEL2, fg=TXT, font=LF).pack(
                     side="left", pady=3, padx=6)
            tk.Label(h2, text="⤢", bg=PANEL2, fg=SUB,
                     font=("Segoe UI",9)).pack(side="right", padx=4)

            fig = plt.Figure(figsize=(4.2, 2.6), facecolor=PANEL)
            cv  = FigureCanvasTkAgg(fig, master=cell)
            cv.get_tk_widget().pack(
                fill="both", expand=True, padx=2, pady=(0,0))
            cv.get_tk_widget().bind(
                "<Double-Button-1>",
                lambda e, k=key, t=title: self._popup_color(k, t))

            # 코멘트 — 선택/복사 가능한 Text 위젯
            cmt_txt = tk.Text(
                cell, height=3, wrap="word",
                bg=CARD2, fg=SUB,
                font=("Segoe UI", 7),
                relief="flat", padx=6, pady=3,
                highlightthickness=0,
                cursor="xterm",
                state="disabled")
            cmt_txt.pack(fill="x", padx=2, pady=(0,2))
            cmt_txt.bind("<Control-c>", lambda e, w=cmt_txt: self._copy_text(w))
            cmt_txt.bind("<Control-C>", lambda e, w=cmt_txt: self._copy_text(w))

            self._color_charts[key] = {
                "fig":fig, "cv":cv,
                "title":title, "cmt_widget":cmt_txt}

        # 하단: 수치 테이블
        tbl_f = tk.Frame(f, bg=PANEL,
                         highlightbackground=BORDER, highlightthickness=1)
        tbl_f.pack(fill="x", padx=4, pady=(0,4))
        tk.Label(tbl_f, text=_L("  📋 지표 수치 테이블","  📋 Metric Data Table"),
                 bg=PANEL2, fg=TXT, font=MFB,
                 highlightbackground=BORDER,
                 highlightthickness=1).pack(
                 fill="x", pady=(0,2))

        col_defs = ("cond","day","lab_b","yi","delta_e",
                    "glcm_con","glcm_eng","glcm_hom")
        self._color_tree = ttk.Treeview(
            tbl_f, columns=col_defs, show="headings", height=5)
        hdrs = [("cond","Cond",80),("day","Day",38),
                ("lab_b","b*",52),("yi","YI",52),
                ("delta_e","ΔE",52),("glcm_con","GLCM-Con",68),
                ("glcm_eng","GLCM-Eng",72),("glcm_hom","GLCM-Hom",72)]
        for cid,lbl2,w in hdrs:
            self._color_tree.heading(cid, text=lbl2)
            self._color_tree.column(cid, width=w, anchor="center")
        self._color_tree.pack(fill="x", padx=4, pady=4)

    def _popup_color(self, key: str, title: str):
        """컬러 분석 차트 확대 팝업"""
        win = tk.Toplevel(self)
        win.title(title); win.configure(bg=PANEL)
        sw=self.winfo_screenwidth(); sh=self.winfo_screenheight()
        ww,wh=int(sw*0.80),int(sh*0.80)
        win.geometry(f"{ww}x{wh}+{(sw-ww)//2}+{(sh-wh)//2}")
        tk.Label(win,text=f"  {title}",bg=PANEL2,fg=TXT,
                 font=("Segoe UI",11,"bold"),
                 highlightbackground=BORDER,
                 highlightthickness=1).pack(fill="x")
        fig=plt.Figure(facecolor=PANEL)
        cv =FigureCanvasTkAgg(fig,master=win)
        cv.get_tk_widget().pack(fill="both",expand=True,padx=4,pady=(4,0))
        tb_f=tk.Frame(win,bg=PANEL2); tb_f.pack(fill="x")
        NavigationToolbar2Tk(cv,tb_f)
        self._draw_color_chart(fig,key,large=True)
        cv.draw()

        # 코멘트 — 팝업에도 표시, 선택/복사 가능
        cell = self._color_charts.get(key, {})
        cmt_text = cell.get("cmt_widget")
        cmt_content = (cmt_text.get("1.0","end").strip()
                       if cmt_text else "")
        if cmt_content:
            popup_cmt = tk.Text(
                win, height=3, wrap="word",
                bg=CARD2, fg=TXT,
                font=("Segoe UI",10),
                relief="flat", padx=12, pady=6,
                highlightbackground=BORDER,
                highlightthickness=1,
                cursor="xterm",
                state="disabled")
            popup_cmt.pack(fill="x", padx=6, pady=(4,0))
            self._set_cmt(popup_cmt, cmt_content)
            popup_cmt.bind("<Control-c>",
                lambda e, w=popup_cmt: self._copy_text(w))
            popup_cmt.bind("<Control-C>",
                lambda e, w=popup_cmt: self._copy_text(w))

        tk.Button(win,text=_L("닫기","Close"),command=win.destroy,
                  bg=BTN,fg=TXT,font=MF,relief="flat",
                  padx=16,pady=4).pack(pady=6)

    def _refresh_color_tab(self):
        """컬러 분석 탭 전체 갱신"""
        an = [img for img in self.images
              if img.get("roi") and
              not np.isnan(img.get("lab",{}).get("b", np.nan))]
        for key,cell in self._color_charts.items():
            cell["fig"].clear()
            self._draw_color_chart(cell["fig"],key,large=False)
            cell["cv"].draw()
            cmt = self._make_chart_comment(key, an)
            self._set_cmt(cell["cmt_widget"], cmt)
        self._refresh_color_table()

    def _refresh_color_table(self):
        """수치 테이블 갱신"""
        self._color_tree.delete(*self._color_tree.get_children())
        an = [img for img in self.images if img.get("roi")
              and not np.isnan(img.get("lab",{}).get("b",np.nan))]
        def df(d):
            try: return float(d)
            except: return 9999
        for img in sorted(an, key=lambda x:(x["cond"],df(x["day"]))):
            b  = img.get("lab",{}).get("b",np.nan)
            yi = img.get("yellowness_idx",np.nan)
            de = img.get("delta_e",np.nan)
            gc = img.get("glcm",{})
            # 색상 태그
            tag = ("g" if (not np.isnan(b) and b>=self.cfg_b_good.get()) else
                   "w" if (not np.isnan(b) and b>=self.cfg_b_warn.get()) else "r")
            self._color_tree.insert("","end",tags=(tag,),values=(
                img["cond"][:14], img["day"],
                f"{b:.1f}"  if not np.isnan(b)  else "-",
                f"{yi:.0f}" if not np.isnan(yi) else "-",
                f"{de:.1f}" if not np.isnan(de) else "-",
                f"{gc.get('contrast',np.nan):.0f}"    if not np.isnan(gc.get('contrast',np.nan))    else "-",
                f"{gc.get('energy',np.nan):.4f}"      if not np.isnan(gc.get('energy',np.nan))      else "-",
                f"{gc.get('homogeneity',np.nan):.4f}" if not np.isnan(gc.get('homogeneity',np.nan)) else "-",
            ))
        self._color_tree.tag_configure("g", foreground=GREEN)
        self._color_tree.tag_configure("w", foreground=AMBER)
        self._color_tree.tag_configure("r", foreground=RED)

    def _draw_color_chart(self, fig: plt.Figure, key: str, large: bool):
        """컬러 분석 개별 차트 그리기"""
        an = [img for img in self.images if img.get("roi")
              and not np.isnan(img.get("lab",{}).get("b",np.nan))]
        fs_t = 12 if large else 9
        fs_a = 10 if large else 8
        fs_k = 9  if large else 7
        lw   = 2.5 if large else 1.8
        ms   = 8   if large else 5
        fig.patch.set_facecolor(PANEL)

        def no_data():
            ax=fig.add_subplot(111); ax.axis("off")
            ax.set_facecolor(PANEL)
            ax.text(0.5,0.5,_L("분석 후 표시된다","Run analysis first"),
                    transform=ax.transAxes,ha="center",va="center",
                    color=SUB,fontsize=10)

        if not an: no_data(); return

        conds = list(dict.fromkeys(img["cond"] for img in an))
        def df(d):
            try: return float(d)
            except: return 9999

        def sa(ax):
            styled_ax(ax, PANEL)
            ax.title.set_color(TXT)
            ax.xaxis.label.set_color(SUB)
            ax.yaxis.label.set_color(SUB)

        # ── Lab b* 추이 ─────────────────────────────
        if key == "lab_b":
            ax=fig.add_subplot(111); sa(ax)
            ax.set_title(_L("Lab b*  (황색도, ↓=산화)",
                            "Lab b*  (yellow, ↓=oxidized)"),
                         fontsize=fs_t,pad=6)
            ax.set_xlabel("Day",fontsize=fs_a)
            ax.set_ylabel("b*",fontsize=fs_a)
            ax.axhspan(30,70,alpha=0.06,color=GOLD,label=_L("황색 정상 범위","Normal range"))
            for ci,cond in enumerate(conds):
                col=COND_COLORS[ci%len(COND_COLORS)]
                pts=sorted([(img["day"],img["lab"]["b"])
                             for img in an if img["cond"]==cond],
                           key=lambda x:df(x[0]))
                if pts:
                    xs=[df(p[0]) for p in pts]
                    ys=[p[1] for p in pts]
                    ax.plot(xs,ys,"o-",color=col,lw=lw,ms=ms,label=cond)
                    for x,y in zip(xs,ys):
                        ax.annotate(f"{y:.1f}",(x,y),
                                    xytext=(0,6),textcoords="offset points",
                                    ha="center",fontsize=fs_k,color=col)
            ax.legend(fontsize=fs_k,framealpha=0.8,edgecolor=BORDER)
            fig.tight_layout(pad=1.2)

        # ── ΔE 색차 ─────────────────────────────────
        elif key == "delta_e":
            ax=fig.add_subplot(111); sa(ax)
            ax.set_title(_L("ΔE 색차  (0일 기준, ↑=변화 큼)",
                            "ΔE Color Diff  (from day0, ↑=more change)"),
                         fontsize=fs_t,pad=6)
            ax.set_xlabel("Day",fontsize=fs_a)
            ax.set_ylabel("ΔE",fontsize=fs_a)
            # 참고선
            ax.axhline(y=3,  color=AMBER,lw=0.8,ls="--",
                       label=_L("ΔE=3 (인지 가능)","ΔE=3 (perceptible)"))
            ax.axhline(y=10, color=RED,  lw=0.8,ls="--",
                       label=_L("ΔE=10 (큰 변화)","ΔE=10 (large change)"))
            for ci,cond in enumerate(conds):
                col=COND_COLORS[ci%len(COND_COLORS)]
                pts=sorted(
                    [(img["day"],img.get("delta_e",np.nan))
                     for img in an
                     if img["cond"]==cond
                     and not np.isnan(img.get("delta_e",np.nan))],
                    key=lambda x:df(x[0]))
                if pts:
                    xs=[df(p[0]) for p in pts]
                    ys=[p[1] for p in pts]
                    ax.plot(xs,ys,"^-",color=col,lw=lw,ms=ms,label=cond)
                    for x,y in zip(xs,ys):
                        ax.annotate(f"{y:.1f}",(x,y),
                                    xytext=(0,6),textcoords="offset points",
                                    ha="center",fontsize=fs_k,color=col)
            ax.legend(fontsize=fs_k,framealpha=0.8,edgecolor=BORDER)
            fig.tight_layout(pad=1.2)

        # ── Lab L* 밝기 ──────────────────────────────
        elif key == "lab_L":
            ax=fig.add_subplot(111); sa(ax)
            ax.set_title(_L("Lab L*  (밝기, ↑=산화→흰색)",
                            "Lab L*  (lightness, ↑=oxidized)"),
                         fontsize=fs_t,pad=6)
            ax.set_xlabel("Day",fontsize=fs_a)
            ax.set_ylabel("L*",fontsize=fs_a)
            for ci,cond in enumerate(conds):
                col=COND_COLORS[ci%len(COND_COLORS)]
                pts=sorted([(img["day"],img["lab"]["L"])
                             for img in an if img["cond"]==cond],
                           key=lambda x:df(x[0]))
                if pts:
                    xs=[df(p[0]) for p in pts]
                    ys=[p[1] for p in pts]
                    ax.plot(xs,ys,"s-",color=col,lw=lw,ms=ms,label=cond)
                    for x,y in zip(xs,ys):
                        ax.annotate(f"{y:.1f}",(x,y),
                                    xytext=(0,6),textcoords="offset points",
                                    ha="center",fontsize=fs_k,color=col)
            ax.legend(fontsize=fs_k,framealpha=0.8,edgecolor=BORDER)
            fig.tight_layout(pad=1.2)

        # ── GLCM 대비 ────────────────────────────────
        elif key == "glcm_con":
            ax=fig.add_subplot(111); sa(ax)
            ax.set_title(_L("GLCM 대비  (산화→균일해짐→감소)",
                            "GLCM Contrast  (↓ as surface oxidizes)"),
                         fontsize=fs_t,pad=6)
            ax.set_xlabel("Day",fontsize=fs_a)
            ax.set_ylabel(_L("대비","Contrast"),fontsize=fs_a)
            for ci,cond in enumerate(conds):
                col=COND_COLORS[ci%len(COND_COLORS)]
                pts=sorted(
                    [(img["day"],img.get("glcm",{}).get("contrast",np.nan))
                     for img in an
                     if img["cond"]==cond
                     and not np.isnan(img.get("glcm",{}).get("contrast",np.nan))],
                    key=lambda x:df(x[0]))
                if pts:
                    xs,ys=zip(*pts)
                    ax.plot(xs,ys,"D-",color=col,lw=lw,ms=ms,label=cond)
            ax.legend(fontsize=fs_k,framealpha=0.8,edgecolor=BORDER)
            fig.tight_layout(pad=1.2)

        # ── GLCM 에너지 ──────────────────────────────
        elif key == "glcm_eng":
            ax=fig.add_subplot(111); sa(ax)
            ax2=ax.twinx()
            ax.set_title(_L("GLCM 에너지 & 균일도  (↑=표면 균일)",
                            "GLCM Energy & Homogeneity"),
                         fontsize=fs_t,pad=6)
            ax.set_xlabel("Day",fontsize=fs_a)
            ax.set_ylabel(_L("에너지","Energy"),fontsize=fs_a)
            ax2.set_ylabel(_L("균일도","Homogeneity"),
                           fontsize=fs_a,color=TEAL)
            ax2.tick_params(colors=TEAL,labelsize=fs_k)
            for ci,cond in enumerate(conds):
                col=COND_COLORS[ci%len(COND_COLORS)]
                pts_e=sorted(
                    [(img["day"],img.get("glcm",{}).get("energy",np.nan))
                     for img in an
                     if img["cond"]==cond
                     and not np.isnan(img.get("glcm",{}).get("energy",np.nan))],
                    key=lambda x:df(x[0]))
                pts_h=sorted(
                    [(img["day"],img.get("glcm",{}).get("homogeneity",np.nan))
                     for img in an
                     if img["cond"]==cond
                     and not np.isnan(img.get("glcm",{}).get("homogeneity",np.nan))],
                    key=lambda x:df(x[0]))
                if pts_e:
                    xs,ys=zip(*pts_e)
                    ax.plot(xs,ys,"o-",color=col,lw=lw,ms=ms,
                            label=f"{cond} E")
                if pts_h:
                    xs,ys=zip(*pts_h)
                    ax2.plot(xs,ys,"s--",color=col,lw=lw,ms=ms,
                             alpha=0.7,label=f"{cond} H")
            lines1,lab1=ax.get_legend_handles_labels()
            lines2,lab2=ax2.get_legend_handles_labels()
            ax.legend(lines1+lines2,lab1+lab2,
                      fontsize=fs_k,framealpha=0.8,
                      edgecolor=BORDER,ncol=2)
            fig.tight_layout(pad=1.2)

        # ── 종합 비교 (정규화) ───────────────────────
        elif key == "all_trend":
            ax=fig.add_subplot(111); sa(ax)
            ax.set_title(_L("지표 종합  (0~1 정규화, 높을수록 미산화(pristine))",
                            "All Metrics Normalized (higher=fresher)"),
                         fontsize=fs_t,pad=6)
            ax.set_xlabel("Day",fontsize=fs_a)
            ax.set_ylabel(_L("정규화 값","Normalized"),fontsize=fs_a)
            ax.set_ylim(-0.05,1.10)

            metric_defs = [
                ("lab_b_norm",   "b*",       "o-",  COND_COLORS[0]),
                ("yi_norm",      "YI",        "s--", COND_COLORS[1]),
                ("delta_e_inv",  "1-ΔE/max", "^:",  COND_COLORS[2]),
                ("S_norm",       "S-ch",      "D-",  COND_COLORS[3]),
            ]

            def normalize(vals):
                v = np.array(vals, dtype=float)
                v = v[~np.isnan(v)]
                if len(v)<2: return None, None, None
                return v.min(), v.max(), True

            for ci,cond in enumerate(conds):
                imgs_c=sorted([i for i in an if i["cond"]==cond],
                               key=lambda x:df(x["day"]))
                if not imgs_c: continue
                col=COND_COLORS[ci%len(COND_COLORS)]
                days_c=[i["day"] for i in imgs_c]

                # b* 정규화
                b_vals=[i["lab"]["b"] for i in imgs_c]
                b_min,b_max=min(b_vals),max(b_vals)
                b_r=b_max-b_min if b_max!=b_min else 1
                b_norm=[(v-b_min)/b_r for v in b_vals]

                # S채널 정규화
                s_vals=[i["s_mean"] for i in imgs_c]
                s_min,s_max=min(s_vals),max(s_vals)
                s_r=s_max-s_min if s_max!=s_min else 1
                s_norm=[(v-s_min)/s_r for v in s_vals]

                ax.plot(days_c,b_norm,"o-",color=col,lw=lw,ms=ms,
                        label=f"{cond[:8]} b*")
                ax.plot(days_c,s_norm,"s--",color=col,lw=lw*0.7,ms=ms*0.7,
                        alpha=0.6,label=f"{cond[:8]} S")

            ax.legend(fontsize=fs_k,framealpha=0.8,
                      edgecolor=BORDER,ncol=2)
            fig.tight_layout(pad=1.2)

    # ─────────────────────────────────────────
    #  3. 차트 탭 — H/S/I 통합 + 더블클릭 확대
    # ─────────────────────────────────────────
    def _build_chart(self):
        f = self._tfs["chart"]

        # ── 임계값 빠른 설정 바 ─────────────────────
        thresh_f = tk.Frame(f, bg=PANEL2,
                            highlightbackground=BORDER, highlightthickness=1)
        thresh_f.pack(fill="x")
        tk.Label(thresh_f,
                 text=_L("  ⚙ 임계값 (즉시 재계산):",
                          "  ⚙ Thresholds (recalc now):"),
                 bg=PANEL2, fg=TXT, font=MFB).pack(side="left", pady=5, padx=8)

        # H 범위
        tk.Label(thresh_f, text=_L("H하한:","H-lo:"),
                 bg=PANEL2, fg=TEAL, font=LF).pack(side="left", padx=(8,1))
        tk.Spinbox(thresh_f, textvariable=self.cfg_h_lo,
                   from_=0, to=90, increment=1, width=4,
                   bg=PANEL, fg=TXT, relief="flat", font=LF).pack(side="left")
        tk.Label(thresh_f, text=_L("H상한:","H-hi:"),
                 bg=PANEL2, fg=TEAL, font=LF).pack(side="left", padx=(6,1))
        tk.Spinbox(thresh_f, textvariable=self.cfg_h_hi,
                   from_=30, to=120, increment=1, width=4,
                   bg=PANEL, fg=TXT, relief="flat", font=LF).pack(side="left")
        tk.Label(thresh_f, text=_L("S임계:","S-thr:"),
                 bg=PANEL2, fg=PURPLE, font=LF).pack(side="left", padx=(6,1))
        tk.Spinbox(thresh_f, textvariable=self.cfg_s_thresh,
                   from_=0.01, to=0.5, increment=0.01, width=5,
                   format="%.2f",
                   bg=PANEL, fg=TXT, relief="flat", font=LF).pack(side="left")

        # S 판정기준
        tk.Label(thresh_f, text=_L("S양호≥:","S-good≥:"),
                 bg=PANEL2, fg=GREEN, font=LF).pack(side="left", padx=(8,1))
        tk.Spinbox(thresh_f, textvariable=self.cfg_s_good,
                   from_=10, to=255, increment=5, width=4,
                   bg=PANEL, fg=TXT, relief="flat", font=LF).pack(side="left")
        tk.Label(thresh_f, text=_L("경고≥:","warn≥:"),
                 bg=PANEL2, fg=AMBER, font=LF).pack(side="left", padx=(4,1))
        tk.Spinbox(thresh_f, textvariable=self.cfg_s_warn,
                   from_=5, to=200, increment=5, width=4,
                   bg=PANEL, fg=TXT, relief="flat", font=LF).pack(side="left")

        # 재계산 버튼
        tk.Button(thresh_f,
                  text=_L("▶ 재계산","▶ Recalc"),
                  command=lambda: self._recalc_from_chart(),
                  bg=ACCENT, fg="white", font=MFB,
                  relief="flat", padx=10, pady=3,
                  cursor="hand2").pack(side="right", padx=8, pady=4)

        tk.Label(f,
                 text=_L("  📈 채널별 차트  —  더블클릭: 크게 보기",
                          "  📈 Channel Charts  —  dbl-click: enlarge"),
                 bg=BG, fg=SUB, font=MFB).pack(anchor="w", pady=(4,2))

        grid = tk.Frame(f, bg=BG)
        grid.pack(fill="both", expand=True)
        for c in range(3): grid.columnconfigure(c, weight=1)
        for r in range(2): grid.rowconfigure(r, weight=1)

        self._charts = {}
        defs = [
            (0, 0, "s_trend",   _L("S채널 평균 추이",    "S-ch Mean")),
            (0, 1, "yr_trend",  _L("황색 잔존 비율 (%)", "Yellow Ratio")),
            (0, 2, "h_trend",   _L("H채널 추이",         "H-ch Mean")),
            (1, 0, "box",       _L("세그먼트 분포",       "Seg Distribution")),
            (1, 1, "hist",      _L("채널 히스토그램",     "Histogram")),
            (1, 2, "decay",     _L("조건별 감소율",       "Decay Rate")),
        ]
        for row,col,key,title in defs:
            cell = tk.Frame(grid, bg=PANEL,
                            highlightbackground=BORDER,
                            highlightthickness=1)
            cell.grid(row=row, column=col,
                      padx=3, pady=3, sticky="nsew")

            hdr2 = tk.Frame(cell, bg=PANEL2,
                            highlightbackground=BORDER,
                            highlightthickness=1)
            hdr2.pack(fill="x")
            tk.Label(hdr2, text=f"  {title}",
                     bg=PANEL2, fg=TXT, font=LF).pack(
                     side="left", pady=3, padx=6)
            tk.Label(hdr2, text="⤢",
                     bg=PANEL2, fg=SUB,
                     font=("Segoe UI",9)).pack(side="right", padx=4)

            fig = plt.Figure(figsize=(4.2, 2.6), facecolor=PANEL)
            cv  = FigureCanvasTkAgg(fig, master=cell)
            cv.get_tk_widget().pack(
                fill="both", expand=True, padx=2, pady=(0,0))
            cv.get_tk_widget().bind(
                "<Double-Button-1>",
                lambda e, k=key, t=title: self._popup(k, t))

            # 코멘트 — 선택/복사 가능한 Text 위젯
            cmt_txt = tk.Text(
                cell, height=3, wrap="word",
                bg=CARD2, fg=SUB,
                font=("Segoe UI", 7),
                relief="flat", padx=6, pady=3,
                highlightthickness=0,
                cursor="xterm",
                state="disabled")
            cmt_txt.pack(fill="x", padx=2, pady=(0,2))
            # Ctrl+C 복사 허용 (disabled 상태에서도)
            cmt_txt.bind("<Control-c>", lambda e, w=cmt_txt: self._copy_text(w))
            cmt_txt.bind("<Control-C>", lambda e, w=cmt_txt: self._copy_text(w))

            self._charts[key] = {
                "fig":fig, "cv":cv,
                "title":title, "cmt_widget":cmt_txt}

    @staticmethod
    def _set_cmt(widget: tk.Text, text: str):
        """코멘트 Text 위젯에 텍스트 설정 (disabled 상태 유지)"""
        widget.configure(state="normal")
        widget.delete("1.0", "end")
        widget.insert("end", text)
        widget.configure(state="disabled")

    def _copy_text(self, widget: tk.Text):
        """선택된 텍스트 or 전체 텍스트를 클립보드에 복사"""
        try:
            sel = widget.get(tk.SEL_FIRST, tk.SEL_LAST)
        except tk.TclError:
            sel = widget.get("1.0", "end").strip()
        self.clipboard_clear()
        self.clipboard_append(sel)

    def _popup(self, key: str, title: str):
        win = tk.Toplevel(self)
        win.title(title)
        win.configure(bg=PANEL)
        sw = self.winfo_screenwidth()
        sh = self.winfo_screenheight()
        ww, wh = int(sw*0.80), int(sh*0.80)
        win.geometry(f"{ww}x{wh}+{(sw-ww)//2}+{(sh-wh)//2}")

        tk.Label(win, text=f"  {title}",
                 bg=PANEL2, fg=TXT,
                 font=("Segoe UI",11,"bold"),
                 highlightbackground=BORDER,
                 highlightthickness=1).pack(fill="x")

        fig = plt.Figure(facecolor=PANEL)
        cv  = FigureCanvasTkAgg(fig, master=win)
        cv.get_tk_widget().pack(fill="both", expand=True, padx=4, pady=(4,0))
        tb_f = tk.Frame(win, bg=PANEL2)
        tb_f.pack(fill="x")
        NavigationToolbar2Tk(cv, tb_f)

        self._draw_chart(fig, key, large=True)
        cv.draw()

        # 코멘트 — 팝업에도 표시, 선택/복사 가능
        cell = self._charts.get(key, {})
        cmt_text = cell.get("cmt_widget")
        cmt_content = (cmt_text.get("1.0","end").strip()
                       if cmt_text else "")
        if cmt_content:
            popup_cmt = tk.Text(
                win, height=3, wrap="word",
                bg=CARD2, fg=TXT,
                font=("Segoe UI", 10),
                relief="flat", padx=12, pady=6,
                highlightbackground=BORDER,
                highlightthickness=1,
                cursor="xterm",
                state="disabled")
            popup_cmt.pack(fill="x", padx=6, pady=(4,0))
            self._set_cmt(popup_cmt, cmt_content)
            popup_cmt.bind("<Control-c>",
                lambda e, w=popup_cmt: self._copy_text(w))
            popup_cmt.bind("<Control-C>",
                lambda e, w=popup_cmt: self._copy_text(w))

        tk.Button(win, text=_L("닫기","Close"), command=win.destroy,
                  bg=BTN, fg=TXT, font=MF,
                  relief="flat", padx=16, pady=4).pack(pady=6)

    # ─────────────────────────────────────────
    #  목록 재빌드
    # ─────────────────────────────────────────
    # 카드 썸네일 사이즈 (사용자 요청 — 더 크게)
    _CARD_THUMB_W = 96
    _CARD_THUMB_H = 80

    def _rebuild_list(self):
        for w in self._lf.winfo_children():
            w.destroy()
        # 카드 frame 참조 캐시 — 선택 변경 시 _refresh_card_border 에 사용
        self._cards_by_idx = {}

        total = len(self.images)
        done  = sum(1 for img in self.images if img.get("roi"))
        self._roi_stat.set(f"ROI {done}/{total}")
        self._roi_ok_lbl.configure(
            fg=GREEN if done==total and total>0 else AMBER)

        # 같은 cond 그룹 내 ROI 비율 일관성 점검 — outlier 마킹
        consistency = check_roi_group_consistency(self.images)

        # ROI 품질 우선순위 정렬 — 문제 있는 카드를 위로
        # (self.images 자체는 변경 안 함; 표시 순서만 변경)
        def _quality_priority(i):
            img = self.images[i]
            flag = img.get("roi_flag")
            inconsistent, _ = consistency.get(i, (False, 1.0))
            if flag == "failed":            return 0  # 시편 검출 실패 (수동 필수)
            if flag == "warn_paper":        return 1  # 흰 배경 포함
            if flag == "warn_small":        return 2  # 면적 작음
            if flag == "warn_off":          return 3  # 가장자리 근접
            if not img.get("roi"):          return 4  # ROI 미설정
            if inconsistent:                return 5  # 그룹 불일치 (PURPLE)
            return 9                                  # good / manual / db OK
        display_order = sorted(range(len(self.images)),
                               key=lambda i: (_quality_priority(i), i))

        for idx in display_order:
            img = self.images[idx]
            is_sel  = (idx == self.sel_idx)
            has_roi = img.get("roi") is not None
            done_a  = not np.isnan(img.get("s_mean", np.nan))
            roi_flag = img.get("roi_flag")
            inconsistent, group_iou = consistency.get(idx, (False, 1.0))

            # 카드 테두리: 선택 > failed > warn > inconsistent > good > 기본
            if is_sel:
                brd = ACCENT
            else:
                brd = _border_color_for_roi(roi_flag, has_roi,
                    {"green": GREEN, "amber": AMBER, "red": RED,
                     "purple": PURPLE, "border": BORDER},
                    inconsistent=inconsistent)
            # warning/failed/inconsistent 는 굵게 강조
            thick = 2 if (is_sel
                          or roi_flag in ("warn_small", "warn_off", "warn_paper", "failed")
                          or (inconsistent and roi_flag != "failed"
                              and roi_flag not in ("warn_small", "warn_off", "warn_paper"))) else 1
            card = tk.Frame(self._lf,
                            bg=CARD2 if is_sel else CARD,
                            highlightbackground=brd,
                            highlightthickness=thick,
                            cursor="hand2")
            card.pack(fill="x", padx=5, pady=3)
            self._cards_by_idx[idx] = card
            # 일관성 점검 결과 카드에 보존 (ROI 라벨에서 사용)
            img["_roi_inconsistent"] = inconsistent
            img["_roi_group_iou"] = group_iou
            # 라만 매칭 ⚛ 뱃지 (있으면 우상단)
            if img.get("raman_id") is not None:
                badge = tk.Label(card, text="⚛", fg="#a78bfa",
                                 bg=card["bg"],
                                 font=("Segoe UI", 11, "bold"))
                badge._is_raman_badge = True
                badge.place(relx=1.0, rely=0.0, anchor="ne", x=-6, y=2)

            # 썸네일 + 필드
            top = tk.Frame(card, bg=card["bg"])
            top.pack(fill="x", padx=5, pady=(5,2))

            th = img.get("thumb")
            if th:
                # PhotoImage 캐시 — 같은 thumb 객체면 재사용 (resize 비용 절감)
                cached = img.get("_tk_thumb_cache")
                if (cached is not None
                        and cached[0] is th
                        and cached[1] == (self._CARD_THUMB_W, self._CARD_THUMB_H)):
                    tk_img = cached[2]
                else:
                    tk_img = ImageTk.PhotoImage(
                        th.resize((self._CARD_THUMB_W, self._CARD_THUMB_H),
                                  Image.LANCZOS))
                    img["_tk_thumb_cache"] = (
                        th, (self._CARD_THUMB_W, self._CARD_THUMB_H), tk_img)
                self._refs[f"li_{idx}"] = tk_img
                lbl = tk.Label(top, image=tk_img,
                               bg=card["bg"], cursor="hand2")
                lbl.pack(side="left", padx=(0,5))
                lbl.bind("<Button-1>", lambda e,i=idx: self._select(i))
            else:
                tk.Label(top, text="📷", bg=card["bg"],
                         fg=SUB, font=("Segoe UI",24),
                         width=5).pack(side="left", padx=(0,5))

            flds = tk.Frame(top, bg=card["bg"])
            flds.pack(side="left", fill="both", expand=True)

            # 파일명
            icon = " ✦" if img.get("auto_parsed") else ""
            tk.Label(flds,
                     text=img["name"][:22]+icon,
                     bg=card["bg"],
                     fg=TEAL if img.get("auto_parsed") else SUB,
                     font=LF, anchor="w").pack(anchor="w")

            # DB 로드 이미지에 tkVar가 없으면 생성
            if "day_var" not in img:
                img["day_var"]  = tk.StringVar(value=img.get("day",""))
            if "cond_var" not in img:
                img["cond_var"] = tk.StringVar(value=img.get("cond",""))

            # day/cond 입력
            for lbl_txt, var_key, w_ in [("day:", "day_var", 7),
                                          ("cond:","cond_var",14)]:
                rr = tk.Frame(flds, bg=card["bg"])
                rr.pack(fill="x", pady=1)
                tk.Label(rr, text=lbl_txt, bg=card["bg"],
                         fg=SUB, font=LF, width=5,
                         anchor="w").pack(side="left")
                ent = tk.Entry(rr, textvariable=img[var_key],
                               width=w_, bg=PANEL2, fg=TXT,
                               insertbackground=TXT,
                               font=MF, relief="flat",
                               highlightbackground=BORDER,
                               highlightthickness=1)
                ent.pack(side="left", padx=2)
                ent.bind("<Return>",
                         lambda e,i=idx: self._sync(i))
                ent.bind("<FocusOut>",
                         lambda e,i=idx: self._sync(i))

            # 빠른 조건 버튼
            qr = tk.Frame(flds, bg=card["bg"])
            qr.pack(fill="x", pady=(1,0))
            for lbl2,col in self._presets:
                short = (lbl2.replace("Native-","N-")
                             .replace("Al₂O₃-","Al-"))
                tk.Button(qr, text=short,
                          command=lambda l=lbl2,i=idx: self._set_cond(i,l),
                          bg=card["bg"], fg=col,
                          font=("Segoe UI",6,"bold"),
                          relief="flat", cursor="hand2",
                          padx=2, pady=1).pack(side="left", padx=1)

            # 하단
            bot = tk.Frame(card, bg=card["bg"])
            bot.pack(fill="x", padx=5, pady=(0,5))

            # ROI 상태 + 자동 추정 flag
            if has_roi:
                sym, ko_lbl, en_lbl = _ROI_FLAG_LABEL.get(
                    roi_flag, ("✔", "ROI", "ROI"))
                roi_txt = (f"{sym} {img['roi'][2]-img['roi'][0]}"
                           f"×{img['roi'][3]-img['roi'][1]}")
                if roi_flag == "failed":
                    roi_col = RED
                elif roi_flag in ("warn_small", "warn_off", "warn_paper"):
                    roi_col = AMBER
                elif roi_flag in ("manual", "good"):
                    roi_col = GREEN
                else:
                    roi_col = GREEN
            else:
                roi_txt = "○ No ROI"
                ko_lbl, en_lbl = "", ""
                roi_col = AMBER
            tk.Label(bot, text=roi_txt, bg=card["bg"],
                     fg=roi_col, font=MFB).pack(side="left")
            if ko_lbl:
                tk.Label(bot, text=" " + _L(ko_lbl, en_lbl),
                         bg=card["bg"], fg=roi_col,
                         font=("Segoe UI", 7)).pack(side="left")
            # 같은 조건 그룹 ROI 일관성 — outlier 표시
            if inconsistent and has_roi:
                tk.Label(bot,
                         text=_L(f" ◇ 조건 불일치 (IoU {group_iou*100:.0f}%)",
                                 f" ◇ Group mismatch (IoU {group_iou*100:.0f}%)"),
                         bg=card["bg"], fg=PURPLE,
                         font=("Segoe UI", 7, "bold")).pack(side="left")

            # 🎯 ROI 선택 버튼
            tk.Button(bot, text="🎯",
                      command=lambda i=idx: self._sel_roi(i),
                      bg=card["bg"], fg=TEAL,
                      font=("Segoe UI",10),
                      relief="flat", cursor="hand2",
                      padx=2).pack(side="left", padx=2)

            # ✦ 동일 조건에 ROI 복사 버튼 (ROI 있을 때만 활성)
            cond_copy_btn = tk.Button(
                bot, text=_L("✦조건복사","✦CopyByC"),
                command=lambda i=idx: self._copy_roi_same_cond(i),
                bg=card["bg"],
                fg=PURPLE if has_roi else BORDER,
                font=("Segoe UI",7,"bold"),
                relief="flat", cursor="hand2" if has_roi else "arrow",
                padx=3, pady=1,
                state="normal" if has_roi else "disabled")
            cond_copy_btn.pack(side="left", padx=2)

            # 분석 결과
            if done_a:
                s  = img.get("s_mean", np.nan)
                yr = img.get("yellow_ratio", np.nan)
                yi = img.get("yellowness_idx", np.nan)
                col = self._s_color(s)
                result_txt = f"S={s:.0f}  Y={yr*100:.0f}%"
                if not np.isnan(yi):
                    result_txt += f"  YI={yi:.0f}"
                tk.Label(bot,
                         text=result_txt,
                         bg=card["bg"], fg=col,
                         font=MFB).pack(side="left", padx=4)

            tk.Button(bot, text="✕",
                      command=lambda i=idx: self._remove(i),
                      bg=card["bg"], fg=RED, font=LF,
                      relief="flat", cursor="hand2",
                      padx=3).pack(side="right")

            card.bind("<Button-1>", lambda e,i=idx: self._select(i))

    # ─────────────────────────────────────────
    #  ROI 관련
    # ─────────────────────────────────────────
    def _open_roi(self):
        idx = self.sel_idx
        if idx < 0 or idx >= len(self.images):
            messagebox.showinfo(_L("알림","Info"),_L("이미지를 먼저 선택한다.","Select an image first."))
            return
        self._do_open_roi(idx)

    def _sel_roi(self, idx):
        self._select(idx)
        self._do_open_roi(idx)

    def _do_open_roi(self, idx):
        def on_confirm(roi, target_img=None):
            """roi 확정 — target_img 지정 시 해당 이미지에, 없으면 idx 이미지에"""
            img = target_img if target_img is not None else self.images[idx]
            img["roi"]   = roi
            img["mask"]  = roi_to_mask(img["rgb"].shape, roi)
            img["roi_flag"] = "manual"
            img["roi_reason"] = "사용자 직접 설정"
            img["stats"] = {}
            img["s_mean"] = np.nan
            img["yellow_ratio"] = np.nan
            img["thumb"] = make_thumb(img["rgb"], self.TW, self.TH, roi)

        def on_confirm_and_refresh(roi, target_img=None):
            on_confirm(roi, target_img)
            self._rebuild_list()
            self._refresh_orig()
            confirmed_img = target_img if target_img is not None else self.images[idx]
            self._set_status(
                f"✔ ROI 확정: {confirmed_img['name']}  "
                f"({roi[2]-roi[0]}×{roi[3]-roi[1]}px)")

        def on_copy_all(roi):
            """ROI 좌표를 ROI 없는 모든 이미지에 초기값으로 복사"""
            self._apply_roi_to_others(idx, roi, same_cond_only=False)
            on_confirm(roi)
            self._rebuild_list()
            self._set_status(
                f'⊕ ' + _L("ROI 전체복사 완료","ROI copied to all") + f"  ({roi[2]-roi[0]}×{roi[3]-roi[1]}px)")

        ROISelector(
            self,
            self.images[idx],
            on_confirm_and_refresh,
            on_copy_all,
            images=self.images,
            current_idx=idx,
        )

    def _apply_roi_to_others(self, src_idx: int, roi: tuple,
                              same_cond_only: bool = False):
        """
        ROI를 다른 이미지에 비율(0~1) 기반으로 적용.
        src 이미지의 ROI 비율을 계산해 각 대상 이미지 크기에 맞게 변환.
        → 이미지 크기/해상도가 달라도 화면상 같은 위치에 ROI가 생김.
        ROI 없는 이미지에만 적용 (이미 있는 것은 유지).
        """
        src_img  = self.images[src_idx]
        src_cond = src_img.get("cond_var", tk.StringVar()).get().strip()

        # 원본 이미지 크기 기준으로 비율 계산
        src_h, src_w = src_img["rgb"].shape[:2]
        x0, y0, x1, y1 = roi
        rx0 = x0 / src_w;  ry0 = y0 / src_h
        rx1 = x1 / src_w;  ry1 = y1 / src_h

        for i, other in enumerate(self.images):
            if i == src_idx:
                continue
            if same_cond_only:
                other_cond = other.get("cond_var", tk.StringVar()).get().strip()
                if other_cond != src_cond:
                    continue
            if other.get("roi") is not None:
                continue   # 이미 ROI 있는 것은 건드리지 않음

            # 대상 이미지 크기에 맞게 비율 → 픽셀 변환
            oh, ow = other["rgb"].shape[:2]
            nx0 = max(0, int(rx0 * ow))
            ny0 = max(0, int(ry0 * oh))
            nx1 = min(ow, int(rx1 * ow))
            ny1 = min(oh, int(ry1 * oh))

            if nx1 > nx0 and ny1 > ny0:
                scaled = (nx0, ny0, nx1, ny1)
                other["roi"]   = scaled
                other["mask"]  = roi_to_mask(other["rgb"].shape, scaled)
                other["roi_flag"] = "manual"
                other["roi_reason"] = "Copy All 적용"
                other["stats"] = {}
                other["s_mean"] = np.nan
                other["yellow_ratio"] = np.nan
                other["thumb"] = make_thumb(
                    other["rgb"], self.TW, self.TH, scaled)

    def _copy_roi_same_cond(self, src_idx: int):
        """카드 버튼 — 같은 조건의 ROI 없는 이미지에만 복사"""
        src_img = self.images[src_idx]
        if src_img.get("roi") is None:
            messagebox.showinfo(_L("알림","Info"),
                _L("ROI가 없다. 먼저 🎯 버튼으로 ROI를 선택한다.","No ROI. Use 🎯 to select ROI first."))
            return
        src_cond = src_img["cond_var"].get().strip()
        if not src_cond:
            messagebox.showinfo(_L("알림","Info"),
                _L("조건(cond)이 비어 있다. 먼저 조건을 입력한다.","Condition (cond) is empty. Enter condition first."))
            return
        self._apply_roi_to_others(src_idx, src_img["roi"],
                                   same_cond_only=True)
        self._rebuild_list()
        self._set_status(
            f"✦ 동일 조건 ROI 복사: '{src_cond}'  "
            f"({src_img['roi'][2]-src_img['roi'][0]}"
            f"×{src_img['roi'][3]-src_img['roi'][1]}px)")

    def _clear_roi(self):
        idx = self.sel_idx
        if idx < 0 or idx >= len(self.images): return
        img = self.images[idx]
        img["roi"] = img["mask"] = None
        img["roi_flag"] = None
        img["roi_reason"] = ""
        img["stats"] = {}
        img["s_mean"] = img["yellow_ratio"] = np.nan
        img["thumb"] = make_thumb(img["rgb"], self.TW, self.TH)
        self._rebuild_list()
        self._refresh_orig()
        self._roi_info.set(_L("ROI 제거됨","ROI cleared"))

    # ─────────────────────────────────────────
    #  원본 뷰
    # ─────────────────────────────────────────
    def _refresh_orig(self):
        self._orig_cv.delete("all")
        idx = self.sel_idx
        if idx < 0 or idx >= len(self.images): return
        img = self.images[idx]
        roi = img.get("roi")

        cw  = self._orig_cv.winfo_width()  or 600
        ch_ = self._orig_cv.winfo_height() or 400

        pil  = Image.fromarray(img["rgb"])
        pil.thumbnail((cw, ch_), Image.LANCZOS)
        scale = pil.width / img["rgb"].shape[1]

        if roi:
            dark = pil.point(lambda p: int(p*0.35))
            x0,y0,x1,y1 = [int(v*scale) for v in roi]
            mask_p = Image.new("L", pil.size, 0)
            ImageDraw.Draw(mask_p).rectangle([x0,y0,x1,y1], fill=255)
            pil = Image.composite(pil, dark, mask_p)
            draw = ImageDraw.Draw(pil)
            draw.rectangle([x0,y0,x1,y1], outline=(59,130,246), width=2)
            sz = 5
            for cx,cy in [(x0,y0),(x1,y0),(x0,y1),(x1,y1)]:
                draw.rectangle([cx-sz,cy-sz,cx+sz,cy+sz],
                               fill=(59,130,246))
            w_ = roi[2]-roi[0]; h_ = roi[3]-roi[1]
            draw.text((x0+5, y0+3),
                      f" ROI  {w_}×{h_}px ",
                      fill=(59,130,246))
            self._roi_info.set(
                f"✔ ({roi[0]},{roi[1]})~({roi[2]},{roi[3]})  "
                f"{w_}×{h_}px")
        else:
            self._roi_info.set(_L("ROI 미선택 — [🎯 ROI 선택] 버튼을 누르세요","No ROI — click [🎯 Set ROI] button"))

        tk_img = ImageTk.PhotoImage(pil)
        self._refs["orig"] = tk_img
        ox = (cw - pil.width)  // 2
        oy = (ch_ - pil.height) // 2
        self._orig_cv.create_image(ox, oy, anchor="nw", image=tk_img)
        if not roi:
            self._orig_cv.create_text(
                cw//2, ch_-22, anchor="center",
                text=_L("[🎯 ROI 선택] 버튼을 눌러 시편 영역을 지정한다","Click [🎯 Set ROI] to define the specimen area"),
                fill=AMBER, font=MF)

    # ─────────────────────────────────────────
    #  HSI 뷰
    # ─────────────────────────────────────────
    def _refresh_hsi(self):
        self._hsi_cv.delete("all")
        idx = self.sel_idx
        if idx < 0 or idx >= len(self.images): return
        img = self.images[idx]
        roi = img.get("roi")
        if not roi: return

        H_ch, S_ch, I_ch = img["hsi"]
        ch = self.ch_var.get()

        # ── 채널 데이터 준비 ───────────────────────
        if ch in ("H","S","I"):
            ch_data = {"H":H_ch,"S":S_ch,"I":I_ch}[ch]
        elif ch == "b*":
            # Lab b* 채널: ROI 크롭 후 계산
            x0,y0,x1,y1 = roi
            crop_rgb = img["rgb"][y0:y1, x0:x1]
            bgr = cv2.cvtColor(crop_rgb, cv2.COLOR_RGB2BGR)
            lab = cv2.cvtColor(bgr, cv2.COLOR_BGR2Lab)
            b_real = lab[:,:,2].astype(np.float32) - 128  # -128~127
            # 0~255로 정규화하여 시각화용 uint8 생성
            ch_data = ((b_real + 128) / 255.0 * 255).clip(0,255).astype(np.uint8)
        elif ch == "L*":
            x0,y0,x1,y1 = roi
            crop_rgb = img["rgb"][y0:y1, x0:x1]
            bgr = cv2.cvtColor(crop_rgb, cv2.COLOR_RGB2BGR)
            lab = cv2.cvtColor(bgr, cv2.COLOR_BGR2Lab)
            L_real = lab[:,:,0].astype(np.float32) / 255.0 * 100  # 0~100
            ch_data = (L_real / 100.0 * 255).clip(0,255).astype(np.uint8)
        elif ch == "ΔE":
            # 0일차 기준 색차 맵 (픽셀별)
            x0,y0,x1,y1 = roi
            crop_rgb = img["rgb"][y0:y1, x0:x1]
            # 같은 조건 0일차 이미지 찾기
            ref_img = None
            for other in self.images:
                try:
                    if (other["cond"] == img["cond"] and
                            float(other["day"]) == 0 and
                            other is not img):
                        ref_img = other; break
                except: pass
            if ref_img is not None:
                rx0,ry0,rx1,ry1 = ref_img.get("roi", roi)
                ref_crop = ref_img["rgb"][ry0:ry1, rx0:rx1]
                # 크기 맞추기
                h_,w_ = crop_rgb.shape[:2]
                ref_crop = cv2.resize(ref_crop, (w_, h_))
                bgr1 = cv2.cvtColor(crop_rgb, cv2.COLOR_RGB2BGR)
                bgr2 = cv2.cvtColor(ref_crop, cv2.COLOR_RGB2BGR)
                lab1 = cv2.cvtColor(bgr1, cv2.COLOR_BGR2Lab).astype(np.float32)
                lab2 = cv2.cvtColor(bgr2, cv2.COLOR_BGR2Lab).astype(np.float32)
                dE = np.sqrt(np.sum((lab1-lab2)**2, axis=2))
                # ΔE 0~50 범위를 0~255로 정규화
                ch_data = (dE / 50.0 * 255).clip(0,255).astype(np.uint8)
            else:
                # 0일차 없으면 S채널로 대체
                ch_data = S_ch[roi[1]:roi[3], roi[0]:roi[2]]
        elif ch == "YI":
            x0,y0,x1,y1 = roi
            crop_rgb = img["rgb"][y0:y1, x0:x1].astype(np.float32)/255.0
            R = crop_rgb[:,:,0]; G = crop_rgb[:,:,1]; B = crop_rgb[:,:,2]
            G_safe = np.where(G > 0.01, G, 0.01)
            yi_map = (100.0*(1.28*R - 1.06*B) / G_safe)
            # YI 0~120 범위를 0~255로 정규화
            ch_data = ((yi_map / 120.0)*255).clip(0,255).astype(np.uint8)
        else:
            ch_data = S_ch

        # ── 크롭 (HSI는 이미 전체 이미지, Lab계열은 이미 크롭됨) ──
        if ch in ("H","S","I"):
            x0,y0,x1,y1 = roi
            crop = ch_data[y0:y1, x0:x1]
        else:
            crop = ch_data   # 이미 크롭됨

        display = self._colorize(crop, ch)

        # 세그먼트 그리드는 HSI 전용 (stats가 있을 때만)
        if ch in ("H","S","I") and img.get("stats"):
            display = self._draw_seg_grid(display, img, roi)

        cw  = self._hsi_cv.winfo_width()  or 400
        ch_ = self._hsi_cv.winfo_height() or 200
        pil = Image.fromarray(display)
        pil.thumbnail((cw, ch_), Image.LANCZOS)
        tk_img = ImageTk.PhotoImage(pil)
        self._refs["hsi"] = tk_img
        self._hsi_cv.create_image(cw//2, ch_//2,
                                  anchor="center", image=tk_img)

        # 상단 수치 표시
        s_  = img.get("s_mean", np.nan)
        yr_ = img.get("yellow_ratio", np.nan)
        b_  = img.get("lab", {}).get("b", np.nan)
        yi_ = img.get("yellowness_idx", np.nan)
        de_ = img.get("delta_e", np.nan)
        ch_val = {"H": f"H={np.mean(H_ch[img['mask'].astype(bool)]):.1f}"
                       if img.get("mask") is not None else "",
                  "S": f"S={s_:.1f}" if not np.isnan(s_) else "",
                  "I": f"I={np.mean(I_ch[img['mask'].astype(bool)]):.1f}"
                       if img.get("mask") is not None else "",
                  "b*": f"b*={b_:.1f}" if not np.isnan(b_) else "b*(N/A)",
                  "L*": f"L*={img.get('lab',{}).get('L',np.nan):.1f}",
                  "ΔE": f"ΔE={de_:.1f}" if not np.isnan(de_) else "ΔE(vs D0)",
                  "YI": f"YI={yi_:.0f}" if not np.isnan(yi_) else "YI(N/A)",
                  }.get(ch, "")
        col = self._s_color(s_)
        self._hsi_cv.create_text(
            6, 6, anchor="nw",
            text=f"[{img['day']}d / {img['cond']}]  {_L('채널','ch')}:{ch}  {ch_val}",
            fill=col, font=LF)

    def _colorize(self, ch_data, ch):
        n = ch_data.astype(np.float32) / 255.0
        if ch == "H":
            hsv = np.zeros((*ch_data.shape,3), np.uint8)
            hsv[:,:,0] = (ch_data.astype(np.float32)*179/255).astype(np.uint8)
            hsv[:,:,1] = 200; hsv[:,:,2] = 220
            return cv2.cvtColor(hsv, cv2.COLOR_HSV2RGB)
        elif ch == "S":
            out = np.zeros((*ch_data.shape,3), np.uint8)
            out[:,:,0] = (n*80+100).clip(0,255).astype(np.uint8)
            out[:,:,1] = (n*100+80).clip(0,255).astype(np.uint8)
            out[:,:,2] = (n*80+160).clip(0,255).astype(np.uint8)
            return out
        elif ch == "I":
            out = np.zeros((*ch_data.shape,3), np.uint8)
            out[:,:,0] = (n*200+55).clip(0,255).astype(np.uint8)
            out[:,:,1] = (n*160+60).clip(0,255).astype(np.uint8)
            out[:,:,2] = (n*40+20).clip(0,255).astype(np.uint8)
            return out
        elif ch == "b*":
            # 황색(높음)=노랑, 무채색(중간)=회색, 청색(낮음)=파랑
            out = np.zeros((*ch_data.shape,3), np.uint8)
            out[:,:,0] = (n*255).clip(0,255).astype(np.uint8)   # R: 높을수록 빨강
            out[:,:,1] = (n*220).clip(0,255).astype(np.uint8)   # G
            out[:,:,2] = ((1-n)*200).clip(0,255).astype(np.uint8)  # B: 낮을수록 파랑
            return out
        elif ch == "L*":
            # 밝기: 낮음=어두운 파랑, 높음=밝은 흰색
            out = np.zeros((*ch_data.shape,3), np.uint8)
            out[:,:,0] = (n*180+60).clip(0,255).astype(np.uint8)
            out[:,:,1] = (n*180+60).clip(0,255).astype(np.uint8)
            out[:,:,2] = (n*200+55).clip(0,255).astype(np.uint8)
            return out
        elif ch == "ΔE":
            # 색차: 낮음(변화없음)=초록, 높음(변화큼)=빨강 (RdYlGn 역)
            out = np.zeros((*ch_data.shape,3), np.uint8)
            out[:,:,0] = (n*255).clip(0,255).astype(np.uint8)   # R 증가
            out[:,:,1] = ((1-n)*200+30).clip(0,255).astype(np.uint8)  # G 감소
            out[:,:,2] = (30*(1-n)).clip(0,255).astype(np.uint8)
            return out
        elif ch == "YI":
            # YI: 높음(미산화)=황금색, 낮음(산화)=청회색
            out = np.zeros((*ch_data.shape,3), np.uint8)
            out[:,:,0] = (n*220+30).clip(0,255).astype(np.uint8)
            out[:,:,1] = (n*180+40).clip(0,255).astype(np.uint8)
            out[:,:,2] = ((1-n)*160+20).clip(0,255).astype(np.uint8)
            return out
        else:
            gray = ch_data[:,:,np.newaxis].repeat(3, axis=2)
            return gray

    def _draw_seg_grid(self, display, img, roi):
        out = display.copy()
        ch  = self.ch_var.get()
        H, W = out.shape[:2]
        rows = self.rows_var.get()
        cols = self.cols_var.get()
        x0,y0,x1,y1 = roi
        sh = (y1-y0)/rows; sw = (x1-x0)/cols

        for r in range(1,rows):
            y = int(r*sh)
            cv2.line(out,(0,y),(W,y),(180,190,210),1)
        for c in range(1,cols):
            x = int(c*sw)
            cv2.line(out,(x,0),(x,H),(180,190,210),1)

        for s in img["stats"].get(ch,[]):
            sy0=s.get("ry0",0)-y0; sy1=s.get("ry1",1)-y0
            sx0=s.get("rx0",0)-x0; sx1=s.get("rx1",1)-x0
            mv=s.get("mean",np.nan)
            if s.get("pixels",0)<10 or np.isnan(mv): continue
            seg_col=((34,160,80) if mv>=self.cfg_s_good.get() else
                     (200,120,20) if mv>=self.cfg_s_warn.get() else (200,50,50))
            ov=out.copy()
            cv2.rectangle(ov,(sx0+1,sy0+1),(sx1-1,sy1-1),seg_col,-1)
            cv2.addWeighted(ov,0.12,out,0.88,0,out)
            cv2.rectangle(out,(sx0+1,sy0+1),(sx1-1,sy1-1),seg_col,1)
            cv2.putText(out,f"#{s['seg']}",(sx0+3,sy0+12),
                        cv2.FONT_HERSHEY_SIMPLEX,0.32,(80,90,120),1,cv2.LINE_AA)
            cv2.putText(out,f"{mv:.0f}",(sx0+3,sy0+24),
                        cv2.FONT_HERSHEY_SIMPLEX,0.36,seg_col,1,cv2.LINE_AA)
        return out

    # ─────────────────────────────────────────
    #  세그먼트 테이블
    # ─────────────────────────────────────────
    def _update_table(self):
        self.tree.delete(*self.tree.get_children())
        idx = self.sel_idx
        if idx < 0 or idx >= len(self.images): return
        img  = self.images[idx]
        ch   = self.ch_var.get()
        mask = img.get("mask")
        rgb  = img["rgb"]

        # S채널 통계는 "S" 키 기준 (채널 선택과 무관하게 항상 표시)
        seg_key = "S" if "S" in img["stats"] else ch
        segs = img["stats"].get(seg_key, [])

        def f(v, fmt=".1f"):
            return format(v, fmt) if not np.isnan(v) else "-"

        for s in segs:
            ry0,ry1 = s.get("ry0",0), s.get("ry1",1)
            rx0,rx1 = s.get("rx0",0), s.get("rx1",1)
            pix     = s.get("pixels", 0)

            if pix < self.cfg_min_pix.get():
                continue

            # S채널 평균/편차 (통계에서)
            s_mean = s.get("mean", np.nan)
            s_std  = s.get("std",  np.nan)

            # ROI 크롭 슬라이스
            seg_rgb  = rgb[ry0:ry1, rx0:rx1]
            seg_mask = mask[ry0:ry1, rx0:rx1] if mask is not None \
                       else np.ones((ry1-ry0, rx0-rx1 if rx1>rx0 else 1), bool)

            if seg_rgb.size == 0 or seg_mask.sum() == 0:
                continue

            # 황색비율
            yr_s = compute_yellow_ratio(seg_rgb, seg_mask,
                                        h_lo_deg=self.cfg_h_lo.get(),
                                        h_hi_deg=self.cfg_h_hi.get(),
                                        s_thresh=self.cfg_s_thresh.get())

            # Lab b* / L*
            lab_m = compute_lab_metrics(seg_rgb, seg_mask)
            b_s   = lab_m["b"]
            L_s   = lab_m["L"]

            # YI
            yi_s  = compute_yellowness_index(seg_rgb, seg_mask)

            # ΔE (같은 조건 0일차와 비교)
            de_s  = np.nan
            try:
                ref_img = next(
                    (o for o in self.images
                     if o["cond"] == img["cond"]
                     and float(o.get("day","99")) == 0
                     and o is not img), None)
                if ref_img is not None:
                    ref_mask = ref_img.get("mask")
                    ref_crop = ref_img["rgb"][ry0:ry1, rx0:rx1]
                    ref_seg_mask = ref_mask[ry0:ry1, rx0:rx1] \
                                   if ref_mask is not None \
                                   else np.ones_like(seg_mask)
                    ref_lab = compute_lab_metrics(ref_crop, ref_seg_mask)
                    de_s    = compute_delta_e(lab_m, ref_lab)
            except Exception:
                pass

            # 색상 태그: b* 기준
            if not np.isnan(b_s):
                tag = ("g" if b_s >= self.cfg_b_good.get()
                       else "w" if b_s >= self.cfg_b_warn.get()
                       else "r")
            elif not np.isnan(s_mean):
                tag = ("g" if s_mean >= self.cfg_s_good.get()
                       else "w" if s_mean >= self.cfg_s_warn.get()
                       else "r")
            else:
                tag = ""

            self.tree.insert("","end", tags=(tag,), values=(
                f"#{s['seg']}",
                f(s_mean),
                f(s_std, ".1f"),
                f(b_s,   ".1f"),
                f(yi_s,  ".0f"),
                f"{yr_s*100:.0f}%" if not np.isnan(yr_s) else "-",
                f(de_s,  ".1f"),
                f(L_s,   ".1f"),
            ))

        self.tree.tag_configure("g", foreground=GREEN)
        self.tree.tag_configure("w", foreground=AMBER)
        self.tree.tag_configure("r", foreground=RED)
        # 지표 요약 패널도 함께 갱신
        self._update_stat_panel()

    def _update_stat_panel(self):
        """ROI 상세 탭 — 지표 요약 (한/영 지원)"""
        idx = self.sel_idx
        if idx < 0 or idx >= len(self.images):
            self._set_cmt(self._stat_text,
                _L("이미지를 선택하면 지표가 표시된다.",
                   "Select an image to view metrics."))
            return

        img = self.images[idx]

        def f(v, fmt=".1f"):
            return format(float(v), fmt) if not np.isnan(float(v)) else "-"

        def judge(v, good, warn, higher_good=True):
            if np.isnan(v):
                return _L("미분석","N/A"), ""
            if higher_good:
                if v >= good:   return _L("✅ 양호","✅ Good"),   f"≥{good}"
                elif v >= warn: return _L("⚠ 경계","⚠ Border"), f"{warn}~{good}"
                else:           return _L("🔴 산화","🔴 Oxid."),  f"<{warn}"
            else:
                if v <= warn:   return _L("✅ 양호","✅ Good"),   f"≤{warn}"
                elif v <= good: return _L("⚠ 경계","⚠ Border"), f"{warn}~{good}"
                else:           return _L("🔴 산화","🔴 Oxid."),  f">{good}"

        s_  = img.get("s_mean",        np.nan)
        yr_ = img.get("yellow_ratio",  np.nan)
        yi_ = img.get("yellowness_idx",np.nan)
        b_  = img.get("lab",{}).get("b", np.nan)
        L_  = img.get("lab",{}).get("L", np.nan)
        a_  = img.get("lab",{}).get("a", np.nan)
        de_ = img.get("delta_e",       np.nan)
        gc_ = img.get("glcm",{}).get("contrast",    np.nan)
        ge_ = img.get("glcm",{}).get("energy",      np.nan)
        gh_ = img.get("glcm",{}).get("homogeneity", np.nan)
        gr_ = img.get("glcm",{}).get("correlation", np.nan)

        sg  = self.cfg_s_good.get();  sw_ = self.cfg_s_warn.get()
        bg  = self.cfg_b_good.get();  bw  = self.cfg_b_warn.get()
        yg  = self.cfg_yi_good.get(); yw  = self.cfg_yi_warn.get()
        thr = self.cfg_s_thresh.get()

        s_j,  s_n  = judge(s_,  sg,  sw_)
        b_j,  b_n  = judge(b_,  bg,  bw)
        yi_j, yi_n = judge(yi_, yg,  yw)
        de_j, de_n = judge(de_, 10,  3, higher_good=False)

        # 산화 진행률
        progress_str = ""
        ref_imgs = [o for o in self.images
                    if o["cond"]==img["cond"]
                    and not np.isnan(o.get("lab",{}).get("b",np.nan))]
        if len(ref_imgs) >= 2:
            def df2(d):
                try: return float(d)
                except: return 9999
            rs = sorted(ref_imgs, key=lambda x: df2(x["day"]))
            b_max = rs[0]["lab"]["b"]
            b_min = rs[-1]["lab"]["b"]
            if b_max > 0:
                prog = max(0, min(100, (b_max-b_)/(b_max-b_min+1e-6)*100))
                bar  = "█"*int(prog/5) + "░"*(20-int(prog/5))
                progress_str = (
                    f"\nOxidation Progress (b* basis):\n"
                    f"  [{bar}] {prog:.0f}%\n"
                    f"  Day0 b*={b_max:.1f} → Current={f(b_)} → Final={b_min:.1f}")

        # 세그먼트 분포
        seg_summary = ""
        segs = img["stats"].get("S", [])
        vs = [(s["seg"], s.get("mean",np.nan)) for s in segs
              if not np.isnan(s.get("mean",np.nan))
              and s.get("pixels",0) >= self.cfg_min_pix.get()]
        if vs:
            ss = sorted(vs, key=lambda x: x[1])
            avg = np.mean([v for _,v in vs])
            seg_summary = (
                f"\nSegment Distribution ({len(vs)} valid):\n"
                f"  Avg={avg:.1f}  Min=#{ss[0][0]}({ss[0][1]:.1f})"
                f"  Max=#{ss[-1][0]}({ss[-1][1]:.1f})")
        yr_v = yr_*100 if not np.isnan(yr_) else np.nan
        lines = [
            f"══ {img['cond']}  /  Day {img['day']} ══",
            f"File: {img['name']}", "",
            "─── HSI Color ────────────────────────────",
            f"  S-ch Mean:    {f(s_):>7}  {s_j}  ({s_n})",
            f"    → Pristine: 150~200 / Oxidizing: 20~60 / Fully oxidized: <20",
            f"    → Lower = more HfO2 (achromatic) replacing yellow HfS2",
            f"  Yellow Ratio: {f(yr_v, '0f'):>6}%",
            f"    → Pixels with H:35~75 deg · S≥{thr:.2f}",
            f"  YI:           {f(yi_, '0f'):>7}  {yi_j}  ({yi_n})",
            f"    → ASTM E313. Pristine: 50~110 / Border: 35~50 / Oxidized: 20~35",
            "",
            "─── Lab Color Space ──────────────────────",
            f"  b* (yel-blue):  {f(b_):>7}  {b_j}  ({b_n})",
            f"    → Key★. +pos=yellow / 0=achromatic / -neg=blue",
            f"    → Pristine: +50~60 / Oxidizing: +15~30 / Fully: +3~10",
            f"    → Robust to lighting; best Raman correlation",
            f"  L* (lightness): {f(L_):>7}",
            f"    → 0=black~100=white. Pristine: 60~75 / Oxidized: 85~95",
            f"  a* (red-green): {f(a_):>7}",
            f"    → Negligible change during HfS2 oxidation",
            f"  Delta-E:        {f(de_):>7}  {de_j}  (vs. day-0)",
            f"    → CIE76. <1:imperceptible / 1~3:subtle / 3~10:clear / >10:major",
        ]
        if progress_str:
            lines += ["", "─── Oxidation Progress ───────────────────", progress_str]
        lines += [
            "", "─── GLCM Texture ─────────────────────────",
            f"  Contrast:    {f(gc_):>8}",
            f"    → Local intensity diff. High=pristine, decreases on oxidation",
            f"  Energy:      {f(ge_,'.5f'):>8}",
            f"    → Pattern repeatability 0~1. Increases on oxidation",
            f"  Homogeneity: {f(gh_,'.5f'):>8}",
            f"    → Pixel similarity 0~1. Increases on oxidation",
            f"  Correlation: {f(gr_,'.4f'):>8}",
            f"    → Linear dependency -1~1. Surface directionality",
        ]
        if seg_summary:
            lines += ["", "─── Segment Distribution ─────────────────", seg_summary]

        self._set_cmt(self._stat_text, "\n".join(lines))



    def _copy_stat_text(self):
        txt = self._stat_text.get("1.0","end").strip()
        self.clipboard_clear(); self.clipboard_append(txt)
        self._set_status(_L("지표 요약 복사 완료","Metrics copied"))

    def _stat_popup(self):
        """지표 요약 확대 팝업"""
        idx = self.sel_idx
        if idx < 0 or idx >= len(self.images): return
        img = self.images[idx]

        win = tk.Toplevel(self)
        win.title(f"지표 요약 — {img['name']}")
        win.configure(bg=PANEL)
        sw = self.winfo_screenwidth(); sh = self.winfo_screenheight()
        ww, wh = 560, 480
        win.geometry(f"{ww}x{wh}+{(sw-ww)//2}+{(sh-wh)//2}")
        tk.Label(win,
                 text=f"  {img['cond']}  /  " + _L(f"{img['day']}일차  —  지표 요약", f"Day {img['day']}  —  Metric Summary"),
                 bg=PANEL2, fg=TXT,
                 font=("Segoe UI",11,"bold"),
                 highlightbackground=BORDER,
                 highlightthickness=1).pack(fill="x")

        txt_content = self._stat_text.get("1.0","end").strip()
        txt_f = tk.Frame(win, bg=PANEL)
        txt_f.pack(fill="both", expand=True, padx=6, pady=6)
        txt_w = tk.Text(txt_f, wrap="word",
                        bg=CARD2, fg=TXT,
                        font=("Segoe UI",11),   # 큰 글씨
                        relief="flat", padx=12, pady=10,
                        highlightbackground=BORDER,
                        highlightthickness=1,
                        cursor="xterm",
                        state="disabled")
        vsb = tk.Scrollbar(txt_f, orient="vertical",
                           command=txt_w.yview)
        txt_w.configure(yscrollcommand=vsb.set)
        vsb.pack(side="right", fill="y")
        txt_w.pack(fill="both", expand=True)
        self._set_cmt(txt_w, txt_content)
        txt_w.bind("<Control-c>", lambda e: self._copy_text(txt_w))
        txt_w.bind("<Control-C>", lambda e: self._copy_text(txt_w))

        btn_f = tk.Frame(win, bg=PANEL)
        btn_f.pack(fill="x", pady=4)
        tk.Button(btn_f, text=_L("📋 전체 복사","📋 Copy All"),
                  command=lambda: (self.clipboard_clear(),
                                   self.clipboard_append(txt_content)),
                  bg=BTN, fg=TXT, font=MF,
                  relief="flat", padx=12, pady=4,
                  cursor="hand2").pack(side="left", padx=8)
        tk.Button(btn_f, text=_L("닫기","Close"), command=win.destroy,
                  bg=BTN, fg=TXT, font=MF,
                  relief="flat", padx=14, pady=4,
                  cursor="hand2").pack(side="right", padx=8)

    def _hsi_popup(self):
        """채널 뷰 더블클릭 확대 팝업"""
        idx = self.sel_idx
        if idx < 0 or idx >= len(self.images): return
        img = self.images[idx]

        win = tk.Toplevel(self)
        ch = self.ch_var.get()
        win.title(f"채널 뷰 [{ch}]  —  {img['name']}")
        win.configure(bg=PANEL)
        sw = self.winfo_screenwidth(); sh = self.winfo_screenheight()
        ww, wh = int(sw*0.75), int(sh*0.75)
        win.geometry(f"{ww}x{wh}+{(sw-ww)//2}+{(sh-wh)//2}")
        tk.Label(win,
                 text=f"  {_L('채널','Ch')}: {ch}  |  {img['cond']}  " + _L(f"{img['day']}일", f"Day {img['day']}"),
                 bg=PANEL2, fg=TXT,
                 font=("Segoe UI",11,"bold"),
                 highlightbackground=BORDER,
                 highlightthickness=1).pack(fill="x")

        cv_ = tk.Canvas(win, bg=PANEL, highlightthickness=0)
        cv_.pack(fill="both", expand=True, padx=4, pady=4)

        def _render_large(event=None):
            cw = cv_.winfo_width()  or ww-20
            ch_ = cv_.winfo_height() or wh-80
            # _refresh_hsi 결과 이미지를 큰 캔버스에 다시 렌더
            ref = self._refs.get("hsi")
            if ref:
                cv_.delete("all")
                cv_.create_image(cw//2, ch_//2,
                                 anchor="center", image=ref)

        cv_.bind("<Configure>", _render_large)
        # 실제 큰 버전 새로 생성
        self.after(100, lambda: self._render_hsi_to_canvas(cv_, ww-20, wh-120))

        tk.Button(win, text=_L("닫기","Close"), command=win.destroy,
                  bg=BTN, fg=TXT, font=MF,
                  relief="flat", padx=14, pady=4,
                  cursor="hand2").pack(pady=6)

    def _render_hsi_to_canvas(self, cv_: tk.Canvas, cw: int, ch_: int):
        """지정 캔버스에 현재 채널 뷰를 큰 크기로 렌더"""
        idx = self.sel_idx
        if idx < 0 or idx >= len(self.images): return
        img = self.images[idx]
        roi = img.get("roi")
        if not roi: return

        H_ch, S_ch, I_ch = img["hsi"]
        ch = self.ch_var.get()

        if ch in ("H","S","I"):
            ch_data = {"H":H_ch,"S":S_ch,"I":I_ch}[ch]
            x0,y0,x1,y1 = roi
            crop = ch_data[y0:y1, x0:x1]
        else:
            # Lab/YI/ΔE는 _refresh_hsi와 동일한 계산
            x0,y0,x1,y1 = roi
            crop_rgb = img["rgb"][y0:y1, x0:x1]
            if ch == "b*":
                bgr = cv2.cvtColor(crop_rgb, cv2.COLOR_RGB2BGR)
                lab = cv2.cvtColor(bgr, cv2.COLOR_BGR2Lab)
                crop = ((lab[:,:,2].astype(np.float32)) / 255.0 * 255).clip(0,255).astype(np.uint8)
            elif ch == "L*":
                bgr = cv2.cvtColor(crop_rgb, cv2.COLOR_RGB2BGR)
                lab = cv2.cvtColor(bgr, cv2.COLOR_BGR2Lab)
                crop = lab[:,:,0]
            elif ch == "YI":
                r = crop_rgb[:,:,0].astype(np.float32)/255
                g = np.where(crop_rgb[:,:,1]>2, crop_rgb[:,:,1].astype(np.float32)/255, 0.01)
                b = crop_rgb[:,:,2].astype(np.float32)/255
                crop = ((100*(1.28*r-1.06*b)/g)/120*255).clip(0,255).astype(np.uint8)
            else:
                _, S_ch2, _ = rgb_to_hsi(crop_rgb)
                crop = S_ch2

        display = self._colorize(crop, ch)
        if ch in ("H","S","I") and img.get("stats"):
            display = self._draw_seg_grid(display, img, roi)

        pil = Image.fromarray(display)
        pil = pil.resize((cw, ch_), Image.LANCZOS)
        tk_img = ImageTk.PhotoImage(pil)
        self._refs["hsi_popup"] = tk_img
        cv_.delete("all")
        cv_.create_image(cw//2, ch_//2, anchor="center", image=tk_img)
    def _update_heatmap(self):
        self.hm_fig.clear()
        metric = self._hm_metric.get()
        n = self.rows_var.get() * self.cols_var.get()

        # 분석 완료 이미지
        an = [img for img in self.images
              if img.get("roi") and
              not np.isnan(img.get("s_mean", np.nan))]
        if not an: return

        def df(d):
            try: return float(d)
            except: return 9999

        # 조건별로 정렬
        an = sorted(an, key=lambda x: (x["cond"], df(x["day"])))
        labels = [f"{img['day']}d/{img['cond'][:6]}" for img in an]

        # 지표별 행렬 구성
        mat  = np.full((len(an), n), np.nan)
        cmap_name = "RdYlGn"
        vmin, vmax = 0, 255
        title_suffix = ""

        if metric == "S":
            # S채널 세그먼트 평균
            for di, img in enumerate(an):
                for s in img["stats"].get("S", []):
                    if s["seg"] < n:
                        mat[di, s["seg"]] = s.get("mean", np.nan)
            vmin, vmax = 0, 255
            cmap_name = "RdYlGn"
            title_suffix = _L("S채널 평균  (높을수록 미산화(pristine))","S-ch Mean  (higher=more pristine)")

        elif metric == "b":
            # Lab b* 세그먼트별 계산
            for di, img in enumerate(an):
                mask = img.get("mask")
                rgb  = img["rgb"]
                for s in img["stats"].get("S", []):
                    if s["seg"] >= n: continue
                    ry0,ry1 = s.get("ry0",0), s.get("ry1",1)
                    rx0,rx1 = s.get("rx0",0), s.get("rx1",1)
                    seg_rgb  = rgb[ry0:ry1, rx0:rx1]
                    seg_mask = mask[ry0:ry1, rx0:rx1] \
                               if mask is not None \
                               else np.ones((ry1-ry0, rx1-rx0), bool)
                    if seg_rgb.size > 0 and seg_mask.sum() > 0:
                        lab = compute_lab_metrics(seg_rgb, seg_mask)
                        mat[di, s["seg"]] = lab["b"]
            vmin, vmax = 0, 70
            cmap_name = "YlOrRd_r"
            title_suffix = _L("Lab b*  (높을수록 황색(미산화))","Lab b*  (higher=more yellow/pristine)")

        elif metric == "YI":
            for di, img in enumerate(an):
                mask = img.get("mask")
                rgb  = img["rgb"]
                for s in img["stats"].get("S", []):
                    if s["seg"] >= n: continue
                    ry0,ry1 = s.get("ry0",0), s.get("ry1",1)
                    rx0,rx1 = s.get("rx0",0), s.get("rx1",1)
                    seg_rgb  = rgb[ry0:ry1, rx0:rx1]
                    seg_mask = mask[ry0:ry1, rx0:rx1] \
                               if mask is not None \
                               else np.ones((ry1-ry0, rx1-rx0), bool)
                    if seg_rgb.size > 0 and seg_mask.sum() > 0:
                        mat[di, s["seg"]] = compute_yellowness_index(
                            seg_rgb, seg_mask)
            vmin, vmax = 0, 120
            cmap_name = "RdYlGn"
            title_suffix = _L("YI  (높을수록 미산화(pristine))","YI  (higher=more pristine)")

        elif metric == "dE":
            # 조건별 0일차 기준 ΔE
            ref_map = {}
            for img in an:
                cond = img["cond"]
                try:
                    if float(img["day"]) == 0:
                        ref_map[cond] = img
                except: pass
            for di, img in enumerate(an):
                ref = ref_map.get(img["cond"])
                if ref is None: continue
                mask  = img.get("mask")
                rmask = ref.get("mask")
                rgb   = img["rgb"]
                rrgb  = ref["rgb"]
                for s in img["stats"].get("S", []):
                    if s["seg"] >= n: continue
                    ry0,ry1 = s.get("ry0",0), s.get("ry1",1)
                    rx0,rx1 = s.get("rx0",0), s.get("rx1",1)
                    seg_rgb  = rgb[ry0:ry1, rx0:rx1]
                    seg_mask = mask[ry0:ry1, rx0:rx1] \
                               if mask is not None \
                               else np.ones((ry1-ry0, rx1-rx0), bool)
                    ref_seg  = rrgb[ry0:ry1, rx0:rx1]
                    ref_mask = rmask[ry0:ry1, rx0:rx1] \
                               if rmask is not None \
                               else np.ones_like(seg_mask)
                    if seg_rgb.size > 0 and seg_mask.sum() > 0:
                        lab1 = compute_lab_metrics(seg_rgb, seg_mask)
                        lab2 = compute_lab_metrics(ref_seg, ref_mask)
                        mat[di, s["seg"]] = compute_delta_e(lab1, lab2)
            vmin, vmax = 0, 30
            cmap_name = "RdYlGn_r"
            title_suffix = _L("ΔE (0일차 기준)  (낮을수록 변화없음)","ΔE (from day-0)  (lower=less change)")

        # 그리기
        ax = self.hm_fig.add_subplot(111)
        ax.set_facecolor(PANEL)
        self.hm_fig.patch.set_facecolor(PANEL)

        im = ax.imshow(mat, aspect="auto",
                       cmap=cmap_name,
                       vmin=vmin, vmax=vmax,
                       interpolation="nearest")

        ax.set_yticks(range(len(labels)))
        ax.set_yticklabels(labels, color=TXT, fontsize=5)
        ax.set_xticks(range(n))
        ax.set_xticklabels([f"#{i}" for i in range(n)],
                           color=TXT, fontsize=5)
        ax.tick_params(colors=SUB, length=2)
        for sp in ax.spines.values(): sp.set_color(BORDER)
        ax.set_title(title_suffix, color=TXT, fontsize=7, pad=4)

        # 셀 값 표시
        for di in range(len(an)):
            for si in range(n):
                v = mat[di, si]
                if not np.isnan(v):
                    ax.text(si, di, f"{v:.0f}",
                            ha="center", va="center",
                            fontsize=4,
                            color="white" if v < (vmin+vmax)*0.5
                                  else "black")

        cb = self.hm_fig.colorbar(im, ax=ax, fraction=0.03, pad=0.01)
        cb.ax.tick_params(colors=SUB, labelsize=5)
        self.hm_fig.tight_layout(pad=0.6)
        self.hm_cv.draw()

    # ─────────────────────────────────────────
    #  비교 그리드
    # ─────────────────────────────────────────
    def _refresh_compare(self):
        for w in self._cmp_fr.winfo_children(): w.destroy()
        an = [img for img in self.images
              if not np.isnan(img.get("s_mean",np.nan))]
        if not an:
            tk.Label(self._cmp_fr, text=_L("분석 완료된 데이터가 없다","No analyzed data"),
                     bg=BG, fg=SUB, font=MF).pack(padx=40, pady=40)
            return

        conds = list(dict.fromkeys(
            img["cond"] for img in an if img["cond"].strip()))
        def df(d):
            try: return float(d)
            except: return 9999
        days = sorted(set(img["day"] for img in an if img["day"].strip()),
                      key=df)
        if not conds: conds=["(none)"]
        if not days:  days=["(none)"]

        CW=self.TW+20; CH=self.TH+100   # 높이 확장 (지표 추가)

        tk.Label(self._cmp_fr,text="",bg=BG,width=14).grid(
            row=0,column=0,padx=2,pady=2)
        for ci,day in enumerate(days):
            col=COND_COLORS[ci%len(COND_COLORS)]
            tk.Label(self._cmp_fr,text=_L(f"  {day}일차  ", f"  Day {day}  "),
                     bg=PANEL2,fg=col,font=MFB,
                     relief="flat",width=14).grid(
                     row=0,column=ci+1,padx=3,pady=3)

        for ri,cond in enumerate(conds):
            rcol=COND_COLORS[ri%len(COND_COLORS)]
            tk.Label(self._cmp_fr,text=f" {cond} ",
                     bg=PANEL2,fg=rcol,font=MFB,
                     anchor="w",wraplength=110).grid(
                     row=ri+1,column=0,padx=3,pady=3,sticky="nsew")

            for ci,day in enumerate(days):
                matched=[img for img in an
                         if img["cond"].strip()==cond
                         and img["day"].strip()==day]
                cell=tk.Frame(self._cmp_fr,bg=CARD,
                              highlightbackground=BORDER,
                              highlightthickness=1,
                              width=CW,height=CH,
                              cursor="hand2")
                cell.grid(row=ri+1,column=ci+1,padx=3,pady=3)
                cell.grid_propagate(False)

                if not matched:
                    tk.Label(cell,text="+",bg=CARD,fg=BORDER,
                             font=("Segoe UI",20)).place(
                             relx=0.5,rely=0.5,anchor="center")
                    continue

                img=matched[0]; ridx=self.images.index(img)
                if img.get("thumb"):
                    tk_img=ImageTk.PhotoImage(img["thumb"])
                    self._refs[f"cmp_{ridx}"]=tk_img
                    lbl=tk.Label(cell,image=tk_img,bg=CARD,cursor="hand2")
                    lbl.place(x=4,y=4)
                    lbl.bind("<Button-1>",lambda e,i=ridx: self._select(i))

                # ── 지표 표시 (S, YR, YI, b*, ΔE) ──
                s_  = img.get("s_mean",      np.nan)
                yr_ = img.get("yellow_ratio", np.nan)
                yi_ = img.get("yellowness_idx",np.nan)
                b_  = img.get("lab",{}).get("b", np.nan)
                de_ = img.get("delta_e",     np.nan)

                y_pos = CH - 90
                for val, prefix, color_fn, fmt in [
                    (s_,  "S",  self._s_color,  "{:.0f}"),
                    (yr_*100 if not np.isnan(yr_) else np.nan,
                           "Y%", lambda v: GOLD,  "{:.0f}%"),
                    (yi_, "YI", self._yi_color,  "{:.0f}"),
                    (b_,  "b*", self._b_color,   "{:.1f}"),
                    (de_, "ΔE", lambda v: (RED if v>10 else AMBER if v>3 else GREEN), "{:.1f}"),
                ]:
                    if not np.isnan(val):
                        col_ = color_fn(val)
                        tk.Label(cell,
                                 text=f"{prefix}={fmt.format(val)}",
                                 bg=CARD, fg=col_,
                                 font=("Segoe UI",7,"bold")).place(x=4, y=y_pos)
                        y_pos += 16

                # 더블클릭 → 셀 상세 팝업
                cell.bind("<Double-Button-1>",
                          lambda e, i=ridx: self._cmp_detail_popup(i))

        self._cmp_cv.configure(scrollregion=self._cmp_cv.bbox("all"))

    def _cmp_detail_popup(self, img_idx: int):
        """비교 그리드 셀 더블클릭 → 해당 이미지 상세 지표 팝업"""
        if img_idx < 0 or img_idx >= len(self.images): return
        img = self.images[img_idx]

        win = tk.Toplevel(self)
        win.title(f"상세 지표  —  {img['name']}")
        win.configure(bg=PANEL)
        sw = self.winfo_screenwidth(); sh = self.winfo_screenheight()
        ww, wh = 520, 600
        win.geometry(f"{ww}x{wh}+{(sw-ww)//2}+{(sh-wh)//2}")

        # 헤더
        tk.Label(win,
                 text=f"  {img['cond']}  /  " + _L(f"{img['day']}일차", f"Day {img['day']}"),
                 bg=PANEL2, fg=TXT,
                 font=("Segoe UI",11,"bold"),
                 highlightbackground=BORDER,
                 highlightthickness=1).pack(fill="x")

        # 썸네일
        th = img.get("thumb")
        if th:
            big = th.resize((200,160), Image.LANCZOS)
            tk_img = ImageTk.PhotoImage(big)
            self._refs[f"popup_th_{img_idx}"] = tk_img
            tk.Label(win, image=tk_img, bg=PANEL).pack(pady=8)

        # 지표 텍스트 (선택/복사 가능)
        def _safe(v): return v if not np.isnan(v) else None

        lines = [
            f"파일명:   {img['name']}",
            f"조건:     {img['cond']}",
            f"날짜:     {img['day']}일",
            "",
            "─── HSI Color ─────────────────────────",
            f"S-ch Mean:      {img.get('s_mean',np.nan):.1f}  "
            f"(pristine 150~200 / oxidizing 20~60 / fully <20)",
            f"Yellow Ratio:   {img.get('yellow_ratio',np.nan)*100:.1f}%  "
            f"(higher = more pristine)",
            f"YI:             {img.get('yellowness_idx',np.nan):.0f}  "
            f"(pristine 50~110 / oxidized 20~35)",
            "",
            "─── Lab Color Space ───────────────────",
            f"Lab b* (yellow): {img.get('lab',{}).get('b',np.nan):.1f}  "
            f"(pristine +50~60 / oxidized +3~10)",
            f"Lab L* (light): {img.get('lab',{}).get('L',np.nan):.1f}  "
            f"(pristine 60~75 / oxidized→white 85~95)",
            f"Lab a* (r-g):   {img.get('lab',{}).get('a',np.nan):.1f}  "
            f"(negligible change in HfS₂)",
            f"ΔE (color):     {img.get('delta_e',np.nan):.1f}  "
            f"(<3 imperceptible / 3~10 clear / >10 major)",
            "",
            "─── GLCM Texture ──────────────────────",
            f"Contrast:           {img.get('glcm',{}).get('contrast',np.nan):.1f}  "
            f"(decreases on oxidation)",
            f"Energy:             {img.get('glcm',{}).get('energy',np.nan):.4f}  "
            f"(increases on oxidation)",
            f"Homogeneity:        {img.get('glcm',{}).get('homogeneity',np.nan):.4f}  "
            f"(increases on oxidation)",
        ]

        # 산화 상태 판정 추가
        b_ = img.get('lab',{}).get('b', np.nan)
        s_ = img.get('s_mean', np.nan)
        if not np.isnan(b_) and not np.isnan(s_):
            if b_ > 40 and s_ > 80:
                status = "✅ Pristine (unoxidized)"
            elif b_ > 20 and s_ > 40:
                status = "⚠️ Boundary (oxidizing)"
            else:
                status = "🔴 Oxidized (HfO₂ forming)"
            lines += ["", "─── Overall Assessment ─────────────", status]

        txt_content = "\n".join(lines)

        txt_f = tk.Frame(win, bg=PANEL)
        txt_f.pack(fill="both", expand=True, padx=6, pady=(0,4))
        txt_w = tk.Text(txt_f, wrap="word",
                        bg=CARD2, fg=TXT,
                        font=("Segoe UI",9),
                        relief="flat", padx=10, pady=8,
                        highlightbackground=BORDER,
                        highlightthickness=1,
                        cursor="xterm",
                        state="disabled")
        vsb = tk.Scrollbar(txt_f, orient="vertical",
                           command=txt_w.yview)
        txt_w.configure(yscrollcommand=vsb.set)
        vsb.pack(side="right", fill="y")
        txt_w.pack(fill="both", expand=True)
        self._set_cmt(txt_w, txt_content)
        txt_w.bind("<Control-c>", lambda e: self._copy_text(txt_w))
        txt_w.bind("<Control-C>", lambda e: self._copy_text(txt_w))

        btn_f = tk.Frame(win, bg=PANEL)
        btn_f.pack(fill="x", pady=4)
        tk.Button(btn_f, text=_L("📋 전체 복사","📋 Copy All"),
                  command=lambda: (self.clipboard_clear(),
                                   self.clipboard_append(txt_content)),
                  bg=BTN, fg=TXT, font=MF,
                  relief="flat", padx=10, pady=4,
                  cursor="hand2").pack(side="left", padx=8)
        tk.Button(btn_f, text=_L("닫기","Close"),
                  command=win.destroy,
                  bg=BTN, fg=TXT, font=MF,
                  relief="flat", padx=14, pady=4,
                  cursor="hand2").pack(side="right", padx=8)

    # ─────────────────────────────────────────
    #  차트 그리기
    # ─────────────────────────────────────────
    # ─────────────────────────────────────────
    #  차트별 자동 코멘트 생성
    # ─────────────────────────────────────────
    def _make_chart_comment(self, key: str, an: list) -> str:
        """
        3단 구성: 📐 축 설명 | 📖 해석 가이드 | 📊 데이터 요약
        한국어/영어 분기 (_KO 전역변수 기준)
        """
        if not an:
            return ""

        def df(d):
            try: return float(d)
            except: return 9999

        def _safe(v): return 0.0 if np.isnan(v) else float(v)

        conds = list(dict.fromkeys(img["cond"] for img in an))

        def trend(get_val, unit="", higher_is="fresh"):
            parts = []
            for cond in conds:
                pts = sorted([(df(img["day"]), _safe(get_val(img)))
                              for img in an if img["cond"]==cond],
                             key=lambda x: x[0])
                pts = [(d,v) for d,v in pts if d != 9999]
                if len(pts) < 2: continue
                v0, vN = pts[0][1], pts[-1][1]
                diff = vN - v0
                pct = abs(diff)/abs(v0)*100 if v0 != 0 else 0
                if abs(diff) < abs(v0)*0.03:
                    direction = _L("→ 변화없음","→ no change")
                elif (diff < 0) == (higher_is == "fresh"):
                    direction = _L(f"↓{abs(diff):.1f}{unit}({pct:.0f}%) 산화↑",
                                   f"↓{abs(diff):.1f}{unit}({pct:.0f}%) oxidized↑")
                else:
                    direction = f"↑{abs(diff):.1f}{unit}({pct:.0f}%)"
                parts.append(f"{cond[:10]}: {v0:.1f}→{vN:.1f}{unit} {direction}")
            return "  /  ".join(parts) if parts else _L("데이터 없음","no data")

        # ── 차트 탭 ──────────────────────────────────────────

        if key == "s_trend":
            data = trend(lambda i: _safe(i["s_mean"]), higher_is="fresh")
            return (
                "📐 X: Day  |  Y: S-channel (HSI Saturation) mean (0~255)\n"
                "📖 S↑=pristine(vivid)  S↓=oxidized(achromatic).  "
                "Pristine: S≈150~200 / Oxidizing: 20~60 / Fully oxidized: <20\n"
                f"📊 {data}")

        elif key == "yr_trend":
            data = trend(lambda i: _safe(i["yellow_ratio"])*100, "%", higher_is="fresh")
            return (
                "📐 X: Day  |  Y(left): Yellow pixel ratio 0~100%  |  Y(right): YI\n"
                "📖 Yellow ratio↓ = loss of pristine pixels = oxidation. "
                "YI: Pristine 50~110 / Border 35~50 / Oxidized 20~35\n"
                f"📊 Yellow ratio: {data}")

        elif key == "h_trend":
            data = trend(
                lambda i: float(np.mean(
                    i["hsi"][0][i["mask"].astype(bool)]
                )) if i.get("mask") is not None else np.nan,
                higher_is=None)
            return (
                "📐 X: Day  |  Y: H-channel (Hue) mean (0~255, 0=red·45=yellow·85=green)\n"
                "📖 Pristine HfS₂: H≈45 (yellow). Oxidation→achromatic, H becomes irregular. "
                "Interpret with S-ch and b*\n"
                f"📊 {data}")

        elif key == "box":
            ch = self.ch_var.get()
            return (
                f"📐 X: Condition·Day groups  |  Y: {ch}-channel value (0~255)\n"
                "📖 Median line, box=IQR, whiskers=min~max, dots=outliers. "
                "Oxidation: median drops + IQR narrows (surface homogenization)\n"
                "📊 Left→Right by date. Descending median = oxidation progressing")

        elif key == "hist":
            ch = self.ch_var.get()
            ch_range = {
                "H": _L("0~255 (색조각도 비례)","0~255 (hue angle proportional)"),
                "S": _L("0~255 (0=무채색, 255=최대 채도)","0~255 (0=achromatic, 255=max saturation)"),
                "I": _L("0~255 (0=검정, 255=흰색)","0~255 (0=black, 255=white)"),
            }
            rng = ch_range.get(ch, "0~255")
            return (
                f"📐 X: {ch}-channel value ({rng})  |  Y: Pixel density (normalized)\n"
                "📖 Solid=initial(day0), dashed=final(last day). "
                "Peak shift left = value decrease = oxidation. Flattening = surface inhomogeneity\n"
                "📊 Compare day0↔last for each condition. Larger shift = greater change")

        elif key == "decay":
            decays = []
            for cond in conds:
                pts = sorted([i for i in an if i["cond"]==cond], key=lambda x: df(x["day"]))
                if len(pts) >= 2 and pts[0]["s_mean"] > 0:
                    rs = (pts[0]["s_mean"]-pts[-1]["s_mean"])/pts[0]["s_mean"]*100
                    ry = (pts[0]["yellow_ratio"]-pts[-1]["yellow_ratio"]
                         )/max(pts[0]["yellow_ratio"],0.001)*100
                    decays.append((cond, rs, ry))
            if decays:
                decays.sort(key=lambda x: x[1], reverse=True)
                data = "  /  ".join(
                    f"{c[:10]}: S↓{rs:.0f}% Y↓{ry:.0f}%" for c,rs,ry in decays)
            else:
                data = _L("데이터 없음","no data")
            return (
                "📐 X: Condition  |  Y: Decay rate % (first→last day)\n"
                "📖 Solid bar=S-ch decay, hatched=Yellow ratio decay. "
                "Both high = fast oxidation. S>YR = overall saturation loss\n"
                f"📊 {data}")

        # ── 컬러 분석 탭 ──────────────────────────────────────

        elif key == "lab_b":
            data = trend(lambda i: _safe(i.get("lab",{}).get("b",0)), higher_is="fresh")
            return (
                "📐 X: Day  |  Y: CIE Lab b* (-128~+127)\n"
                "📖 b*>0=yellow, b*≈0=achromatic, b*<0=blue. "
                "Pristine: b*≈+50~60 / Oxidizing: +15~30 / Fully: +3~10. "
                "Lighting-robust key metric★\n"
                f"📊 {data}")

        elif key == "delta_e":
            data = trend(lambda i: _safe(i.get("delta_e",0)), higher_is="oxidized")
            return (
                "📐 X: Day  |  Y: ΔE (CIE76 color difference, 0=no change)\n"
                "📖 Color change from day-0 reference. "
                "ΔE<1:imperceptible / 1~3:subtle / 3~10:noticeable / >10:major\n"
                f"📊 {data}")

        elif key == "lab_L":
            data = trend(lambda i: _safe(i.get("lab",{}).get("L",0)), higher_is="oxidized")
            return (
                "📐 X: Day  |  Y: Lab L* lightness (0=black~100=white)\n"
                "📖 Pristine: L*≈60~75 (slightly dark yellow) / Oxidized→white: 85~95. "
                "Caution: may confuse with white substrate. Use with b* and S\n"
                f"📊 {data}")

        elif key == "glcm_con":
            data = trend(lambda i: _safe(i.get("glcm",{}).get("contrast",0)),
                         higher_is="fresh")
            return (
                "📐 X: Day  |  Y: GLCM Contrast (0=fully uniform)\n"
                "📖 Mean local intensity difference. High=pristine, drops on oxidation (surface homogenization). "
                "Sharp drop = oxidation acceleration point\n"
                f"📊 {data}")

        elif key == "glcm_eng":
            data_e = trend(lambda i: _safe(i.get("glcm",{}).get("energy",0)),
                           higher_is="oxidized")
            return (
                "📐 X: Day  |  Y(left): GLCM Energy(0~1)  |  Y(right): Homogeneity(0~1)\n"
                "📖 Energy=pattern repeatability (higher=simpler surface). "
                "Homogeneity=pixel similarity. Both increase on oxidation\n"
                f"📊 Energy: {data_e}")

        elif key == "all_trend":
            data = trend(lambda i: _safe(i.get("lab",{}).get("b",0)), higher_is="fresh")
            return (
                "📐 X: Day  |  Y: 0~1 normalized (day-0 = 1.0 for each metric)\n"
                "📖 b*(solid)·S-ch(dashed) normalized comparison. "
                "Both declining = high-confidence oxidation signal. "
                "Only one declining = check for lighting change or noise\n"
                f"📊 b* basis: {data}")

        return ""


    def _refresh_charts(self):
        an = [img for img in self.images
              if not np.isnan(img.get("s_mean", np.nan))]
        for key, cell in self._charts.items():
            cell["fig"].clear()
            self._draw_chart(cell["fig"], key, large=False)
            cell["cv"].draw()
            # 코멘트 업데이트
            cmt = self._make_chart_comment(key, an)
            self._set_cmt(cell["cmt_widget"], cmt)

    def _draw_chart(self, fig: plt.Figure, key: str, large: bool):
        an = [img for img in self.images
              if not np.isnan(img.get("s_mean",np.nan))]
        bg_ = PANEL
        fs_t = 12 if large else 9
        fs_a = 10 if large else 8
        fs_k = 9  if large else 7
        lw   = 2.5 if large else 1.8
        ms   = 8   if large else 5
        fig.patch.set_facecolor(bg_)

        def no_data():
            ax=fig.add_subplot(111); ax.axis("off")
            ax.text(0.5,0.5,_L("데이터 없음","No data"),
                    transform=ax.transAxes,
                    ha="center",va="center",color=SUB,fontsize=10)

        if not an: no_data(); return
        conds = list(dict.fromkeys(img["cond"] for img in an))
        def df(d):
            try: return float(d)
            except: return 9999

        def sa(ax):
            styled_ax(ax, bg_)
            ax.title.set_color(TXT)
            ax.xaxis.label.set_color(SUB)
            ax.yaxis.label.set_color(SUB)

        # S채널 추이
        if key == "s_trend":
            ax=fig.add_subplot(111); sa(ax)
            ax.set_title(_L("S채널 평균  (채도↓ = 산화↑)",
                            "S-ch Mean  (lower=oxidized)"),
                         fontsize=fs_t, pad=6)
            ax.set_xlabel("Day",fontsize=fs_a)
            ax.set_ylabel(_L("S 평균","S Mean"),fontsize=fs_a)
            for ci,cond in enumerate(conds):
                col=COND_COLORS[ci%len(COND_COLORS)]
                pts=sorted([(img["day"],img["s_mean"])
                             for img in an if img["cond"]==cond],
                           key=lambda x:df(x[0]))
                if pts:
                    xs=[df(p[0]) for p in pts]
                    ys=[p[1] for p in pts]
                    ax.plot(xs,ys,"o-",color=col,lw=lw,ms=ms,label=cond)
                    for x,y in zip(xs,ys):
                        ax.annotate(f"{y:.1f}",(x,y),
                                    xytext=(0,6),
                                    textcoords="offset points",
                                    ha="center",fontsize=fs_k,color=col)
            ax.legend(fontsize=fs_k,framealpha=0.8,
                      edgecolor=BORDER)
            fig.tight_layout(pad=1.2)

        # 황색 잔존 비율 + YI 듀얼 축
        elif key == "yr_trend":
            ax=fig.add_subplot(111); sa(ax)
            ax.set_title(_L("황색 잔존 비율 & YI  (↓ = 산화 진행)",
                            "Yellow Ratio & YI  (lower=oxidized)"),
                         fontsize=fs_t, pad=6)
            ax.set_xlabel("Day",fontsize=fs_a)
            ax.set_ylabel(_L("황색 비율 (%)","Yellow Ratio (%)"),
                          fontsize=fs_a)

            ax2 = ax.twinx()
            ax2.set_ylabel(_L("YI (Yellowness Index)","YI"),
                           fontsize=fs_a, color=PURPLE)
            ax2.tick_params(colors=PURPLE, labelsize=fs_k)

            for ci,cond in enumerate(conds):
                col=COND_COLORS[ci%len(COND_COLORS)]
                # 황색비율 (왼쪽 y축, 실선)
                pts_yr=sorted([(img["day"],img["yellow_ratio"]*100)
                                for img in an if img["cond"]==cond],
                               key=lambda x:df(x[0]))
                if pts_yr:
                    xs=[df(p[0]) for p in pts_yr]
                    ys=[p[1] for p in pts_yr]
                    ax.plot(xs,ys,"s-",color=col,lw=lw,ms=ms,
                            label=f"{cond} Y%")
                    for x,y in zip(xs,ys):
                        ax.annotate(f"{y:.0f}%",(x,y),
                                    xytext=(0,6),
                                    textcoords="offset points",
                                    ha="center",fontsize=fs_k,color=col)
                # YI (오른쪽 y축, 점선)
                pts_yi=sorted(
                    [(img["day"], img.get("yellowness_idx",np.nan))
                     for img in an
                     if img["cond"]==cond
                     and not np.isnan(img.get("yellowness_idx",np.nan))],
                    key=lambda x:df(x[0]))
                if pts_yi:
                    xs2=[df(p[0]) for p in pts_yi]
                    ys2=[p[1] for p in pts_yi]
                    ax2.plot(xs2,ys2,"^--",color=col,lw=lw,ms=ms,
                             alpha=0.7, label=f"{cond} YI")
                    for x,y in zip(xs2,ys2):
                        ax2.annotate(f"{y:.0f}",(x,y),
                                     xytext=(0,-12),
                                     textcoords="offset points",
                                     ha="center",fontsize=fs_k,
                                     color=col, alpha=0.8)

            lines1,labels1 = ax.get_legend_handles_labels()
            lines2,labels2 = ax2.get_legend_handles_labels()
            ax.legend(lines1+lines2, labels1+labels2,
                      fontsize=fs_k, framealpha=0.8,
                      edgecolor=BORDER, ncol=2)
            fig.tight_layout(pad=1.2)

        # H채널 추이
        elif key == "h_trend":
            ax=fig.add_subplot(111); sa(ax)
            ax.set_title(_L("H채널 평균  (색조 변화)",
                            "H-ch Mean  (hue shift)"),
                         fontsize=fs_t, pad=6)
            ax.set_xlabel("Day",fontsize=fs_a)
            ax.set_ylabel(_L("H 평균","H Mean"),fontsize=fs_a)
            for ci,cond in enumerate(conds):
                col=COND_COLORS[ci%len(COND_COLORS)]
                pts=[]
                for img in an:
                    if img["cond"]!=cond: continue
                    H_ch,_,_=img["hsi"]
                    msk=img.get("mask")
                    if msk is not None:
                        pts.append((img["day"],float(np.mean(H_ch[msk]))))
                pts.sort(key=lambda x:df(x[0]))
                if pts:
                    xs,ys=zip(*pts)
                    ax.plot(xs,ys,"^-",color=col,lw=lw,ms=ms,label=cond)
                    for x,y in zip(xs,ys):
                        ax.annotate(f"{y:.1f}",(x,y),
                                    xytext=(0,6),
                                    textcoords="offset points",
                                    ha="center",fontsize=fs_k,color=col)
            ax.legend(fontsize=fs_k,framealpha=0.8,
                      edgecolor=BORDER)
            fig.tight_layout(pad=1.2)

        # 세그먼트 박스플롯
        elif key == "box":
            ch=self.ch_var.get()
            ax=fig.add_subplot(111); sa(ax)
            ax.set_title(_L(f"세그먼트별 {ch}채널 분포",
                            f"Seg {ch}-ch Distribution"),
                         fontsize=fs_t, pad=6)
            ax.set_ylabel(_L(f"{ch} 값",f"{ch} Value"),fontsize=fs_a)
            all_data,all_labels,all_colors=[],[],[]
            for ci,cond in enumerate(conds):
                col=COND_COLORS[ci%len(COND_COLORS)]
                for img in sorted([i for i in an if i["cond"]==cond],
                                  key=lambda x:df(x["day"])):
                    vals=[s["mean"] for s in img["stats"].get(ch,[])
                          if not np.isnan(s.get("mean",np.nan))]
                    if vals:
                        all_data.append(vals)
                        all_labels.append(f"{cond[:5]}\n{img['day']}d")
                        all_colors.append(col)
            if all_data:
                bp=ax.boxplot(all_data,patch_artist=True,
                              medianprops=dict(color=TXT,lw=1.5))
                for patch,c in zip(bp["boxes"],all_colors):
                    patch.set_facecolor(c); patch.set_alpha(0.5)
                ax.set_xticks(range(1,len(all_labels)+1))
                ax.set_xticklabels(all_labels,fontsize=fs_k-1,
                                   color=TXT,rotation=30,ha="right")
            fig.tight_layout(pad=1.2)

        # 히스토그램
        elif key == "hist":
            ch=self.ch_var.get()
            ch_i={"H":0,"S":1,"I":2}[ch]
            ax=fig.add_subplot(111); sa(ax)
            ax.set_title(_L(f"{ch}채널 히스토그램  (실선=초기, 점선=최후)",
                            f"{ch}-ch Histogram (solid=first, dash=last)"),
                         fontsize=fs_t, pad=6)
            ax.set_xlabel(_L(f"{ch} 값","Value"),fontsize=fs_a)
            ax.set_ylabel(_L("밀도","Density"),fontsize=fs_a)
            for ci,cond in enumerate(conds):
                col=COND_COLORS[ci%len(COND_COLORS)]
                imgs_c=sorted([i for i in an if i["cond"]==cond],
                               key=lambda x:df(x["day"]))
                pairs=([imgs_c[0],imgs_c[-1]] if len(imgs_c)>1
                       else imgs_c)
                for li,img in enumerate(pairs):
                    msk=img.get("mask")
                    cd=img["hsi"][ch_i]
                    if msk is None: continue
                    vals=cd[msk].astype(np.float32)
                    hist,edges=np.histogram(vals,bins=64,
                                            range=(0,255),density=True)
                    centers=(edges[:-1]+edges[1:])/2
                    ax.plot(centers,hist,"-" if li==0 else "--",
                            color=col,
                            alpha=0.9 if li==0 else 0.55,
                            lw=lw,
                            label=f"{cond[:7]} {img['day']}d")
            ax.legend(fontsize=fs_k,framealpha=0.8,
                      edgecolor=BORDER,ncol=2)
            fig.tight_layout(pad=1.2)

        # 감소율 바차트
        elif key == "decay":
            ax=fig.add_subplot(111); sa(ax)
            ax.set_title(_L("조건별 S채널·황색비율 감소율",
                            "Decay Rate by Condition"),
                         fontsize=fs_t, pad=6)
            ax.set_ylabel(_L("감소율 (%)","Decay (%)"),fontsize=fs_a)
            bl,bv,by,bc=[],[],[],[]
            for ci,cond in enumerate(conds):
                pts=sorted([i for i in an if i["cond"]==cond],
                           key=lambda x:df(x["day"]))
                if len(pts)>=2:
                    s0=pts[0]["s_mean"]; sN=pts[-1]["s_mean"]
                    y0=pts[0]["yellow_ratio"]*100
                    yN=pts[-1]["yellow_ratio"]*100
                    if s0>0:
                        bl.append(cond[:14])
                        bv.append(max((s0-sN)/s0*100,0))
                        by.append(max(y0-yN,0))
                        bc.append(COND_COLORS[ci%len(COND_COLORS)])
            if bv:
                xp=np.arange(len(bl)); w=0.35
                b1=ax.bar(xp-w/2,bv,w,color=bc,alpha=0.85,
                          label=_L("S채널 감소율","S Decay"))
                b2=ax.bar(xp+w/2,by,w,color=bc,alpha=0.4,
                          hatch="//",
                          label=_L("황색비율 감소","Yellow Decay"))
                for bar,val in zip(b1,bv):
                    ax.text(bar.get_x()+bar.get_width()/2,
                            bar.get_height()+0.5,f"{val:.1f}%",
                            ha="center",va="bottom",
                            color=TXT,fontsize=fs_k,fontweight="bold")
                for bar,val in zip(b2,by):
                    ax.text(bar.get_x()+bar.get_width()/2,
                            bar.get_height()+0.5,f"{val:.1f}%",
                            ha="center",va="bottom",
                            color=TXT,fontsize=fs_k)
                ax.set_xticks(xp)
                ax.set_xticklabels(bl,color=TXT,
                                   fontsize=fs_k,rotation=15,ha="right")
                ax.set_ylim(0,max(max(bv),max(by) if by else 0)*1.25+5)
                ax.legend(fontsize=fs_k,framealpha=0.8,edgecolor=BORDER)
            else:
                ax.text(0.5,0.5,
                        _L("조건당 2개 이상 필요",
                           "Need ≥2 days per condition"),
                        transform=ax.transAxes,
                        ha="center",va="center",color=SUB,fontsize=10)
            fig.tight_layout(pad=1.2)

    # ─────────────────────────────────────────
    #  전체 분석
    # ─────────────────────────────────────────
    def _run_all(self):
        if not self.images:
            messagebox.showwarning(_L("주의","Warning"),"이미지를 추가한다."); return

        no_roi=[img["name"] for img in self.images if not img.get("roi")]
        if no_roi:
            if not messagebox.askyesno(
                _L("ROI 미완료","ROI Incomplete"),
                f"No ROI ({len(no_roi)} images):\n"
                + "\n".join(f"  • {n}" for n in no_roi[:5])
                + ("\n  ..." if len(no_roi)>5 else "")
                + "\n\nAnalyze only images with ROI set?"):
                return


        for img in self.images:
            img["day"]  = img["day_var"].get().strip() \
                          if "day_var" in img else img.get("day","")
            img["cond"] = img["cond_var"].get().strip() \
                          if "cond_var" in img else img.get("cond","")
            if not img.get("roi"): continue

            roi=img["roi"]
            if img.get("mask") is None:
                img["mask"]=roi_to_mask(img["rgb"].shape,roi)
            H_ch,S_ch,I_ch=img["hsi"]
            mask=img["mask"]

            for ch_name,ch_data in [("H",H_ch),("S",S_ch),("I",I_ch)]:
                rows = self.rows_var.get()
                cols = self.cols_var.get()
                img["stats"][ch_name]=seg_stats(
                    ch_data, mask, roi, rows, cols,
                    min_pix=self.cfg_min_pix.get())

            img["s_mean"]         = compute_s_mean(img["rgb"], mask)
            img["yellow_ratio"]   = compute_yellow_ratio(
                img["rgb"], mask,
                h_lo_deg=self.cfg_h_lo.get(),
                h_hi_deg=self.cfg_h_hi.get(),
                s_thresh=self.cfg_s_thresh.get())
            img["yellowness_idx"] = compute_yellowness_index(img["rgb"], mask)
            img["lab"]            = compute_lab_metrics(img["rgb"], mask)
            img["glcm"]           = compute_glcm_metrics(img["rgb"], mask)
            img["thumb"]          = make_thumb(img["rgb"],self.TW,self.TH,roi)

        # ΔE 계산 — 조건별로 가장 이른 날짜를 기준(ref)으로
        def _day_f(d):
            try: return float(d)
            except: return 9999

        conds_found = list(dict.fromkeys(
            img["cond"] for img in self.images if img.get("roi")))
        for cond in conds_found:
            imgs_c = sorted(
                [img for img in self.images
                 if img.get("roi") and img["cond"] == cond],
                key=lambda x: _day_f(x["day"]))
            if not imgs_c: continue
            ref_lab = imgs_c[0].get("lab")
            for img in imgs_c:
                img["delta_e"] = compute_delta_e(
                    img.get("lab", {}), ref_lab) if ref_lab else np.nan

        self._rebuild_list()
        self._refresh_orig()
        self._refresh_hsi()
        self._update_table()
        self._update_heatmap()
        self._refresh_compare()
        self._refresh_charts()
        self._refresh_color_tab()
        self._refresh_raman_tab()
        self._set_status(_L(f"분석 완료 — {len(self.images)}개", f"Analysis done — {len(self.images)} images"))

    # ─────────────────────────────────────────
    #  이미지 추가
    # ─────────────────────────────────────────
    def _sort_images_by_cond_day(self):
        """self.images 를 cond → day(숫자) → name 으로 정렬. sel_idx 보정.

        이미지 추가/일괄 로드 후 호출 — 그래프 X축이 입력 순서가 아닌
        실험 조건+날짜 순서로 그려지도록 보장.
        """
        if not self.images:
            return
        cur_obj = (self.images[self.sel_idx]
                   if 0 <= self.sel_idx < len(self.images) else None)
        def _key(img):
            cond = (img.get("cond") or "").strip()
            try:
                day = float(img.get("day") or 9999)
            except (ValueError, TypeError):
                day = 9999.0
            return (cond, day, img.get("name") or "")
        self.images.sort(key=_key)
        if cur_obj is not None:
            try:
                self.sel_idx = self.images.index(cur_obj)
            except ValueError:
                pass

    def _add_image(self, pil_img, name, defer_rebuild=False):
        """이미지 1장을 self.images 에 추가.

        defer_rebuild=True 이면 _rebuild_list / 상태바 갱신을 호출자가 마지막에
        한 번만 하도록 위임 — N장 일괄 추가 시 O(N²) → O(N).
        """
        pil_rgb = pil_img.convert("RGB")
        rgb = np.array(pil_rgb)
        H_ch,S_ch,I_ch = rgb_to_hsi(rgb)
        auto_day, auto_cond = parse_filename_tags(name)
        auto_parsed = bool(auto_day or auto_cond)
        # 자동 ROI 추정 — paper 분리 + bounding box + 품질 플래그
        auto_roi, roi_flag, roi_reason = auto_detect_roi(rgb, cond=auto_cond)
        entry = {
            "name":name, "rgb":rgb,
            "hsi":(H_ch,S_ch,I_ch),
            "mask": roi_to_mask(rgb.shape, auto_roi),
            "roi": auto_roi,
            "roi_flag": roi_flag,
            "roi_reason": roi_reason,
            "stats":{}, "thumb":None,
            "day_var":  tk.StringVar(value=auto_day),
            "cond_var": tk.StringVar(value=auto_cond),
            "day":auto_day, "cond":auto_cond,
            "auto_parsed":auto_parsed,
            "s_mean":np.nan, "yellow_ratio":np.nan, "yellowness_idx":np.nan,
            "lab": {"L":np.nan,"a":np.nan,"b":np.nan},
            "glcm": {"contrast":np.nan,"energy":np.nan,
                     "homogeneity":np.nan,"correlation":np.nan},
            "delta_e": np.nan,
        }
        entry["thumb"] = make_thumb(rgb, self.TW, self.TH, auto_roi)
        self.images.append(entry)
        if len(self.images)==1: self.sel_idx=0
        if defer_rebuild:
            return
        self._sort_images_by_cond_day()
        self._rebuild_list()
        sym, ko, en = _ROI_FLAG_LABEL.get(roi_flag, ("•", roi_flag, roi_flag))
        self._set_status(
            _L(f"추가: {name}", f"Added: {name}")
            + (f"  →  day={auto_day}, cond={auto_cond}" if auto_parsed else "")
            + _L(f"  | 자동 ROI {sym} {ko}", f"  | Auto ROI {sym} {en}"))

    def _add_images_bulk(self, items):
        """이미지 여러 장 일괄 추가. items = [(pil_img, name), ...].

        무거운 작업(PIL decode, rgb_to_hsi, auto_detect_roi, make_thumb)을
        ThreadPoolExecutor 로 병렬 처리한 뒤 메인 스레드에서 tk.StringVar
        를 만들고 dict 를 조립해 self.images 에 append. 마지막에 _rebuild_list
        를 단 한 번만 호출.

        반환: 추가 성공한 이미지 수, 실패 메시지 리스트.
        """
        if not items:
            return 0, []
        from concurrent.futures import ThreadPoolExecutor

        TW, TH = self.TW, self.TH

        def _prep(pil_img, name):
            try:
                pil_rgb = pil_img.convert("RGB")
                rgb = np.array(pil_rgb)
                hsi = rgb_to_hsi(rgb)
                day, cond = parse_filename_tags(name)
                roi, flag, reason = auto_detect_roi(rgb, cond=cond)
                thumb = make_thumb(rgb, TW, TH, roi)
                mask = roi_to_mask(rgb.shape, roi)
                return ("ok", name, rgb, hsi, day, cond,
                        roi, flag, reason, thumb, mask)
            except Exception as ex:
                return ("err", name, str(ex))

        max_workers = max(2, min(8, (os.cpu_count() or 4)))
        results = []
        with ThreadPoolExecutor(max_workers=max_workers) as ex:
            for r in ex.map(lambda it: _prep(*it), items):
                results.append(r)

        errors = []
        added = 0
        for r in results:
            if r[0] != "ok":
                errors.append(f"{r[1]}: {r[2]}")
                continue
            (_, name, rgb, hsi, day, cond,
             roi, flag, reason, thumb, mask) = r
            entry = {
                "name": name, "rgb": rgb, "hsi": hsi,
                "mask": mask, "roi": roi,
                "roi_flag": flag, "roi_reason": reason,
                "stats": {}, "thumb": thumb,
                "day_var":  tk.StringVar(value=day),
                "cond_var": tk.StringVar(value=cond),
                "day": day, "cond": cond,
                "auto_parsed": bool(day or cond),
                "s_mean": np.nan, "yellow_ratio": np.nan,
                "yellowness_idx": np.nan,
                "lab": {"L": np.nan, "a": np.nan, "b": np.nan},
                "glcm": {"contrast": np.nan, "energy": np.nan,
                         "homogeneity": np.nan, "correlation": np.nan},
                "delta_e": np.nan,
            }
            self.images.append(entry)
            added += 1

        if added > 0 and self.sel_idx < 0:
            self.sel_idx = 0
        if added > 0:
            self._sort_images_by_cond_day()
            self._rebuild_list()
        return added, errors

    def _load(self):
        paths=filedialog.askopenfilenames(
            title=_L("이미지 선택","Select Images"),
            filetypes=[(_L("이미지","Image"),"*.png *.jpg *.jpeg *.bmp *.tiff"),
                       (_L("전체","All"),"*.*")])
        if not paths: return
        items = []
        for p in paths:
            try:
                items.append((Image.open(p), os.path.basename(p)))
            except Exception as ex:
                messagebox.showerror(_L("오류","Error"), f"{p}\n{ex}")
        added, errors = self._add_images_bulk(items)
        if errors:
            messagebox.showwarning(
                _L("일부 실패","Partial failure"),
                "\n".join(errors[:8]))
        self._set_status(
            _L(f"✓ {added}개 추가", f"✓ {added} added"))

    def _paste(self):
        pil=None
        try:
            from PIL import ImageGrab
            pil=ImageGrab.grabclipboard()
        except Exception: pass
        if pil is None:
            for cmd in (["xclip","-selection","clipboard","-t","image/png","-o"],
                        ["xsel","--clipboard","--output"]):
                try:
                    import subprocess
                    r=subprocess.run(cmd,capture_output=True,timeout=3)
                    if r.returncode==0 and r.stdout:
                        pil=Image.open(io.BytesIO(r.stdout)); break
                except Exception: pass
        if pil is None:
            messagebox.showwarning(_L("클립보드","Clipboard"),_L("이미지가 없다.","No image found.")); return
        if isinstance(pil,list): pil=Image.open(pil[0]) if pil else None
        if pil is None: return
        self._add_image(pil,
            f"clipboard_{datetime.datetime.now().strftime('%H%M%S')}.png")

    def _on_drop(self, event):
        self.configure(bg=BG)
        paths=parse_drop_paths(event.data)
        imgs=[p for p in paths
              if os.path.splitext(p)[1].lower() in _IMG_EXTS]
        if not imgs:
            self.sv.set(_L("⚠ 이미지 파일 없음","⚠ No image files")); return
        # ── Evaluation 탭 활성화 시: 평가대상으로 추가 ──
        try:
            if self._atab.get() == "predict":
                cap = PRED_MAX_TARGETS - len(self._pred_targets)
                if cap <= 0:
                    messagebox.showwarning(
                        _L("평가대상 가득", "Targets full"),
                        _L(f"이미 {PRED_MAX_TARGETS}개입니다. 기존 항목을 삭제 후 다시 시도하세요.",
                           f"Already at {PRED_MAX_TARGETS}. Remove some first."))
                    return
                accepted = imgs[:cap]
                skipped  = len(imgs) - cap
                added    = 0
                for p in accepted:
                    try:
                        self._pred_add_target(Image.open(p),
                                              os.path.basename(p))
                        added += 1
                    except Exception as ex:
                        messagebox.showerror(_L("오류","Error"),
                                             f"{p}\n{ex}")
                msg = _L(f"✓ 평가대상 {added}개 추가",
                         f"✓ {added} target(s) added")
                if skipped > 0:
                    msg += _L(f"  (한도 초과로 {skipped}개 스킵)",
                              f"  ({skipped} skipped — over limit)")
                self._set_status(msg)
                return
        except Exception:
            pass
        items = []
        for p in imgs:
            try:
                items.append((Image.open(p), os.path.basename(p)))
            except Exception as ex:
                messagebox.showerror(_L("오류","Error"), f"{p}\n{ex}")
        added, errors = self._add_images_bulk(items)
        if errors:
            messagebox.showwarning(
                _L("일부 실패","Partial failure"),
                "\n".join(errors[:8]))
        self._set_status(_L(f"✓ {added}개 추가 — ROI를 선택한다",
                            f"✓ {added} added — set ROI"))

    # ─────────────────────────────────────────
    #  선택/동기화/유틸
    # ─────────────────────────────────────────
    def _refresh_card_border(self, idx):
        """카드 테두리/배경만 갱신 — 선택 변경 시 사용 (전체 재빌드 회피)."""
        cards = getattr(self, "_cards_by_idx", None)
        if not cards or idx is None or idx not in cards:
            return
        card = cards[idx]
        try:
            if not card.winfo_exists():
                return
        except tk.TclError:
            return
        if idx < 0 or idx >= len(self.images):
            return
        img = self.images[idx]
        is_sel = (idx == self.sel_idx)
        has_roi = img.get("roi") is not None
        roi_flag = img.get("roi_flag")
        inconsistent = bool(img.get("_roi_inconsistent", False))
        if is_sel:
            brd = ACCENT
        else:
            brd = _border_color_for_roi(roi_flag, has_roi,
                {"green": GREEN, "amber": AMBER, "red": RED,
                 "purple": PURPLE, "border": BORDER},
                inconsistent=inconsistent)
        thick = 2 if (is_sel
                      or roi_flag in ("warn_small", "warn_off", "warn_paper", "failed")
                      or (inconsistent and roi_flag != "failed"
                          and roi_flag not in ("warn_small", "warn_off", "warn_paper"))) else 1
        new_bg = CARD2 if is_sel else CARD
        try:
            card.configure(bg=new_bg, highlightbackground=brd,
                           highlightthickness=thick)
            # 자식 위젯의 bg 도 카드와 일치하도록 재귀 (CARD/CARD2 만 교체)
            def _update_bg(w):
                try:
                    cur = w.cget("bg")
                    if cur in (CARD, CARD2):
                        w.configure(bg=new_bg)
                except tk.TclError:
                    pass
                for c in w.winfo_children():
                    _update_bg(c)
            _update_bg(card)
        except tk.TclError:
            pass

    def _select(self, idx):
        old = self.sel_idx
        self.sel_idx = idx
        # 전체 재빌드 대신 두 카드(이전/현재) 의 테두리만 갱신 — 빠름
        try:
            if old != idx and old is not None and old >= 0:
                self._refresh_card_border(old)
            self._refresh_card_border(idx)
        except Exception:
            # 카드 캐시 없거나 stale 이면 안전 폴백
            self._rebuild_list()
        self._refresh_orig()
        self._refresh_hsi()
        self._update_table()
        if self._atab.get() not in ("detail",):
            self._switch("detail")

    def _sync(self, idx):
        img=self.images[idx]
        img["day"] =img["day_var"].get().strip()
        img["cond"]=img["cond_var"].get().strip()

    def _set_cond(self, idx, label):
        img=self.images[idx]
        img["cond_var"].set(label); img["cond"]=label
        self._rebuild_list()

    def _apply_preset(self, label):
        if self.sel_idx<0 or self.sel_idx>=len(self.images):
            messagebox.showinfo(_L("알림","Info"),_L("이미지를 먼저 선택한다.","Select an image first.")); return
        self._set_cond(self.sel_idx, label)

    def _remove(self, idx):
        self.images.pop(idx)
        if self.sel_idx>=len(self.images):
            self.sel_idx=len(self.images)-1
        self._rebuild_list(); self._refresh_orig()

    def _auto_roi_all_unmanual(self):
        """사용자 manual 또는 DB 로드 ROI 는 보호하고 나머지에 자동 ROI 재추정."""
        if not self.images:
            self._set_status(_L("이미지 없음", "No images"))
            return
        targets = [i for i, img in enumerate(self.images)
                   if img.get("roi_flag") != "manual"
                   and img.get("roi_source") != "db"]
        skipped = len(self.images) - len(targets)
        if not targets:
            self._set_status(_L(f"전부 사용자 설정 ROI — 자동 적용 대상 없음 (보호된 {skipped}장)",
                                f"All ROIs are manual — none to auto-apply ({skipped} protected)"))
            return
        if not messagebox.askyesno(
                _L("자동 ROI 재실행", "Re-run Auto ROI"),
                _L(f"이미지 {len(targets)}장에 자동 ROI 재추정합니다.\n"
                   f"(직접 설정한 {skipped}장은 보호됨)\n\n계속할까요?",
                   f"Re-run auto ROI on {len(targets)} images.\n"
                   f"({skipped} manual ROIs protected)\n\nContinue?")):
            return
        flag_cnt = {"good": 0, "warn_small": 0, "warn_off": 0,
                    "warn_paper": 0, "failed": 0}
        for i in targets:
            img = self.images[i]
            new_roi, new_flag, new_reason = auto_detect_roi(
                img["rgb"], cond=img.get("cond"))
            img["roi"] = new_roi
            img["mask"] = roi_to_mask(img["rgb"].shape, new_roi)
            img["roi_flag"] = new_flag
            img["roi_reason"] = new_reason
            img["thumb"] = make_thumb(img["rgb"], self.TW, self.TH, new_roi)
            # 분석 결과는 무효화 (ROI 변경됨)
            img["stats"] = {}
            img["s_mean"] = np.nan
            img["yellow_ratio"] = np.nan
            flag_cnt[new_flag] = flag_cnt.get(new_flag, 0) + 1
        self._rebuild_list()
        try: self._refresh_orig()
        except Exception: pass
        self._set_status(_L(
            f"✓ 자동 ROI 완료 — OK {flag_cnt['good']} | 경고 {flag_cnt['warn_small']+flag_cnt['warn_off']+flag_cnt['warn_paper']} | 실패 {flag_cnt['failed']} | 보호 {skipped}",
            f"✓ Auto ROI done — OK {flag_cnt['good']} | Warn {flag_cnt['warn_small']+flag_cnt['warn_off']+flag_cnt['warn_paper']} | Fail {flag_cnt['failed']} | Protected {skipped}"))

    def _delete_all_images(self):
        """이미지 목록만 전체 삭제 (Raman 데이터는 보존)."""
        if not self.images:
            self._set_status(_L("삭제할 이미지 없음", "No images to delete"))
            return
        n = len(self.images)
        if not messagebox.askyesno(
                _L("확인", "Confirm"),
                _L(f"이미지 {n}장을 모두 삭제할까요?\n(Raman 데이터는 유지됩니다)",
                   f"Delete all {n} images?\n(Raman data will be kept)")):
            return
        self.images.clear()
        self.sel_idx = -1
        self._refs.clear()
        self._rebuild_list()
        # 이미지 관련 캔버스/차트 초기화 (Raman 은 건드리지 않음)
        try: self._orig_cv.delete("all")
        except Exception: pass
        try: self._hsi_cv.delete("all")
        except Exception: pass
        try: self.tree.delete(*self.tree.get_children())
        except Exception: pass
        try:
            self.hm_fig.clear(); self.hm_cv.draw()
        except Exception: pass
        for c in self._charts.values():
            try: c["fig"].clear(); c["cv"].draw()
            except Exception: pass
        for c in self._color_charts.values():
            try: c["fig"].clear(); c["cv"].draw()
            except Exception: pass
        self._set_status(_L(f"이미지 {n}장 전체 삭제 완료", f"Deleted {n} images"))

    def _clear(self):
        if self.images and not messagebox.askyesno(
                _L("확인","Confirm"),_L("전체 초기화할까요?","Clear all data?")): return
        self.images.clear(); self.sel_idx=-1; self._refs.clear()
        self._rebuild_list()
        self._orig_cv.delete("all")
        self._hsi_cv.delete("all")
        self.tree.delete(*self.tree.get_children())
        self.hm_fig.clear(); self.hm_cv.draw()
        for c in self._charts.values():
            c["fig"].clear(); c["cv"].draw()
        for c in self._color_charts.values():
            c["fig"].clear(); c["cv"].draw()
        for c in self._raman_charts.values():
            c["fig"].clear(); c["cv"].draw()
        self._set_status(_L("초기화 완료","Cleared"))

    def _export_csv(self):
        if not any(img["stats"] for img in self.images):
            messagebox.showwarning(_L("주의","Warning"),"먼저 분석을 실행한다."); return
        path=filedialog.asksaveasfilename(
            defaultextension=".csv",filetypes=[("CSV","*.csv")])
        if not path: return
        rows_out=[]
        for img in self.images:
            for ch in ("H","S","I"):
                for s in img["stats"].get(ch,[]):
                    rows_out.append({
                        "file":img["name"],"day":img["day"],
                        "condition":img["cond"],"channel":ch,
                        "seg":s["seg"],"row":s["row"],"col":s["col"],
                        "pixels":s["pixels"],
                        "mean":round(s.get("mean",np.nan),3),
                        "std": round(s.get("std",np.nan),3),
                        "s_mean_roi":round(img.get("s_mean",np.nan),3),
                        "yellow_ratio":round(img.get("yellow_ratio",np.nan),4),
                        "yellowness_idx":round(img.get("yellowness_idx",np.nan),2),
                        "lab_L":  round(img.get("lab",{}).get("L",np.nan),2),
                        "lab_a":  round(img.get("lab",{}).get("a",np.nan),2),
                        "lab_b":  round(img.get("lab",{}).get("b",np.nan),2),
                        "delta_e":round(img.get("delta_e",np.nan),2),
                        "glcm_contrast":   round(img.get("glcm",{}).get("contrast",np.nan),3),
                        "glcm_energy":     round(img.get("glcm",{}).get("energy",np.nan),6),
                        "glcm_homogeneity":round(img.get("glcm",{}).get("homogeneity",np.nan),4),
                    })
        if not rows_out:
            messagebox.showinfo(_L("알림","Info"),"데이터 없음"); return
        with open(path,"w",newline="",encoding="utf-8-sig") as f:
            w=csv.DictWriter(f,fieldnames=list(rows_out[0].keys()))
            w.writeheader(); w.writerows(rows_out)
        messagebox.showinfo(_L("완료","Done"),f"저장:\n{path}")

    def _set_lang(self, ko: bool):
        """UI 언어 전환 — 모든 탭 텍스트·차트·코멘트 즉시 갱신"""
        set_lang(ko)

        # 언어 버튼 색상
        self._lang_btn_ko.configure(
            bg=ACCENT if ko else BTN, fg="white" if ko else TXT)
        self._lang_btn_en.configure(
            bg=BTN if ko else ACCENT, fg=TXT if ko else "white")

        # 탭 버튼 텍스트
        tab_labels = {
            "detail":   _L("🔍 ROI·상세",   "🔍 ROI·Detail"),
            "compare":  _L("🔲 조건 비교",   "🔲 Condition Grid"),
            "chart":    _L("📈 차트",          "📈 Charts"),
            "color":    _L("🔬 컬러 분석",    "🔬 Color Analysis"),
            "raman":    _L("📡 Raman 분석",   "📡 Raman Analysis"),
            "predict":  _L("🎯 평가 대상",    "🎯 Evaluation"),
            "settings": _L("⚙ 설정",           "⚙ Settings"),
        }
        for key, lbl in tab_labels.items():
            if key in self._tbts:
                self._tbts[key].configure(text=lbl)

        # 현재 탭 콘텐츠 즉시 갱신
        active = self._atab.get()
        if active == "detail":
            self._refresh_orig()
            self._refresh_hsi()
            self._update_stat_panel()
        elif active == "compare":
            self._refresh_compare()
        elif active == "chart":
            self._refresh_charts()
        elif active == "color":
            self._refresh_color_tab()
        elif active == "raman":
            self._refresh_raman_tab()

        # 비활성 탭 코멘트도 100ms 후 백그라운드 갱신
        self.after(100, self._refresh_all_comments)

        self._set_status(_L("한국어로 전환되었다.",
                             "Switched to English."))

    def _refresh_all_comments(self):
        """언어 전환 시 모든 차트 코멘트 갱신"""
        an_s = [img for img in self.images
                if not np.isnan(img.get("s_mean", np.nan))]
        an_c = [img for img in self.images
                if img.get("roi") and
                not np.isnan(img.get("lab",{}).get("b", np.nan))]
        for key, cell in self._charts.items():
            self._set_cmt(cell["cmt_widget"],
                          self._make_chart_comment(key, an_s))
        for key, cell in self._color_charts.items():
            self._set_cmt(cell["cmt_widget"],
                          self._make_chart_comment(key, an_c))
        self._update_stat_panel()


    def _set_status(self, msg):
        self.sv.set(
            f"[{datetime.datetime.now().strftime('%H:%M:%S')}]  {msg}")
        self.update_idletasks()


# ══════════════════════════════════════════════
if __name__ == "__main__":
    App().mainloop()
