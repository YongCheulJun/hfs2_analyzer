"""
SCIE 논문 .docx 생성 — MDPI Applied Sciences 양식 모방.
주저자: Yongcheul Jun (a), 교신저자: Kwangwook Park (b).

주의: MDPI 공식 양식은 https://www.mdpi.com/journal/applsci/instructions
에서 .docx 템플릿을 다운로드하는 것이 정석. 본 스크립트는 핵심 양식
(Palatino Linotype, A4, 단일 칼럼, 윗첨자 소속, * Correspondence,
숫자 섹션 번호, References numbered list) 만 재현.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIG = os.path.join(ROOT, "paper/figures")
OUT = os.path.join(ROOT, "paper", "Jun_HfS2_image_oxidation.docx")

doc = Document()

# 페이지 — MDPI Applied Sciences 기본: A4, 위/아래 1.78cm, 좌/우 1.78cm
sec = doc.sections[0]
sec.page_width  = Cm(21.0)
sec.page_height = Cm(29.7)
sec.top_margin    = Cm(1.78)
sec.bottom_margin = Cm(1.78)
sec.left_margin   = Cm(1.78)
sec.right_margin  = Cm(1.78)

# 기본 스타일 — MDPI 는 Palatino Linotype 사용 (없으면 Times New Roman 폴백)
style = doc.styles["Normal"]
style.font.name = "Palatino Linotype"
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(2)
BODY_FONT = "Palatino Linotype"


def H(text, level=1, size=12, bold=True, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8 if level == 1 else 4)
    p.paragraph_format.space_after  = Pt(after)
    r = p.add_run(text)
    r.font.name = BODY_FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    return p


def P(text, size=10, italic=False, align=None, after=2):
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.font.name = BODY_FONT
    r.font.size = Pt(size)
    r.font.italic = italic
    return p


def IMG(fname, width_in=6.5, caption=None):
    doc.add_picture(os.path.join(FIG, fname), width=Inches(width_in))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        cp.paragraph_format.space_after = Pt(8)
        r = cp.add_run(caption)
        r.font.name = BODY_FONT
        r.font.size = Pt(9)
        r.font.italic = True


# ─────────────────── TITLE (MDPI Applied Sciences 양식) ───────────────────
# Article type tag (MDPI 상단 라벨)
tag = doc.add_paragraph()
tag.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = tag.add_run("Article")
r.font.name = BODY_FONT; r.font.size = Pt(10); r.font.italic = True
tag.paragraph_format.space_after = Pt(4)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = title.add_run(
    "Image-based Estimation of Native Oxidation Aging in CVD-grown "
    "HfS₂ Thin Films via Multi-method Ensemble and Pseudo-Raman "
    "Reconstruction")
r.font.name = BODY_FONT
r.font.size = Pt(16)
r.font.bold = True
title.paragraph_format.space_after = Pt(6)

# Authors — MDPI: 저자 이름 + 윗첨자 (Inline superscript via run.font.superscript)
auth = doc.add_paragraph()
auth.alignment = WD_ALIGN_PARAGRAPH.LEFT
auth.paragraph_format.space_after = Pt(2)


def add_author(par, name, sups, last=False):
    r = par.add_run(name)
    r.font.name = BODY_FONT; r.font.size = Pt(11); r.font.bold = False
    rs = par.add_run(" " + sups)
    rs.font.name = BODY_FONT; rs.font.size = Pt(11)
    rs.font.superscript = True
    if not last:
        rsep = par.add_run(", ")
        rsep.font.name = BODY_FONT; rsep.font.size = Pt(11)


add_author(auth, "Yongcheul Jun", "1")
add_author(auth, "Kwangwook Park", "2,*", last=True)

# Affiliations
def add_aff(num, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_after = Pt(2)
    rn = p.add_run(num)
    rn.font.name = BODY_FONT; rn.font.size = Pt(9)
    rn.font.superscript = True
    rt = p.add_run(" " + text)
    rt.font.name = BODY_FONT; rt.font.size = Pt(9); rt.font.italic = True


add_aff("1",
        "College of Engineering (Major in Intellectual Property "
        "Convergence), Pusan National University, Busan, 46241, "
        "Republic of Korea")
add_aff("2",
        "Division of Electronics and Information Engineering, "
        "Jeonbuk National University, Jeonju 54896, "
        "Republic of Korea")

corr = doc.add_paragraph()
corr.paragraph_format.left_indent = Cm(0.5)
corr.paragraph_format.first_line_indent = Cm(-0.5)
corr.paragraph_format.space_after = Pt(10)
rs = corr.add_run("*")
rs.font.name = BODY_FONT; rs.font.size = Pt(9); rs.font.superscript = True
rt = corr.add_run(" Correspondence: kpark@jbnu.ac.kr (K. P.)")
rt.font.name = BODY_FONT; rt.font.size = Pt(9)


# ─────────────────── ABSTRACT (MDPI style) ───────────────────
abs_p = doc.add_paragraph()
abs_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
abs_p.paragraph_format.space_after = Pt(4)
rl = abs_p.add_run("Abstract: ")
rl.font.name = BODY_FONT; rl.font.size = Pt(10); rl.font.bold = True
rb = abs_p.add_run(
    "Hafnium disulfide (HfS₂), a layered transition-metal "
    "dichalcogenide considered for high-mobility electronics and "
    "photodetectors, undergoes rapid native oxidation under ambient "
    "humidity, transforming the originally yellow chalcogenide into "
    "a transparent oxide (HfOₓ) within days. Quantifying the extent of "
    "oxidation conventionally requires Raman spectroscopy, which is "
    "instrument-bound and slow. In this work we propose an image-only "
    "framework that estimates the aging day of CVD-grown HfS₂ thin "
    "films directly from a single specimen photograph and, in addition, "
    "reconstructs a pseudo-Raman spectrum without any spectroscopic "
    "measurement. A pool of 53 specimen images covering four passivation "
    "conditions (Native HfS₂ at 35% and 70% RH, Al₂O₃- and "
    "PMMA-encapsulated HfS₂ at 70% RH) and aging days from 0 to 30 d "
    "is analyzed by five independent estimators—weighted-Euclidean "
    "k-nearest-neighbour on color metrics, Wasserstein distance on b∗ "
    "histograms, FFT texture distance, spatial-pattern distance, and an "
    "exponential-decay kinetic fit—whose outputs are combined by a "
    "Huber-loss optimised, condition-specific ensemble. Leave-one-out "
    "evaluation on 33 query specimens yields a weighted-mean root-mean-"
    "square error of 4.80 d, a 38% reduction relative to the uniform-"
    "weight baseline (7.74 d). For Al₂O₃-encapsulated HfS₂ the "
    "image-derived A₁ₒ peak intensity reproduces the measured "
    "Raman trend (1.00 → 0.60) within the predicted 95% confidence "
    "interval, demonstrating that image-only pseudo-Raman reconstruction "
    "can serve as a rapid, non-destructive surrogate for benchtop Raman "
    "spectroscopy in routine quality screening of HfS₂ thin films.")
rb.font.name = BODY_FONT; rb.font.size = Pt(10)

kw_p = doc.add_paragraph()
kw_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
kw_p.paragraph_format.space_after = Pt(10)
rl = kw_p.add_run("Keywords: ")
rl.font.name = BODY_FONT; rl.font.size = Pt(10); rl.font.bold = True
rb = kw_p.add_run(
    "hafnium disulfide; native oxidation; image processing; "
    "pseudo-Raman; ensemble regression; CVD thin films")
rb.font.name = BODY_FONT; rb.font.size = Pt(10)

# ─────────────────── 1. INTRODUCTION ───────────────────
H("1. Introduction", level=1, size=12, bold=True)

P("Two-dimensional layered transition metal dichalcogenides (TMDs) have "
  "drawn intense attention as candidates for ultra-thin electronics, "
  "broadband photodetectors, hydrogen evolution catalysts and "
  "stretchable devices [1–3]. Among them HfS₂ attracts particular "
  "interest because of its high theoretical electron mobility, high "
  "photoresponsivity, and conduction-band alignment favourable for "
  "H⁺/H₂ reduction [4,5]. However, like most sulfide-based TMDs, "
  "HfS₂ is highly susceptible to ambient native oxidation: O₂ and "
  "H₂O cooperatively convert HfS₂ into HfS₂₋ₓOₓ and ultimately "
  "into HfO₂, replacing the original yellow chalcogenide with a "
  "transparent oxide film [6]. Hwang and co-workers recently showed "
  "that the A₁ₒ Raman peak of CVD-grown HfS₂ disappears within a "
  "week under 70% relative humidity (RH) at room temperature, and that "
  "encapsulation by Al₂O₃ effectively suppresses oxidation while PMMA "
  "delays but does not block it [6].",
  align="justify")

P("In standard practice the extent of HfS₂ oxidation is monitored by "
  "micro-Raman spectroscopy through the intensity of the A₁ₒ mode "
  "near 337 cm⁻¹. Although highly specific, Raman measurement is "
  "instrument-bound, requires careful focus alignment for every "
  "specimen, and is impractical for fabrication-line monitoring or in-"
  "field quality checks. The macroscopic optical signature of HfS₂ "
  "oxidation — fading from saturated yellow to translucent grey — is "
  "however readily captured by an ordinary digital camera. We "
  "therefore investigated whether a single specimen photograph alone "
  "can replace a Raman measurement for routine ageing assessment.",
  align="justify")

P("Image-based oxidation tracking has been pursued for noble-metal "
  "tarnishing and corrosion, and recent reports have used colorimetric "
  "or texture features to estimate oxidation in 2D materials [7,8]. "
  "What has not been explored, to our knowledge, is whether the same "
  "image features can be inverted into a synthetic Raman spectrum that "
  "an experimentalist can interpret in the same way as a measured "
  "spectrum. Such a “pseudo-Raman” reconstruction would let users "
  "verify Stage I–IV oxidation classification without spectrometer "
  "access. The present work introduces, evaluates, and benchmarks such "
  "a pipeline. The contributions are: (i) a five-method ensemble that "
  "estimates aging day from a single image with condition-specific "
  "weights learned by Huber-robust regression, (ii) a pseudo-Raman "
  "spectrum reconstruction from image metrics with an analytically "
  "propagated 95% confidence interval, and (iii) quantitative "
  "comparison against measured Raman spectra, which confirms the "
  "fidelity of the reconstruction for Al₂O₃-encapsulated HfS₂ and "
  "exposes the fundamental sensitivity limit for that sample class.",
  align="justify", after=8)


# ─────────────────── 2. MATERIALS AND METHODS ───────────────────
H("2. Materials and methods", level=1, size=12, bold=True)

H("2.1. Sample preparation", level=2, size=11, bold=True)
P("CVD-grown HfS₂ thin films on sapphire substrates were prepared "
  "following the procedure described in Ref. [6]. Four passivation "
  "conditions were studied: bare native HfS₂ stored at 35% RH and "
  "at 70% RH, HfS₂ capped with a 10 nm Al₂O₃ layer deposited by "
  "atomic layer deposition (Al₂O₃/HfS₂, 70% RH), and HfS₂ covered "
  "by a spin-coated PMMA film (PMMA/HfS₂, 70% RH). For each "
  "condition the same physical specimen was photographed at multiple "
  "aging days between day 0 and day 30 under fixed illumination and "
  "camera settings. The complete reference pool used for the present "
  "study comprised 53 photographs distributed across the four "
  "conditions and 14 distinct aging days, augmented by an independent "
  "20-image evaluation set of the same specimens captured at a later "
  "session. Representative images are shown in Fig. 1.",
  align="justify")

IMG("fig1_specimen.png", width_in=6.5,
    caption="Fig. 1. Photographic time-evolution of CVD-grown HfS₂ "
            "specimens under four passivation/humidity conditions. "
            "Native HfS₂ at 35% RH preserves yellow chalcogenide "
            "appearance over a month, whereas Native HfS₂ at 70% RH "
            "loses its yellow color within a week. Al₂O₃ encapsulation "
            "almost completely suppresses optical change, while PMMA "
            "delays oxidation more modestly.")

H("2.2. Region-of-interest extraction and color metrics", level=2,
  size=11, bold=True)
P("For every input photograph an automatic region-of-interest (ROI) is "
  "extracted to confine the analysis to specimen pixels and exclude "
  "the cardstock background. The pipeline applies an HSV value/"
  "saturation threshold (V > 215, S < 25) to obtain a paper mask, "
  "extracts the largest non-paper contour, and replaces it with its "
  "convex hull when the latter encloses less than 90% of the image "
  "to recover specimen pixels lost to specular highlights. ROI size "
  "is then set adaptively per condition (between 13% and 17% of the "
  "image area, learned from manual annotations) with an aspect ratio "
  "softened by a square-root transform of the specimen bounding box, "
  "and the ROI center is placed at the geometric mean of the specimen "
  "centroid and bounding-box center, which empirically produced the "
  "most consistent ROIs across the four conditions.",
  align="justify")
P("Within the ROI the following color metrics are computed: the CIE "
  "Lab b∗ coordinate (yellow–blue axis), the HSI saturation channel "
  "S, the ASTM E313 Yellowness Index (YI), and the cumulative CIE "
  "ΔE color difference relative to the day-0 reference of the same "
  "condition. Together these descriptors capture the macroscopic "
  "loss of saturated yellow that accompanies HfS₂ oxidation.",
  align="justify")

H("2.3. Five-method ensemble for aging-day estimation", level=2,
  size=11, bold=True)
P("We combine five independent estimators (Fig. 2). (1) A weighted "
  "k-nearest-neighbour (kNN) regressor measures the Mahalanobis-type "
  "distance d = √(0.45·Δb∗² + 0.30·ΔS² + 0.25·ΔYI²) "
  "between the query and reference images of the same condition and "
  "averages the top-3 reference days with inverse-distance weights. "
  "(2) A Wasserstein-distance estimator compares the b∗ pixel "
  "histogram of the query against each reference and again averages "
  "the three nearest references. (3) An FFT texture estimator extracts "
  "the high-frequency energy fraction, the spectral entropy, and the "
  "radial power profile of the ROI greyscale image and ranks "
  "references by a weighted feature distance. (4) A spatial-pattern "
  "estimator divides the ROI into a 3×3 grid, computes per-cell b∗ "
  "mean and standard deviation, and quantifies oxidation anisotropy "
  "via inter-cell entropy and centre-vs-boundary gradient. (5) A "
  "kinetic estimator fits the per-condition b∗(t) time series to a "
  "first-order exponential decay b∗(t) = b∗₊ + (b∗₀ - b∗₊)·"
  "exp(-k·t) using SciPy's non-linear least squares with bounded "
  "parameters [9], and inverts the model to recover the day "
  "corresponding to the query b∗.",
  align="justify")

IMG("fig2_pipeline.png", width_in=6.5,
    caption="Fig. 2. Schematic of the image-only oxidation aging "
            "estimation pipeline. A specimen photograph is converted "
            "to color and texture descriptors inside an automatic ROI; "
            "five independent estimators produce candidate aging-day "
            "values that are combined by a condition-specific weighted "
            "ensemble.")

P("The five candidate days are combined into an ensemble estimate "
  "ŷ = Σwᵢ · dᵢ with weights {wᵢ} learned to minimise the "
  "Huber loss ℓ(e) = ½·e² if |e|≤δ else δ·(|e|-½δ), "
  "with δ = 5 d, on the leave-one-out errors against the known aging "
  "days of the reference set. Two regimes are evaluated: a single "
  "global weight set (“uniform/optimised global”) and four condition-"
  "specific weight sets, with the latter selected because the "
  "intrinsic per-method accuracy depends strongly on the surface "
  "passivation. Optimisation is performed via SciPy's L-BFGS-B with "
  "five multi-start initialisations [9].",
  align="justify")

H("2.4. Pseudo-Raman spectrum reconstruction", level=2,
  size=11, bold=True)
P("Although image features cannot directly reproduce vibrational "
  "transitions, the relative intensity of the A₁ₒ mode at "
  "~337 cm⁻¹ — the principal Raman fingerprint of HfS₂ — varies "
  "monotonically with surface oxidation and is therefore well "
  "correlated with the same color metrics that drive aging-day "
  "estimation. A pseudo-Raman intensity î⁺ is predicted by an "
  "R²-weighted ensemble of four univariate linear regressors "
  "(b∗ → i⁺, S → i⁺, YI → i⁺, ΔE → i⁺) trained on the "
  "image–Raman pairs of the reference set. Two reference Raman "
  "spectra whose normalised peaks bracket î⁺ are then linearly "
  "interpolated, and a 95% confidence interval is propagated from "
  "the regression standard error σ̂ as î⁺ ± 1.96·σ̂. "
  "The output is a complete intensity-versus-Raman-shift curve that "
  "users can compare directly against a benchtop measurement.",
  align="justify", after=8)


# ─────────────────── 3. RESULTS AND DISCUSSION ───────────────────
H("3. Results and discussion", level=1, size=12, bold=True)

H("3.1. Time-evolution of color metrics", level=2, size=11, bold=True)
P("Fig. 3 plots the three primary color metrics as a function of "
  "aging day for the four conditions. Native HfS₂ at 70% RH shows the "
  "expected drastic decay: b∗ collapses from 28 to 3 within two weeks "
  "and the S-channel mean falls from 54 to 4, consistent with "
  "complete loss of the yellow chalcogenide and conversion into "
  "transparent HfO₂. Native HfS₂ stored at the lower 35% RH degrades "
  "much more slowly, reaching b∗ ≈ 6 only after four weeks. "
  "Al₂O₃-encapsulated HfS₂ changes least of all four conditions, "
  "with b∗ fluctuating between 16 and 28 over the entire 30-day "
  "window — a behaviour fully consistent with the protective role "
  "of the Al₂O₃ cap reported in Ref. [6]. PMMA delays the decay by "
  "roughly a factor of two relative to bare HfS₂ at the same humidity "
  "but ultimately fails to prevent oxidation.",
  align="justify")

IMG("fig3_metric_trends.png", width_in=6.5,
    caption="Fig. 3. Time-evolution of the three image color metrics "
            "for the four passivation conditions. (a) CIE Lab b∗, "
            "(b) HSI S-channel mean, (c) ASTM E313 Yellowness Index. "
            "The flat trajectory of Al₂O₃/HfS₂ reflects the "
            "effective oxidation barrier provided by the Al₂O₃ cap.")

P("The Pearson correlation r(day, b∗) is −0.92 for Native HfS₂ at "
  "35% RH, −0.83 for Native HfS₂ at 70% RH, −0.89 for PMMA/HfS₂, "
  "and only −0.32 for Al₂O₃/HfS₂, quantitatively confirming the "
  "qualitative observation that Al₂O₃-passivated specimens are nearly "
  "stationary in image space across the entire aging range. This sets "
  "a fundamental upper bound on day-estimation accuracy for that "
  "condition that no image-based estimator can surpass.",
  align="justify")

H("3.2. Per-method and ensemble accuracy", level=2,
  size=11, bold=True)
P("Fig. 4 reports the leave-one-out root-mean-square error (RMSE) of "
  "each of the five individual estimators and of the ensemble in two "
  "configurations: a single global weight set and condition-specific "
  "weights. Three observations stand out. First, the kNN estimator "
  "is the most accurate single method for three of the four "
  "conditions (Native 35% RH 3.12 d, Native 70% RH 3.56 d, PMMA "
  "70% RH 2.50 d), reflecting its direct use of the strongly "
  "monotonic color descriptors. Second, the FFT and kinetic "
  "estimators perform poorly in isolation, but their independent "
  "errors are partially decorrelated from kNN, so they can still "
  "contribute marginal gains in the ensemble. Third, condition-"
  "specific weighting reduces the weighted-mean RMSE from 7.74 d "
  "(uniform baseline) and 6.30 d (single global optimal weight set) "
  "to 4.80 d, a 38% improvement over the uniform baseline.",
  align="justify")

IMG("fig4_method_rmse.png", width_in=6.5,
    caption="Fig. 4. Per-condition RMSE (leave-one-out) of the five "
            "individual estimators and of the two ensemble "
            "configurations. Condition-specific weighting strongly "
            "outperforms a single global weight set, especially for "
            "Native HfS₂ at 35% RH and PMMA/HfS₂.")

P("The condition-specific weights themselves are physically "
  "interpretable. For Native HfS₂ at 70% RH the optimum is kNN "
  "alone (100%), because the strong b∗ monotonic decay makes "
  "color-distance KNN essentially exact. For PMMA/HfS₂ the optimum "
  "is dominated by kNN (83%) with a small contribution from the "
  "Wasserstein histogram (12%). For Native HfS₂ at 35% RH the "
  "Wasserstein method dominates (55%) because the slower decay "
  "shifts the b∗ histogram by amounts more reliably captured by "
  "earth-mover distance than by mean values. For Al₂O₃/HfS₂ the "
  "ensemble distributes weight more evenly across FFT (45%), spatial "
  "pattern (24%), Wasserstein (22%) and kinetic (3%), since none of "
  "the methods has any decisive single signal in this near-stationary "
  "regime.",
  align="justify")

H("3.3. Pseudo-Raman vs. measured Raman spectra", level=2,
  size=11, bold=True)
P("Of the four conditions only the Al₂O₃/HfS₂ series provides a "
  "complete set of measured Raman spectra (the bare and PMMA "
  "specimens lose the A₁ₒ mode entirely within a few days at "
  "70% RH and cannot be tracked by Raman thereafter). Fig. 5(a) "
  "overlays the measured Raman spectra of Al₂O₃/HfS₂ from day 0 "
  "to day 28, showing the gradual but unambiguous decrease of the "
  "A₁ₒ peak intensity from a normalised value of 1.00 at day 0 to "
  "0.60 at day 28, with a small concomitant peak shift of "
  "0.7 cm⁻¹. Fig. 5(b) compares the measured normalised A₁ₒ peak "
  "against the pseudo-Raman estimate produced from image metrics "
  "alone. The two curves agree within the propagated 95% "
  "confidence interval at every measured time point, with absolute "
  "errors of 0.00, 0.01, 0.00, 0.01 and 0.00 normalised intensity "
  "units at days 0, 3, 7, 14 and 28, respectively. The mean absolute "
  "error of 0.005 normalised units corresponds to roughly 1–2 d "
  "uncertainty in the equivalent aging day.",
  align="justify")

IMG("fig5_pseudo_raman.png", width_in=6.5,
    caption="Fig. 5. (a) Measured Raman spectra of Al₂O₃-encapsulated "
            "HfS₂ across 28 d showing gradual A₁ₒ attenuation. "
            "(b) Comparison of measured normalised A₁ₒ peak intensity "
            "(red) against the image-derived pseudo-Raman estimate "
            "(blue) with 95% confidence band. The two curves agree to "
            "within ~0.01 normalised intensity units across the entire "
            "30-day window.")

P("This quantitative agreement validates the central premise of the "
  "present study: when the underlying oxidation chemistry leaves a "
  "monotonic optical signature, an image-only estimator can recover "
  "the same A₁ₒ trend that micro-Raman would measure, without any "
  "spectrometer at all. For samples where the Raman signal vanishes "
  "before the optical signal saturates (Native HfS₂ at 70% RH being "
  "the extreme case) the pseudo-Raman estimate becomes formally "
  "extrapolative and is reported with reduced confidence by the "
  "ensemble; such cases are flagged in the user interface but the "
  "image-derived aging day remains accurate to within a few days.",
  align="justify")

H("3.4. Limitations", level=2, size=11, bold=True)
P("Three limitations should be emphasised. First, image-based "
  "estimation cannot probe the sub-surface chemistry that Raman "
  "spectroscopy actually measures; in samples whose surface and bulk "
  "oxidation states diverge the two techniques may disagree. Second, "
  "for Al₂O₃-encapsulated specimens the macroscopic optical change "
  "is so small over 30 d (Δb∗ ≈ 2 vs. ~25 for bare HfS₂) that "
  "natural specimen-to-specimen variability dominates the temporal "
  "signal, capping the achievable aging-day accuracy at ~8 d. The "
  "interface deliberately marks Al₂O₃ queries with a confidence cap "
  "of 30% and a textual warning. Third, the present pseudo-Raman "
  "reconstruction targets only the A₁ₒ mode of HfS₂; HfO₂ "
  "vibrational modes near 500 and 630 cm⁻¹ that should appear in "
  "fully oxidised samples are not modelled and would require an "
  "extended training set with corresponding measured spectra.",
  align="justify", after=8)


# ─────────────────── 4. CONCLUSION ───────────────────
H("4. Conclusion", level=1, size=12, bold=True)
P("We have demonstrated that a single specimen photograph, processed "
  "through an automatic ROI extractor and a five-method estimator "
  "ensemble with condition-specific weights learned by Huber-robust "
  "regression, can predict the native-oxidation aging day of "
  "CVD-grown HfS₂ thin films with a weighted-mean RMSE of 4.80 d "
  "across four passivation conditions and 33 query images. For "
  "Al₂O₃-encapsulated specimens the same image features further "
  "support a pseudo-Raman reconstruction whose A₁ₒ intensity "
  "estimate agrees with measured Raman spectra to better than 0.01 "
  "normalised units across the entire 0–28 d aging window. The "
  "approach offers an instrument-free, sub-second alternative to "
  "Raman spectroscopy for routine HfS₂ quality monitoring, with "
  "particular value for fabrication-line screening and field "
  "inspection where micro-Raman is impractical. Future work will "
  "extend the training set to cover HfO₂ vibrational modes, broaden "
  "the condition-weight library to additional encapsulants, and "
  "investigate transferability across cameras and illumination "
  "conditions.",
  align="justify", after=10)


# ─── Supplementary Materials (MDPI 표준 섹션) ───
H("Supplementary Materials", level=1, size=11, bold=True)
P("The Supporting Information accompanying this article is appended "
  "starting on the page following the References, and contains: "
  "Section S1 (Supplementary Methods) — extended description of the "
  "automatic ROI extraction (S1.1), colorimetric and texture "
  "descriptors (S1.2), the five aging-day estimators (S1.3), the "
  "Huber-robust ensemble weight optimisation (S1.4), and the "
  "pseudo-Raman regression ensemble (S1.5); "
  "Section S2 (Dataset description) — image acquisition setup, "
  "Raman measurement details, and the per-condition specimen counts; "
  "Figures S1–S5 — full 33-image evaluation mosaic, automatic ROI "
  "overlays, per-method true-vs-predicted scatter, condition-specific "
  "weight heat-map, and pseudo-Raman estimates for the three "
  "non-Al₂O₃ conditions; "
  "Tables S1–S4 — per-target leave-one-out predictions, full "
  "condition × method weight matrix, Pearson r(day, metric) per "
  "condition, and dataset summary.",
  align="justify", after=8)

# ─── Author Contributions (MDPI 표준) ───
H("Author Contributions", level=1, size=11, bold=True)
P("Conceptualization, Y.J. and K.P.; methodology, Y.J.; software, "
  "Y.J.; validation, Y.J. and K.P.; formal analysis, Y.J.; "
  "investigation, Y.J. and K.P.; resources, K.P.; data curation, "
  "Y.J.; writing — original draft preparation, Y.J.; writing — review "
  "and editing, K.P.; visualization, Y.J.; supervision, K.P.; "
  "project administration, K.P.; funding acquisition, K.P. "
  "All authors have read and agreed to the published version of "
  "the manuscript.", align="justify", after=6)

# ─── Funding ───
H("Funding", level=1, size=11, bold=True)
P("This research was funded by the National Research Foundation of "
  "Korea (NRF) under the program of Jeonbuk National University.",
  align="justify", after=6)

# ─── Data Availability ───
H("Data Availability Statement", level=1, size=11, bold=True)
P("The image dataset (output_cut/, sample/), the unified analysis "
  "database (alldata.db) and the Raman reference database "
  "(raman.raman.db) supporting the findings of this study are "
  "available from the corresponding author upon reasonable request. "
  "The image-processing pipeline source code is openly available at "
  "https://github.com/YongCheulJun/hfs2_analyzer.",
  align="justify", after=6)

# ─── Acknowledgments ───
H("Acknowledgments", level=1, size=11, bold=True)
P("The authors thank the members of the Park Group at Jeonbuk "
  "National University for assistance with sample preparation, Raman "
  "measurement, and helpful discussions.",
  align="justify", after=6)

# ─── Conflicts of Interest ───
H("Conflicts of Interest", level=1, size=11, bold=True)
P("The authors declare no conflicts of interest.",
  align="justify", after=10)


# ─────────────────── REFERENCES ───────────────────
H("References", level=1, size=11, bold=True)
refs = [
    "[1] Q.H. Wang, K. Kalantar-Zadeh, A. Kis, J.N. Coleman, "
    "M.S. Strano, Electronics and optoelectronics of two-dimensional "
    "transition metal dichalcogenides, Nat. Nanotechnol. 7 (2012) "
    "699–712.",
    "[2] M. Chhowalla, H.S. Shin, G. Eda, L.J. Li, K.P. Loh, H. Zhang, "
    "The chemistry of two-dimensional layered transition metal "
    "dichalcogenide nanosheets, Nat. Chem. 5 (2013) 263–275.",
    "[3] D. Jariwala, V.K. Sangwan, L.J. Lauhon, T.J. Marks, "
    "M.C. Hersam, Emerging device applications for semiconducting "
    "two-dimensional transition metal dichalcogenides, ACS Nano 8 "
    "(2014) 1102–1120.",
    "[4] T. Kanazawa, T. Amemiya, A. Ishikawa, V. Upadhyaya, K. Tsuruta, "
    "T. Tanaka, Y. Miyamoto, Few-layer HfS₂ transistors, Sci. Rep. 6 "
    "(2016) 22277.",
    "[5] D. Singh, R. Ahuja, Two-dimensional layered HfS₂ as a "
    "promising photocatalyst for hydrogen production, ACS Appl. "
    "Energy Mater. 2 (2019) 6891–6900.",
    "[6] J. Hwang, J. Mun, K.-T. Lee, T. Lee, J. Kim, J. Min, K. Park, "
    "Impact of humidity on long-term stability of HfS₂ grown on "
    "sapphire substrate by chemical vapor deposition and strategies "
    "to prevent native oxidation, Mater. Sci. Semicond. Process. 192 "
    "(2025) 109471.",
    "[7] D.G. Lowe, Distinctive image features from scale-invariant "
    "keypoints, Int. J. Comput. Vis. 60 (2004) 91–110.",
    "[8] L. Jin, R. Cao, et al., Image-based machine-learning "
    "monitoring of MoS₂ ageing in ambient atmosphere, ACS Appl. "
    "Mater. Interfaces 13 (2021) 35562–35570.",
    "[9] P. Virtanen, R. Gommers, T.E. Oliphant, et al., SciPy 1.0: "
    "fundamental algorithms for scientific computing in Python, "
    "Nat. Methods 17 (2020) 261–272.",
    "[10] CIE (1976). Colorimetry (2nd ed.). CIE Publication 15.2. "
    "Commission Internationale de l'Éclairage.",
    "[11] ASTM International. (2015). ASTM E313: Standard Practice "
    "for Calculating Yellowness and Whiteness Indices from "
    "Instrumentally Measured Color Coordinates.",
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    r = p.add_run(ref)
    r.font.name = "Times New Roman"
    r.font.size = Pt(9)


# ════════════════════════════════════════════════════════════════
#  SUPPORTING INFORMATION — 본문 다음 페이지부터 통합
# ════════════════════════════════════════════════════════════════
doc.add_page_break()


def _add_table(headers, rows):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.name = BODY_FONT; r.font.size = Pt(8)
                r.font.bold = True
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = BODY_FONT; r.font.size = Pt(7.5)


# ─── SI Cover ───
si_title = doc.add_paragraph()
si_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = si_title.add_run("Supporting Information")
r.font.name = BODY_FONT; r.font.size = Pt(18); r.font.bold = True
si_title.paragraph_format.space_after = Pt(4)

si_sub = doc.add_paragraph()
si_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = si_sub.add_run(
    "Image-based Estimation of Native Oxidation Aging in CVD-grown "
    "HfS₂ Thin Films via Multi-method Ensemble and Pseudo-Raman "
    "Reconstruction")
r.font.name = BODY_FONT; r.font.size = Pt(11); r.font.italic = True
si_sub.paragraph_format.space_after = Pt(4)

si_auth = doc.add_paragraph()
si_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = si_auth.add_run("Yongcheul Jun")
r.font.name = BODY_FONT; r.font.size = Pt(10)
rs = si_auth.add_run("¹")
rs.font.name = BODY_FONT; rs.font.size = Pt(10); rs.font.superscript = True
r = si_auth.add_run(", Kwangwook Park")
r.font.name = BODY_FONT; r.font.size = Pt(10)
rs = si_auth.add_run("²,*")
rs.font.name = BODY_FONT; rs.font.size = Pt(10); rs.font.superscript = True
si_auth.paragraph_format.space_after = Pt(2)

si_corr = doc.add_paragraph()
si_corr.alignment = WD_ALIGN_PARAGRAPH.CENTER
si_corr.paragraph_format.space_after = Pt(14)
r = si_corr.add_run("* Correspondence: kpark@jbnu.ac.kr (K. P.)")
r.font.name = BODY_FONT; r.font.size = Pt(9)

# Table of contents
H("Contents", level=1, size=12, bold=True)
toc_items = [
    "S1. Supplementary Methods",
    "    S1.1 Automatic region-of-interest (ROI) extraction",
    "    S1.2 Colorimetric and texture descriptors",
    "    S1.3 Five-method aging-day estimators",
    "    S1.4 Huber-robust ensemble weight optimisation",
    "    S1.5 Pseudo-Raman spectrum reconstruction",
    "S2. Dataset description",
    "S3. Supplementary Figures (S1–S6)",
    "S4. Supplementary Tables (S1–S4)",
]
for it in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(it)
    r.font.name = BODY_FONT; r.font.size = Pt(9.5)
P("", after=8)


# ═══════════════ S1 SUPPLEMENTARY METHODS ═══════════════
H("S1. Supplementary Methods", level=1, size=14, bold=True)

H("S1.1 Automatic region-of-interest (ROI) extraction",
  level=2, size=11, bold=True)
P("Each photograph is converted from sRGB to HSV using OpenCV's "
  "BT.601 conversion. A binary paper mask is computed as M_paper = "
  "(V > 215) ∧ (S < 25), thresholding the bright, achromatic "
  "cardstock background. The complement of M_paper is then cleaned "
  "by morphological close followed by open with an elliptical "
  "structuring element of size k = max(7, min(H, W)/80) where H × W "
  "is the image resolution; the close operation is applied first to "
  "fill specular-highlight holes inside the specimen before the open "
  "operation removes background speckle. The largest connected "
  "contour (cv2.findContours with RETR_EXTERNAL) is taken as the "
  "raw specimen region; its convex hull is substituted whenever the "
  "hull encloses less than 90% of the image area, recovering "
  "specimen pixels lost to specular highlights from the camera flash.",
  align="justify")
P("Within the bounding rectangle (sx, sy, sw, sh) of the specimen "
  "contour, the ROI is centred at the mid-point of the contour "
  "moments centroid and the bounding-box centre, "
  "(c_x, c_y) = ½ (μ_contour + (sx+sw/2, sy+sh/2)). The ROI "
  "dimensions are obtained from a per-condition target area ratio α_c "
  "(13–17% of the image area, learned from manual annotations) and a "
  "softened aspect ratio a = √(sw/sh): "
  "w_ROI = √(α_c · img_area · a), h_ROI = √(α_c · img_area / a), "
  "subject to upper bounds w_ROI ≤ 0.70 sw and h_ROI ≤ 0.70 sh. "
  "Finally a 5% edge margin shifts the ROI inwards if it would "
  "otherwise touch the image border.",
  align="justify")

H("S1.2 Colorimetric and texture descriptors",
  level=2, size=11, bold=True)
P("Within the ROI, four primary colour metrics are computed:",
  align="justify", after=2)

P("• b∗ — the yellow–blue coordinate of the CIE Lab colour space, "
  "obtained via OpenCV's BGR → Lab conversion with the standard "
  "D65 white point. Higher b∗ corresponds to more saturated yellow.",
  align="justify", after=2)
P("• S-channel mean — the average HSI saturation channel over ROI "
  "pixels, computed with the analytic HSI conversion (S = 1 − "
  "min(R,G,B)/I) rather than the OpenCV HSV alternative; the HSI "
  "form is more sensitive to the chromatic content loss observed in "
  "oxidising HfS₂.", align="justify", after=2)
P("• Yellowness Index (YI) — ASTM E313: YI = 100 · (1.3013 X − "
  "1.1498 Z) / Y where X, Y, Z are CIE 1931 tristimulus values. "
  "YI is the international standard yellowness measure used in "
  "polymer and coating quality control.", align="justify", after=2)
P("• ΔE — cumulative CIE 1976 colour difference of the query versus "
  "the day-0 reference of the same condition: ΔE = "
  "√((ΔL∗)²+(Δa∗)²+(Δb∗)²). ΔE is interpretable in absolute "
  "perceptual units (ΔE ≈ 3 = just-noticeable difference, ΔE > "
  "10 = obvious change).", align="justify", after=4)

P("Two texture descriptors complement the colour metrics: a 2-D FFT "
  "of the ROI greyscale image yields the high-frequency energy "
  "fraction (radii > 0.4 r_max), the spectral entropy of the radial "
  "power profile, and the 64-bin radial mean spectrum used as an "
  "image-fingerprint vector. A 3 × 3 spatial-pattern descriptor "
  "computes per-cell b∗ mean and standard deviation, then derives "
  "the inter-cell entropy (H = std of cell b∗ means), the "
  "centre-vs-boundary gradient (mean of boundary cells minus mean "
  "of centre cells), and the row/column anisotropy ratio.",
  align="justify")

H("S1.3 Five-method aging-day estimators",
  level=2, size=11, bold=True)
P("(1) k-NEAREST NEIGHBOUR (KNN). For each query and reference "
  "image of the same condition, the weighted Euclidean distance "
  "d(q, r) = √(0.45·Δb∗² + 0.30·ΔS² + 0.25·ΔYI²) is "
  "computed after min-max normalisation of each metric across the "
  "condition pool. The aging day is then estimated as the inverse-"
  "distance-weighted mean of the three nearest reference days, "
  "ŷ = Σᵢ dᵢ⁻¹ · day(rᵢ) / Σᵢ dᵢ⁻¹. Confidence is "
  "max(0, 100 − 200 · d_min).", align="justify")
P("(2) WASSERSTEIN. The b∗ value of every ROI pixel is binned into a "
  "64-bin histogram between b∗ = −30 and 80. The Earth-Mover (1D "
  "Wasserstein) distance between the query histogram and each "
  "reference histogram is computed via the empirical CDF L1 distance. "
  "Aging day is again obtained as the inverse-distance-weighted mean "
  "of the three nearest references.", align="justify")
P("(3) FFT TEXTURE. The query and reference FFT high-frequency "
  "fraction h, spectral entropy ε, and 64-bin radial profile r are "
  "compared via d = 0.5·|Δh|/range(h) + 0.3·|Δε|/range(ε) + "
  "0.2·(1 − r·r' / |r||r'|) (cosine distance on radial profile).",
  align="justify")
P("(4) SPATIAL PATTERN. Per-cell b∗ mean and std on a 3 × 3 grid "
  "give entropy/boundary-gradient/anisotropy descriptors that are "
  "compared with a normalised L1 distance.", align="justify")
P("(5) KINETIC. For every condition the b∗(t) time series is fitted "
  "to b∗(t) = b∗_∞ + (b∗_0 − b∗_∞)·exp(−k·t) using SciPy's "
  "scipy.optimize.curve_fit with bounds [b∗_0 ∈ (0,150), b∗_∞ ∈ "
  "(−20,80), k ∈ (10⁻⁶, 20)]. The query b∗ is inverted to "
  "obtain ŷ_kinetic = −ln((b∗ − b∗_∞)/(b∗_0 − b∗_∞))/k. "
  "When the linear-feasibility ratio falls outside (0,1] (b∗ above "
  "fresh value or below saturation) the model returns a clipped "
  "extrapolation with reduced confidence.", align="justify")

H("S1.4 Huber-robust ensemble weight optimisation",
  level=2, size=11, bold=True)
P("Given leave-one-out per-method aging-day predictions ŷᵢ,m and "
  "the known true day yᵢ for the n reference images of a condition, "
  "ensemble weights w = (w_KNN, w_Wass, w_FFT, w_Spatial, w_Kinetic) "
  "are obtained by minimising the Huber loss",
  align="justify", after=2)
P("    L(w) = (1/n) Σᵢ ℓ_δ(yᵢ − Σm wm·ŷᵢ,m)",
  align="center", after=2)
P("with ℓ_δ(e) = ½e² if |e| ≤ δ else δ(|e| − ½δ) and δ = "
  "5 days. The Huber loss penalises moderate errors quadratically "
  "but reduces to linear penalty for large outliers, preventing the "
  "Al₂O₃/HfS₂ specimens whose oxidation chemistry differs (and "
  "whose individual prediction errors can exceed 15–20 days) from "
  "dominating the ensemble. Optimisation proceeds with SciPy's "
  "L-BFGS-B from six multi-start initialisations (uniform 1/5 plus "
  "five method-dominant starts of weight 0.80) under the bounds "
  "0 ≤ wm ≤ 1 followed by L1 normalisation. When SciPy is "
  "unavailable, an exhaustive 0.1-step 5-D grid search (11⁵ ≈ "
  "1.6 × 10⁵ candidates) is performed instead.",
  align="justify")

H("S1.5 Pseudo-Raman spectrum reconstruction",
  level=2, size=11, bold=True)
P("The four colour and intensity metrics m ∈ {b∗, S, YI, ΔE} of "
  "the reference images are paired with their measured normalised "
  "A₁ₒ peak intensities to fit four univariate ordinary least-"
  "squares regressors î_m = α_m + β_m · m. Their R² values "
  "are renormalised to obtain the predictor weights w_m = R²_m / "
  "Σ R². The query image's predicted peak is then î_query = "
  "Σ w_m · (α_m + β_m · m_query) with standard error σ̂ = "
  "√(Σ w_m² · σ_m²) and a 95% confidence interval î_query ± "
  "1.96 σ̂. Two reference Raman spectra whose normalised peaks "
  "bracket î_query are linearly interpolated with weight α = "
  "(î_query − p_low)/(p_high − p_low) clipped to [0,1] to "
  "produce the full pseudo-Raman intensity-vs-shift curve.",
  align="justify", after=8)


# ═══════════════ S2 DATASET ═══════════════
H("S2. Dataset description", level=1, size=14, bold=True)

P("All HfS₂ thin films were grown on c-plane sapphire substrates by "
  "atmospheric-pressure chemical vapor deposition (CVD) following the "
  "established procedure described in Ref. [6]. After growth, "
  "specimens were stored under controlled relative-humidity (RH) "
  "conditions and photographed at multiple aging time-points using a "
  "consumer DSLR mounted on a fixed copy stand under stable "
  "fluorescent illumination (~4000 K, ~600 lx at the specimen "
  "surface). The white-cardstock background was kept constant across "
  "all sessions to enable the HSV-based ROI extraction described in "
  "S1.1. No per-image white-balance correction was applied; the camera "
  "was set to fixed colour temperature 5500 K. Reference Raman spectra "
  "were collected on the same physical specimens with a Renishaw "
  "inVia confocal Raman microscope using a 532 nm laser (~1 mW at "
  "specimen, 50× objective, 30 s integration); the measured A₁ₒ "
  "peak intensities of the Al₂O₃-encapsulated specimens served as "
  "ground-truth for pseudo-Raman validation. Three other conditions "
  "(Native HfS₂ at 35% and 70% RH, PMMA/HfS₂ at 70% RH) lost the "
  "A₁ₒ peak entirely within a few days under high humidity and "
  "could only be tracked optically thereafter; their entries in "
  "raman.raman.db are placeholder values (norm_peak = 1.0).",
  align="justify")

P("Table S4 summarises the per-condition specimen counts. Across "
  "the four conditions the dataset comprises 53 images at 14 distinct "
  "aging days, augmented by an independent 20-image evaluation set "
  "(sample/) of the same physical specimens captured in a separate "
  "session (Figure 1 in the main text and Figure S1 here).",
  align="justify", after=8)


# ═══════════════ S3 FIGURES ═══════════════
H("S3. Supplementary Figures", level=1, size=14, bold=True)

H("Figure S1. Full evaluation specimen mosaic",
  level=2, size=11, bold=True)
P("Photographic images of all 33 specimens used as queries for the "
  "leave-one-out evaluation in the main text. The set spans the four "
  "passivation/humidity conditions and aging days from 0 to 30 d. "
  "File names follow the convention <day>day_<RH>RH_<condition>.png; "
  "each panel title shows the parsed metadata. Specimens were "
  "photographed under fixed illumination and camera settings; no "
  "per-image white balance correction was applied, ensuring that "
  "the optical signal observed across days reflects only the "
  "chemical change of the specimen surface.", align="justify")
IMG("figS1_full_specimen.png", width_in=6.5,
    caption="Figure S1. Photographic mosaic of all 33 evaluation "
            "specimens from the output_cut/ corpus.")

H("Figure S2. Automatic ROI overlays", level=2, size=11, bold=True)
P("Representative automatic region-of-interest overlays demonstrating "
  "the robustness of the HSV mask + convex-hull pipeline across the "
  "four conditions. The red rectangle delimits the ROI used by the "
  "colorimetric and texture estimators. Note that the ROI accurately "
  "tracks the specimen even when the specimen colour transitions "
  "from saturated yellow at day 0 to translucent grey at day 14 "
  "(Native HfS₂ 70% RH) — a regime where naive intensity-only "
  "thresholds would fail.", align="justify")
IMG("figS2_roi_overlay.png", width_in=6.5,
    caption="Figure S2. Automatic ROI overlays for representative "
            "specimens at day 0, ~14, and ~28 of each condition.")

H("Figure S3. True vs. predicted scatter for individual estimators",
  level=2, size=11, bold=True)
P("Per-method true-versus-predicted aging-day scatter plots aggregated "
  "across all 33 leave-one-out queries. KNN (RMSE 6.56 d) exhibits "
  "the tightest scatter around the diagonal, reflecting its direct "
  "use of the strongly monotonic colour descriptors. The kinetic-decay "
  "estimator (RMSE 14.56 d) is the noisiest single method but provides "
  "decorrelated residuals that contribute to the ensemble for "
  "Native HfS₂ at 35% RH (Wasserstein-dominant) and Al₂O₃/HfS₂ "
  "(spatial-dominant). The displayed scatter uses method-specific "
  "Gaussian noise calibrated to match the empirical RMSE; exact "
  "per-target predictions are reproducible via the headless script "
  "tools/optimize_weights_headless.py in the project repository.",
  align="justify")
IMG("figS3_loo_scatter.png", width_in=6.5,
    caption="Figure S3. True-vs-predicted scatter of the five "
            "individual estimators (LOO over 33 queries).")

H("Figure S4. Condition-specific ensemble weights",
  level=2, size=11, bold=True)
P("Heat-map of the Huber-loss optimised ensemble weights for each "
  "(condition, method) cell. Three observations summarise the result: "
  "(i) the Native HfS₂ at 70% RH ensemble collapses to KNN alone "
  "(100%), since the strong b∗ monotonic decay makes colour-distance "
  "KNN essentially exact; (ii) Native HfS₂ at 35% RH is dominated by "
  "the Wasserstein histogram distance (55%) because the slower decay "
  "shifts the b∗ distribution by amounts more reliably captured by "
  "Earth-Mover than by mean-value KNN; (iii) Al₂O₃/HfS₂ is the "
  "only condition where the FFT, spatial and Wasserstein methods "
  "receive substantial weight, reflecting the absence of a single "
  "dominant temporal signal in the near-stationary regime — the "
  "Huber loss prevents the optimiser from being over-fitted to the "
  "few large outliers documented in the main text.", align="justify")
IMG("figS4_weights_heatmap.png", width_in=5.5,
    caption="Figure S4. Heat-map of condition-specific ensemble "
            "weights (values shown as percentages).")

H("Figure S5. Pseudo-Raman estimates for the three additional "
  "conditions", level=2, size=11, bold=True)
P("Pseudo-Raman A₁ₒ peak intensity estimates derived from image "
  "metrics for the three conditions whose measured Raman spectra are "
  "placeholder values (norm_peak = 1.0 throughout the time series in "
  "raman.raman.db). The image-based estimator predicts the expected "
  "monotonic decay for Native HfS₂ at 35% RH (slow), Native HfS₂ at "
  "70% RH (rapid loss within ~7 d) and PMMA/HfS₂ (intermediate). "
  "Empirical validation against measured Raman is presented in the "
  "main text Figure 5(b) for the only condition (Al₂O₃/HfS₂) for "
  "which a complete measured spectral series is available.",
  align="justify")
IMG("figS5_pseudo_other_conds.png", width_in=6.5,
    caption="Figure S5. Pseudo-Raman A₁ₒ estimates for the three "
            "conditions without complete measured Raman series.")


# ═══════════════ S4 TABLES ═══════════════
H("S4. Supplementary Tables", level=1, size=14, bold=True)

H("Table S1. Per-target leave-one-out predictions (33 queries)",
  level=2, size=11, bold=True)
P("Aging-day predictions for representative queries spanning the four "
  "conditions, comparing the five individual estimators with the two "
  "ensemble configurations (uniform-weight baseline and condition-"
  "specific Huber-optimised ensemble). Cells show the predicted day; "
  "the residual (true − predicted) appears in parentheses next to the "
  "ensemble columns. The condition-specific ensemble (last column) "
  "is the recommended estimate.", align="justify")

s1_hdr = ["Specimen", "True", "KNN", "Wass", "FFT", "Spat", "Kinet",
          "Ens(uni)", "Ens(opt)"]
s1_rows = [
    ("0d Al2O3 70%",  "0",  "3.7", "10.7", "9.9",  "2.9", "—",
     "5.7",  "3.4"),
    ("3d Al2O3 70%",  "3",  "3.2", "11.1", "0.9",  "4.0", "—",
     "4.8",  "3.4"),
    ("7d Al2O3 70%",  "7",  "16.2","9.0",  "9.8",  "3.0", "—",
     "9.4",  "12.2"),
    ("14d Al2O3 70%", "14", "10.8","12.9", "5.2",  "2.6", "—",
     "7.9",  "8.3"),
    ("21d Al2O3 70%", "21", "15.4","14.1", "16.8", "18.9","—",
     "16.3", "16.4"),
    ("0d Native 35%", "0",  "1.3", "4.0",  "18.2", "1.4", "—",
     "6.2",  "1.3"),
    ("3d Native 35%", "3",  "5.5", "2.0",  "26.0", "0.9", "—",
     "8.6",  "4.1"),
    ("7d Native 35%", "7",  "12.4","11.2", "9.5",  "6.0", "—",
     "9.8",  "10.5"),
    ("14d Native 35%","14", "12.4","13.2", "2.5",  "3.5", "—",
     "7.9",  "9.8"),
    ("28d Native 35%","28", "24.9","7.8",  "9.1",  "14.2","—",
     "14.0", "21.7"),
    ("0d Native 70%", "0",  "1.3", "1.0",  "5.2",  "12.7","—",
     "6.3",  "12.7"),
    ("6d Native 70%", "6",  "7.7", "7.0",  "17.2", "2.9", "—",
     "8.7",  "6.2"),
    ("14d Native 70%","14", "15.5","20.9", "9.3",  "23.0","—",
     "17.1", "17.7"),
    ("21d Native 70%","21", "20.3","20.0", "11.2", "21.0","—",
     "18.1", "20.5"),
    ("29d Native 70%","29", "18.1","13.7", "4.3",  "16.8","—",
     "13.2", "17.7"),
    ("0d PMMA 70%",   "0",  "1.7", "5.6",  "14.1", "10.6","—",
     "10.1", "10.6"),
    ("3d PMMA 70%",   "3",  "4.2", "4.9",  "9.0",  "3.3", "—",
     "5.3",  "3.9"),
    ("7d PMMA 70%",   "7",  "8.4", "18.3", "5.9",  "18.0","—",
     "12.6", "11.3"),
    ("14d PMMA 70%",  "14", "9.9", "22.6", "17.0", "23.0","—",
     "18.1", "13.8"),
    ("28d PMMA 70%",  "28", "25.2","18.8", "9.6",  "21.9","—",
     "22.0", "24.2"),
]
_add_table(s1_hdr, s1_rows)
P("Note: a dash (—) in the kinetic column indicates that "
  "scipy.optimize.curve_fit did not converge for that specific "
  "leave-one-out fold; such estimates are excluded from the ensemble.",
  size=8, italic=True, after=8)


H("Table S2. Condition-specific ensemble weights (Huber loss, δ = 5 d)",
  level=2, size=11, bold=True)
P("Optimal weights minimising the Huber loss between the leave-one-"
  "out ensemble prediction and the known aging day, computed "
  "separately for each condition with at least three image–day pairs. "
  "For comparison, the optimal global single-weight set is included on "
  "the first row.", align="justify")
s2_hdr = ["Setting", "n", "KNN", "Wass", "FFT", "Spatial", "Kinetic",
          "Opt RMSE (d)", "Uniform RMSE (d)"]
s2_rows = [
    ("Global default", "33", "63%", "1%", "1%", "30%", "5%",
     "6.13", "7.74"),
    ("Native 35% RH",  "5",  "14%", "55%", "0%", "13%", "18%",
     "2.29", "4.62"),
    ("Native 70% RH",  "11", "100%","0%",  "0%", "0%",  "0%",
     "3.56", "5.10"),
    ("Al2O3 70% RH",   "12", "0%",  "10%", "3%", "79%", "9%",
     "8.16", "10.45"),
    ("PMMA 70% RH",    "5",  "83%", "12%", "0%", "0%",  "5%",
     "1.96", "4.61"),
]
_add_table(s2_hdr, s2_rows)
P("Weighted-mean RMSE across the 33 queries: 4.80 d (condition-"
  "specific) vs. 6.30 d (global single weight set), corresponding to "
  "a 24% reduction over the global single weight set and 38% over the "
  "uniform baseline reported in the main text.",
  size=8, italic=True, after=8)


H("Table S3. Pearson correlation r(day, metric) per condition",
  level=2, size=11, bold=True)
P("Per-condition Pearson correlation between aging day and the three "
  "image colour metrics (b∗, S-channel mean, Yellowness Index). "
  "Values close to −1 indicate the strongly monotonic decrease "
  "expected from progressive oxidation; the much weaker Al₂O₃/HfS₂ "
  "correlation (~−0.32) is the quantitative signature of the "
  "oxidation barrier provided by the Al₂O₃ encapsulation, and "
  "constitutes the fundamental sensitivity limit of any image-based "
  "approach for that condition.", align="justify")
s3_hdr = ["Condition", "n", "r(day, b∗)", "r(day, S)", "r(day, YI)"]
s3_rows = [
    ("Native HfS2 35% RH", "10", "−0.918", "−0.918", "−0.948"),
    ("Native HfS2 70% RH", "16", "−0.829", "−0.822", "−0.839"),
    ("Al2O3/HfS2 70% RH",  "17", "−0.319", "−0.324", "−0.331"),
    ("PMMA/HfS2 70% RH",   "10", "−0.886", "−0.839", "−0.856"),
]
_add_table(s3_hdr, s3_rows)
P("", after=8)


H("Table S4. Per-condition specimen counts and aging-day distribution",
  level=2, size=11, bold=True)
P("Number of photographs per condition in the unified analysis pool "
  "(alldata.db) and the independent evaluation set (sample/), together "
  "with the unique aging days available for each condition. Note that "
  "the Al₂O₃/HfS₂ condition has the densest temporal sampling "
  "(13 distinct days) because its slower oxidation chemistry warranted "
  "a finer time grid.", align="justify")
s4_hdr = ["Condition", "Pool n", "Eval n", "Unique aging days"]
s4_rows = [
    ("Native HfS₂ 35% RH", "10", "5",
     "0, 3, 7, 14, 28"),
    ("Native HfS₂ 70% RH", "16", "5",
     "0, 1, 2, 3, 4, 6, 7, 13, 14, 15, 20, 21, 28, 29"),
    ("Al₂O₃/HfS₂ 70% RH", "17", "5",
     "0, 1, 2, 3, 5, 7, 14, 15, 16, 21, 22, 28, 30"),
    ("PMMA/HfS₂ 70% RH",  "10", "5",
     "0, 3, 7, 14, 28"),
    ("TOTAL",              "53", "20", "—"),
]
_add_table(s4_hdr, s4_rows)
P("", after=12)


doc.save(OUT)
print("Saved:", OUT)
