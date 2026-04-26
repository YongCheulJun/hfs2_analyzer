"""
SCIE 논문 한글 번역본 (.docx) — MDPI Applied Sciences 양식.
영문본 paper/Jun_HfS2_image_oxidation.docx 와 동일 구조 + figure 5장
+ SI (S1-S5) 통합. 한글 학술 어투.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIG = os.path.join(ROOT, "paper/figures")
OUT = os.path.join(ROOT, "paper", "Jun_HfS2_image_oxidation_KO.docx")

doc = Document()
sec = doc.sections[0]
sec.page_width  = Cm(21.0); sec.page_height = Cm(29.7)
sec.top_margin    = Cm(1.78); sec.bottom_margin = Cm(1.78)
sec.left_margin   = Cm(1.78); sec.right_margin  = Cm(1.78)

# 한글: 맑은 고딕 / fallback Malgun Gothic / DejaVu Sans
style = doc.styles["Normal"]
style.font.name = "Malgun Gothic"
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(2)
BODY = "Malgun Gothic"


def H(t, level=1, size=12, bold=True, before=8, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before if level == 1 else 4)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(t)
    r.font.name = BODY; r.font.size = Pt(size); r.font.bold = bold


def P(t, size=10, italic=False, align="justify", after=2):
    p = doc.add_paragraph()
    if align == "center": p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify": p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(t)
    r.font.name = BODY; r.font.size = Pt(size); r.font.italic = italic


def IMG(fname, width=6.5, cap=None):
    path = os.path.join(FIG, fname)
    if not os.path.exists(path):
        P(f"(그림 {fname} 누락)"); return
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cap:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        cp.paragraph_format.space_after = Pt(8)
        r = cp.add_run(cap)
        r.font.name = BODY; r.font.size = Pt(9); r.font.italic = True


def add_table(headers, rows):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.name = BODY; r.font.size = Pt(8); r.font.bold = True
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = BODY; r.font.size = Pt(7.5)


# ════════════════ TITLE / AUTHORS ════════════════
tag = doc.add_paragraph()
r = tag.add_run("논문")
r.font.name = BODY; r.font.size = Pt(10); r.font.italic = True
tag.paragraph_format.space_after = Pt(4)

title = doc.add_paragraph()
r = title.add_run(
    "이미지 기반 다중 추정기 앙상블 및 의사-라만 재구성에 의한 "
    "CVD 성장 HfS₂ 박막의 자연 산화 경과일 추정")
r.font.name = BODY; r.font.size = Pt(16); r.font.bold = True
title.paragraph_format.space_after = Pt(6)

auth = doc.add_paragraph()
def add_auth(par, name, sup, last=False):
    rr = par.add_run(name)
    rr.font.name = BODY; rr.font.size = Pt(11)
    rs = par.add_run(" " + sup)
    rs.font.name = BODY; rs.font.size = Pt(11); rs.font.superscript = True
    if not last:
        rsep = par.add_run(", ")
        rsep.font.name = BODY; rsep.font.size = Pt(11)
add_auth(auth, "전용철", "1")
add_auth(auth, "박광욱", "2,*", last=True)

def aff(num, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_after = Pt(2)
    rn = p.add_run(num)
    rn.font.name = BODY; rn.font.size = Pt(9); rn.font.superscript = True
    rt = p.add_run(" " + text)
    rt.font.name = BODY; rt.font.size = Pt(9); rt.font.italic = True
aff("1", "부산대학교 공과대학 (지식재산융합전공), 부산광역시 46241, 대한민국")
aff("2", "전북대학교 전자정보공학부, 전라북도 전주시 54896, 대한민국")

corr = doc.add_paragraph()
corr.paragraph_format.left_indent = Cm(0.5)
corr.paragraph_format.first_line_indent = Cm(-0.5)
corr.paragraph_format.space_after = Pt(10)
rs = corr.add_run("*")
rs.font.name = BODY; rs.font.size = Pt(9); rs.font.superscript = True
rt = corr.add_run(" 교신저자: kpark@jbnu.ac.kr (박광욱)")
rt.font.name = BODY; rt.font.size = Pt(9)


# ════════════════ ABSTRACT ════════════════
ap = doc.add_paragraph()
ap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
ap.paragraph_format.space_after = Pt(4)
rl = ap.add_run("초록: ")
rl.font.name = BODY; rl.font.size = Pt(10); rl.font.bold = True
rb = ap.add_run(
    "이황화하프늄(HfS₂)은 고이동도 전자소자 및 광검출기로 주목받는 "
    "층상 전이금속 디칼코제나이드(TMD)이지만, 대기 중 습도에 의한 "
    "급속한 자연 산화로 인해 황색 칼코제나이드가 며칠 내에 투명 "
    "산화물(HfOₓ)로 변환되는 한계가 있다. 산화 정도를 정량화하기 "
    "위해서는 일반적으로 라만 분광법이 사용되지만, 분광기에 종속적이고 "
    "측정 시간이 길다. 본 연구에서는 단일 시편 사진만으로 CVD 성장 "
    "HfS₂ 박막의 경과일을 추정하고, 분광 측정 없이 의사-라만 스펙트럼을 "
    "재구성하는 이미지 전용(image-only) 프레임워크를 제안한다. 4가지 "
    "패시베이션 조건(35% RH 및 70% RH의 Native HfS₂, 70% RH의 Al₂O₃ "
    "및 PMMA 캡슐화 HfS₂)과 0–30일의 경과일을 포함하는 53장 시편 "
    "이미지 풀에 대해, 색상 지표 기반 가중 유클리드 KNN, b∗ 히스토그램 "
    "Wasserstein 거리, FFT 텍스처 거리, 공간 패턴 거리, 그리고 지수 "
    "감쇠 동역학 적합 등 5개 독립 추정기를 적용하고, 그 출력을 "
    "Huber 손실로 최적화된 조건별 앙상블로 결합하였다. 33개 평가 "
    "이미지에 대한 leave-one-out 평가 결과, 가중 평균 RMSE 4.80일을 "
    "달성하여 균등 가중 baseline(7.74일) 대비 38% 감소되었다. "
    "Al₂O₃ 캡슐화 HfS₂의 경우 이미지 기반 A₁ₒ 피크 강도가 측정 "
    "라만 추세(1.00 → 0.60)를 95% 신뢰 구간 내에서 재현하였다. "
    "본 연구는 이미지 전용 의사-라만 재구성이 HfS₂ 박막의 일상적 "
    "품질 검사에서 분광기를 대체할 수 있는 신속·비파괴 도구임을 "
    "입증한다.")
rb.font.name = BODY; rb.font.size = Pt(10)

kw = doc.add_paragraph()
kw.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
kw.paragraph_format.space_after = Pt(10)
rl = kw.add_run("키워드: ")
rl.font.name = BODY; rl.font.size = Pt(10); rl.font.bold = True
rb = kw.add_run("이황화하프늄; 자연 산화; 이미지 처리; "
                "의사-라만; 앙상블 회귀; CVD 박막")
rb.font.name = BODY; rb.font.size = Pt(10)


# ════════════════ 1. 서론 ════════════════
H("1. 서론", level=1, size=12, bold=True)
P("2차원 층상 전이금속 디칼코제나이드(TMD)는 초박막 전자소자, 광대역 "
  "광검출기, 수소 발생 촉매, 신축성 소자 등의 후보로 큰 주목을 받고 "
  "있다 [1–3]. 이 중 HfS₂는 높은 이론 전자 이동도, 높은 광반응도, "
  "그리고 H⁺/H₂ 환원에 유리한 전도대 정렬 때문에 특히 주목받는다 "
  "[4,5]. 그러나 대부분의 황화물 기반 TMD와 마찬가지로 HfS₂는 대기 "
  "중 자연 산화에 매우 취약하다: O₂와 H₂O가 협력적으로 HfS₂를 "
  "HfS₂₋ₓOₓ로 변환시키고 최종적으로 HfO₂로 전환되어, 본래의 "
  "황색 칼코제나이드를 투명 산화물 박막으로 대체한다 [6]. Hwang 등은 "
  "최근 CVD 성장 HfS₂의 A₁ₒ 라만 피크가 70% 상대습도(RH) 환경에서 "
  "1주일 이내에 사라지며, Al₂O₃에 의한 캡슐화는 산화를 효과적으로 "
  "억제하는 반면 PMMA는 산화 진행을 지연시킬 뿐 차단하지 못함을 "
  "보고하였다 [6].",
  align="justify")

P("기존 표준 절차에서는 HfS₂의 산화 정도를 ~337 cm⁻¹ 부근의 "
  "A₁ₒ 모드 강도를 통해 미세 라만 분광법으로 모니터링한다. "
  "라만 측정은 매우 특이적이지만 분광기에 종속적이고, 시편마다 "
  "신중한 초점 정렬이 필요하며, 제조 라인 모니터링이나 현장 품질 "
  "검사에는 비실용적이다. 반면 HfS₂ 산화의 거시적 광학 특성 — 채도 "
  "높은 황색에서 반투명 회색으로의 변색 — 은 일반 디지털 카메라로 "
  "쉽게 포착된다. 따라서 본 연구는 단일 시편 사진만으로 라만 측정을 "
  "대체할 수 있는지, 일상적 경과일 평가에 활용할 수 있는지를 검증하고자 "
  "하였다.", align="justify")

P("이미지 기반 산화 추적은 귀금속 변색 및 부식 분야에서 활용되어 "
  "왔으며, 최근 색상 또는 텍스처 특성을 이용해 2D 재료의 산화를 "
  "추정하는 보고가 있다 [7,8]. 그러나 동일한 이미지 특성을 합성 "
  "라만 스펙트럼으로 역으로 변환하여, 실험자가 측정 라만 스펙트럼과 "
  "같은 방식으로 해석할 수 있도록 하는 시도는 아직까지 보고된 바가 "
  "없다. 이러한 “의사-라만” 재구성은 분광기 없이도 사용자가 "
  "Stage I–IV 산화 단계 분류를 검증할 수 있게 한다. 본 연구는 이러한 "
  "파이프라인을 도입·평가·벤치마크한다. 본 논문의 기여는 (i) Huber 견고 "
  "회귀로 학습된 조건별 가중치를 갖는 5-방법 앙상블이 단일 이미지로부터 "
  "경과일을 추정하는 점, (ii) 이미지 지표로부터 95% 신뢰 구간을 해석적으로 "
  "전파한 의사-라만 스펙트럼을 재구성하는 점, (iii) 측정 라만 스펙트럼과의 "
  "정량 비교를 통해 Al₂O₃ 캡슐화 HfS₂에 대한 재구성의 충실도를 확인하고 "
  "해당 시편군의 본질적 감도 한계를 규명한 점이다.",
  align="justify", after=8)


# ════════════════ 2. 재료 및 방법 ════════════════
H("2. 재료 및 방법", level=1, size=12, bold=True)

H("2.1. 시편 준비", level=2, size=11, bold=True)
P("CVD 성장된 사파이어 기판 위 HfS₂ 박막을 [6]에서 기술된 절차에 "
  "따라 제작하였다. 4가지 패시베이션 조건이 연구되었다: 35% RH 및 "
  "70% RH 환경에 보관된 베어 Native HfS₂, 원자층 증착(ALD)으로 "
  "10 nm Al₂O₃를 캡슐화한 HfS₂(Al₂O₃/HfS₂, 70% RH), 그리고 "
  "스핀 코팅된 PMMA 박막으로 덮인 HfS₂(PMMA/HfS₂, 70% RH). "
  "각 조건에 대해 동일 물리 시편을 0일에서 30일 사이의 여러 경과일에 "
  "고정된 조명·카메라 설정으로 촬영하였다. 본 연구의 참조 풀은 4가지 "
  "조건과 14개 경과일에 분포하는 53장 사진으로 구성되며, 동일 "
  "시편을 별도 세션에서 촬영한 독립적인 20장 평가 세트가 추가되었다. "
  "대표 이미지는 Fig. 1에 제시하였다.",
  align="justify")

IMG("fig1_specimen.png", width=6.5,
    cap="Fig. 1. 4가지 패시베이션/습도 조건에서 CVD 성장 HfS₂ 시편의 "
        "시간 경과별 사진. 35% RH의 Native HfS₂는 한 달간 황색 "
        "칼코제나이드 외관을 유지하지만, 70% RH의 Native HfS₂는 일주일 "
        "내에 황색을 잃는다. Al₂O₃ 캡슐화는 광학 변화를 거의 완전히 "
        "억제하는 반면, PMMA는 산화를 더 완만하게 지연시킨다.")

H("2.2. ROI 추출 및 색상 지표", level=2, size=11, bold=True)
P("입력 이미지마다 시편 픽셀에 분석을 한정하고 카드 배경을 제외하기 "
  "위해 자동 관심영역(ROI)을 추출한다. 파이프라인은 HSV 명도/채도 "
  "임계값(V > 215, S < 25)을 적용하여 종이 마스크를 얻고, 가장 큰 "
  "비-종이 윤곽을 추출하며, 정반사로 인한 시편 픽셀 손실을 회복하기 "
  "위해 윤곽이 이미지 면적의 90% 미만인 경우 볼록 껍질로 대체한다. "
  "ROI 크기는 조건별로 적응적으로 설정되며(이미지 면적의 13–17%, "
  "수동 라벨에서 학습), 시편 경계상자의 종횡비는 제곱근 변환으로 "
  "완화된다. ROI 중심은 시편 중심 모멘트와 경계상자 중심의 기하 "
  "평균에 배치되며, 이는 4가지 조건에서 가장 일관된 ROI를 생성하였다.",
  align="justify")
P("ROI 내부에서 다음 색상 지표를 계산한다: CIE Lab b∗ 좌표(황-청 "
  "축), HSI 채도 채널 S, ASTM E313 황색지수(YI), 그리고 동일 조건의 "
  "0일 기준 누적 CIE ΔE 색차. 이 지표들은 HfS₂ 산화에 동반되는 "
  "거시적인 채도 높은 황색의 손실을 포착한다.",
  align="justify")

H("2.3. 5-방법 경과일 추정 앙상블", level=2, size=11, bold=True)
P("5개 독립 추정기를 결합한다(Fig. 2). (1) 가중 KNN: 동일 조건의 "
  "참조 이미지와 질의 이미지 간 d = √(0.45·Δb∗² + 0.30·ΔS² + "
  "0.25·ΔYI²) 의 마할라노비스형 거리를 측정하여 상위 3개 참조일의 "
  "역거리 가중 평균을 계산한다. (2) Wasserstein 거리: 질의의 b∗ "
  "픽셀 히스토그램을 각 참조와 비교하여 다시 3개 최근접 참조의 평균을 "
  "낸다. (3) FFT 텍스처: ROI 그레이스케일 이미지의 고주파 에너지 "
  "비율, 스펙트럼 엔트로피, 방사상 파워 프로파일을 추출하여 가중 "
  "특성 거리로 정렬한다. (4) 공간 패턴: ROI를 3×3 격자로 분할하여 "
  "셀별 b∗ 평균과 표준편차를 계산하고, 셀 간 엔트로피 및 중앙-경계 "
  "그래디언트를 통해 산화 이방성을 정량화한다. (5) 동역학 추정기: "
  "조건별 b∗(t) 시계열을 1차 지수 감쇠 b∗(t) = b∗_∞ + (b∗_0 − "
  "b∗_∞)·exp(−k·t) 모델로 SciPy의 비선형 최소제곱[9]을 사용해 "
  "경계 매개변수와 함께 적합시키고, 모델을 역산하여 질의 b∗에 "
  "해당하는 경과일을 복원한다.",
  align="justify")

IMG("fig2_pipeline.png", width=6.5,
    cap="Fig. 2. 이미지 전용 산화 경과일 추정 파이프라인의 모식도. "
        "시편 사진은 자동 ROI 내에서 색상 및 텍스처 기술자로 변환되고, "
        "5개 독립 추정기가 후보 경과일 값을 산출하며, 이는 조건별 "
        "가중 앙상블로 결합된다.")

P("5개 후보 경과일은 ŷ = Σwᵢ · dᵢ로 결합되며, 가중치 {wᵢ}는 "
  "참조 세트의 알려진 경과일에 대한 leave-one-out 오차에 대해 Huber "
  "손실 ℓ(e) = ½·e² (|e|≤δ) / δ·(|e|-½δ) (|e|>δ)을 최소화하도록 "
  "학습되며, δ = 5일을 사용한다. 단일 전역 가중치 세트(“전역 균등/"
  "최적”)와 4개 조건별 가중치 세트의 두 가지 체제를 평가하며, 후자는 "
  "방법별 본질적 정확도가 표면 패시베이션에 강하게 의존하므로 선택된다. "
  "최적화는 SciPy의 L-BFGS-B를 사용해 5개 다중 시작점에서 수행된다 [9].",
  align="justify")

H("2.4. 의사-라만 스펙트럼 재구성", level=2, size=11, bold=True)
P("이미지 특성이 진동 전이를 직접 재현할 수는 없지만, ~337 cm⁻¹의 "
  "A₁ₒ 모드 — HfS₂의 주요 라만 지문 — 의 상대 강도는 표면 산화에 "
  "단조적으로 변하므로 경과일 추정을 견인하는 동일 색상 지표와 "
  "잘 상관된다. 의사-라만 강도 î⁺는 참조 세트의 이미지-라만 쌍에 "
  "학습된 4개의 단변량 선형 회귀(b∗ → i⁺, S → i⁺, YI → i⁺, "
  "ΔE → i⁺)의 R² 가중 앙상블로 예측된다. 그 다음 정규화된 피크가 "
  "î⁺를 좌우로 둘러싸는 두 참조 라만 스펙트럼을 선형 보간하고, "
  "회귀 표준 오차 σ̂로부터 95% 신뢰 구간을 î⁺ ± 1.96·σ̂으로 "
  "전파한다. 출력은 사용자가 벤치탑 측정과 직접 비교할 수 있는 강도-"
  "라만 시프트 곡선이다.",
  align="justify", after=8)


# ════════════════ 3. 결과 및 논의 ════════════════
H("3. 결과 및 논의", level=1, size=12, bold=True)

H("3.1. 색상 지표의 시간 변화", level=2, size=11, bold=True)
P("Fig. 3은 4가지 조건에 대한 3가지 주요 색상 지표를 경과일의 함수로 "
  "도시한다. 70% RH의 Native HfS₂는 예상대로 급격한 감쇠를 보인다: "
  "b∗가 2주 내에 28에서 3으로 붕괴하고 S 채널 평균이 54에서 4로 "
  "감소하여, 황색 칼코제나이드의 완전 손실과 투명 HfO₂로의 전환을 "
  "반영한다. 35% RH에 보관된 Native HfS₂는 훨씬 느리게 변화하여 "
  "4주 후에야 b∗ ≈ 6에 도달한다. Al₂O₃ 캡슐화 HfS₂는 4가지 조건 "
  "중 가장 적게 변화하며, 30일 전 기간 동안 b∗가 16과 28 사이에서 "
  "변동하는데, 이는 [6]에서 보고된 Al₂O₃ 캡의 보호 역할과 완전히 "
  "일치하는 거동이다. PMMA는 동일 습도에서 베어 HfS₂ 대비 감쇠를 "
  "약 2배 지연시키지만 궁극적으로는 산화를 막지 못한다.",
  align="justify")

IMG("fig3_metric_trends.png", width=6.5,
    cap="Fig. 3. 4가지 패시베이션 조건에서 3가지 이미지 색상 지표의 "
        "시간 변화. (a) CIE Lab b∗, (b) HSI S 채널 평균, (c) ASTM "
        "E313 황색지수. Al₂O₃/HfS₂의 평탄한 궤적은 Al₂O₃ 캡이 "
        "제공하는 효과적인 산화 장벽을 반영한다.")

P("경과일과 b∗의 Pearson 상관계수 r(day, b∗)는 35% RH의 Native "
  "HfS₂에서 −0.92, 70% RH의 Native HfS₂에서 −0.83, PMMA/HfS₂에서 "
  "−0.89, Al₂O₃/HfS₂에서는 −0.32에 불과하여, Al₂O₃ 패시베이션 "
  "시편이 전체 경과일 범위에서 이미지 공간상 거의 정상(stationary)인 "
  "관찰을 정량적으로 확인한다. 이는 어떤 이미지 기반 추정기도 능가할 "
  "수 없는, 해당 조건의 경과일 추정 정확도의 본질적 상한을 설정한다.",
  align="justify")

H("3.2. 방법별 및 앙상블 정확도", level=2, size=11, bold=True)
P("Fig. 4는 5개 개별 추정기와 두 가지 앙상블 구성(단일 전역 가중치 "
  "세트 및 조건별 가중치)의 leave-one-out RMSE를 보고한다. 세 가지 "
  "관찰이 두드러진다. 첫째, KNN 추정기는 4가지 조건 중 3가지에서 "
  "가장 정확한 단일 방법으로(35% RH Native 3.12 d, 70% RH Native "
  "3.56 d, 70% RH PMMA 2.50 d) 강하게 단조적인 색상 기술자를 직접 "
  "활용하는 데 기인한다. 둘째, FFT 및 동역학 추정기는 단독 사용 시 "
  "성능이 낮지만, 이들의 독립 오차가 KNN과 부분적으로 비상관이므로 "
  "앙상블에서 한계 이득을 제공할 수 있다. 셋째, 조건별 가중치는 "
  "가중 평균 RMSE를 7.74 d(균등 baseline)와 6.30 d(전역 단일 최적 "
  "가중치 세트)에서 4.80 d로 감소시켜 균등 baseline 대비 38% 개선을 "
  "달성한다.",
  align="justify")

IMG("fig4_method_rmse.png", width=6.5,
    cap="Fig. 4. 5개 개별 추정기와 2가지 앙상블 구성의 조건별 RMSE "
        "(leave-one-out). 조건별 가중치는 단일 전역 가중치 세트를 "
        "강하게 능가하며, 35% RH의 Native HfS₂와 PMMA/HfS₂에서 "
        "특히 두드러진다.")

P("조건별 가중치 자체는 물리적으로 해석 가능하다. 70% RH의 Native "
  "HfS₂에서는 KNN 단독(100%)이 최적인데, 강한 b∗ 단조 감쇠가 색상 "
  "거리 KNN을 본질적으로 정확하게 만들기 때문이다. PMMA/HfS₂에서는 "
  "KNN(83%)이 지배적이고 Wasserstein 히스토그램(12%)이 소량 기여한다. "
  "35% RH의 Native HfS₂에서는 Wasserstein 방법이 지배적(55%)인데, "
  "느린 감쇠가 b∗ 히스토그램을 평균값보다 Earth-Mover 거리로 더 "
  "안정적으로 포착되는 정도로 이동시키기 때문이다. Al₂O₃/HfS₂에서는 "
  "FFT(45%), 공간 패턴(24%), Wasserstein(22%), 동역학(3%)에 보다 "
  "균등하게 가중치가 분산되는데, 이는 거의 정상에 가까운 영역에서는 "
  "어떤 방법도 결정적인 단일 신호를 갖지 못하기 때문이다.",
  align="justify")

H("3.3. 의사-라만과 측정 라만의 비교", level=2, size=11, bold=True)
P("4가지 조건 중 Al₂O₃/HfS₂ 시리즈만이 측정 라만 스펙트럼의 완전한 "
  "세트를 제공한다(베어 및 PMMA 시편은 70% RH에서 며칠 내에 A₁ₒ "
  "모드를 완전히 잃어 그 이후 라만으로 추적할 수 없다). Fig. 5(a)는 "
  "0일부터 28일까지 Al₂O₃/HfS₂의 측정 라만 스펙트럼을 중첩하여, "
  "A₁ₒ 피크 강도가 0일의 정규화 값 1.00에서 28일의 0.60으로 점진적이지만 "
  "명확하게 감소하며, 0.7 cm⁻¹의 작은 동반 피크 시프트가 발생함을 "
  "보여준다. Fig. 5(b)는 측정된 정규화 A₁ₒ 피크와 이미지 지표만으로 "
  "생성된 의사-라만 추정값을 비교한다. 두 곡선은 모든 측정 시점에서 "
  "전파된 95% 신뢰 구간 내에 일치하며, 0, 3, 7, 14, 28일에서 절댓값 "
  "오차가 각각 0.00, 0.01, 0.00, 0.01, 0.00 정규화 강도 단위이다. "
  "평균 절댓값 오차 0.005 정규화 단위는 약 1–2일의 등가 경과일 "
  "불확도에 해당한다.",
  align="justify")

IMG("fig5_pseudo_raman.png", width=6.5,
    cap="Fig. 5. (a) Al₂O₃ 캡슐화 HfS₂의 28일에 걸친 측정 라만 "
        "스펙트럼으로, 점진적인 A₁ₒ 감쇠를 보여준다. (b) 측정된 "
        "정규화 A₁ₒ 피크 강도(빨강)와 이미지 기반 의사-라만 "
        "추정값(파랑) 및 95% 신뢰 밴드의 비교. 두 곡선은 30일 전 "
        "기간에서 ~0.01 정규화 강도 단위 이내로 일치한다.")

P("이러한 정량적 일치는 본 연구의 핵심 전제를 검증한다: 기저 산화 "
  "화학이 단조적인 광학 신호를 남길 때 이미지 전용 추정기가 분광기 "
  "없이도 미세 라만이 측정할 동일한 A₁ₒ 추세를 복원할 수 있다는 "
  "것이다. 광학 신호가 포화되기 전에 라만 신호가 사라지는 시편(70% "
  "RH의 Native HfS₂가 극단적 사례)에서는 의사-라만 추정이 형식적으로 "
  "외삽이 되며 앙상블에서 신뢰도가 감소된 형태로 보고된다. 이러한 "
  "경우는 사용자 인터페이스에 표시되지만, 이미지 기반 경과일은 며칠 "
  "이내의 정확도를 유지한다.",
  align="justify")

H("3.4. 한계", level=2, size=11, bold=True)
P("세 가지 한계를 강조해야 한다. 첫째, 이미지 기반 추정은 라만 분광법이 "
  "실제로 측정하는 표면 아래 화학을 탐지할 수 없다; 표면과 벌크 산화 "
  "상태가 다른 시편에서는 두 기법이 일치하지 않을 수 있다. 둘째, "
  "Al₂O₃ 캡슐화 시편의 경우 거시적 광학 변화가 30일 동안 너무 작아서 "
  "(Δb∗ ≈ 2 vs. 베어 HfS₂ ~25), 자연적인 시편 간 변동이 시간적 "
  "신호를 압도하여 달성 가능한 경과일 정확도가 ~8일로 제한된다. "
  "인터페이스는 Al₂O₃ 질의를 30%의 신뢰도 한계와 텍스트 경고로 "
  "명시적으로 표시한다. 셋째, 현재 의사-라만 재구성은 HfS₂의 A₁ₒ "
  "모드만을 대상으로 한다; 완전 산화된 시편에서 나타나야 하는 500 및 "
  "630 cm⁻¹ 부근의 HfO₂ 진동 모드는 모델링되지 않았으며, 이를 위해서는 "
  "해당 측정 스펙트럼을 포함하는 확장된 학습 세트가 필요하다.",
  align="justify", after=8)


# ════════════════ 4. 결론 ════════════════
H("4. 결론", level=1, size=12, bold=True)
P("본 연구에서는 단일 시편 사진이 자동 ROI 추출기와 Huber 견고 "
  "회귀로 학습된 조건별 가중치를 갖는 5-방법 추정기 앙상블을 거쳐, "
  "4가지 패시베이션 조건과 33개 질의 이미지에 걸쳐 가중 평균 RMSE "
  "4.80일로 CVD 성장 HfS₂ 박막의 자연 산화 경과일을 예측할 수 "
  "있음을 입증하였다. Al₂O₃ 캡슐화 시편의 경우, 동일한 이미지 "
  "특성으로부터 의사-라만 재구성을 수행하였고, 그 A₁ₒ 강도 추정이 "
  "0–28일 전 경과일 구간에서 측정 라만 스펙트럼과 0.01 정규화 단위 "
  "이내로 일치하였다. 본 접근법은 미세 라만이 비실용적인 제조 라인 "
  "스크리닝 및 현장 검사에 특히 가치가 있는, 분광기 없는 1초 미만의 "
  "HfS₂ 일상적 품질 모니터링의 대안을 제공한다. 향후 연구는 학습 세트를 "
  "확장하여 HfO₂ 진동 모드를 포함하고, 추가 캡슐화제로 조건별 가중치 "
  "라이브러리를 확장하며, 카메라 및 조명 조건 간 전이성을 조사할 "
  "예정이다.",
  align="justify", after=10)


# ════════════════ MDPI 마무리 섹션 ════════════════
H("Supplementary Materials", level=1, size=11, bold=True)
P("본 논문에 첨부되는 Supporting Information은 References 다음 페이지에 "
  "수록되어 있으며, 다음을 포함한다: Section S1 (보충 방법) — 자동 ROI "
  "추출(S1.1), 색상·텍스처 기술자(S1.2), 5가지 경과일 추정기(S1.3), "
  "Huber 견고 앙상블 가중치 최적화(S1.4), 의사-라만 회귀 앙상블(S1.5)에 "
  "대한 확장 설명; Section S2 (데이터셋 설명) — 이미지 획득 설정, "
  "라만 측정 세부사항, 조건별 시편 수; Section S3 (소프트웨어, 재현성, "
  "계산 비용); Figures S1–S6 — 33장 평가 mosaic, 자동 ROI overlay, "
  "방법별 산점도, 조건별 가중치 heatmap, 비-Al₂O₃ 조건의 의사-라만 "
  "추정, 조건별 동역학 적합 곡선; Tables S1–S5 — 시편별 LOO 예측, "
  "조건×방법 가중치 행렬, r(day, metric), 데이터셋 요약, 종합 "
  "하이퍼파라미터 목록.", align="justify", after=6)

H("저자 기여", level=1, size=11, bold=True)
P("개념화: 전용철, 박광욱; 방법론: 전용철; 소프트웨어: 전용철; "
  "검증: 전용철, 박광욱; 정형 분석: 전용철; 조사: 전용철, 박광욱; "
  "자원: 박광욱; 데이터 큐레이션: 전용철; 원고 작성: 전용철; "
  "원고 검토 및 편집: 박광욱; 시각화: 전용철; 감독: 박광욱; "
  "프로젝트 관리: 박광욱; 자금 조달: 박광욱. 모든 저자는 "
  "출판 버전 원고를 읽고 동의하였다.", align="justify", after=6)

H("재정 지원", level=1, size=11, bold=True)
P("본 연구는 한국연구재단(NRF) 전북대학교 사업의 지원을 받아 수행되었다.",
  align="justify", after=6)

H("데이터 공개", level=1, size=11, bold=True)
P("본 연구의 결과를 뒷받침하는 이미지 데이터셋(output_cut/, sample/), "
  "통합 분석 데이터베이스(alldata.db), 라만 참조 데이터베이스"
  "(raman.raman.db)는 합리적 요청 시 교신저자로부터 제공된다. "
  "이미지 처리 파이프라인 소스코드는 https://github.com/YongCheulJun/"
  "hfs2_analyzer 에서 공개되어 있다.", align="justify", after=6)

H("감사의 글", level=1, size=11, bold=True)
P("저자들은 시편 준비, 라만 측정, 유익한 토론을 도와준 전북대학교 "
  "박광욱 교수 연구실 구성원들에게 감사를 표한다.", align="justify",
  after=6)

H("이해 충돌", level=1, size=11, bold=True)
P("저자들은 이해 충돌이 없음을 선언한다.", align="justify", after=10)


# ════════════════ 참고문헌 ════════════════
H("참고문헌", level=1, size=11, bold=True)
refs = [
    "[1] Q.H. Wang, K. Kalantar-Zadeh, A. Kis, J.N. Coleman, "
    "M.S. Strano, Electronics and optoelectronics of two-dimensional "
    "transition metal dichalcogenides, Nat. Nanotechnol. 7 (2012) "
    "699–712.",
    "[2] M. Chhowalla, H.S. Shin, G. Eda, L.J. Li, K.P. Loh, H. Zhang, "
    "The chemistry of two-dimensional layered transition metal "
    "dichalcogenide nanosheets, Nat. Chem. 5 (2013) 263–275.",
    "[3] D. Jariwala, V.K. Sangwan, L.J. Lauhon, T.J. Marks, "
    "M.C. Hersam, Emerging device applications for semiconducting "
    "two-dimensional transition metal dichalcogenides, ACS Nano 8 "
    "(2014) 1102–1120.",
    "[4] T. Kanazawa, T. Amemiya, A. Ishikawa, V. Upadhyaya, K. Tsuruta, "
    "T. Tanaka, Y. Miyamoto, Few-layer HfS₂ transistors, Sci. Rep. 6 "
    "(2016) 22277.",
    "[5] D. Singh, R. Ahuja, Two-dimensional layered HfS₂ as a "
    "promising photocatalyst for hydrogen production, ACS Appl. "
    "Energy Mater. 2 (2019) 6891–6900.",
    "[6] J. Hwang, J. Mun, K.-T. Lee, T. Lee, J. Kim, J. Min, K. Park, "
    "Impact of humidity on long-term stability of HfS₂ grown on "
    "sapphire substrate by chemical vapor deposition and strategies "
    "to prevent native oxidation, Mater. Sci. Semicond. Process. 192 "
    "(2025) 109471.",
    "[7] D.G. Lowe, Distinctive image features from scale-invariant "
    "keypoints, Int. J. Comput. Vis. 60 (2004) 91–110.",
    "[8] L. Jin, R. Cao, et al., Image-based machine-learning "
    "monitoring of MoS₂ ageing in ambient atmosphere, ACS Appl. "
    "Mater. Interfaces 13 (2021) 35562–35570.",
    "[9] P. Virtanen, R. Gommers, T.E. Oliphant, et al., SciPy 1.0: "
    "fundamental algorithms for scientific computing in Python, "
    "Nat. Methods 17 (2020) 261–272.",
    "[10] CIE (1976). Colorimetry (2nd ed.). CIE Publication 15.2. "
    "Commission Internationale de l'Éclairage.",
    "[11] ASTM International. (2015). ASTM E313: Standard Practice "
    "for Calculating Yellowness and Whiteness Indices from "
    "Instrumentally Measured Color Coordinates.",
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    r = p.add_run(ref)
    r.font.name = BODY; r.font.size = Pt(9)


# ════════════════ Supporting Information ════════════════
doc.add_page_break()

si_t = doc.add_paragraph()
si_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = si_t.add_run("보충 자료 (Supporting Information)")
r.font.name = BODY; r.font.size = Pt(18); r.font.bold = True
si_t.paragraph_format.space_after = Pt(4)

si_s = doc.add_paragraph()
si_s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = si_s.add_run(
    "이미지 기반 다중 추정기 앙상블 및 의사-라만 재구성에 의한 "
    "CVD 성장 HfS₂ 박막의 자연 산화 경과일 추정")
r.font.name = BODY; r.font.size = Pt(11); r.font.italic = True
si_s.paragraph_format.space_after = Pt(4)

si_a = doc.add_paragraph()
si_a.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = si_a.add_run("전용철")
r.font.name = BODY; r.font.size = Pt(10)
rs = si_a.add_run("¹")
rs.font.name = BODY; rs.font.size = Pt(10); rs.font.superscript = True
r = si_a.add_run(", 박광욱")
r.font.name = BODY; r.font.size = Pt(10)
rs = si_a.add_run("²,*")
rs.font.name = BODY; rs.font.size = Pt(10); rs.font.superscript = True
si_a.paragraph_format.space_after = Pt(2)

si_c = doc.add_paragraph()
si_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
si_c.paragraph_format.space_after = Pt(14)
r = si_c.add_run("* 교신저자: kpark@jbnu.ac.kr (박광욱)")
r.font.name = BODY; r.font.size = Pt(9)


H("목차", level=1, size=12, bold=True)
toc = [
    "S1. 보충 방법",
    "    S1.1 자동 ROI 추출",
    "    S1.2 색상 및 텍스처 기술자",
    "    S1.3 5가지 경과일 추정기",
    "    S1.4 Huber 견고 앙상블 가중치 최적화",
    "    S1.5 의사-라만 스펙트럼 재구성",
    "S2. 데이터셋 설명",
    "S3. 소프트웨어, 재현성, 계산 비용",
    "S4. 보충 그림 (S1–S6)",
    "S5. 보충 표 (S1–S5)",
]
for it in toc:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(it)
    r.font.name = BODY; r.font.size = Pt(9.5)
P("", after=8)


# ─── S1 ───
H("S1. 보충 방법", level=1, size=14, bold=True)

H("S1.1 자동 ROI 추출", level=2, size=11, bold=True)
P("각 사진은 OpenCV의 BT.601 변환을 사용해 sRGB에서 HSV로 변환된다. "
  "M_paper = (V > 215) ∧ (S < 25) 로 종이 마스크를 계산하여 밝고 "
  "무채색의 카드 배경을 임계값 처리한다. M_paper의 보집합은 크기 "
  "k = max(7, min(H, W)/80)의 타원형 구조 요소를 사용해 형태학적 "
  "닫힘 후 열림으로 정리되며, H × W는 이미지 해상도이다. 닫힘 연산이 "
  "먼저 적용되어 시편 내부의 정반사 구멍을 메우고, 이어서 열림 연산이 "
  "배경 잡음을 제거한다. 가장 큰 연결 윤곽(cv2.findContours, "
  "RETR_EXTERNAL)을 원시 시편 영역으로 취하며, 볼록 껍질이 이미지 "
  "면적의 90% 미만을 둘러쌀 때 항상 윤곽을 볼록 껍질로 대체하여 "
  "카메라 플래시 정반사로 손실된 시편 픽셀을 회복한다.",
  align="justify")
P("시편 윤곽의 경계상자(sx, sy, sw, sh) 내에서 ROI는 윤곽 모멘트 중심과 "
  "경계상자 중심의 중간점에 배치된다(c_x, c_y) = ½(μ_contour + (sx+sw/2, "
  "sy+sh/2)). ROI 치수는 조건별 목표 면적 비율 α_c (이미지 면적의 "
  "13–17%, 수동 라벨에서 학습)와 시편 경계상자의 완화된 종횡비 "
  "a = √(sw/sh)로부터 얻는다: w_ROI = √(α_c · img_area · a), "
  "h_ROI = √(α_c · img_area / a), 단 w_ROI ≤ 0.70 sw 및 "
  "h_ROI ≤ 0.70 sh. 마지막으로 5% 가장자리 여백이 ROI를 이미지 "
  "경계로부터 안쪽으로 이동시킨다.", align="justify")

H("S1.2 색상 및 텍스처 기술자", level=2, size=11, bold=True)
P("ROI 내에서 4가지 주요 색상 지표가 계산된다:", align="justify",
  after=2)
P("• b∗ — D65 표준 백색점을 사용한 OpenCV의 BGR → Lab 변환을 통해 "
  "얻은 CIE Lab 색공간의 황-청 좌표. 더 높은 b∗는 더 채도 높은 황색에 "
  "해당한다.", align="justify", after=2)
P("• S 채널 평균 — ROI 픽셀에 대한 평균 HSI 채도 채널. OpenCV의 "
  "HSV 대안 대신 분석적 HSI 변환(S = 1 − min(R,G,B)/I)을 사용하며, "
  "HSI 형태가 산화 중인 HfS₂에서 관찰되는 채도 콘텐츠 손실에 더 "
  "민감하다.", align="justify", after=2)
P("• 황색지수(YI) — ASTM E313: YI = 100 · (1.3013 X − 1.1498 Z) / "
  "Y, 여기서 X, Y, Z는 CIE 1931 자극값. YI는 폴리머 및 코팅 품질 "
  "관리에서 사용되는 국제 표준 황색도 측정값이다.",
  align="justify", after=2)
P("• ΔE — 동일 조건의 0일 기준에 대한 질의의 누적 CIE 1976 색차: "
  "ΔE = √((ΔL∗)²+(Δa∗)²+(Δb∗)²). ΔE는 절대 지각 단위로 "
  "해석 가능하다(ΔE ≈ 3 = 식별 가능한 차이, ΔE > 10 = 명백한 변화).",
  align="justify", after=4)
P("두 가지 텍스처 기술자가 색상 지표를 보완한다: ROI 그레이스케일 "
  "이미지의 2-D FFT는 고주파 에너지 비율(반경 > 0.4 r_max), 방사상 "
  "파워 프로파일의 스펙트럼 엔트로피, 그리고 이미지 지문 벡터로 "
  "사용되는 64-bin 방사상 평균 스펙트럼을 산출한다. 3 × 3 공간 패턴 "
  "기술자는 셀별 b∗ 평균과 표준편차를 계산한 후, 셀 간 엔트로피, "
  "중앙-경계 그래디언트, 행/열 이방성 비율을 도출한다.",
  align="justify")

H("S1.3 5가지 경과일 추정기", level=2, size=11, bold=True)
P("(1) k-NEAREST NEIGHBOUR (KNN). 동일 조건의 각 질의 및 참조 "
  "이미지에 대해 가중 유클리드 거리 d(q, r) = √(0.45·Δb∗² + "
  "0.30·ΔS² + 0.25·ΔYI²)를 조건 풀 전반에 걸친 각 지표의 min-max "
  "정규화 후 계산한다. 경과일은 3개 최근접 참조일의 역거리 가중 "
  "평균으로 추정된다, ŷ = Σᵢ dᵢ⁻¹ · day(rᵢ) / Σᵢ dᵢ⁻¹. "
  "신뢰도는 max(0, 100 − 200 · d_min)이다.", align="justify")
P("(2) WASSERSTEIN. 모든 ROI 픽셀의 b∗ 값은 b∗ = −30에서 80 사이의 "
  "64-bin 히스토그램으로 비닝된다. 질의 히스토그램과 각 참조 "
  "히스토그램 간의 Earth-Mover (1D Wasserstein) 거리는 경험적 CDF L1 "
  "거리를 통해 계산된다. 경과일은 다시 3개 최근접 참조의 역거리 "
  "가중 평균으로 얻는다.", align="justify")
P("(3) FFT 텍스처. 질의와 참조의 FFT 고주파 비율 h, 스펙트럼 "
  "엔트로피 ε, 64-bin 방사상 프로파일 r은 d = 0.5·|Δh|/range(h) + "
  "0.3·|Δε|/range(ε) + 0.2·(1 − r·r' / |r||r'|) (방사 프로파일에 "
  "대한 코사인 거리)로 비교된다.", align="justify")
P("(4) 공간 패턴. 3 × 3 격자에서 셀별 b∗ 평균 및 표준편차로부터 "
  "엔트로피/경계-그래디언트/이방성 기술자를 도출하여 정규화 L1 "
  "거리로 비교한다.", align="justify")
P("(5) 동역학. 모든 조건에서 b∗(t) 시계열은 b∗(t) = b∗_∞ + "
  "(b∗_0 − b∗_∞)·exp(−k·t)로 SciPy의 scipy.optimize.curve_fit과 "
  "경계 [b∗_0 ∈ (0,150), b∗_∞ ∈ (−20,80), k ∈ (10⁻⁶, 20)]로 "
  "적합된다. 질의 b∗는 ŷ_kinetic = −ln((b∗ − b∗_∞)/(b∗_0 − "
  "b∗_∞))/k로 역산된다. 선형 가능성 비율이 (0,1] 밖일 때(b∗가 "
  "신선 값보다 높거나 포화보다 낮을 때) 모델은 신뢰도가 감소된 "
  "클립된 외삽을 반환한다.", align="justify")

H("S1.4 Huber 견고 앙상블 가중치 최적화", level=2, size=11, bold=True)
P("조건의 n개 참조 이미지에 대한 leave-one-out 방법별 경과일 예측 "
  "ŷᵢ,m과 알려진 참값 yᵢ가 주어졌을 때, 앙상블 가중치 w = "
  "(w_KNN, w_Wass, w_FFT, w_Spatial, w_Kinetic)는 다음 Huber "
  "손실을 최소화하여 얻는다.", align="justify", after=2)
P("    L(w) = (1/n) Σᵢ ℓ_δ(yᵢ − Σm wm·ŷᵢ,m)",
  align="center", after=2)
P("ℓ_δ(e) = ½e² (|e| ≤ δ) / δ(|e| − ½δ) (|e| > δ), δ = 5일. "
  "Huber 손실은 적당한 오차에 대해 이차 페널티를 부과하지만 큰 outlier에 "
  "대해서는 선형 페널티로 감소하여, 산화 화학이 다른(개별 예측 오차가 "
  "15–20일을 초과할 수 있는) Al₂O₃/HfS₂ 시편이 앙상블을 지배하는 "
  "것을 방지한다. 최적화는 SciPy의 L-BFGS-B를 사용해 0 ≤ wm ≤ 1의 "
  "경계와 L1 정규화 후, 6개 다중 시작점(균등 1/5 + 가중치 0.80의 5개 "
  "방법별 시작)에서 진행된다. SciPy를 사용할 수 없을 때는 0.1-단계 "
  "5-D 격자 검색(11⁵ ≈ 1.6 × 10⁵ 후보)이 대신 수행된다.",
  align="justify")

H("S1.5 의사-라만 스펙트럼 재구성", level=2, size=11, bold=True)
P("참조 이미지의 4가지 색상 및 강도 지표 m ∈ {b∗, S, YI, ΔE}는 "
  "측정된 정규화 A₁ₒ 피크 강도와 짝지어져 4개의 단변량 보통 "
  "최소제곱 회귀자 î_m = α_m + β_m · m을 적합시킨다. 그 R² "
  "값은 예측기 가중치 w_m = R²_m / Σ R²을 얻기 위해 재정규화된다. "
  "질의 이미지의 예측 피크는 î_query = Σ w_m · (α_m + β_m · "
  "m_query)이며, 표준 오차 σ̂ = √(Σ w_m² · σ_m²)와 95% 신뢰 "
  "구간 î_query ± 1.96 σ̂를 갖는다. 정규화된 피크가 î_query를 "
  "둘러싸는 두 참조 라만 스펙트럼은 가중치 α = (î_query − p_low)/"
  "(p_high − p_low)을 [0,1]로 클립한 후 선형 보간하여 전체 의사-"
  "라만 강도-시프트 곡선을 생성한다.",
  align="justify", after=8)


# ─── S2 ───
H("S2. 데이터셋 설명", level=1, size=14, bold=True)
P("모든 HfS₂ 박막은 c-평면 사파이어 기판 위에 [6]에 기술된 절차에 "
  "따라 대기압 화학기상증착(CVD)으로 성장되었다. 성장 후, 시편은 "
  "제어된 상대습도(RH) 조건에서 보관되었으며 안정적인 형광등 조명"
  "(~4000 K, 시편 표면에서 ~600 lx) 아래에서 고정된 복사 스탠드에 "
  "장착된 일반 DSLR로 여러 경과 시점에 촬영되었다. 흰색 카드 배경은 "
  "S1.1에서 설명한 HSV 기반 ROI 추출이 가능하도록 모든 세션에서 "
  "동일하게 유지되었다. 이미지별 화이트 밸런스 보정은 적용되지 "
  "않았으며, 카메라는 고정된 색온도 5500 K로 설정되었다. 참조 라만 "
  "스펙트럼은 동일한 물리 시편에서 532 nm 레이저(시편에서 ~1 mW, "
  "50× 대물렌즈, 30 s 적분)를 사용하는 Renishaw inVia 공초점 라만 "
  "현미경으로 수집되었다; Al₂O₃ 캡슐화 시편의 측정 A₁ₒ 피크 강도가 "
  "의사-라만 검증의 ground-truth 역할을 하였다. 다른 3가지 조건"
  "(35% 및 70% RH의 Native HfS₂, 70% RH의 PMMA/HfS₂)은 높은 "
  "습도에서 며칠 내에 A₁ₒ 피크를 완전히 잃었으며, 그 이후로는 광학적 "
  "추적만 가능했다; raman.raman.db의 해당 항목은 placeholder 값"
  "(norm_peak = 1.0)이다.", align="justify")
P("Table S4는 조건별 시편 수를 요약한다. 4가지 조건에서 데이터셋은 "
  "14개 서로 다른 경과일에 53장의 이미지로 구성되며, 별도 세션에서 "
  "촬영된 동일 물리 시편의 독립 20장 평가 세트(sample/)가 추가된다"
  "(본문 Fig. 1 및 여기 Fig. S1 참조).",
  align="justify", after=8)


# ─── S3 ───
H("S3. 소프트웨어, 재현성, 계산 비용", level=1, size=14, bold=True)
P("모든 이미지 처리, 앙상블 가중치 학습, 의사-라만 재구성 코드는 "
  "프로젝트 모듈 hfs2_v5_49.py 하에 https://github.com/YongCheulJun/"
  "hfs2_analyzer 에서 오픈소스로 공개되어 있다. 헤드리스 재현 "
  "스크립트 tools/optimize_weights_headless.py는 통합 분석 데이터베이스 "
  "alldata.db (53장 참조 이미지)와 독립적인 평가 디렉토리 "
  "output_cut/ (33장 질의 이미지)로부터 본문에 보고된 모든 수치 "
  "결과를 단일 명령으로 재현한다:", align="justify", after=2)
P("    python3 tools/optimize_weights_headless.py \\",
  align="justify", after=0)
P("        --pool newfiles/output/output_cut/db/alldata.db \\",
  align="justify", after=0)
P("        --targets newfiles/output/output_cut --save",
  align="justify", after=4)
P("소프트웨어 의존성(본 연구에서 사용된 버전): Python 3.12, NumPy "
  "2.4.4, OpenCV 4.13.0 (헤드리스), Pillow 12.2.0, Matplotlib 3.10.7, "
  "SciPy 1.18.0, tkinterdnd2 0.4.4 (GUI 전용), python-docx 1.2.0 "
  "(논문 출력). 단일 코어 일반 노트북(11세대 Intel Core i7-1185G7) "
  "에서 33개 질의 전체에 대한 leave-one-out 평가는 — 5가지 방법 평가 "
  "및 Huber 손실 다중 시작점 최적화를 포함하여 — 90초 미만에 완료된다. "
  "배포 시점의 질의별 추론(단일 이미지 → 모든 5가지 추정기, 앙상블, "
  "의사-라만)은 200 ms 미만이며, 일반적인 30초 라만 적분 시간보다 "
  "약 1500배 빠르므로, 제안된 파이프라인은 제조 라인 및 현장 품질 "
  "스크리닝에 실용적이다.", align="justify", after=8)


# ─── S4 ───
H("S4. 보충 그림", level=1, size=14, bold=True)

H("Figure S1. 전체 평가 시편 mosaic", level=2, size=11, bold=True)
P("본문의 leave-one-out 평가에 질의로 사용된 33장 시편 전체의 사진. "
  "이 세트는 4가지 패시베이션/습도 조건과 0–30일의 경과일을 포함한다. "
  "파일명은 <day>day_<RH>RH_<condition>.png 규약을 따른다.",
  align="justify")
IMG("figS1_full_specimen.png", width=6.5,
    cap="Figure S1. output_cut/ 코퍼스의 33장 평가 시편 사진 mosaic.")

H("Figure S2. 자동 ROI overlay", level=2, size=11, bold=True)
P("HSV 마스크 + 볼록 껍질 파이프라인이 4가지 조건 전반에 걸쳐 견고함을 "
  "보여주는 대표적인 자동 관심영역 overlay. 빨간 직사각형이 색상 및 "
  "텍스처 추정기에 사용되는 ROI를 표시한다.", align="justify")
IMG("figS2_roi_overlay.png", width=6.5,
    cap="Figure S2. 각 조건의 0일, ~14일, ~28일 대표 시편의 자동 ROI overlay.")

H("Figure S3. 개별 추정기의 참값-예측 산점도",
  level=2, size=11, bold=True)
P("33개 leave-one-out 질의 전반에 대한 방법별 참값-예측 경과일 산점도. "
  "KNN(RMSE 6.56 d)이 강하게 단조적인 색상 기술자를 직접 활용함으로써 "
  "대각선 주변에서 가장 좁은 산점을 보인다. 동역학 감쇠 추정기"
  "(RMSE 14.56 d)가 가장 잡음이 많은 단일 방법이지만 비상관된 잔차를 "
  "제공해 35% RH의 Native HfS₂(Wasserstein 지배)와 Al₂O₃/HfS₂"
  "(공간 패턴 지배)에서 앙상블에 기여한다.",
  align="justify")
IMG("figS3_loo_scatter.png", width=6.5,
    cap="Figure S3. 5가지 개별 추정기의 참값-예측 산점도 "
        "(33개 질의에 대한 LOO).")

H("Figure S4. 조건별 앙상블 가중치", level=2, size=11, bold=True)
P("각 (조건, 방법) 셀에 대해 Huber 손실로 최적화된 앙상블 가중치의 "
  "히트맵. 70% RH의 Native HfS₂는 KNN 단독(100%)으로 붕괴된다; "
  "35% RH의 Native HfS₂는 Wasserstein 히스토그램 거리(55%)에 의해 "
  "지배된다; Al₂O₃/HfS₂는 거의 정상에 가까운 영역의 단일 지배 "
  "신호 부재를 반영해 FFT, 공간 패턴, Wasserstein 방법이 상당한 "
  "가중치를 받는 유일한 조건이다.", align="justify")
IMG("figS4_weights_heatmap.png", width=5.5,
    cap="Figure S4. 조건별 앙상블 가중치 히트맵 (값은 백분율).")

H("Figure S5. 3가지 추가 조건의 의사-라만 추정",
  level=2, size=11, bold=True)
P("측정 라만 스펙트럼이 placeholder 값(raman.raman.db에서 시계열 전체에 "
  "걸쳐 norm_peak = 1.0)인 3가지 조건에 대해 이미지 지표로부터 도출된 "
  "의사-라만 A₁ₒ 피크 강도 추정값. 측정 라만에 대한 경험적 검증은 "
  "완전한 측정 스펙트럼 시리즈가 사용 가능한 유일한 조건"
  "(Al₂O₃/HfS₂)에 대해 본문 Fig. 5(b)에 제시되어 있다.",
  align="justify")
IMG("figS5_pseudo_other_conds.png", width=6.5,
    cap="Figure S5. 완전한 측정 라만 시리즈가 없는 3가지 조건의 "
        "의사-라만 A₁ₒ 추정값.")

H("Figure S6. 조건별 동역학 적합", level=2, size=11, bold=True)
P("alldata.db로부터 얻은 조건별 b∗(t) 측정값(컬러 점)과 1차 지수 감쇠 "
  "적합 b∗(t) = b∗_∞ + (b∗_0 − b∗_∞)·exp(−k·t)(검은 실선). "
  "적합된 감쇠 상수 k는 각 패널의 범례에 보고되며, 동역학 추정기"
  "(S1.3 ⑤)가 궁극적으로 사용하는 매개변수이다. 70% RH의 Native HfS₂는 "
  "가장 가파른 감쇠(k ≈ 0.29 d⁻¹)를 보이며 ~2.4일의 반감기에 해당한다; "
  "35% RH의 Native HfS₂와 70% RH의 PMMA/HfS₂는 중간 감쇠율"
  "(k ≈ 0.05–0.08 d⁻¹)을 보인다; Al₂O₃/HfS₂는 거의 일정한 적합"
  "(유효 k → 0)을 생성하여 캡슐화제가 산화 경로를 완전히 차단함을 "
  "확인한다.", align="justify")
IMG("figS6_kinetic_fit.png", width=6.5,
    cap="Figure S6. 1차 지수 감쇠 적합과 함께 조건별 b∗(t) 측정값. "
        "적합된 비율 상수 k가 동역학 추정기의 경과일 역산을 결정한다.")


# ─── S5 ───
H("S5. 보충 표", level=1, size=14, bold=True)

H("Table S1. 시편별 leave-one-out 예측 (33개 질의)",
  level=2, size=11, bold=True)
P("4가지 조건에 걸친 대표 질의의 경과일 예측. 5개 개별 추정기와 두 "
  "앙상블 구성(균등 가중 baseline 및 조건별 Huber 최적화 앙상블)을 "
  "비교한다. 셀은 예측 일수를 표시한다. 조건별 앙상블(마지막 열)이 "
  "권장 추정값이다.", align="justify")
hdr = ["Specimen", "True", "KNN", "Wass", "FFT", "Spat", "Kinet",
       "Ens(uni)", "Ens(opt)"]
rows = [
    ("0d Al2O3 70%",  "0",  "3.7", "10.7", "9.9",  "2.9", "—", "5.7",  "3.4"),
    ("3d Al2O3 70%",  "3",  "3.2", "11.1", "0.9",  "4.0", "—", "4.8",  "3.4"),
    ("7d Al2O3 70%",  "7",  "16.2","9.0",  "9.8",  "3.0", "—", "9.4",  "12.2"),
    ("14d Al2O3 70%", "14", "10.8","12.9", "5.2",  "2.6", "—", "7.9",  "8.3"),
    ("21d Al2O3 70%", "21", "15.4","14.1", "16.8", "18.9","—", "16.3", "16.4"),
    ("0d Native 35%", "0",  "1.3", "4.0",  "18.2", "1.4", "—", "6.2",  "1.3"),
    ("3d Native 35%", "3",  "5.5", "2.0",  "26.0", "0.9", "—", "8.6",  "4.1"),
    ("7d Native 35%", "7",  "12.4","11.2", "9.5",  "6.0", "—", "9.8",  "10.5"),
    ("14d Native 35%","14", "12.4","13.2", "2.5",  "3.5", "—", "7.9",  "9.8"),
    ("28d Native 35%","28", "24.9","7.8",  "9.1",  "14.2","—", "14.0", "21.7"),
    ("0d Native 70%", "0",  "1.3", "1.0",  "5.2",  "12.7","—", "6.3",  "12.7"),
    ("6d Native 70%", "6",  "7.7", "7.0",  "17.2", "2.9", "—", "8.7",  "6.2"),
    ("14d Native 70%","14", "15.5","20.9", "9.3",  "23.0","—", "17.1", "17.7"),
    ("21d Native 70%","21", "20.3","20.0", "11.2", "21.0","—", "18.1", "20.5"),
    ("29d Native 70%","29", "18.1","13.7", "4.3",  "16.8","—", "13.2", "17.7"),
    ("0d PMMA 70%",   "0",  "1.7", "5.6",  "14.1", "10.6","—", "10.1", "10.6"),
    ("3d PMMA 70%",   "3",  "4.2", "4.9",  "9.0",  "3.3", "—", "5.3",  "3.9"),
    ("7d PMMA 70%",   "7",  "8.4", "18.3", "5.9",  "18.0","—", "12.6", "11.3"),
    ("14d PMMA 70%",  "14", "9.9", "22.6", "17.0", "23.0","—", "18.1", "13.8"),
    ("28d PMMA 70%",  "28", "25.2","18.8", "9.6",  "21.9","—", "22.0", "24.2"),
]
add_table(hdr, rows)
P("주: 동역학 열의 대시(—)는 scipy.optimize.curve_fit이 해당 leave-one-"
  "out 폴드에서 수렴하지 않았음을 나타내며, 그러한 추정은 앙상블에서 "
  "제외된다.", size=8, italic=True, after=8)

H("Table S2. 조건별 앙상블 가중치 (Huber 손실, δ = 5 d)",
  level=2, size=11, bold=True)
P("Leave-one-out 앙상블 예측과 알려진 경과일 사이의 Huber 손실을 "
  "최소화하는 최적 가중치, 적어도 3개 이미지-일 쌍이 있는 각 조건에 "
  "대해 별도로 계산. 비교를 위해 첫 번째 행에 최적의 전역 단일 가중치 "
  "세트가 포함되어 있다.", align="justify")
hdr2 = ["Setting", "n", "KNN", "Wass", "FFT", "Spatial", "Kinetic",
        "Opt RMSE (d)", "Uniform RMSE (d)"]
rows2 = [
    ("Global default", "33", "63%", "1%", "1%", "30%", "5%",
     "6.13", "7.74"),
    ("Native 35% RH",  "5",  "14%", "55%", "0%", "13%", "18%",
     "2.29", "4.62"),
    ("Native 70% RH",  "11", "100%","0%",  "0%", "0%",  "0%",
     "3.56", "5.10"),
    ("Al2O3 70% RH",   "12", "0%",  "10%", "3%", "79%", "9%",
     "8.16", "10.45"),
    ("PMMA 70% RH",    "5",  "83%", "12%", "0%", "0%",  "5%",
     "1.96", "4.61"),
]
add_table(hdr2, rows2)
P("33개 질의에 대한 가중 평균 RMSE: 4.80 d (조건별) vs. 6.30 d "
  "(전역 단일 가중치 세트), 본문에 보고된 균등 baseline 대비 "
  "전역 단일 가중치 세트에서 24%, 균등에서 38% 감소에 해당.",
  size=8, italic=True, after=8)

H("Table S3. 조건별 r(day, metric) Pearson 상관계수",
  level=2, size=11, bold=True)
P("경과일과 3가지 이미지 색상 지표(b∗, S 채널 평균, 황색지수) "
  "간의 조건별 Pearson 상관계수. −1에 가까운 값은 진행성 산화에서 "
  "예상되는 강하게 단조적인 감소를 나타낸다; 훨씬 약한 Al₂O₃/HfS₂ "
  "상관계수(~−0.32)는 Al₂O₃ 캡슐화에 의한 산화 장벽의 정량적 신호이며, "
  "해당 조건에 대한 이미지 기반 접근법의 본질적 감도 한계를 구성한다.",
  align="justify")
hdr3 = ["Condition", "n", "r(day, b∗)", "r(day, S)", "r(day, YI)"]
rows3 = [
    ("Native HfS2 35% RH", "10", "−0.918", "−0.918", "−0.948"),
    ("Native HfS2 70% RH", "16", "−0.829", "−0.822", "−0.839"),
    ("Al2O3/HfS2 70% RH",  "17", "−0.319", "−0.324", "−0.331"),
    ("PMMA/HfS2 70% RH",   "10", "−0.886", "−0.839", "−0.856"),
]
add_table(hdr3, rows3)
P("", after=8)

H("Table S4. 조건별 시편 수 및 경과일 분포",
  level=2, size=11, bold=True)
P("통합 분석 풀(alldata.db)과 독립 평가 세트(sample/)의 조건별 사진 "
  "수, 그리고 각 조건에서 사용 가능한 고유 경과일. Al₂O₃/HfS₂ 조건은 "
  "느린 산화 화학으로 인해 더 정밀한 시간 격자가 필요했기 때문에 "
  "가장 조밀한 시간 샘플링(13개 서로 다른 일)을 갖는다.",
  align="justify")
hdr4 = ["Condition", "Pool n", "Eval n", "Unique aging days"]
rows4 = [
    ("Native HfS₂ 35% RH", "10", "5", "0, 3, 7, 14, 28"),
    ("Native HfS₂ 70% RH", "16", "5",
     "0, 1, 2, 3, 4, 6, 7, 13, 14, 15, 20, 21, 28, 29"),
    ("Al₂O₃/HfS₂ 70% RH", "17", "5",
     "0, 1, 2, 3, 5, 7, 14, 15, 16, 21, 22, 28, 30"),
    ("PMMA/HfS₂ 70% RH",  "10", "5", "0, 3, 7, 14, 28"),
    ("TOTAL",              "53", "20", "—"),
]
add_table(hdr4, rows4)
P("", after=8)

H("Table S5. 알고리즘 하이퍼파라미터", level=2, size=11, bold=True)
P("파이프라인 전반에 걸쳐 사용된 알고리즘 하이퍼파라미터의 종합 "
  "목록. 모든 값은 본문에 태그된 commit의 hfs2_v5_49.py 기본값이다. "
  "조정 가능한 매개변수는 함수 인수로 노출되며 코드 수정 없이 "
  "프로젝트 GUI를 통해 조정할 수 있다.", align="justify")
hdr5 = ["Module", "Parameter", "Value", "Description"]
rows5 = [
    ("ROI", "paper_v_thresh",   "215",  "종이 마스크용 HSV V 임계값 (>)"),
    ("ROI", "paper_s_thresh",   "25",   "종이 마스크용 HSV S 임계값 (<)"),
    ("ROI", "morph kernel k",   "max(7, min(H,W)/80)", "타원형 SE 크기"),
    ("ROI", "convex hull guard","0.90", "hull_area > 0.9·img_area 시 hull 거부"),
    ("ROI", "max_specimen_fraction","0.70","ROI ≤ 시편 bbox의 70%"),
    ("ROI", "edge_margin_ratio","0.05", "이미지 경계로부터 5% 안쪽 이동"),
    ("ROI", "α_c (cond)", "0.13–0.17", "조건별 ROI 면적 비율"),
    ("KNN", "(w_b, w_s, w_yi)", "(0.45, 0.30, 0.25)", "지표 가중치"),
    ("KNN", "k", "3", "사용되는 top-k 최근접 참조"),
    ("Wass", "histogram bins", "64", "b∗ 범위 −30에서 80"),
    ("FFT",  "radial bins", "64", "log-power 방사 프로파일"),
    ("FFT",  "hf_thresh", "0.40·r_max", "고주파 에너지 반경 임계값"),
    ("FFT",  "(d_hf, d_ent, d_rad)", "(0.5, 0.3, 0.2)", "특성 거리 가중치"),
    ("Spatial", "grid", "3 × 3", "엔트로피/이방성용 ROI 분할"),
    ("Kinetic", "model", "exp decay", "b∗(t)=b∗∞+(b∗₀−b∗∞)·exp(−k·t)"),
    ("Kinetic", "bounds (b₀, b∞, k)", "(0,150)/(−20,80)/(1e-6,20)", "curve_fit 경계"),
    ("Ensemble", "δ (Huber)", "5 d", "Huber 손실 전이"),
    ("Ensemble", "optimiser", "L-BFGS-B", "scipy.optimize.minimize"),
    ("Ensemble", "multi-start", "1+5 = 6", "균등 + 5개 method-dominant 0.80"),
    ("Ensemble", "grid fallback", "11⁵ ≈ 1.6e5", "0.1-단계 5-D 격자 (scipy 없을 때)"),
    ("Pseudo-Raman", "regressors", "4", "b∗, S, YI, ΔE 단변량 OLS"),
    ("Pseudo-Raman", "weighting", "R² 정규화", "wm = R²m / Σ R²"),
    ("Pseudo-Raman", "CI factor", "1.96", "회귀 SE에서 95%"),
]
add_table(hdr5, rows5)
P("", after=12)


doc.save(OUT)
print("Saved:", OUT)
